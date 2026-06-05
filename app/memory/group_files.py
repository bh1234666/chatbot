"""
QQ 群文件同步：
  1. 快速索引：拿到 NapCat 文件列表 → 立即创建 KB 节点（download_status="pending"）
  2. 后台下载：异步下载 + 内容提取 → 更新 KB 节点（download_status="done"）
  3. 按需提取：fetch_group_file() 将已下载文件拷贝/链接到当前工作区

智能体看到「群组文件」区域中带「下载中」标记的条目时，
可按人设自然告知用户"还在下载，稍等一下"。
"""
from __future__ import annotations

import json
import logging
import os
import shutil
import time
from dataclasses import dataclass, field
from datetime import datetime, timezone, timedelta
from pathlib import Path

import httpx
import ulid

from app.config import settings
from app.db.pool import pool
from app.core import debug
from app.core.sanitize import sanitize_headline, sanitize_summary
from app.core.bg_tasks import schedule
from app.llm.tools.workspace import (
    _get_workspace_root,
    _BINARY_DETECT_BYTES,
    _FILE_TYPE_TABLE,
)

log = logging.getLogger(__name__)

NAPCAT_URL = settings.napcat_url
MAX_DOWNLOAD_SIZE = 50 * 1024 * 1024  # 50 MB
MAX_CONTENT_EXTRACT_SIZE = 100 * 1024  # 100 KB
MAX_EXTRACT_CHARS = 8000
SYNC_COOLDOWN_SEC = 120  # 两次同步之间的最小间隔
MAX_BG_DOWNLOADS = 3       # 每次 sync 最多启动几个后台下载

_TEXT_CATEGORIES = {"text", "text-structured"}

# Patch 12 (2026-05-02): 改用 (archive_id, group_id) tuple key,
# 多人格群下每个 archive 独立 cooldown,不互相饿死。
_last_sync_time: dict[tuple[str, str], float] = {}
_GROUP_FILE_SYNC_ONLINE_SINCE = int(time.time())
_GROUP_FILE_SYNC_ONLINE_GRACE_SEC = 2

# (archive_id, file_id) 当前正在 _bg_download_and_index 中的文件，防止 heal 重复排队
_in_flight_downloads: dict[tuple[str, str], float] = {}
_IN_FLIGHT_TIMEOUT = 300.0  # 5 分钟无进展即认为已死，允许重试

# 2026-05-15 P85: URL 拿不到的物理文件签名记忆 — 避免重复尝试
# 病因(实测 22:59-23:03 trace): NapCat file_id 不稳定, 每次 poll 同一物理文件 file_id 变,
#   existing_ids 命中不到 → 进 Phase 1 创建新 cold_node → bg download → URL 拿不到 →
#   permanent=True 删 cold_node (但 synced_files 保留) → 下次 sync 重复同样过程。
#   实测 30 警告 / 20 文件 / 4 分钟内, 同一 cpp_join_bench.cpp 反复尝试 4 次。
# 修法: 记 (archive_id, group_id, file_name, file_size) → last_failure_time。
#   sync_group_files Phase 1 跳过近期失败过的物理签名; _download_file 只写后台
#   debug 记录,避免 NapCat 不返回下载 URL 时刷控制台。
_recent_url_failures: dict[tuple[str, str, int], float] = {}
_URL_FAILURE_COOLDOWN = 600.0  # 10 分钟内不再重试同物理签名


@dataclass
class GroupFileItem:
    file_id: str
    file_name: str
    file_size: int
    upload_time: int
    uploader_uin: int = 0
    uploader_name: str = ""
    busid: int = 0
    raw: dict = field(default_factory=dict)


# ── NapCat API 调用 ────────────────────────────────────────────

class NapCatUnavailable(Exception):
    """NapCat HTTP 调用本身失败(网络/超时/非 200);区别于"调用成功但返回空"。

    用于 sync_group_files 区分两种 case:
      - 调用成功 → items=[] (群里真没文件): 设 cooldown,避免每条消息都打
      - 调用失败 → 抛 NapCatUnavailable: 不设 cooldown,下条消息会重试,
        但 fire-and-forget 调用会吞下异常不影响主流程。
    """


async def fetch_group_files(group_id: str) -> list[GroupFileItem]:
    """从 NapCat 拉取群文件列表（含根目录 + 子文件夹递归）。

    成功返回 list(可为空)。NapCat 不可达时抛 NapCatUnavailable,调用方根据
    异常类型决定是否设置 cooldown。
    """
    items: list[GroupFileItem] = []

    async def _collect(folder_id: str | None = None) -> None:
        if folder_id is not None:
            r = await _napcat("get_group_file_list", {
                "group_id": int(group_id),
                "folder_id": folder_id,
            })
        else:
            r = await _napcat("get_group_root_files", {
                "group_id": int(group_id),
            })
        # _napcat 失败时已 raise NapCatUnavailable,这里 r 必非 None
        # NapCat 在文件夹空、被删除或权限不足时会返回 {"data": null},
        # dict.get 的默认值只在 key 缺失时生效——key 存在但 value=None 还是会拿到 None。
        # 用 `or {}` 兜住这个 case,然后下面 .get("files", []) 才安全。
        data = r.get("data") or {}
        if not data:
            log.debug("napcat returned empty data for folder=%s group=%s", folder_id, group_id)
            return
        for f in data.get("files", []) or []:
            try:
                upload_time = int(f.get("upload_time", 0))
                items.append(GroupFileItem(
                    file_id=str(f["file_id"]),
                    file_name=str(f.get("file_name", "")),
                    file_size=int(f.get("file_size", 0)),
                    upload_time=upload_time,
                    uploader_uin=int(f.get("uploader_uin", 0)),
                    uploader_name=str(f.get("uploader_name", "")),
                    busid=int(f.get("busid", 0)),
                    raw=f,
                ))
            except (KeyError, ValueError, TypeError):
                continue
        for folder in data.get("folders", []) or []:
            fid = str(folder.get("folder_id", ""))
            if fid:
                await _collect(fid)

    await _collect()   # NapCatUnavailable 异常透传给调用方
    return items


async def _napcat(action: str, params: dict) -> dict:
    """调用 NapCat,失败抛 NapCatUnavailable。"""
    try:
        async with httpx.AsyncClient(timeout=15.0) as client:
            r = await client.post(f"{NAPCAT_URL}/{action}", json=params)
            if r.status_code == 200:
                return r.json()
            raise NapCatUnavailable(f"napcat {action}: HTTP {r.status_code}")
    except NapCatUnavailable:
        raise
    except Exception as e:
        raise NapCatUnavailable(f"napcat {action}: {type(e).__name__}: {e}") from e


# ── 同步主逻辑（两阶段）────────────────────────────────────────

async def sync_group_files(archive_id: str, group_id: str) -> int:
    """同步群新文件。

    阶段 0（恢复）：扫描所有 download_status="pending" 的 cold_nodes，
      对其中文件已在磁盘的直接修复为 "done"，没在磁盘的重新调度后台下载。
      这一步不触发 NapCat 调用，所以不受 cooldown 约束。
    阶段 1（同步、毫秒级）：遍历 NapCat 文件列表，对每个新文件立即创建
      KB 节点（download_status="pending"）并写入 synced_files。
    阶段 2（异步、后台）：下载文件 → 提取内容 → 用 lite 模型更新 KB 节点
      摘要（download_status="done"）。
    """
    # 阶段 0：恢复卡住的 pending 节点（不受 cooldown 约束）
    healed = await _heal_pending_nodes(archive_id, group_id)

    now = time.time()
    # Patch 12: 用 (archive_id, group_id) tuple,多人格群独立 cooldown
    cooldown_key = (archive_id, group_id)
    last = _last_sync_time.get(cooldown_key, 0)
    if now - last < SYNC_COOLDOWN_SEC:
        return healed

    # 2026-05-01 修: 区分"调用成功但群里没文件"和"NapCat 不可达"。
    # 旧版两种情况都返回 [] 然后 if not items: return,导致空文件群每条
    # 消息都重新打 NapCat 一次,SYNC_COOLDOWN_SEC 形同虚设。
    # 新版: 成功(无论空/非空)都设 cooldown;NapCat 不可达不设 cooldown,
    # 让下次消息能尝试重连,但当前调用直接返回 healed 不影响主流程。
    try:
        items = await fetch_group_files(group_id)
    except NapCatUnavailable as e:
        log.debug("fetch_group_files unavailable, skip this round: %s", e)
        return healed

    # 调用成功 → 设 cooldown(无论 items 是否为空)
    _last_sync_time[cooldown_key] = now
    if not items:
        return healed

    online_cutoff = _GROUP_FILE_SYNC_ONLINE_SINCE - _GROUP_FILE_SYNC_ONLINE_GRACE_SEC
    recent_items = [it for it in items if (it.upload_time or 0) >= online_cutoff]
    skipped_before_online = len(items) - len(recent_items)
    if skipped_before_online:
        debug.log(
            "group_files.sync.skip_before_online",
            f"archive={archive_id} group={group_id} "
            f"skipped {skipped_before_online} files uploaded before process online cutoff "
            f"({online_cutoff}); using upload_time as high-watermark",
        )
    items = recent_items
    if not items:
        return healed

    # 2026-05-15 重设计 (用户明确要求: "不应该去重,给旧的打上时间戳。qq文件本来就能重名"):
    # QQ 群文件**合法允许重名** —— 不同人传同名,或同人前后传两份同名但内容不同。
    # 旧版用 (file_name, file_size) 当唯一指纹 → 第二份合法上传被静默丢弃。
    #
    # **新策略:只用 file_id 去重**。
    # NapCat 给每次上传的 file_id 是 32 字符十六进制(从生产 log 观察像内容哈希,
    # 例 "8064fc99e3d14cfa903f7316d0fcec6c"),对同一物理文件应当稳定。
    # 即使在 file_id-不稳定的情况下退化也安全:
    #   - 下面 INSERT 用 ON CONFLICT (PK) DO NOTHING, 同 file_id 重复入会被静默忽略
    #   - 撞同名但不同 file_id → 走重命名分支(下方),把老的 file_name 加时间戳
    # 旧的 (name, size) 早期签名彻底不再使用 (idx_synced_files_dedup 索引已在
    # main._ensure_post_baseline_schema 里 DROP 掉)。
    # 配套:`fetch_group_file` 等下游读取已改用 kb_node_id 稳定 UUID 而非 file_name 做 JOIN。
    # 2026-05-15 P71: 二次去重防 file_id 不稳定导致的 1000+ 次 rename 风暴。
    # 病因(实测 trace 16:25-19:17 累计 1350 次 synced_files.rename):
    #   `cpp_join_bench.cpp` 1 个文件被 rename 75 次, 5 个文件各 50 次, 36 个文件各 25 次。
    #   原因: NapCat 给同一物理文件每次 poll 返回不同的 file_id (注释说"应当稳定"但实际不稳定)。
    #   主 file_id 去重失效 → 每次 poll 整批 N 个文件全部"新"→ N 次 rename 风暴。
    # 修法: existing_ids 失败时, 加二次签名 (file_name, file_size, upload_time) 匹配。
    #   匹配上视为同一物理文件, 跳过 INSERT / rename。
    #   NapCat file_id 真稳定时此分支没开销; 不稳定时大量节省 DB 写。
    existing_ids: set[str] = set()
    existing_sigs: set[tuple[str, int, int]] = set()  # (file_name, file_size, upload_time)
    async with pool().acquire() as conn:
        rows = await conn.fetch(
            "SELECT file_id, file_name, file_size, upload_time FROM synced_files "
            "WHERE archive_id = $1 AND group_id = $2",
            archive_id, group_id,
        )
        for r in rows:
            existing_ids.add(r["file_id"])
            # 二次签名: 物理内容指纹 (允许重名 - QQ 群文件可重名 - 但 name+size+ts 三元组
            # 极少撞车; 真撞了顶多漏掉一条少见 case, 远比 1000+ 次 rename 风暴划算)
            existing_sigs.add((r["file_name"], r["file_size"] or 0, r["upload_time"] or 0))

    # P71: 双层过滤 — 先看 file_id, 再看 (name,size,ts) 三元组
    # 2026-05-15 P85 加一层: URL 失败 cooldown 内的物理签名(name, size)也跳过 —
    #   防止 NapCat file_id 不稳定 + URL 拿不到的双重失败循环 (实测同物理文件
    #   反复尝试 4 次警告刷屏)。
    new_items = []
    _suppressed_dup_count = 0
    _suppressed_url_fail = 0
    _suppressed_bot_generated = 0
    _now_p85 = time.time()
    # 2026-05-17 Round 14f: 过滤 bot 自己生成的 wav 文件
    # 实测 (trace 08:28): bot 发 tts_xxx_yyy.wav 到群, NapCat 返时当群文件,
    # sync_group_files 不知道是 bot 自己发的, 添 KB placeholder (永久污染).
    # 模式: tts_<trace 8 字>_<13 位 ts>.wav 是 _handle_tts (Round 11) 输出;
    #       _voice_<trace 8 字>_<13 位 ts>.wav 是自动 voice 流程输出.
    import re as _re_bot
    _BOT_FILE_PATTERN = _re_bot.compile(
        r'^(tts|_voice|_tts)_[a-f0-9]{6,16}_\d{10,16}(?:_\d+)?\.(wav|mp3|m4a)$',
        _re_bot.IGNORECASE,
    )
    for it in items:
        if _BOT_FILE_PATTERN.match(it.file_name or ""):
            _suppressed_bot_generated += 1
            continue  # bot 自己生成的 → 不入 KB
        if it.file_id in existing_ids:
            continue  # file_id 主键匹配
        _sig = (it.file_name, it.file_size or 0, it.upload_time or 0)
        if _sig in existing_sigs:
            _suppressed_dup_count += 1
            continue  # 三元组匹配, 视为同一文件 (file_id 是 NapCat 重新生成的)
        # P85: 看 URL 失败 memo (按 group_id+name+size 物理签名, 忽略 upload_time)
        _fail_sig = (group_id, it.file_name, it.file_size or 0)
        _last_fail = _recent_url_failures.get(_fail_sig)
        if _last_fail and (_now_p85 - _last_fail) < _URL_FAILURE_COOLDOWN:
            _suppressed_url_fail += 1
            continue  # 近期 URL 拿不到, NapCat 还在不稳定 → 不浪费再创节点
        new_items.append(it)
    if _suppressed_bot_generated:
        debug.log(
            "group_files.sync.skip_bot_generated",
            f"archive={archive_id} group={group_id} "
            f"suppressed {_suppressed_bot_generated} bot-generated wav/mp3 files "
            f"(pattern: tts_/voice_<trace>_<ts>) — not indexing as KB",
        )
    if _suppressed_dup_count:
        debug.log(
            "group_files.sync.dedup_secondary",
            f"archive={archive_id} group={group_id} "
            f"suppressed {_suppressed_dup_count} dup uploads via (name,size,ts) signature "
            f"(NapCat file_id appears unstable for this group)",
        )
    if _suppressed_url_fail:
        debug.log(
            "group_files.sync.skip_failed_url",
            f"P85: archive={archive_id} group={group_id} "
            f"suppressed {_suppressed_url_fail} files whose URL fetch failed recently "
            f"(within {_URL_FAILURE_COOLDOWN}s cooldown) — NapCat unstable for these",
        )
    if not new_items:
        return 0

    # 确保工作区 group_files 目录存在
    ws_dir = _get_workspace_root() / archive_id / group_id / "group_files"
    ws_dir.mkdir(parents=True, exist_ok=True)

    synced = 0
    bg_downloads = 0

    for item in new_items:
        # 阶段 1：快速创建 KB 节点 + synced_files 记录
        kb_node_id = await _create_pending_kb_node(
            archive_id=archive_id,
            group_id=group_id,
            item=item,
        )

        workspace_rel = ""
        if kb_node_id:
            ts = item.upload_time if item.upload_time > 0 else int(time.time())
            ts_str = datetime.fromtimestamp(ts, tz=timezone(timedelta(hours=8))).strftime(
                "%Y%m%d_%H%M%S"
            )
            safe_name = _safe_filename(item.file_name)
            workspace_rel = f"group_files/{ts_str}_{safe_name}"

        async with pool().acquire() as conn:
            # 2026-05-15 重设计:撞名 → 把老的 file_name 加时间戳保留访问 (而不是
            # 静默丢弃新上传)。
            # 步骤:
            #   1. SELECT 同 (archive, group, file_name) 但不同 file_id 的旧行
            #   2. 对每一行 UPDATE file_name = "{stem}_{old_upload_ts}{ext}"
            #      (workspace_path 不动 — 磁盘上的文件名已含时间戳前缀,不需要移动)
            #   3. INSERT 新行,使用原 file_name
            # 路由保护:fetch_group_file 的 JOIN 必须改成用 kb_node_id (稳定 UUID)
            # 而不是 (file_name, file_size) — 老的 KB 节点的 file_metadata 里
            # 存的还是原 filename,改名后旧 KB 节点 → synced_files JOIN 才不会断。
            # 那条 JOIN 已经在 fetch_group_file 里同步改好。
            async with conn.transaction():
                colliding = await conn.fetch(
                    """
                    SELECT file_id, upload_time FROM synced_files
                    WHERE archive_id = $1 AND group_id = $2
                      AND file_name = $3 AND file_id != $4
                    """,
                    archive_id, group_id, item.file_name, item.file_id,
                )
                tz = timezone(timedelta(hours=8))
                for old in colliding:
                    old_ts_val = old["upload_time"] or int(time.time())
                    old_ts_str = datetime.fromtimestamp(old_ts_val, tz=tz).strftime(
                        "%Y%m%d_%H%M%S"
                    )
                    stem, ext = os.path.splitext(item.file_name)
                    new_old_name = f"{stem}_{old_ts_str}{ext}"
                    await conn.execute(
                        """
                        UPDATE synced_files SET file_name = $1
                        WHERE archive_id = $2 AND group_id = $3 AND file_id = $4
                        """,
                        new_old_name, archive_id, group_id, old["file_id"],
                    )
                    # 2026-05-15: 用 debug.log(file-only structured trace)而不是 log.info
                    # 控制台不再被这条刷屏(per-poll 一次能触发几十条同名重传),
                    # 但事件全文照常进 debug 日志文件,需要审计 grep 即可。
                    debug.log(
                        "synced_files.rename",
                        f"{item.file_name!r} → {new_old_name!r} "
                        f"(archive={archive_id} group={group_id} file_id={old['file_id']}) "
                        f"due to new upload with same name",
                    )

                # 现在可以安全插入新行:同 file_name 的旧行已改名,
                # 同 file_id 的话 (再次轮询同一物理文件) PK 冲突 → DO NOTHING。
                await conn.execute(
                    """
                    INSERT INTO synced_files
                        (archive_id, group_id, file_id, file_name, file_size,
                         upload_time, uploader_uin, uploader_name, busid,
                         workspace_path, kb_node_id)
                    VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, $10, $11)
                    ON CONFLICT (archive_id, group_id, file_id) DO NOTHING
                    """,
                    archive_id, group_id, item.file_id, item.file_name, item.file_size,
                    item.upload_time, item.uploader_uin, item.uploader_name, item.busid,
                    workspace_rel, kb_node_id,
                )
        synced += 1

        # 阶段 2：后台下载（限制并发数）
        if kb_node_id and workspace_rel and bg_downloads < MAX_BG_DOWNLOADS:
            # 预注册 in-flight：在 task 真正进入函数体之前关掉 heal 的 race window，
            # 否则 create_task 后到 task 内部首行注册之间，下一轮 heal 可能重复排队。
            _in_flight_downloads[(archive_id, item.file_id)] = time.time()
            schedule(
                _bg_download_and_index(
                    archive_id=archive_id,
                    group_id=group_id,
                    kb_node_id=kb_node_id,
                    item=item,
                    workspace_rel=workspace_rel,
                    ws_dir=str(ws_dir),
                ),
                name=f"group_file_download:{archive_id}:{group_id}:{item.file_id}",
            )
            bg_downloads += 1

    if synced:
        debug.log(
            "group_files.sync.phase1",
            f"archive={archive_id} group={group_id} new={synced} "
            f"total={len(items)} bg_downloads={bg_downloads}",
        )
    return synced


# ── 阶段 1：快速索引（无 LLM，无下载）─────────────────────────

async def _create_pending_kb_node(
    *,
    archive_id: str,
    group_id: str,
    item: GroupFileItem,
) -> str | None:
    """用模板立即创建 KB 节点（不调 LLM），标记 download_status="pending"。"""
    tz = timezone(timedelta(hours=8))
    ts_val = item.upload_time if item.upload_time > 0 else int(time.time())
    ts = datetime.fromtimestamp(ts_val, tz=tz).strftime("%Y-%m-%d %H:%M")
    size_s = _fmt_size(item.file_size)

    headline = f"{item.uploader_name} uploaded {item.file_name}"
    content = (
        f"{item.uploader_name} uploaded `{item.file_name}` at {ts} ({size_s}). "
        "The file is still being downloaded and indexed in the background; content summary will appear after indexing finishes. "
        "Treat this node as pending and wait for a ready workspace path before fetching it.\n\n"
        "文件仍在后台下载和索引；未 ready 前只按 pending 状态处理。"
    )

    cid = f"c_{ulid.ULID()}"
    file_meta = json.dumps({
        "filename": item.file_name,
        "workspace_path": "",
        "archive_id": archive_id,
        "group_id": group_id,
        "upload_time": item.upload_time,
        "uploader_name": item.uploader_name,
        "file_size": item.file_size,
        "napcat_file_id": item.file_id,
        "busid": item.busid,
        "download_status": "pending",
    }, ensure_ascii=False)

    try:
        async with pool().acquire() as conn:
            await conn.execute(
                """
                INSERT INTO cold_nodes
                    (id, archive_id, group_id, user_id, scope,
                     node_type, headline, content,
                     salience, source_refs, file_metadata)
                VALUES ($1, $2, $3, NULL, 'kb', 'file', $4, $5, $6, $7::jsonb, $8)
                """,
                cid, archive_id, group_id,
                sanitize_headline(headline),
                sanitize_summary(content),
                0.5,
                json.dumps([]),
                file_meta,
            )
        log.debug("pending kb node: %s file=%s", cid, item.file_name)
        # 2026-05-16 Dream: emit 文件上传事件 (触发 D15-D18 可寻性增强)
        try:
            from app.core.dream import event_bus
            import os
            ext = os.path.splitext(item.file_name)[1].lower()
            await event_bus.emit("file_uploaded",
                                  archive_id=archive_id, group_id=group_id,
                                  file_id=cid, file_name=item.file_name, ext=ext)
        except Exception:
            pass
        return cid
    except Exception:
        log.exception("create_pending_kb_node failed for %s", item.file_name)
        return None


# ── 阶段 0：恢复卡住的 pending 节点 ─────────────────────────────

def _to_workspace_rel(local_path: str, archive_id: str, group_id: str) -> str:
    """把绝对的 local_path 转成 synced_files.workspace_path 应该存的相对形式
    （相对于 ws_root/archive_id/group_id/）。失败返回空字符串。"""
    try:
        ws_root = _get_workspace_root()
        base = (ws_root / archive_id / group_id).resolve()
        rel = os.path.relpath(Path(local_path).resolve(), base)
        # 跨平台统一用正斜杠（DB 里之前存的就是正斜杠形式）
        return rel.replace("\\", "/")
    except (ValueError, OSError):
        return ""


def _find_existing_local_file(
    archive_id: str, group_id: str, workspace_rel: str, file_name: str
) -> str | None:
    """在工作区里查找文件是否已存在。返回绝对路径，找不到返回 None。

    依次尝试：
      1. 精确匹配 workspace_rel
      2. group_files/ 目录下任何以 _<file_name> 结尾的文件（带时间戳前缀）
      3. group_files/ 目录下以 _<file_name>_1 等重名后缀结尾的文件
    """
    ws_root = _get_workspace_root()
    base = ws_root / archive_id / group_id

    # 1. 精确匹配
    if workspace_rel:
        p = base / workspace_rel
        if p.exists() and p.is_file():
            return str(p)

    # 2 / 3. 同目录下扫描重名变体
    if not file_name:
        return None
    safe_name = _safe_filename(file_name)
    if not safe_name:
        return None
    gf_dir = base / "group_files"
    if not gf_dir.exists():
        return None
    stem, ext = os.path.splitext(safe_name)
    try:
        candidates = list(gf_dir.iterdir())
    except OSError:
        return None
    # 匹配 *_<safe_name> 或 *_<stem>_<n><ext>
    for p in candidates:
        if not p.is_file():
            continue
        nm = p.name
        if nm.endswith("_" + safe_name):
            return str(p)
        # 匹配 *_<stem>_<digit><ext>（_download_file 的重名后缀）
        if ext and nm.endswith(ext):
            mid = nm[:-len(ext)] if ext else nm
            if "_" + stem + "_" in mid:
                return str(p)
    return None


async def _heal_pending_nodes(archive_id: str, group_id: str) -> int:
    """扫描所有 pending 状态的 file 节点，能修就修，修不了就重排队。

    场景：
      - 服务重启/迁移导致 Phase 2 中断 → 文件在磁盘但 cold_nodes 没更新
      - MAX_BG_DOWNLOADS 溢出导致 Phase 2 从未启动 → 文件不在磁盘需要重下
      - workspace_rel 路径预测错位 → 文件在磁盘但 DB 路径对不上

    返回修复（done）+ 重排队（rescheduled）的总数。
    """
    async with pool().acquire() as conn:
        rows = await conn.fetch(
            """
            SELECT cn.id AS node_id, cn.file_metadata,
                   sf.file_id, sf.file_name, sf.file_size, sf.upload_time,
                   sf.uploader_uin, sf.uploader_name, sf.busid,
                   sf.workspace_path
            FROM cold_nodes cn
            LEFT JOIN synced_files sf
              ON sf.archive_id = cn.archive_id
             AND sf.group_id = cn.group_id
             AND sf.kb_node_id = cn.id
            WHERE cn.archive_id = $1 AND cn.group_id = $2
              AND cn.scope = 'kb' AND cn.node_type = 'file'
            """,
            archive_id, group_id,
        )

    pending_rows: list[dict] = []
    for r in rows:
        meta = {}
        if r["file_metadata"]:
            try:
                raw = r["file_metadata"]
                meta = json.loads(raw) if isinstance(raw, str) else (raw or {})
            except (json.JSONDecodeError, TypeError):
                meta = {}
        # 2026-05-16: 跳过已被标记 deleted 的节点 (D20 / _mark_download_failed
        # permanent=True 都会标记). 这些文件确认不存在, 不再反复尝试下载,
        # 避免重启后控制台刷一批 "no download url" 警告。
        if meta.get("deleted"):
            continue
        if meta.get("download_status") == "pending":
            pending_rows.append({"row": r, "meta": meta})

    if not pending_rows:
        return 0

    # 确保工作区目录存在（可能整个 archive 工作区都没建过）
    ws_dir = _get_workspace_root() / archive_id / group_id / "group_files"
    ws_dir.mkdir(parents=True, exist_ok=True)

    healed = 0
    rescheduled = 0
    # bg_slots 要扣掉已经在跑的 in-flight 数量，避免和 sync_group_files 已启动的
    # task 叠加导致并发突破 MAX_BG_DOWNLOADS。
    bg_slots = max(0, MAX_BG_DOWNLOADS - len(_in_flight_downloads))

    for entry in pending_rows:
        r = entry["row"]
        meta = entry["meta"]
        node_id = r["node_id"]
        file_name = (
            r["file_name"] or meta.get("filename", "") or ""
        )
        workspace_rel = r["workspace_path"] or meta.get("workspace_path", "") or ""

        # 1. 优先看磁盘上是不是已经有这个文件
        local = _find_existing_local_file(
            archive_id, group_id, workspace_rel, file_name,
        )
        if local:
            actual_rel = _to_workspace_rel(local, archive_id, group_id) or workspace_rel
            new_meta = dict(meta)
            new_meta["download_status"] = "done"
            new_meta["workspace_path"] = actual_rel
            new_meta.pop("download_error", None)
            try:
                async with pool().acquire() as conn:
                    async with conn.transaction():
                        await conn.execute(
                            "UPDATE cold_nodes SET file_metadata = $1, "
                            "updated_at = NOW() WHERE id = $2 AND archive_id = $3",
                            json.dumps(new_meta, ensure_ascii=False),
                            node_id, archive_id,
                        )
                        if r["file_id"] is not None:
                            await conn.execute(
                                "UPDATE synced_files SET workspace_path = $1 "
                                "WHERE archive_id = $2 AND file_id = $3",
                                actual_rel, archive_id, r["file_id"],
                            )
                healed += 1
                debug.log(
                    "group_files.heal.file",
                    f"node={node_id} file={file_name} path={actual_rel}",
                )
            except Exception:
                log.exception("heal pending failed for node=%s", node_id)
            continue

        # 已经有 bg 任务在跑这个文件 → 跳过，等它完成
        in_flight_key = (archive_id, r["file_id"])
        in_flight_at = _in_flight_downloads.get(in_flight_key)
        if in_flight_at is not None and (time.time() - in_flight_at) < _IN_FLIGHT_TIMEOUT:
            continue
        if not r["file_id"]:
            # 没有 file_id 没法重下,标记失败让模型可以告诉用户
            # 2026-05-09 Patch 43:这是永久不可恢复(NapCat file_id 每次重启都变,
            # 旧 file_id 丢失就是死)→ 直接删除节点,避免长期污染 KB
            await _mark_download_failed(
                node_id, "丢失 NapCat file_id,无法重新下载", permanent=True,
            )
            continue
        # 2026-05-15 P88: heal 也查 P85 URL 失败 memo
        # 病因(实测 23:00-23:06 trace): group 333637049 有 125 pending 节点,
        #   每 3-5s 一次 heal 重排 3 个 bg 下载, 每个都 NapCat URL 拿不到永久失败。
        #   P85 已在 sync_group_files Phase 1 跳过这些 sig, 但 heal 是另一路径,
        #   仍然反复创建 bg 任务浪费 NapCat HTTP + asyncio 调度。
        # 修法: heal 检查物理签名 (group_id, name, size) 是否在 cooldown 内失败过,
        #   是 → 不重排, 直接 _mark_download_failed(permanent=True) 删节点。
        _fail_sig_p88 = (group_id, file_name, int(r["file_size"] or meta.get("file_size", 0) or 0))
        _last_fail_p88 = _recent_url_failures.get(_fail_sig_p88)
        if _last_fail_p88 and (time.time() - _last_fail_p88) < _URL_FAILURE_COOLDOWN:
            # 这文件物理签名近期 URL 拿不到 — 别再创 bg 任务, 直接清节点
            await _mark_download_failed(
                node_id, "P88: URL 失败 cooldown 内, 跳过重排", permanent=True,
            )
            continue
        # 受 MAX_BG_DOWNLOADS 限速：剩下的 pending 留给下次同步处理
        if bg_slots <= 0:
            continue
        item = GroupFileItem(
            file_id=r["file_id"],
            file_name=file_name,
            file_size=int(r["file_size"] or meta.get("file_size", 0) or 0),
            upload_time=int(r["upload_time"] or meta.get("upload_time", 0) or 0),
            uploader_uin=int(r["uploader_uin"] or 0),
            uploader_name=r["uploader_name"] or meta.get("uploader_name", "") or "",
            busid=int(r["busid"] or meta.get("busid", 0) or 0),
        )
        # 预注册 in-flight，关掉 create_task 到 task 入口注册之间的 race window
        _in_flight_downloads[in_flight_key] = time.time()
        schedule(
            _bg_download_and_index(
                archive_id=archive_id,
                group_id=group_id,
                kb_node_id=node_id,
                item=item,
                workspace_rel=workspace_rel or f"group_files/{_safe_filename(file_name)}",
                ws_dir=str(ws_dir),
            ),
            name=f"group_file_download:{archive_id}:{group_id}:{item.file_id}",
        )
        bg_slots -= 1
        rescheduled += 1

    if healed or rescheduled:
        debug.log(
            "group_files.heal.batch",
            f"archive={archive_id} group={group_id} "
            f"healed={healed} rescheduled={rescheduled} (of {len(pending_rows)} pending)",
        )
    return healed + rescheduled


# ── 阶段 2：后台下载 + LLM 更新 ─────────────────────────────────

async def _bg_download_and_index(
    *,
    archive_id: str,
    group_id: str,
    kb_node_id: str,
    item: GroupFileItem,
    workspace_rel: str,
    ws_dir: str,
) -> None:
    """后台任务：下载文件 → 提取内容 → LLM 摘要 → 更新 KB 节点。"""
    in_flight_key = (archive_id, item.file_id)
    _in_flight_downloads[in_flight_key] = time.time()
    try:
        # 1. 下载
        local_path = await _download_file(item, group_id, ws_dir)
        if not local_path:
            # 2026-05-10 Patch 49: "无法获取下载链接" 是永久失败 — NapCat 重启后
            # 旧 file_id 全死, 下次重试也拿不到 URL.
            # 2026-05-16 修订: permanent=True 现在改为加 [已删除] 标记
            # (而非 DELETE), 保留历史索引让模型知道这个文件曾存在但已不可下载.
            await _mark_download_failed(kb_node_id, "无法获取下载链接", permanent=True)
            return

        # 用真实落盘路径覆盖 Phase 1 预测的 workspace_rel：当 upload_time=0 时
        # Phase 1 和 _download_file 各自调 time.time()，时间戳可能差几秒；
        # 加上重名 _1 后缀，预测路径常和真实路径对不上。
        actual_rel = _to_workspace_rel(local_path, archive_id, group_id)
        if actual_rel:
            workspace_rel = actual_rel

        # 更新 workspace_path（用真实路径）
        async with pool().acquire() as conn:
            await conn.execute(
                "UPDATE synced_files SET workspace_path = $1 "
                "WHERE archive_id = $2 AND file_id = $3",
                workspace_rel, archive_id, item.file_id,
            )

        # 2. 内容提取
        snippet = ""
        if _should_extract_content(local_path, item.file_size):
            snippet = _extract_text(local_path)

        # 3. LLM 摘要（system prompt 与 kb.index_group_file 共享同一个常量，
        #    避免两处文案漂移）
        # 2026-05-15 P67: 本地短路 — 非文本/无内容文件不调 LLM。
        # 病因(实测 16:25-19:19 trace): 群文件 heal 后台批处理时, 每个图片/视频/无文本 PDF
        # 都送 LLM 生成几乎一模一样的"未提取到文字内容, 可能是扫描件/图像化PDF/加密文档,
        # 需OCR或专门工具打开才能判断具体内容"摘要 (90+ 次 llm.json.input deepseek-v4-flash
        # 调用); LLM 输出对实际检索价值为 0。
        # 修法: snippet 为空 → 按扩展名分类生成确定性 headline+content 文本, 跳过 LLM。
        from app.llm import client as llm
        from app.memory.kb import _GROUP_FILE_INDEX_SYSTEM
        size_s = _fmt_size(item.file_size)
        tz = timezone(timedelta(hours=8))
        ts_val = item.upload_time if item.upload_time > 0 else int(time.time())
        ts = datetime.fromtimestamp(ts_val, tz=tz).strftime("%Y-%m-%d %H:%M")

        _local_short_circuit = False
        if not snippet:
            # 决定是否本地短路: 看扩展名分类
            _ext = os.path.splitext(item.file_name)[1].lower()
            _cat = None
            if _ext in _FILE_TYPE_TABLE:
                _cat, _, _, _ = _FILE_TYPE_TABLE[_ext]
            # 这些类别 100% 没文本内容, 无需 LLM:
            _NO_TEXT_CATS = {"image", "media", "archive", "binary"}
            if _cat in _NO_TEXT_CATS:
                _local_short_circuit = True
                _cat_label_map = {
                    "image": "image", "media": "media file",
                    "archive": "archive", "binary": "binary file",
                }
                _cat_label = _cat_label_map.get(_cat, _cat)
                headline = f"{_cat_label} {item.file_name}"
                content = (
                    f"{item.uploader_name} uploaded `{item.file_name}` at {ts} ({size_s}). "
                    f"The detected file category is { _cat_label }. No plain text was extracted. "
                    "Use the matching image, OCR, media, or archive tool when content inspection is needed.\n\n"
                    "非纯文本文件；需要内容时使用匹配的图片、OCR、媒体或归档工具。"
                )
            elif _cat in {"document", "presentation", "spreadsheet"}:
                # PDF/DOCX/PPTX 提不出文字 → 多半是扫描件/加密/复杂图表, 本地短路
                _local_short_circuit = True
                headline = f"no extracted text for {item.file_name}"
                content = (
                    f"{item.uploader_name} uploaded `{item.file_name}` at {ts} ({size_s}). "
                    "The local extractor did not obtain body text. The file may be scanned, image-heavy, encrypted, or structurally complex. "
                    "Use OCR or a dedicated Office/PDF inspection path when its content matters.\n\n"
                    "本地未提取到正文；需要内容时使用 OCR 或专用 Office/PDF 检查。"
                )

        if _local_short_circuit:
            debug.log(
                "group_files.bg_index.local_summary",
                f"node={kb_node_id} file={item.file_name} cat={_cat} "
                f"(skipped LLM, deterministic template)",
            )
            headline = sanitize_headline(headline)
            content = sanitize_summary(content)

        try:
            if not _local_short_circuit:
                snippet_block = ""
                if snippet:
                    snippet_block = f"## Extracted Content Snippet\n```\n{snippet[:8000]}\n```\n\n内容片段。"
                else:
                    snippet_block = "## Extracted Content Snippet\nNo file content was extracted because the file was too large, non-text, or unavailable.\n\n未提取到文件正文。"

                raw = await llm.chat_json(
                    [
                        {"role": "system", "content": _GROUP_FILE_INDEX_SYSTEM},
                        {"role": "user", "content": (
                            f"## Shared File Metadata\n"
                            f"filename: {item.file_name}\n"
                            f"size: {size_s}\n"
                            f"uploaded_at: {ts}\n"
                            f"uploader: {item.uploader_name}\n\n"
                            f"{snippet_block}"
                        )},
                    ],
                    reasoning="disabled",
                    lite=True,
                    metrics_tag="json.group_file_index",
                )
                if isinstance(raw, dict):
                    headline = sanitize_headline(str(raw.get("headline", "")))
                    content = sanitize_summary(str(raw.get("content", "")))
                else:
                    raise ValueError("non-dict response")
        except Exception:
            log.exception("LLM summary failed for %s, using template", item.file_name)
            headline = f"{item.uploader_name} uploaded {item.file_name}"
            content = (
                f"{item.uploader_name} uploaded `{item.file_name}` at {ts} ({size_s}). "
                "The generated summary failed; use file inspection when content matters.\n\n"
                "文件索引摘要兜底；需要内容时检查文件。"
            )

        # 4. 更新 KB 节点
        file_meta = json.dumps({
            "filename": item.file_name,
            "workspace_path": workspace_rel,
            "archive_id": archive_id,
            "group_id": group_id,
            "upload_time": item.upload_time,
            "uploader_name": item.uploader_name,
            "file_size": item.file_size,
            "napcat_file_id": item.file_id,
            "busid": item.busid,
            "download_status": "done",
        }, ensure_ascii=False)

        async with pool().acquire() as conn:
            await conn.execute(
                """
                UPDATE cold_nodes
                SET headline = $1, content = $2, file_metadata = $3, updated_at = NOW()
                WHERE id = $4 AND archive_id = $5
                """,
                headline, content, file_meta, kb_node_id, archive_id,
            )
        debug.log(
            "group_files.bg_index.done",
            f"node={kb_node_id} file={item.file_name}",
        )

    except Exception:
        log.exception("_bg_download_and_index failed for %s", item.file_name)
        await _mark_download_failed(kb_node_id, "下载或分析过程出错")
    finally:
        _in_flight_downloads.pop(in_flight_key, None)


async def _mark_download_failed(
    kb_node_id: str, reason: str, *, permanent: bool = False,
) -> None:
    """标记下载失败,后续智能体看到后可以告知用户。

    2026-05-09 Patch 43: 加 permanent 参数。
      - permanent=False (默认):仅标记 metadata.download_status="failed",节点保留
        (临时网络错误等可能下次重试成功的场景)
      - permanent=True:永久不可恢复(file_id 丢失等)
    
    2026-05-16 修订: P43 原版直接 DELETE 节点, 违反用户"保留节点 + [已删除]
    标记"原则。改为跟 dream D20 一致的方式:
      - headline 加 [已删除] 前缀
      - content 头部加说明 (让模型知道这文件曾经存在)
      - salience 适度降 (0.2, 不归零)
      - file_metadata.deleted=True (含 deleted_at, deleted_reason)
    
    这样: 主线程 search 仍能命中 → 模型看到 [已删除] → 知道曾经存在但已删,
    不会以为这文件能打开/下载, 也不会反复尝试。
    """
    if permanent:
        try:
            async with pool().acquire() as conn:
                row = await conn.fetchrow(
                    "SELECT archive_id, headline, content, file_metadata FROM cold_nodes WHERE id = $1",
                    kb_node_id,
                )
                if not row:
                    return  # 节点已被其他流程删了
                
                # 解析现有 metadata
                cur_meta = row["file_metadata"]
                meta = {}
                if cur_meta:
                    try:
                        meta = json.loads(cur_meta) if isinstance(cur_meta, str) else (cur_meta if isinstance(cur_meta, dict) else {})
                    except (json.JSONDecodeError, TypeError):
                        meta = {}
                
                # 已标记过 → 跳过 (幂等)
                if meta.get("deleted"):
                    return
                
                old_headline = row["headline"] or ""
                old_content = row["content"] or ""
                
                # 加 [已删除] 前缀
                new_headline = old_headline if old_headline.startswith("[已删除]") else f"[已删除] {old_headline[:25]}"
                
                # content 头部加说明
                now_ts = time.time()
                from datetime import datetime as _dt, timezone as _tz, timedelta as _td
                _human_time = _dt.fromtimestamp(now_ts, tz=_tz(_td(hours=8))).strftime("%Y-%m-%d %H:%M")
                deleted_note = (
                    f"[文件已删除, 索引保留供历史检索]\n"
                    f"删除时间: {_human_time}\n"
                    f"删除原因: {reason[:80]}\n"
                    f"\n[原内容]\n"
                )
                new_content = old_content if old_content.startswith("[文件已删除") else deleted_note + old_content
                
                # 更新 metadata
                meta.update({
                    "deleted": True,
                    "deleted_at": now_ts,
                    "deleted_reason": reason[:200],
                    "deleted_original_path": meta.get("workspace_path", ""),
                    "download_status": "failed",
                })
                
                await conn.execute(
                    """UPDATE cold_nodes
                       SET headline = $1, content = $2, file_metadata = $3,
                           salience = LEAST(salience, 0.2), updated_at = NOW()
                       WHERE id = $4""",
                    new_headline, new_content, json.dumps(meta, ensure_ascii=False),
                    kb_node_id,
                )
                debug.log(
                    "group_files.bg_index.marked_deleted",
                    f"marked KB file node {kb_node_id} as [已删除] "
                    f"(reason={reason[:80]})",
                )
            return  # 已处理完成, 不走 permanent=False 分支
        except Exception:
            log.exception(
                "mark_deleted failed for %s, falling back to non-permanent mark",
                kb_node_id,
            )
            permanent = False  # 降级标记 download_status

    if not permanent:
        try:
            async with pool().acquire() as conn:
                cur = await conn.fetchval(
                    "SELECT file_metadata FROM cold_nodes WHERE id = $1", kb_node_id,
                )
                meta = {}
                if cur:
                    try:
                        meta = json.loads(cur) if isinstance(cur, str) else (cur if isinstance(cur, dict) else {})
                    except (json.JSONDecodeError, TypeError):
                        pass
                meta["download_status"] = "failed"
                meta["download_error"] = reason
                await conn.execute(
                    "UPDATE cold_nodes SET file_metadata = $1, updated_at = NOW() WHERE id = $2",
                    json.dumps(meta, ensure_ascii=False), kb_node_id,
                )
        except Exception:
            log.exception("mark_download_failed failed for %s", kb_node_id)


# ── 按需提取到工作区 ───────────────────────────────────────────

async def fetch_group_file(
    kb_node_id: str, archive_id: str, group_id: str, workspace_dir: str,
) -> dict:
    """将已下载的群文件提取到当前工作区。返回 {ok, path?, error?}。

    智能体调用此工具后可直接 read_file / inspect_file 操作。
    """
    async with pool().acquire() as conn:
        # 先拿 cold_nodes 记录
        cn_row = await conn.fetchrow(
            "SELECT file_metadata FROM cold_nodes WHERE id = $1 AND archive_id = $2",
            kb_node_id, archive_id,
        )
        if not cn_row:
            return {"ok": False, "error": f"文件节点不存在：{kb_node_id}"}

        cn_meta = {}
        if cn_row["file_metadata"]:
            try:
                raw = cn_row["file_metadata"]
                cn_meta = json.loads(raw) if isinstance(raw, str) else (raw or {})
            except (json.JSONDecodeError, TypeError):
                pass
        cn_filename = cn_meta.get("filename", "") or ""
        cn_filesize = cn_meta.get("file_size", 0) or 0

        # 2026-05-15 重设计:JOIN key 改用稳定的 kb_node_id (它在 synced_files 里就存着)。
        # 旧版用 (file_name, file_size) 作 join 条件,有两个问题:
        #   1. 群文件合法重名时,(name, size) 不再唯一,可能命中错误的物理文件
        #   2. 新加的"撞名 → 老的 file_name 改成带时间戳后缀"逻辑会让旧 KB 节点
        #      (file_metadata 里还是原 filename) 通过原名 join 不到对应 synced_files
        #      行,fetch 直接返回"文件节点不存在"。
        # 用 kb_node_id 这两个问题都消失:KB 节点和 synced_files 行是 1:1 直接绑定,
        # 文件名怎么改、磁盘上文件怎么挪都不影响这条路径的稳定性。
        row = await conn.fetchrow(
            """
            SELECT sf.workspace_path, sf.file_id, sf.busid, sf.file_name
            FROM synced_files sf
            WHERE sf.archive_id = $1
              AND sf.group_id = $2
              AND sf.kb_node_id = $3
            """,
            archive_id, group_id, kb_node_id,
        )
    if not row:
        return {"ok": False, "error": f"文件节点不存在：{kb_node_id}"}

    meta = cn_meta
    download_status = meta.get("download_status", "unknown")
    filename = meta.get("filename") or row["file_name"] or "unknown"

    # 找到源文件路径（先精确，再扫描重名变体）
    ws_root = _get_workspace_root()
    workspace_rel = row["workspace_path"] or meta.get("workspace_path", "") or ""
    src_path = ws_root / archive_id / group_id / workspace_rel
    if not src_path.exists():
        alt = _find_existing_local_file(archive_id, group_id, workspace_rel, filename)
        if alt:
            src_path = Path(alt)

    file_on_disk = src_path.exists() and src_path.is_file()

    # pending：如果磁盘上其实已经有文件，就走 healing 路径，否则才告诉用户没下完
    if download_status == "pending" and not file_on_disk:
        return {
            "ok": False,
            "error": (
                f"Shared file {filename} is not downloaded yet (status: {download_status}). "
                "Retry later or ask for another available source file.\n"
                "共享文件尚未下载完成，稍后重试或改用可用来源。"
            ),
        }

    if download_status == "failed" and not file_on_disk:
        reason = meta.get("download_error", "未知错误")
        return {
            "ok": False,
            "error": (
                f"Shared file {filename} download failed: {reason}. "
                "Use another available file or retry after the shared-file backend recovers.\n"
                "共享文件下载失败，改用其它可用文件或等待后台恢复后重试。"
            ),
        }

    # 文件不在磁盘且状态正常 → 重新下载
    if not file_on_disk:
        # 文件可能已被清理，尝试重新下载
        if not row["file_id"]:
            return {
                "ok": False,
                "error": (
                    f"Shared file {filename} has no local copy and cannot be re-downloaded because file_id is missing. "
                    "Use another available source, or ask the user to upload/provide the file again.\n"
                    "共享文件缺少本地副本和 file_id，需要其它来源或用户重新提供。"
                ),
            }
        local = await _download_file(
            GroupFileItem(
                file_id=row["file_id"],
                file_name=filename,
                file_size=meta.get("file_size", 0),
                upload_time=meta.get("upload_time", 0),
                uploader_name=meta.get("uploader_name", ""),
                busid=meta.get("busid", 0),
            ),
            group_id,
            str(src_path.parent),
        )
        if not local:
            return {
                "ok": False,
                "error": (
                    f"Shared file {filename} could not be re-downloaded. "
                    "Use another available source, or ask the user to upload/provide the file again.\n"
                    "共享文件重新下载失败，需要其它来源或用户重新提供。"
                ),
            }
        src_path = Path(local)
        file_on_disk = True

    # 顺手把 pending → done（healing），把对错路径校正
    if download_status in ("pending", "failed"):
        try:
            actual_rel = _to_workspace_rel(str(src_path), archive_id, group_id) or workspace_rel
            new_meta = dict(meta)
            new_meta["download_status"] = "done"
            new_meta["workspace_path"] = actual_rel
            new_meta.pop("download_error", None)
            async with pool().acquire() as conn:
                async with conn.transaction():
                    await conn.execute(
                        "UPDATE cold_nodes SET file_metadata = $1, "
                        "updated_at = NOW() WHERE id = $2 AND archive_id = $3",
                        json.dumps(new_meta, ensure_ascii=False),
                        kb_node_id, archive_id,
                    )
                    if row["file_id"]:
                        await conn.execute(
                            "UPDATE synced_files SET workspace_path = $1 "
                            "WHERE archive_id = $2 AND file_id = $3",
                            actual_rel, archive_id, row["file_id"],
                        )
            debug.log(
                "group_files.fetch.heal",
                f"fetch_group_file healed pending → done: node={kb_node_id} file={filename}",
            )
        except Exception:
            log.exception("fetch_group_file: heal failed for node=%s", kb_node_id)

    # 拷贝到工作区
    # 幂等性：如果工作区里已经有目标文件且和源文件大小一致,认为之前 fetch 过,
    # 直接返回现有路径,避免每次重复调用同 kb_node_id 产生 hello.c → hello_from_group.c
    # → hello_from_group_from_group.c 这种无意义副本链。
    primary_dst = os.path.join(workspace_dir, filename)
    try:
        if os.path.isfile(primary_dst) and os.path.getsize(primary_dst) == os.path.getsize(str(src_path)):
            debug.log(
                "group_files.fetch.reuse",
                f"fetch_group_file: {filename} already in workspace, reusing",
            )
            return {
                "ok": True, "path": filename, "filename": filename,
                "note": "already in workspace, reused existing copy",
            }
    except OSError:
        pass

    dst_filename = filename
    dst = primary_dst
    # 冲突检测：同名但内容不同 → 加后缀
    if os.path.exists(dst):
        stem, ext = os.path.splitext(filename)
        # 优先尝试 _from_group 后缀
        candidate = f"{stem}_from_group{ext}"
        candidate_path = os.path.join(workspace_dir, candidate)
        # 二级幂等：_from_group 版本已存在且大小一致 → 复用
        try:
            if os.path.isfile(candidate_path) and os.path.getsize(candidate_path) == os.path.getsize(str(src_path)):
                debug.log(
                    "group_files.fetch.reuse",
                    f"fetch_group_file: {filename} already in workspace as {candidate}, reusing",
                )
                return {
                    "ok": True, "path": candidate, "filename": filename,
                    "note": "already in workspace, reused existing copy",
                }
        except OSError:
            pass
        # 序号化：_from_group / _from_group_2 / _from_group_3 ... 找第一个空位
        # 避免 3 个并发 fetch 同名文件时第 2、3 个互相覆盖。
        if os.path.exists(candidate_path):
            for i in range(2, 100):
                candidate = f"{stem}_from_group_{i}{ext}"
                candidate_path = os.path.join(workspace_dir, candidate)
                if not os.path.exists(candidate_path):
                    break
        dst_filename = candidate
        dst = candidate_path
    shutil.copy2(str(src_path), dst)

    # 2026-05-10 Patch 57 v3: 记录用户 fetch 的文件 basename,fork 时无条件放行
    # 防止 P57 _looks_like_helper_artifact 把这种文件误识别为 helper artifact 跳过
    try:
        from app.llm.tools.workspace import add_user_fetched_basename
        add_user_fetched_basename(workspace_dir, dst_filename)
    except Exception:
        pass  # 写失败不阻塞 fetch

    log.debug("fetch_group_file: %s → %s", filename, dst)
    # 返回工作区相对路径——下游 read_file / edit_file / inspect_file 都拒绝绝对路径
    # （沙箱安全机制）。直接传文件名即可，因为 dst 就在 workspace_dir 根下。
    return {"ok": True, "path": dst_filename, "filename": filename}


# ── 下载工具函数 ───────────────────────────────────────────────

async def _download_file(item: GroupFileItem, group_id: str, dst_dir: str) -> str | None:
    """下载单个群文件。返回本地绝对路径，失败返回 None。"""
    if item.file_size > MAX_DOWNLOAD_SIZE:
        debug.log(
            "group_files.skip.too_large",
            f"file too large, skip: {item.file_name} ({item.file_size} bytes)",
        )
        return None

    # _napcat 现在失败时会 raise NapCatUnavailable(替代旧版返回 None)。
    # 下载是可选的后台任务,失败时不要抛到外面影响 KB 节点状态——节点会被
    # _heal_pending_nodes 在下次同步时重新调度。
    try:
        url_data = await _napcat("get_group_file_url", {
            "group_id": int(group_id),
            "file_id": item.file_id,
            "busid": item.busid,
        })
    except NapCatUnavailable as e:
        log.debug("get_group_file_url unavailable: %s", e)
        return None
    download_url = ""
    if url_data and isinstance(url_data.get("data"), dict):
        download_url = url_data["data"].get("url", "")
    if not download_url:
        # P85: 记物理签名失败 + 降级 warning 为 debug 如果近期已警过
        # 物理签名 (group_id, file_name, file_size) — NapCat file_id 不稳定时
        # 同物理文件每次 file_id 不一样, 但 name+size 稳定。
        _sig = (group_id, item.file_name, item.file_size)
        _now = time.time()
        _prev_fail = _recent_url_failures.get(_sig)
        _recent_url_failures[_sig] = _now
        # 周期性 GC: dict 长度超 200 时清掉 cooldown 外的条目
        if len(_recent_url_failures) > 200:
            _expired = [k for k, t in _recent_url_failures.items()
                        if _now - t > _URL_FAILURE_COOLDOWN]
            for k in _expired:
                _recent_url_failures.pop(k, None)
        # 2026-05-16: 降级为后台 debug
        # NapCat 能列出文件但不给 URL 时,这些文件无法缓存/索引,但不影响启动或聊天。
        # 控制台 WARN 无用户价值;保留详细后台记录方便后续排查 file_id/busid 问题。
        elapsed = 0.0 if _prev_fail is None else _now - _prev_fail
        debug.log(
            "group_files.download.no_url",
            f"no download url for file={item.file_name} "
            f"group={group_id} file_id={item.file_id} busid={item.busid} "
            f"size={item.file_size} repeat_after={elapsed:.0f}s",
        )
        log.debug(
            "no download url for file=%s group=%s file_id=%s busid=%s size=%s repeat_after=%.0fs",
            item.file_name,
            group_id,
            item.file_id,
            item.busid,
            item.file_size,
            elapsed,
        )
        return None

    safe_name = _safe_filename(item.file_name)
    ts = item.upload_time if item.upload_time > 0 else int(time.time())
    ts_str = datetime.fromtimestamp(ts, tz=timezone(timedelta(hours=8))).strftime(
        "%Y%m%d_%H%M%S"
    )
    local_name = f"{ts_str}_{safe_name}"
    local_path = os.path.join(dst_dir, local_name)

    if os.path.exists(local_path):
        stem, ext = os.path.splitext(local_name)
        local_name = f"{stem}_1{ext}"
        local_path = os.path.join(dst_dir, local_name)

    try:
        async with httpx.AsyncClient(timeout=120.0, follow_redirects=True) as client:
            r = await client.get(download_url)
            if r.status_code == 200:
                content = r.content
                if len(content) > MAX_DOWNLOAD_SIZE:
                    log.warning("download exceeded cap: %s (%s bytes)", item.file_name, len(content))
                    return None
                with open(local_path, "wb") as f:
                    f.write(content)
                log.debug("group file downloaded: %s → %s", item.file_name, local_path)
                return local_path
            else:
                log.warning("download failed for %s: status=%s", item.file_name, r.status_code)
    except Exception:
        log.exception("download error for %s", item.file_name)
    return None


# ── 内容提取 ───────────────────────────────────────────────────

_BINARY_DOC_EXTS = {".pdf", ".docx", ".pptx", ".xlsx"}


def _should_extract_content(local_path: str, file_size: int) -> bool:
    if file_size > MAX_CONTENT_EXTRACT_SIZE:
        return False
    ext = os.path.splitext(local_path)[1].lower()
    # 2026-05-15 PDF/DOCX/PPTX/XLSX bug fix:旧逻辑用「头部 N 字节有无 \x00」
    # 决定是否提取,所有二进制类型(包括 PDF)直接跳过 → _extract_text 拿不到内容
    # → 分类器只能瞎编一段"可能包含 X / 典型章节结构"的填充语,污染检索索引。
    # 修法:扩展名为已知二进制文档时**绕过 NUL 探测**,交给 _extract_text 按
    # 扩展名分派到 pypdf / python-docx 等专门提取器。其他未知扩展仍走旧的
    # NUL 探测(防止把可执行/压缩等真正不可读的文件喂给文本提取)。
    if ext in _BINARY_DOC_EXTS:
        return True
    try:
        with open(local_path, "rb") as f:
            head = f.read(_BINARY_DETECT_BYTES)
        if b"\x00" in head:
            return False
    except OSError:
        return False
    if ext in _FILE_TYPE_TABLE:
        category, _, _, _ = _FILE_TYPE_TABLE[ext]
        if category not in _TEXT_CATEGORIES:
            return False
    return True


def _extract_text(local_path: str, max_chars: int = MAX_EXTRACT_CHARS) -> str:
    """提取群文件文本片段供分类器使用。

    2026-05-15 PDF/DOCX 等二进制文档提取支持(Item: PDF 索引 bug):
      旧版对所有文件都按 utf-8 文本打开 → PDF/DOCX 出来一堆乱码 → 分类器拿不到
      内容 → 编一段"可能包含 X / 典型 Y" 的瞎话。修法:按扩展名分派到专门提取器。
      pypdf / python-docx 是 lazy import — 没装时回退到 "" (分类器会走"未提取"
      诚实分支,不会瞎编)。
      图像化 PDF (扫描件 / 截图导出的 PDF) pypdf 提不出文字 → 返回 "" →
      分类器输出"内容以图像形式存储,需 OCR" → bot 后续可直接 `ocr(image_path=PDF)`
      因为 ocr_bridge 内部已经支持 PDF 输入(_render_pdf_to_images + 逐页 OCR)。
    """
    ext = os.path.splitext(local_path)[1].lower()

    if ext == ".pdf":
        return _extract_pdf_text(local_path, max_chars)
    if ext == ".docx":
        return _extract_docx_text(local_path, max_chars)
    # .pptx / .xlsx 暂时不提取(用 office 工具语义比 raw text 更准;
    # 让分类器走"未提取"诚实分支)。需要时再加 _extract_pptx_text 等。

    # 纯文本路径(包括 .txt .md .py .c .json .csv 等):utf-8 读
    try:
        with open(local_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read(max_chars)
        return content.strip()
    except Exception:
        return ""


def _extract_pdf_text(local_path: str, max_chars: int) -> str:
    """pypdf 抽 PDF 文本层。图像化 PDF (无文本层) 返回 ""。"""
    try:
        import pypdf  # type: ignore
    except ImportError:
        log.warning("pypdf 未安装, 跳过 PDF 文本提取: %s", local_path)
        return ""
    try:
        reader = pypdf.PdfReader(local_path)
    except Exception as e:
        log.warning("pypdf 打开失败 (%s): %s", local_path, e)
        return ""

    parts: list[str] = []
    total = 0
    # 最多读前 10 页;太多页 + 长正文容易超 max_chars,且对索引摘要无增量价值
    for i, page in enumerate(reader.pages[:10]):
        try:
            t = page.extract_text() or ""
        except Exception:
            continue
        t = t.strip()
        if not t:
            continue
        parts.append(t)
        total += len(t)
        if total >= max_chars:
            break
    text = "\n\n".join(parts).strip()
    # 图像化 PDF 的特征:即使文件大几百 KB,pypdf 提出的文字也只是若干页码 +
    # 偶尔一两个字符。这种情况返回 "" 比返回 "1 2 3 ..." 干净 (分类器
    # 看到 "" 会走"未提取"诚实分支,而不是把页码当作"目录"瞎编)。
    if len(text) < 30:
        return ""
    return text[:max_chars]


def _extract_docx_text(local_path: str, max_chars: int) -> str:
    """python-docx 抽 DOCX 段落 + 表格文字。
    
    2026-05-16 修订: 之前只提取 doc.paragraphs, 论文/报告类文档大量文字在表格里
    会被完全遗漏 → 走 deterministic template "未提取文字的 X.docx" 误导模型.
    现在也遍历 doc.tables 提取单元格文字.
    """
    try:
        import docx  # type: ignore
    except ImportError:
        log.warning("python-docx 未安装, 跳过 DOCX 文本提取: %s", local_path)
        return ""
    try:
        doc = docx.Document(local_path)
    except Exception as e:
        log.warning("python-docx 打开失败 (%s): %s", local_path, e)
        return ""

    parts: list[str] = []
    total = 0
    
    # 1. 段落文字
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        parts.append(t)
        total += len(t)
        if total >= max_chars:
            text = "\n".join(parts).strip()
            return text[:max_chars]
    
    # 2. 表格文字 (单元格按 | 分隔, 行换行)
    try:
        for tbl in doc.tables:
            for row in tbl.rows:
                cells_text = []
                for cell in row.cells:
                    ct = (cell.text or "").strip()
                    if ct:
                        cells_text.append(ct)
                if cells_text:
                    line = " | ".join(cells_text)
                    parts.append(line)
                    total += len(line)
                    if total >= max_chars:
                        text = "\n".join(parts).strip()
                        return text[:max_chars]
    except Exception as e:
        log.warning("python-docx 表格提取失败 (%s): %s", local_path, e)
    
    # 2026-05-16 Round 14e: 空提取时检测嵌入图片
    # 实测 trace: paper.docx 走 _extract_docx_text 返空 → KB 入"未提取文字" template
    # 模型完全不知道这是论文(可能含 OCR-able 内容). 检测嵌入图片, 给个有意义的 placeholder
    # 而非 "未提取文字" — 让模型知道这是纯图文档, 待 D17 后台 OCR.
    text = "\n".join(parts).strip()
    if not text:
        img_count = 0
        try:
            # 用 zipfile 数 word/media/ 内图片 (.png/.jpg/.jpeg 等)
            import zipfile
            with zipfile.ZipFile(local_path) as z:
                media_files = [
                    n for n in z.namelist()
                    if n.startswith("word/media/")
                       and n.lower().endswith((".png", ".jpg", ".jpeg", ".bmp", ".gif", ".tiff"))
                ]
                img_count = len(media_files)
        except Exception:
            pass
        if img_count > 0:
            # 返个 placeholder 让 _local_summary 看到非空 (不走 "未提取文字" template)
            # 内容暗示 D17 OCR 后会有真内容
            return (
                f"[此 DOCX 文档主体为图片 ({img_count} 张嵌入图片), "
                f"文字内容待后台 OCR 提取 — 当前仅有图像内容, "
                f"完整内容索引由 D17/D18 任务异步补全]"
            )
    return text[:max_chars]


def _safe_filename(name: str) -> str:
    return "".join(c for c in name if c not in r'<>:"/\|?*' and ord(c) >= 32)


def _fmt_size(size: int) -> str:
    if size < 1024:
        return f"{size:,}B"
    elif size < 1024 * 1024:
        return f"{size / 1024:.0f}KB"
    else:
        return f"{size / 1024 / 1024:.1f}MB"
