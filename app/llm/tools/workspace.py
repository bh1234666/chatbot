"""
工作区工具：mkdir / write / run / read / inspect。

安全设计：
- 所有写入操作限定在 workspace_root 内，拒绝路径遍历（..）
- 沙箱内可读写当前 archive/group 的持久工作区与 .temp 临时区
- 沙箱外默认只读，禁止修改/删除操作
- 命令执行先经过风险策略、路径归属分析与 AST 扫描
- 高危系统命令始终拦截
- 持久工作区保存交付产物，.temp 保存当前会话临时文件，.prev 保存上一轮临时快照
"""
from __future__ import annotations

import asyncio
import ast
import hashlib
import json
import logging
import os
import re
import shutil
import signal
import subprocess  # ── Bug #23 修复:taskkill 杀进程树 ──
import sys
import tempfile
import time
import uuid
from contextvars import ContextVar
from pathlib import Path

from app.config import settings

# 2026-05-06 §A4: FIX_HINT 重复追踪从 module-level dict 改 ContextVar。
# 旧 module-level dict 在多个 helper 并行时互相覆盖 — helper-A 的 reset 清掉 helper-B 的计数。
# ContextVar 确保每个 asyncio task 有独立 dict,PEP 567 自动跨 await propagate。
from contextvars import ContextVar

_fix_hint_counts_var: ContextVar[dict[int, int]] = ContextVar(
    "_fix_hint_counts", default=None,
)


def _get_fix_hint_dict() -> dict[int, int]:
    d = _fix_hint_counts_var.get()
    if d is None:
        d = {}
        _fix_hint_counts_var.set(d)
    return d


def reset_fix_hint_counts() -> None:
    """每个 tool iteration 开始时清空 hint 重复计数。由 client.py 调用。"""
    _fix_hint_counts_var.set({})


# 2026-05-15 P70: bash 命令重复失败追踪
# 病因(实测 05-15 comp_bench): 同一 helper 里 1026 次 bash + 66 失败, 多个相同 gcc
# 编译命令各失败 3-4 次, 但 fix_hint 计数只在 hint 文本同时出现时触发。对于退出码非 0
# 但无明确 hint 的情况(很多 "cd ... && gcc ..." 复合命令), 重复失败没被检测到。
# 修法: 直接按 command 字符串 hash 计失败次数, ≥3 次返 _repeated_command_failure 字段。
_bash_failure_counts_var: ContextVar[dict[int, int] | None] = ContextVar(
    "bash_failure_counts", default=None
)


def _get_bash_failure_dict() -> dict[int, int]:
    d = _bash_failure_counts_var.get()
    if d is None:
        d = {}
        _bash_failure_counts_var.set(d)
    return d


def reset_bash_failure_counts() -> None:
    """每个 helper iteration 开始时清空(可选, 默认随 helper 生命周期保留)。"""
    _bash_failure_counts_var.set({})
from app.core import debug
from app.core.core_processes import registry as proc_registry, current_owner
from app.llm.tools.command_risk import analyze_command

log = logging.getLogger(__name__)

# ── 2026-05-06 §SEC1: bash 速率限制 ──
# helper 可以无节制连续发 bash 命令,60 次/分钟的 per-owner 上限让 helper 收到
# 明确反馈而非被静默忽略。
_BASH_RATE_LIMIT = 60  # commands per minute per owner
_bash_call_log: dict[str, list[float]] = {}
_bash_log_lock = asyncio.Lock()


async def _check_bash_rate(owner: str) -> bool:
    """返回 True 如果 owner 仍未超速。"""
    async with _bash_log_lock:
        now = time.time()
        log_entries = _bash_call_log.setdefault(owner, [])
        cutoff = now - 60
        while log_entries and log_entries[0] < cutoff:
            log_entries.pop(0)
        if len(log_entries) >= _BASH_RATE_LIMIT:
            return False
        log_entries.append(now)
        return True


_MAX_OUTPUT = 64 * 1024  # 64 KB

# ─── git-bash 检测(2026-05-03 修):Bug 1 修复 ────────────────────────
# 历史问题(trace 7e6629f228c84e78):BASH_TOOL_SCHEMA 和 helper system prompt
# 都告诉 LLM "bash 工具有 git-bash/MSYS2 环境,unix 命令可用",但实际上
# `bash` 工具走的是 cmd /c <cmd> 路径(create_subprocess_shell 在 Windows 上
# 默认 cmd.exe),unix 语法 `2>/dev/null` `head -5` `find -name` 全部失败。
# 模型按 prompt 写 → 必失败 → stuck detector 杀 helper(sbt 32s 就被杀)。
#
# 修法: 启动时检测 git-bash(or MSYS2 bash)是否可用,公开 has_unix_shell()
# 让 prompt 描述真实反映实际能力;handle_run 加 prefer_unix_shell 参数,
# `bash` 工具通过它走真 bash.exe -c <cmd>(不再 cmd /c 包装)。

# 进程/平台 helper 已抽离到 process_utils.py(2026-05-20 重构);re-export 兼容。
from app.llm.tools.process_utils import (  # noqa: E402,F401
    _detect_git_bash,
    _detect_asan_support,
    _kill_process_tree,
)


# 模块导入时检测一次(避免每次工具调用重新查 PATH)。
_GIT_BASH_EXE: str | None = _detect_git_bash()
# 2026-05-21: ASan 可用性同样探测一次(实际编译最小程序,MinGW 默认缺 libasan)。
_ASAN_OK: bool = _detect_asan_support()


def has_asan() -> bool:
    """当前环境 gcc 能否真正用 -fsanitize=address 编译+链接。
    供 prompt 编译时据实反映:不可用时不要诱导 helper 用 ASan(否则 cannot find -lasan)。
    """
    return _ASAN_OK


def has_unix_shell() -> bool:
    """Linux/macOS 永远 True;Windows 看 git-bash 是否检测到。
    供 registry.py / delegate.py 在编译 prompt 时实时反映平台真实能力。
    """
    return sys.platform != "win32" or _GIT_BASH_EXE is not None


def git_bash_path() -> str | None:
    """返回检测到的 bash 可执行路径(主要供 debug log / 诊断用)。"""
    return _GIT_BASH_EXE


# ─── workspace.run timeout 策略(2026-05-01 重构) ───────────────────────
# **设计哲学**:LLM 最懂自己要跑什么,最该自决 timeout。
# 默认值故意设很低(1s),强制 LLM 显式传 timeout_sec —— 没传 = 按 1s 跑。
#
# 这样设计的理由:
#   1. 避免 ML/Python 启发式误判(实测:gcc 编译被识别成 ML 给 300s)
#   2. 强迫 LLM 思考 "这命令该跑多久" —— 是健康的认知负担
#   3. 启发式默认 90/300 看似贴心,实际让 LLM 偷懒不传值,失败时不知道为什么
#   4. helper 已无时间硬墙,但单条 workspace.run 卡 5min 仍会让心跳变 stale,
#      触发主进程误判 helper 死了,所以单条命令时长仍应有节制
#
# LLM 没传时给 1s,失败错误会清晰说明 "你没传 timeout_sec,默认 1s,下次显式传"。
_RUN_TIMEOUT_DEFAULT = 1     # ← 默认值。LLM 不传 timeout_sec 时用,故意很低
# ── 2026-05-04 修改:300 → 7200 ──
# 旧版 300s 硬上限导致大型 benchmark/ML 训练/长编译被强制截断,丢失有效输出。
# 绝大多数进程能跑就不该停 — LLM 通过 timeout_sec 自己管理预期,
# 超时后 kill 并在错误信息里告知,让 LLM 根据 stderr/FIX_HINT 决定重试还是换方向。
_RUN_TIMEOUT_HARD_CAP = 7200  # 硬上限 2h,防 LLM 失误传 999999
_RUN_TIMEOUT_MIN = 1         # 最小允许值

# ═══════════════════════════════════════════════════════════════
# Layer 1：高危关键字（始终拦截，不区分语境）
# ═══════════════════════════════════════════════════════════════
from app.llm.tools.workspace_run_checks import (  # noqa: E402,F401
    _DANGEROUS_KEYWORDS,
    _CMD_DESTRUCTIVE_OPS,
    _WIN_BUILTIN_CMDS,
    _WIN_SHELL_OPERATORS,
    _ALREADY_CMD_RE,
    _GCC_BLOCKED_FLAGS,
    _PATH_RE,
    _REDIRECT_RE,
    _compact_repeating_lines,
    _diagnose_build_failure,
    _security_check,
    _check_cmd,
    _check_gcc,
    _extract_paths,
    _is_abs_outside,
    _has_redirect_to_outside,
    _touches_prev_or_outside,
    set_main_thread_provider as _set_workspace_run_main_thread_provider,
)

def copy_workspace_contents(
    src_dir: str, dst_dir: str, *, use_hardlink: bool = True,
    allowed_prefixes: frozenset[str] | set[str] | None = None,
    allowed_basenames: frozenset[str] | set[str] | None = None,
    include_downloaded_media: bool = False,
    include_environment_files: bool = False,
) -> int:
    """Copy all files from src workspace to dst workspace. Returns file count.

    保留 size filter 但门槛放宽 — 主修是 helper 完成时立即删工作区(在 delegate.py)。
    这里只防"极端大文件被反复复制"(如 100MB+ 二进制),不限制正常 benchmark 数据。

    B8 修复: 同时通过 _should_copy_for_fork 跳过历史 helper artifact
    (含 ≥2 层 task_id 前缀的污染文件),防止 fork 把上一轮的污染再带过去。

    Patch 57(2026-05-10):新增 allowed_prefixes 参数 — 当前会话已 spawn 的
    helper task_id 集合,前缀命中的文件不被识别为"历史污染"。
    病因(trace b430c4f228eb40c7):gen_charts 完成后产物 gen_charts_*.png 被 fork
    误识别为 artifact 跳过,embed_charts 沙箱里没图,任务失败,文件没推送给用户。

    2026-05-02 part10 (A9):**特殊目录递归 fork**。
    主区里 `_shared/` 和 `_helpers_shared/` 子目录(主线程或前序 helper 写的共享代码/数据)
    会**完整递归复制**到 helper 工作区,helper 直接 import/read 用。
    其他子目录不递归(避免误带 `_delegate_<task>/` 等内部目录)。

    2026-05-23: read helper 需要访问用户上传图片。调用方可设置
    include_downloaded_media=True 递归复制 `_downloaded_media/`。

    L6-1 (2026-05-09): use_hardlink=True 时对安全扩展名(.csv/.png/.json 等)
    用 os.link() 创建硬链接,零 IO 零额外磁盘空间。跨盘/网络盘自动 fallback 到 copy。
    """
    if not src_dir or not dst_dir or not os.path.isdir(src_dir):
        return 0
    os.makedirs(dst_dir, exist_ok=True)
    count = 0
    hardlinked = 0
    skipped_huge = 0
    skipped_bytes = 0
    skipped_artifacts = 0
    # 2026-05-02 part10 A9:特殊共享目录白名单
    _SHARED_DIRS = {"_shared", "_helpers_shared"}
    _SPECIAL_DIRS = set(_SHARED_DIRS)
    if include_downloaded_media:
        _SPECIAL_DIRS.add("_downloaded_media")
    if include_environment_files:
        _SPECIAL_DIRS.add("_env")
    for name in os.listdir(src_dir):
        src = os.path.join(src_dir, name)
        dst = os.path.join(dst_dir, name)
        # 处理特殊共享目录:递归复制,跳过大小检查(共享脚手架通常很小)
        if os.path.isdir(src) and name in _SPECIAL_DIRS:
            try:
                shutil.copytree(src, dst, dirs_exist_ok=True)
                # 统计复制了几个文件(只算 immediate children + 一层子目录)
                for _root, _dirs, _files in os.walk(dst):
                    count += len(_files)
            except OSError:
                debug.log(
                    "workspace.copy_shared_failed",
                    f"failed to copy shared dir {name}; skipping",
                )
            continue
        if not os.path.isfile(src):
            continue
        try:
            size = os.path.getsize(src)
        except OSError:
            continue
        # 走综合过滤 — _should_copy_to_helper(size + helper artifact 启发式)
        if not _should_copy_to_helper(
            src, size,
            allowed_prefixes=allowed_prefixes,
            allowed_basenames=allowed_basenames,
        ):
            if size > _FORK_COPY_HARD_LIMIT:
                skipped_huge += 1
                skipped_bytes += size
            else:
                skipped_artifacts += 1
            continue

        # L6-1: 对安全扩展名尝试 hardlink
        if use_hardlink:
            ext = os.path.splitext(name)[1].lower()
            if ext in _HARDLINK_SAFE_EXTS:
                try:
                    os.link(src, dst)
                    hardlinked += 1
                    count += 1
                    continue
                except (OSError, NotImplementedError):
                    pass  # 跨盘/网络盘/ReFS 不支持 hardlink,fallback 到 copy

        shutil.copy2(src, dst)
        count += 1
    if skipped_huge or skipped_artifacts:
        debug.log(
            "workspace.copy_filtered",
            f"copied {count} files (hardlinked={hardlinked}), "
            f"skipped {skipped_huge} >1GB "
            f"(saved {skipped_bytes / 1024 / 1024:.1f} MB), "
            f"skipped {skipped_artifacts} historical helper artifacts (multi-prefixed) — "
            f"main workspace fork is now clean",
        )
    return count


# ─── Workspace fork file size cap(防极端大文件污染)──────────────────────
# 设计哲学(Phase 5++ v3 修订, 用户反馈调整):
#   - 真正的修复是 helper 完成时立即删工作区(delegate.py)
#   - 这里的 size cap 仅作"安全网" — 防 helper 写出 GB 级别的极端大文件
#     被反复复制(典型场景: helper 跑训练吐出 weights.bin 几个 GB)
#   - 测试数据文件(几百 MB)放行,因为算法测试本就需要大数据
#
# 用户反馈: "大文件可以进入,不要超过 1GB 就行"
# v3 阈值: 1GB — 只挡真异常大文件(单文件超过 1GB 几乎肯定是 bug)
_FORK_COPY_HARD_LIMIT = 1024 * 1024 * 1024    # 1GB — 软门槛,只挡极端大文件

# L6-1 (2026-05-09): Hardlink-safe file extensions.
# These file types are treated as "read-only baseline" — helper won't in-place modify them.
# .docx/.pptx/.xlsx are included because office.* write paths use copy-on-write semantics.
_HARDLINK_SAFE_EXTS = frozenset({
    ".csv", ".tsv", ".json", ".png", ".jpg", ".jpeg",
    ".gif", ".pdf", ".zip", ".docx", ".pptx", ".xlsx",
})

# B8 修复: 历史 helper 加前缀的产物不应再带进新 helper 的工作区。
# 这些文件对新 helper 没有任何价值,反而会被 _copy_results_to_main 再次
# 加新前缀复制回主区,造成多层 task_id 前缀的指数膨胀。
# 启发式: 文件名前部含至少 2 个看起来像 task_id 的下划线分段
# (4-20 字符,纯小写字母+数字+下划线,且不是常见原始词)。
import re as _re_helper_artifact
_HELPER_ARTIFACT_PATTERN = _re_helper_artifact.compile(
    # 形如 abc_def_realfile.png / first_fit_acmc_novel_chart.png 等
    # 至少 2 个 task_id 段(每段 3-25 字母/数字/_,后跟下划线),后面才是真正名字
    r"^(?:[a-z][a-z0-9]{2,24}_){2,}",
)

# 已知的"非 task_id 前缀"白名单 — 这些是合法源码命名风格,不是 helper artifact
_LEGIT_PREFIX_NAMES = frozenset({
    "first_fit", "best_fit", "next_fit", "worst_fit",
    "heap_sort", "merge_sort", "quick_sort", "radix_sort", "tim_sort",
    "std_sort", "shell_sort", "bubble_sort", "insertion_sort", "selection_sort",
    "test_", "spec_", "src_", "main_", "lib_", "utils_", "core_",
})


def _looks_like_helper_artifact(
    name: str,
    *,
    allowed_prefixes: frozenset[str] | set[str] | None = None,
    allowed_basenames: frozenset[str] | set[str] | None = None,
) -> bool:
    """启发式: 判断文件名是否是历史 helper 加前缀的产物。"""
    lower = name.lower()
    # 合法源码命名 - 跳过此过滤
    for legit in _LEGIT_PREFIX_NAMES:
        if lower.startswith(legit):
            return False
    # P57 v3: 用户 fetch 的文件 - 完整 basename 匹配
    if allowed_basenames and name in allowed_basenames:
        return False
    # Patch 86: 同名 task 产生的数据文件仍是当前可用交付物。
    # 例如 bench_runner_assemble_bench_out.csv 是 task_id=bench_runner_assemble 的基准数据，
    # 旧规则会把 bench_runner_assemble_ 误判成多层 helper 前缀，导致后续 draw helper 拿不到 CSV。
    if allowed_prefixes:
        for prefix in allowed_prefixes:
            if not prefix:
                continue
            if lower.startswith(prefix.lower() + "_"):
                return False

    stem = os.path.splitext(lower)[0]
    segments = stem.split("_")
    if len(segments) < 3:
        return False

    if len(segments) >= 4:
        for i in range(2, len(segments) - 1):
            left = "_".join(segments[:i])
            right = "_".join(segments[i:-1])
            if left and right and left == right:
                return True

    return False


def _should_copy_for_fork(path: str, size: int) -> bool:
    """是否应将此文件复制到 helper fork 工作区(纯 size 门槛 — 接口稳定)。

    门槛:仅挡 >1GB 的极端大文件。测试数据(几百 MB)正常通过。
    real cleanup happens at helper completion (see delegate.py handle_delegate).

    NOTE: 历史 helper artifact 的过滤(B8 修复)放在 copy_workspace_contents 里
    单独做,不混入这个函数 — 这样保持简单稳定的 size-only 接口给老测试。
    """
    return size <= _FORK_COPY_HARD_LIMIT


def _should_copy_to_helper(
    path: str, size: int,
    *,
    allowed_prefixes: frozenset[str] | set[str] | None = None,
    allowed_basenames: frozenset[str] | set[str] | None = None,
) -> bool:
    """B8 + Patch 57 + Patch 57 v3: 综合判断,是否应将此文件带入 helper fork 工作区。

    在 _should_copy_for_fork(size 检查)基础上增加历史 artifact 过滤:
    含 ≥2 层 task_id 前缀的非源码文件不再复制 — 这是 B8 修复防止污染累积。

    Patch 57(2026-05-10):allowed_prefixes 当前会话 helper task_id,前缀命中放行。
    Patch 57 v3(2026-05-10):allowed_basenames 当前会话用户 fetch 的群文件 / 上传文件
    完整 basename,精确匹配放行。修复 trace f973df37 暴露的 fetch_group_file 拉来的
    docx 被 fork 误识别为 artifact 跳过的问题。
    """
    if not _should_copy_for_fork(path, size):
        return False
    name = os.path.basename(path)
    ext = os.path.splitext(name)[1].lower()
    # 编译中间产物 / 缓存:不带过去(2026-05-02 part15)
    # .o/.obj 是编译中间结果 helper 必须重编;.pyc 同理
    # .exe 例外:有些场景 helper 想直接调用前编好的 binary(如 主线程编译好的 testdata),
    #   但绝大多数场景都是过期的,统一跳过让 helper 自己 re-link
    if ext in {".o", ".obj", ".pyc", ".pyo", ".lib", ".a", ".so"}:
        return False
    if name in {"__pycache__", ".pytest_cache", "node_modules", ".git"}:
        return False
    # 2026-05-03 v18:helper 内部元数据 .json 永不带过去(它们属于上一个 helper 的状态,
    # 新 helper 不应继承——尤其 .read_history.json 会让重读限制误生效)
    # 2026-05-10 Patch 58:.helpers_displayed_name.json 是 P58 的推送层 metadata,
    # 不该被 helper 看到(否则 helper 可能误以为是任务相关数据)
    if (
        name in _METADATA_BASENAMES
        or name.startswith(".helper_")
        or name.endswith(_NON_DELIVERABLE_SUFFIXES)
    ):
        return False
    # 源码总是 fork(LLM 可能要读它们)
    if ext in {".py", ".c", ".cpp", ".h", ".hpp", ".js", ".ts", ".sh",
               ".bat", ".cmd", ".ps1", ".md", ".txt", ".json", ".yaml", ".yml"}:
        return True
    # 非源码 + 看起来像历史 helper artifact → 跳过
    if _looks_like_helper_artifact(name, allowed_prefixes=allowed_prefixes):
        return False
    return True


def clean_workspace_dir(ws_dir: str) -> bool:
    """Remove all files and subdirectories from a workspace dir.

    ⚠ SAFETY (Phase 5++ v3): 拒绝清除永久主工作区。只允许清 helper / 临时区。
    用户反馈: "永久工作区的内容不要有任何变动,不要删除"
    主工作区路径必须含 _delegate_ 才允许调用本函数(双保险)。
    """
    if not ws_dir or not os.path.isdir(ws_dir):
        return False
    # ── 安全 guard: 拒绝清主工作区 ──
    if not _is_safe_to_wipe(ws_dir):
        log.error(
            "clean_workspace_dir REFUSED to clean %s "
            "(safety: only paths containing _delegate_ are allowed)",
            ws_dir,
        )
        return False
    for name in os.listdir(ws_dir):
        path = os.path.join(ws_dir, name)
        try:
            if os.path.isfile(path) or os.path.islink(path):
                os.unlink(path)
            elif os.path.isdir(path):
                shutil.rmtree(path, ignore_errors=True)
        except OSError:
            pass
    return True


# 工作区路径安全/可回收判定已抽离到 workspace_paths.py(2026-05-20 重构);re-export 兼容。
from app.llm.tools.workspace_paths import (  # noqa: E402,F401
    _safe_resolve,
    _is_safe_to_wipe,
    _is_known_reclaimable_workspace_dir,
    _is_known_reclaimable_workspace_file,
)


# ── 工作区生命周期 ────────────────────────────────────────────
# Registry: group_key -> active workspace dirs (oldest..newest, for download APIs).
# A single group can have concurrent per-user turns; keeping only one path lets
# a later turn overwrite the current download workspace for an earlier turn.
_workspace_registry: dict[str, list[str]] = {}

# 注意:这里曾经有个 _SOURCE_EXTENSIONS 集合(.py/.js/.ts/.sh/.bat/.cmd/.ps1),
# 用于在 list_generated_files 里把"源码"过滤掉。
# 但这导致 plan.deliverables=["qlearning.py"] 时,文件永远找不到、永远无法交付,
# 也永远不会被 kb.index_generated_files 记录。trace d4896644 实测:
#   plan.deliverables=["environment.py","qlearning.py","dqn.py",...]
#   workspace.missing 报: {'qlearning.py','environment.py','dqn.py','benchmark_rl.py'}
#   ↑ 全是 .py,被这里过滤掉了
# 删除该过滤——是否交付应该完全由 plan.deliverables 决定,
# orchestrator 的交付候选收集已经会过滤中间脚本前缀
# 来处理"中间脚本"的 fallback,不需要在这里做。


def _get_workspace_root() -> Path:
    """Return root for persistent workspaces. Configurable via settings.workspace_root."""
    if settings.workspace_root:
        return Path(settings.workspace_root)
    project_root = Path(__file__).resolve().parent.parent.parent.parent
    return project_root / "data" / "workspaces"


# 2026-05-04 v19.2:matplotlib 中文字体配置兜底
# 实测 trace 2661da1f:LLM 写画图脚本时常忘记设
# `plt.rcParams['font.sans-serif']`,导致中文显示成方框。
# 容器层在 MPLCONFIGDIR 写一个 matplotlibrc 全局配中文字体优先 — 即便
# LLM 不显式设置,matplotlib 启动也会读取此文件。
#
# 内容覆盖:
#  - font.sans-serif: 中文字体优先(Microsoft YaHei / SimHei),
#    再 fallback 到 DejaVu Sans(matplotlib 自带,含部分上下标 Unicode)
#  - axes.unicode_minus: False(避免负号方框)
#
# 不覆盖用户已有 matplotlibrc(检测文件存在就跳过)— 用户可能有自定义。
_DEFAULT_MATPLOTLIBRC_TEXT = """\
# Auto-generated by workspace.py (v19.2) — 中文字体兜底配置
# 用户可手动覆盖此文件,框架不会强制重写(只在文件不存在时初始化一次)。

# 中文字体优先级(Windows 下 Microsoft YaHei / SimHei 必有,
# Mac/Linux 下 fallback 到 DejaVu Sans)
font.sans-serif:  Microsoft YaHei, SimHei, Arial Unicode MS, DejaVu Sans, sans-serif
font.family:      sans-serif

# 负号正常显示(中文字体的减号不是 Unicode minus,要关闭这个)
axes.unicode_minus: False

# 默认 dpi 略提高,导出图清晰
figure.dpi:       100
savefig.dpi:      150

# 数学符号字体(LaTeX 风格 mathtext)— 用 STIX 替代 cm,
# 因为 cm 字体某些 Linux 系统缺失。STIX 是 matplotlib 自带的兜底字体。
mathtext.fontset: stix
"""


def _ensure_matplotlibrc(mpl_config_dir: str) -> None:
    """确保 MPLCONFIGDIR 下有 matplotlibrc 配置文件。幂等,文件已存在则跳过。"""
    try:
        os.makedirs(mpl_config_dir, exist_ok=True)
        rc_path = os.path.join(mpl_config_dir, "matplotlibrc")
        if os.path.isfile(rc_path):
            # 用户/上次启动已写,不覆盖(尊重定制)
            return
        with open(rc_path, "w", encoding="utf-8") as f:
            f.write(_DEFAULT_MATPLOTLIBRC_TEXT)
        log.info("workspace: matplotlibrc initialized at %s", rc_path)
        debug.log(
            "workspace.matplotlibrc.init",
            f"wrote default matplotlibrc to {rc_path}"
        )
    except Exception as e:
        # 失败不能阻塞主流程 — 退而求其次依赖 LLM prompt 引导
        log.warning("workspace: failed to init matplotlibrc: %s", e)


def create_workspace(archive_id: str = "", group_id: str = "") -> str:
    if archive_id and group_id:
        ws = _get_workspace_root() / archive_id / group_id
    else:
        base = Path(tempfile.gettempdir()) / "chatbot_workspaces"
        base.mkdir(parents=True, exist_ok=True)
        ws = base / f"ws_{uuid.uuid4().hex[:12]}"
    ws.mkdir(parents=True, exist_ok=True)
    log.info("workspace created: %s", ws)
    debug.log("workspace.create", str(ws))
    return str(ws)


def get_persistent_workspace_path(archive_id: str, group_id: str) -> str:
    """Return the expected persistent workspace path regardless of existence."""
    return str(_get_workspace_root() / archive_id / group_id)


def cleanup_archive_workspace(archive_id: str) -> bool:
    """Delete the persistent workspace for an archive.

    Called when the archive itself is deleted. Unlike cleanup_workspace(),
    this is explicitly allowed to remove the main workspace (the archive is
    gone, so the persistent data should go too).

    Returns True if a directory was actually removed.
    """
    root = _get_workspace_root()
    target = root / archive_id
    target_str = str(target)
    # Safety: must be directly under workspace root
    norm_target = os.path.normpath(os.path.abspath(target_str))
    norm_root = os.path.normpath(os.path.abspath(str(root)))
    if not norm_target.startswith(norm_root + os.sep):
        log.error("cleanup_archive_workspace REFUSED: %s not under workspace root %s", target_str, norm_root)
        return False
    # Safety: archive_id must match expected pattern
    if not archive_id.startswith("arch_"):
        log.error("cleanup_archive_workspace REFUSED: %s does not match arch_ prefix", archive_id)
        return False
    try:
        if os.path.isdir(target_str):
            shutil.rmtree(target_str, ignore_errors=True)
            log.info("archive workspace cleaned: %s", target_str)
            debug.log("workspace.archive_cleanup", target_str)
            return True
    except Exception:
        log.exception("archive workspace cleanup failed: %s", target_str)
    return False


def cleanup_workspace(ws_dir: str) -> None:
    """Remove an entire workspace directory.

    ⚠ SAFETY (Phase 5++ v3): 拒绝删除永久主工作区。
    用户反馈: "永久工作区的内容不要有任何变动,不要删除。包括下载的群文件"
    """
    if not _is_safe_to_wipe(ws_dir):
        log.error(
            "cleanup_workspace REFUSED to remove %s "
            "(safety: persistent main workspace must be preserved)",
            ws_dir,
        )
        return
    try:
        if os.path.isdir(ws_dir):
            shutil.rmtree(ws_dir, ignore_errors=True)
            log.info("workspace cleaned: %s", ws_dir)
            debug.log("workspace.cleanup", ws_dir)
    except Exception:
        log.exception("workspace cleanup failed: %s", ws_dir)


def cleanup_archive_temp(archive_id: str) -> dict:
    """清理存档下所有临时工作区，保留主工作区永久文件。

    删除内容：
      - <group>/.temp/ 整个临时工作区目录树
      - <group>/_delegate_*/ 泄露到主区的 helper 沙箱
    保留内容：
      - 用户上传的群文件
      - 已提交的核心产物
      - .pause_state.json

    Returns: {"groups_scanned": N, "dirs_removed": N, "bytes_freed": N}
    """
    root = _get_workspace_root()
    archive_dir = root / archive_id
    stats = {"groups_scanned": 0, "dirs_removed": 0, "bytes_freed": 0}

    if not archive_dir.is_dir():
        return stats

    # Safety: archive_id must match expected pattern
    if not archive_id.startswith("arch_"):
        log.error("cleanup_archive_temp REFUSED: %s does not match arch_ prefix", archive_id)
        return stats

    for group_dir in sorted(archive_dir.iterdir()):
        if not group_dir.is_dir():
            continue
        stats["groups_scanned"] += 1

        # 1. 删除 .temp/ 整个临时工作区
        temp_dir = group_dir / ".temp"
        if temp_dir.is_dir():
            try:
                temp_size = sum(
                    f.stat().st_size for f in temp_dir.rglob("*") if f.is_file()
                )
                shutil.rmtree(str(temp_dir), ignore_errors=True)
                if not temp_dir.exists():
                    stats["dirs_removed"] += 1
                    stats["bytes_freed"] += temp_size
                    debug.log("workspace.temp_cleanup", str(temp_dir))
            except Exception:
                log.exception("failed to remove .temp: %s", temp_dir)

        # 2. 删除泄露到主区的 _delegate_* 子目录
        for name in sorted(os.listdir(str(group_dir))):
            if name.startswith("_delegate_"):
                target = group_dir / name
                if target.is_dir():
                    try:
                        target_size = sum(
                            f.stat().st_size for f in target.rglob("*") if f.is_file()
                        )
                        shutil.rmtree(str(target), ignore_errors=True)
                        if not target.exists():
                            stats["dirs_removed"] += 1
                            stats["bytes_freed"] += target_size
                    except Exception:
                        log.exception("failed to remove delegate dir: %s", target)

    if stats["dirs_removed"]:
        log.info(
            "archive temp cleaned: %s — %d groups, %d dirs, %.1f MB freed",
            archive_id,
            stats["groups_scanned"],
            stats["dirs_removed"],
            stats["bytes_freed"] / 1024 / 1024,
        )
    return stats


# ─── 双层工作区(2026-05-03 加,Bug E)─────────────────────────────────
# 历史问题:所有中间产物、helper 沙箱、调试输出都堆在 <archive>/<group>/
# 单一目录里,日积月累工作区像垃圾场,helper 启动时复制 60+ 个文件耗时,主区
# 列文件给用户也是噪声(实测 trace 7e6629 一次 abort 跳过 47 个非 deliverables
# 的中间文件)。
#
# 新设计:
#   <archive>/<group>/                  ← **主工作区**(干净的核心产物)
#   ├── <user 上传的群文件>
#   ├── <经 commit/提升的核心成果(deliverables)>
#   ├── .pause_state.json               ← 持久化的暂停状态
#   └── .temp/                          ← **临时工作区**(当前对话的所有工作)
#       ├── (chat 开始时 sync 自主区的副本)
#       ├── *.c / *.py / *.json         ← Round 2 / helper 产出的所有中间文件
#       ├── _delegate_<tag>_<task>/     ← helper 沙箱(从 .temp 复制)
#       └── ...
#
# 流程:
# 1. chat 开始 → ensure_temp_workspace(main_ws) 把主区当前内容 sync 到 .temp/
#    (保留 .temp 已有的 _delegate_* helper 沙箱,只覆写普通文件)
# 2. Round 2 / helper 全部在 .temp/ 工作
# 3. helper spawn 时,从 .temp/ 复制 → .temp/_delegate_*/(已有的 fork 逻辑)
# 4. chat 成功结束 → promote_to_main(plan.deliverables) 把核心成果回写主区
# 5. 模型可主动 commit_to_main 提升中途阶段成果(新工具)
#
# 这让主区永远只装"对话核心成果",中间过程不污染。

_TEMP_DIRNAME = ".temp"
_PREV_DIRNAME = ".prev"
_SHARED_DIRS = ("_shared", "_helpers_shared")
_rotation_done = False  # 兼容旧测试/调试开关;实际轮转按 main_ws 维度记录
_rotated_main_workspaces: set[str] = set()


def get_main_workspace(archive_id: str = "", group_id: str = "") -> str:
    """主工作区路径(干净区,只放核心成果)。"""
    if archive_id and group_id:
        return str(_get_workspace_root() / archive_id / group_id)
    base = Path(tempfile.gettempdir()) / "chatbot_workspaces"
    return str(base / f"ws_{uuid.uuid4().hex[:12]}")


def get_temp_workspace(main_ws: str) -> str:
    """临时工作区路径(scratch,Round 2 / helper 都在此干活)。"""
    return os.path.join(main_ws, _TEMP_DIRNAME)


def _session_temp_slug(session_tag: str) -> str:
    tag = str(session_tag or "").strip()
    if not tag:
        tag = uuid.uuid4().hex
    return "s_" + hashlib.sha1(tag.encode("utf-8", errors="ignore")).hexdigest()[:16]


def get_session_temp_workspace(main_ws: str, session_tag: str) -> str:
    """Return an isolated temp workspace path for one chat turn."""
    return os.path.join(get_temp_workspace(main_ws), "_sessions", _session_temp_slug(session_tag))


def get_prev_workspace(main_ws: str) -> str:
    """上一轮临时工作区的快照路径(只读历史)。"""
    return os.path.join(main_ws, _PREV_DIRNAME)


def rotate_temp_to_prev(main_ws: str) -> str | None:
    """会话启动时将 .temp/ 轮转为 .prev/ 快照。

    旧的 .prev/ 若存在则先删除。旧的 .temp/ 若存在则 rename 为 .prev/。
    返回新的 .temp/ 路径，或 None(无旧 .temp/ 可轮转)。

    安全：仅当 main_ws 看起来是合法的 workspace 路径时才操作。
    """
    if not main_ws or not os.path.isdir(main_ws):
        return None

    prev_ws = get_prev_workspace(main_ws)
    temp_ws = get_temp_workspace(main_ws)

    # 清理旧 .prev/
    if os.path.isdir(prev_ws):
        shutil.rmtree(prev_ws, ignore_errors=True)
        log.info("rotate_temp_to_prev: removed old .prev/")

    # 轮转: .temp/ → .prev/
    if os.path.isdir(temp_ws):
        os.rename(temp_ws, prev_ws)
        log.info("rotate_temp_to_prev: .temp/ → .prev/")
        debug.log("workspace.rotate", ".temp/ → .prev/")

    # 创建新的空 .temp/
    os.makedirs(temp_ws, exist_ok=True)

    # 2026-05-12 P46-B: 从永久根 sync 上轮 helper 产物到新 .temp/
    # 病因(实测 22:44 trace): 主线程 cwd 在 .temp/, 看不到永久根的上轮 helper 产物.
    # 双层架构原本是 helper push 到 .temp(短期), maintenance promote 到永久根(长期).
    # 但 promote_to_main 只 promote plan.deliverables, helper 中间产物 (helper_*.c)
    # 不被 promote, 跨会话全丢 → 主线程下轮看到空 .temp/ → 重头开始派 helper.
    # 修法 (维护层职责): 启动新 .temp/ 时, 把永久根的 helper_* / _helpers_shared/ 等
    # 已有产物复制到 .temp/, 主线程能直接看到。配合 P46-A (push 时双写永久根) 完整。
    #
    # 跳过列表:
    #   - .temp / .prev: 临时目录自己
    #   - .helper_*_summary 等 . 开头: 上轮元数据, 各自管理
    #   - archive/: archive_stale_artifacts 归档的 14 天前老制品, 不该污染当前会话
    _synced_count = 0
    _SYNC_SKIP_DIRS = {".temp", ".prev", "archive"}
    try:
        for _entry in os.listdir(main_ws):
            if _entry in _SYNC_SKIP_DIRS or _entry.startswith("."):
                continue
            _src = os.path.join(main_ws, _entry)
            _dst = os.path.join(temp_ws, _entry)
            try:
                if os.path.isfile(_src):
                    shutil.copy2(_src, _dst)
                    _synced_count += 1
                elif os.path.isdir(_src):
                    # 子目录 (如 _helpers_shared/, _downloaded_media/) 递归复制
                    shutil.copytree(_src, _dst, dirs_exist_ok=True)
                    _synced_count += 1
            except OSError:
                pass
    except OSError:
        pass
    if _synced_count:
        debug.log(
            "workspace.temp.p46_sync",
            f"P46-B: 从永久根 sync {_synced_count} 个文件/目录到新 .temp/, "
            f"主线程能看到上轮后台工作",
        )

    _write_session_manifest(temp_ws, main_ws)
    debug.log("workspace.temp.fresh", "new .temp/ created with manifest")

    # 注:P58 / P57 v3 / P67 加的 metadata 文件(.helpers_displayed_name.json /
    # .user_fetched_files.json / .rewrite_count.json)都写在 .temp/ 内,随上方
    # .temp/ → .prev/ rotate 自然进入 .prev/(被新空 .temp/ 取代)。无需单独清理。

    return temp_ws


# 会话清单/原子写/编辑计数已抽离到 session_manifest.py(2026-05-20 重构);re-export 兼容。
from app.llm.tools.session_manifest import (  # noqa: E402,F401
    _atomic_write_json,
    _write_session_manifest,
    _peek_edit_count,
)


def fetch_to_temp(
    main_ws: str,
    temp_ws: str,
    *,
    paths: list[str],
    source: str = "main",
) -> tuple[list[str], list[str]]:
    """将文件从永久区或历史快照复制到当前临时工作区。

    三层隔离模型的核心原语：不能直接访问永久区和 .prev/，必须通过此函数复制。

    Args:
        main_ws: 主工作区路径
        temp_ws: 当前临时工作区路径
        paths: 要复制的相对路径列表
        source: "main"(永久区) 或 "prev"(上一轮 .prev/ 快照)

    Returns:
        (copied, skipped): 成功复制的路径 / 跳过的路径(不存在/越界)
    """
    if not main_ws or not temp_ws or not paths:
        return [], list(paths or [])

    if source == "prev":
        src_root = get_prev_workspace(main_ws)
    else:
        src_root = main_ws

    if not os.path.isdir(src_root):
        return [], list(paths)

    copied: list[str] = []
    skipped: list[str] = []
    remap = load_displayed_name_remap(main_ws) if source == "main" else {}

    for p in paths:
        if not p:
            continue
        requested_path = p
        normalized_request = str(p).replace("\\", "/").strip().strip('"').strip("'")
        if normalized_request in {"", ".", "./"}:
            skipped.append(requested_path)
            continue
        try:
            src = _safe_resolve(src_root, p)
        except ValueError:
            skipped.append(requested_path)
            continue
        if not os.path.exists(src) and remap:
            remapped = remap.get(p)
            if not remapped and os.path.basename(p) == p:
                for main_name, displayed_name in remap.items():
                    if displayed_name == p:
                        remapped = main_name
                        break
            if isinstance(remapped, str) and remapped:
                try:
                    remapped_src = _safe_resolve(src_root, remapped)
                    if os.path.exists(remapped_src):
                        src = remapped_src
                except ValueError:
                    pass
        if not os.path.exists(src):
            skipped.append(requested_path)
            continue
        rel = os.path.relpath(src, src_root).replace("\\", "/")
        if source == "prev" and rel.startswith("_delegate_"):
            skipped.append(requested_path)
            continue
        dst_rel = requested_path.replace("\\", "/") if source == "main" else rel
        try:
            _safe_resolve(temp_ws, dst_rel)
        except ValueError:
            dst_rel = os.path.basename(requested_path)
        dst = os.path.join(temp_ws, dst_rel)
        try:
            src_abs = os.path.abspath(src)
            dst_abs = os.path.abspath(dst)
            temp_abs = os.path.abspath(temp_ws)
            if os.path.isdir(src_abs):
                if os.path.samefile(src_abs, os.path.abspath(src_root)):
                    skipped.append(requested_path)
                    continue
                try:
                    if os.path.commonpath([src_abs, temp_abs]) == src_abs:
                        skipped.append(requested_path)
                        continue
                except ValueError:
                    pass
                try:
                    if os.path.commonpath([src_abs, dst_abs]) == src_abs:
                        skipped.append(requested_path)
                        continue
                except ValueError:
                    pass
            os.makedirs(os.path.dirname(dst) or temp_ws, exist_ok=True)
            if os.path.isfile(src):
                shutil.copy2(src, dst)
            elif os.path.isdir(src):
                def _ignore_internal(_dir, names):
                    return [
                        name for name in names
                        if name.startswith("_delegate_")
                        or name in {".temp", ".prev", ".write_backups", "__pycache__"}
                    ]
                shutil.copytree(src, dst, dirs_exist_ok=True, ignore=_ignore_internal)
            copied.append(dst_rel)
        except OSError as e:
            log.warning("fetch_to_temp: failed to copy %s: %s", dst_rel, e)
            skipped.append(requested_path)

    if copied:
        debug.log(
            "workspace.fetch_to_temp",
            f"source={source}: copied {len(copied)}/{len(paths)} files",
            {"copied": copied[:10], "skipped": skipped[:5]},
        )
    return copied, skipped


def ensure_temp_workspace(
    main_ws: str,
    *,
    force_resync: bool = False,
    session_tag: str = "",
    isolate_session: bool = False,
) -> str:
    """确保 .temp/ 临时工作区存在（v2 三层隔离架构）。

    会话首次调用时自动 rotate: .temp/ → .prev/ → 新建 .temp/。
    不再 auto-sync 主区文件到 temp — 改用 fetch_to_temp() 按需复制。
    force_resync: 强制重置 temp（删除全部内容重建，保留 _delegate_*）。
    session_tag: 2026-05-07 Bug 4 fix — 用于跨任务 stale _shared/ 检测。
        每次会话传入不同的 tag（如 user_message[:80] 的 hash），
        若与 _shared/ 里记录的 tag 不同，跳过复制（旧会话残留）。

    Args:
        main_ws: 主工作区路径
        force_resync: True 时强制重置 temp
        session_tag: 当前会话的短标识，用于跨任务检测
        isolate_session: True 时在 .temp/_sessions/<hash>/ 下创建本轮工作区，
            避免同群多用户并发时把“当前产物”混到同一个 .temp 根目录。

    Returns:
        temp_ws 路径
    """
    if not main_ws:
        return ""
    os.makedirs(main_ws, exist_ok=True)

    temp_root = get_temp_workspace(main_ws)
    temp_ws = temp_root

    # 每次进程启动(新会话)首次调用时强制轮转: .temp/ → .prev/ → 新建 .temp/
    # 根因修复: 旧逻辑 `if not os.path.isdir(temp_root)` 只在 .temp/ 不存在时轮转,
    # 若上次会话残留 .temp/(含 PDF/DOCX 等污染文件),新会话直接复用 → helper 被污染 → 卡 loop。
    # 现在用进程级 flag 保证每次新会话都拿到干净的 .temp/。
    global _rotation_done
    _main_key = os.path.realpath(main_ws)
    if not _rotation_done:
        _rotated_main_workspaces.clear()
        _rotation_done = True
    if _main_key not in _rotated_main_workspaces:
        rotate_temp_to_prev(main_ws)
        _rotated_main_workspaces.add(_main_key)
    else:
        os.makedirs(temp_root, exist_ok=True)

    if isolate_session:
        temp_ws = get_session_temp_workspace(main_ws, session_tag)
        os.makedirs(temp_ws, exist_ok=True)
        debug.log("workspace.temp.session", f"isolated temp workspace: {temp_ws}")

    if force_resync:
        # 仅删除非 _delegate_* 内容
        for entry in os.listdir(temp_ws):
            if entry.startswith("_delegate_"):
                continue
            target = os.path.join(temp_ws, entry)
            try:
                if os.path.isfile(target) or os.path.islink(target):
                    os.remove(target)
                elif os.path.isdir(target):
                    shutil.rmtree(target, ignore_errors=True)
            except OSError:
                pass
        _write_session_manifest(temp_ws, main_ws)

    # 2026-05-10 Patch 77: 清理过期的 .helper_*_full_report.txt
    # 病因(trace debug_20260510_134607):log 显示 10+ 个 .helper_xxx_full_report.txt
    # 跨 chat session 持续累积(`pre-existing files` 从 0 涨到 56),主线程在新会话
    # 反复试图 read_file 这些上次 helper 的摘要 → P67 重读拒绝触发 30 次。
    # 这些是 P32 加的 helper 内部摘要,本会话用完该清理(下次会话 helper 会写新的)。
    # 不动 helper sandbox 内的 .helper_summary.txt(供 resume 用)。
    #
    # 策略:启动新 chat session 时清理 main_ws 内 mtime > 1 小时的 .helper_*_full_report.txt
    # 1 小时是经验阈值:helper 完成 30 min 后基本不会再被主线程查;留给"刚 abort 几分钟内 resume"。
    try:
        _now = time.time()
        _STALE_HELPER_REPORT_AGE = 3600.0  # 1 小时
        _cleaned = 0
        for entry in os.listdir(temp_ws):
            if not entry.startswith(".helper_") or not entry.endswith("_full_report.txt"):
                continue
            target = os.path.join(temp_ws, entry)
            try:
                if not os.path.isfile(target):
                    continue
                _age = _now - os.path.getmtime(target)
                if _age > _STALE_HELPER_REPORT_AGE:
                    os.remove(target)
                    _cleaned += 1
            except OSError:
                pass
        if _cleaned > 0:
            debug.log(
                "workspace.helper_reports.cleaned",
                f"P77: 清理了 {_cleaned} 个过期 helper 摘要(.helper_*_full_report.txt > 1h)",
            )
    except OSError:
        pass  # 清理失败不阻塞会话启动

    # v2 default: seed .temp from the permanent main workspace so helpers start with
    # the current durable artifacts, while still keeping helper sandboxes isolated.
    staged_files = 0
    skipped_stage_artifacts = 0
    for entry in os.listdir(main_ws):
        if entry in (_TEMP_DIRNAME, _PREV_DIRNAME, "archive") or entry.startswith("_delegate_"):
            continue
        if entry in _SHARED_DIRS or entry == "_downloaded_media":
            continue
        src = os.path.join(main_ws, entry)
        dst = os.path.join(temp_ws, entry)
        if not os.path.isfile(src):
            continue
        try:
            size = os.path.getsize(src)
        except OSError:
            continue
        if not _should_copy_to_helper(src, size):
            skipped_stage_artifacts += 1
            continue
        try:
            if os.path.exists(dst):
                src_st = os.stat(src)
                dst_st = os.stat(dst)
                if dst_st.st_size == src_st.st_size and abs(dst_st.st_mtime - src_st.st_mtime) < 1.0:
                    continue
            shutil.copy2(src, dst)
            staged_files += 1
        except OSError:
            pass
    if staged_files or skipped_stage_artifacts:
        debug.log(
            "workspace.temp.seeded_from_main",
            f"staged {staged_files} durable file(s) into .temp; "
            f"skipped {skipped_stage_artifacts} filtered artifact(s)",
        )

    # v2: shared dirs are copied separately with session-tag stale checks.
    # 2026-05-07: 清理 main_ws/_helpers_shared/ 中上一会话残留的 helper 共享文件。
    # 若未做此清理，残留文件会被复制到新 .temp/ → helper merge 时合并回主区 → 污染。
    # 2026-05-08 Fix 3: 删除前先保存到 .prev/，防止"工作区空重做"。
    # main_ws/_helpers_shared/ 在 helper 完成时被 merge 更新，
    # 但 .temp/_helpers_shared/ 只在 session 开始复制一次 → .prev/ 缺这些新文件。
    # 删除前先 merge 到 .prev/_helpers_shared/，fetch_to_temp(source='prev') 可找回。
    _hsh_path = os.path.join(main_ws, "_helpers_shared")
    if os.path.isdir(_hsh_path) and session_tag:
        _hsh_tag_file = os.path.join(_hsh_path, ".session_tag")
        _hsh_prev_tag = ""
        try:
            if os.path.isfile(_hsh_tag_file):
                with open(_hsh_tag_file, "r", encoding="utf-8") as _tf:
                    _hsh_prev_tag = _tf.read().strip()
        except OSError:
            pass
        if _hsh_prev_tag != session_tag:
            # 保存到 .prev/ 再清空（防止 work 丢失）
            _prev_ws = get_prev_workspace(main_ws)
            if _prev_ws and os.path.isdir(_prev_ws):
                _prev_hsh = os.path.join(_prev_ws, "_helpers_shared")
                os.makedirs(_prev_hsh, exist_ok=True)
                _preserved = 0
                for _f in os.listdir(_hsh_path):
                    if _f == ".session_tag":
                        continue
                    _fp = os.path.join(_hsh_path, _f)
                    _dp = os.path.join(_prev_hsh, _f)
                    try:
                        if os.path.isfile(_fp) or os.path.islink(_fp):
                            shutil.copy2(_fp, _dp)
                            _preserved += 1
                        elif os.path.isdir(_fp):
                            if os.path.isdir(_dp):
                                shutil.rmtree(_dp, ignore_errors=True)
                            shutil.copytree(_fp, _dp)
                            _preserved += sum(1 for _ in Path(_dp).rglob("*") if _.is_file())
                    except OSError:
                        pass
                if _preserved:
                    debug.log(
                        "workspace.helpers_shared.preserved",
                        f"preserved {_preserved} file(s) from _helpers_shared/ to .prev/",
                    )
            # tag 不匹配或无 tag → 清空旧文件（防止下一段代码复制到新 .temp/）
            for _f in os.listdir(_hsh_path):
                if _f == ".session_tag":
                    continue
                _fp = os.path.join(_hsh_path, _f)
                try:
                    if os.path.isfile(_fp) or os.path.islink(_fp):
                        os.remove(_fp)
                    elif os.path.isdir(_fp):
                        shutil.rmtree(_fp, ignore_errors=True)
                except OSError:
                    pass
            # 写入当前 session_tag
            try:
                with open(_hsh_tag_file, "w", encoding="utf-8") as _tf:
                    _tf.write(session_tag)
            except OSError:
                pass
    copied_shared = 0
    discarded_shared = 0
    for entry in os.listdir(main_ws):
        if entry in (_TEMP_DIRNAME, _PREV_DIRNAME, "archive"):
            continue
        src = os.path.join(main_ws, entry)
        dst = os.path.join(temp_ws, entry)
        if os.path.isdir(src) and entry in _SHARED_DIRS:
            # 2026-05-07 Bug 4: 跨任务 stale _shared/ 检测
            # 主区 _shared/ 可能残留上次会话的脚手架代码（如上次写 AVL tree，
            # 本次写 email），带入 helper 会污染上下文和编译。
            _tag_file = os.path.join(src, ".session_tag")
            _prev_tag = ""
            try:
                if os.path.isfile(_tag_file):
                    with open(_tag_file, "r", encoding="utf-8") as _tf:
                        _prev_tag = _tf.read().strip()
            except OSError:
                pass
            if _prev_tag and session_tag and _prev_tag != session_tag:
                discarded_shared += 1
                debug.log(
                    "workspace.shared.discarded",
                    f"task mismatch, removed {entry}/ "
                    f"(prev_tag={_prev_tag[:20]}... cur_tag={session_tag[:20]}...)",
                )
                continue
            # 2026-05-07: _helpers_shared 是当前会话 transient 共享区,
            # main_ws 中的副本来自上次 helper 合并。若没有 session_tag,
            # 则一定是旧的代码残留(旧版未写 tag 或手动残留),不应带进新会话。
            if entry == "_helpers_shared" and not _prev_tag:
                discarded_shared += 1
                debug.log(
                    "workspace.shared.discarded",
                    f"helpers_shared without session_tag, removed (stale carryover)",
                )
                continue
            try:
                shutil.copytree(src, dst, dirs_exist_ok=True)
                copied_shared += 1
                # 写入当前 session_tag 供下次检测
                if session_tag:
                    try:
                        with open(os.path.join(dst, ".session_tag"), "w", encoding="utf-8") as _tf:
                            _tf.write(session_tag)
                    except OSError:
                        pass
            except OSError as e:
                log.warning("ensure_temp_workspace: failed to copy shared dir %s: %s", entry, e)

    if copied_shared:
        debug.log(
            "workspace.temp.shared",
            f"copied {copied_shared} shared dir(s) to .temp"
            + (f" (discarded {discarded_shared} stale)" if discarded_shared else ""),
        )

    # ── 入站媒体: 把 main_ws/_downloaded_media/ 复制到 .temp/ ──
    # bridge 在收到 QQ 图片/语音/视频时立即下载到 main_ws/_downloaded_media/,
    # 但 LLM 的工作区在 .temp/ 内。这里把已下载的媒体复制到 .temp/ 供 LLM 使用。
    _media_src = os.path.join(main_ws, "_downloaded_media")
    if os.path.isdir(_media_src):
        _media_dst = os.path.join(temp_ws, "_downloaded_media")
        _copied_media = 0
        try:
            os.makedirs(_media_dst, exist_ok=True)
            for _f in os.listdir(_media_src):
                _fp = os.path.join(_media_src, _f)
                _dp = os.path.join(_media_dst, _f)
                if os.path.isfile(_fp) and not os.path.exists(_dp):
                    shutil.copy2(_fp, _dp)
                    _copied_media += 1
            if _copied_media:
                debug.log(
                    "workspace.temp.media",
                    f"copied {_copied_media} downloaded media file(s) to .temp/_downloaded_media/",
                )
        except OSError:
            pass

    return temp_ws


def promote_to_main(
    main_ws: str, temp_ws: str, deliverables: list[str],
) -> tuple[list[str], list[str], dict[str, str]]:
    """从 temp 提升 deliverable 文件到主工作区。

    L5-2 (2026-05-09): 有限前缀匹配解决"plan 写裸名,主区是带前缀名"的 mismatch。
    旧函数只做精确路径匹配,helper 产出 `paper_pptx_xlsx_paper.docx` 时
    plan 写 `paper.docx` → 全 missing。

    2026-06-08: 不再扫描 `_delegate_*` helper 内部沙箱,也不做 `*{basename}` 首个
    命中的宽泛兜底。helper 产物应先通过 copyback 进入 temp/main 可见区; promote 只
    提升当前 temp 中的显式文件或唯一的 `<prefix>_<basename>` 文件。歧义/缺失作为事实
    记录到 skipped,交给 LLM 重新决定。

    Args:
        main_ws: 主工作区路径
        temp_ws: 临时工作区路径 (.temp/)
        deliverables: deliverable 文件名列表(可能裸名,如 "paper.docx")

    Returns:
        (promoted, skipped, name_remap)
        name_remap: deliverable_name → actual_promoted_filename 映射
    """
    import fnmatch

    if not main_ws or not temp_ws:
        return [], list(deliverables), {}
    promoted: list[str] = []
    skipped: list[str] = []
    name_remap: dict[str, str] = {}

    # 构建 temp 工作区文件索引。跳过 helper 内部沙箱;这些不是可直接交付边界。
    _temp_files: dict[str, str] = {}  # basename → full path
    for dirpath, dirnames, filenames in os.walk(temp_ws):
        rel_dir = os.path.relpath(dirpath, temp_ws)
        if rel_dir == ".":
            pass
        elif rel_dir.startswith("_delegate_"):
            dirnames.clear()  # 不深层递归 delegate 子目录
            continue
        dirnames[:] = [
            d for d in dirnames
            if not d.startswith("_delegate_") and d not in {"__pycache__", ".git", ".pytest_cache"}
        ]
        for fname in filenames:
            if fname.startswith(".") or fname in _METADATA_BASENAMES or fname.endswith(_NON_DELIVERABLE_SUFFIXES):
                continue
            full = os.path.join(dirpath, fname)
            _temp_files[fname] = full
            # 也按相对于 temp 的路径索引
            _temp_files[os.path.join(rel_dir, fname).replace("\\", "/")] = full

    for dlv in deliverables:
        if not dlv:
            continue
        src = None
        matched_key = None

        # Level 1: 精确匹配 — basename 或相对路径原样找到
        if dlv in _temp_files:
            src = _temp_files[dlv]
            matched_key = dlv
        elif os.path.basename(dlv) in _temp_files:
            src = _temp_files[os.path.basename(dlv)]
            matched_key = os.path.basename(dlv)

        # Level 2: 有限前缀匹配 — 只接受唯一 `<prefix>_<deliverable>` 文件。
        if src is None:
            dlv_basename = os.path.basename(dlv)
            prefixed_matches = [
                (fname, fpath)
                for fname, fpath in _temp_files.items()
                if "/" not in fname and fnmatch.fnmatch(fname, f"*_{dlv_basename}")
            ]
            # Deduplicate basename/path dual index matches.
            unique_matches: dict[str, str] = {}
            for fname, fpath in prefixed_matches:
                unique_matches[fpath] = fname
            if len(unique_matches) == 1:
                fpath, fname = next(iter(unique_matches.items()))
                src = fpath
                matched_key = fname
            elif len(unique_matches) > 1:
                debug.log(
                    "workspace.promote.ambiguous",
                    f"deliverable {dlv!r} matched multiple prefixed files; leaving for LLM decision: "
                    f"{list(unique_matches.values())[:10]}",
                )

        if src is None or not os.path.isfile(src):
            skipped.append(dlv)
            continue

        # 目标: 用实际匹配到的文件名放到 main_ws
        actual_name = os.path.basename(src)
        dst = os.path.join(main_ws, actual_name)
        try:
            os.makedirs(main_ws, exist_ok=True)
            shutil.copy2(src, dst)
            promoted.append(actual_name)
            name_remap[dlv] = actual_name
        except OSError as e:
            log.warning("promote_to_main: failed to copy %s (matched from %s) → main: %s", dlv, actual_name, e)
            skipped.append(dlv)

    if promoted:
        debug.log(
            "workspace.promote",
            f".temp → main: promoted {len(promoted)} file(s): {promoted[:5]}"
            + (f" ... ({len(promoted)-5} more)" if len(promoted) > 5 else ""),
        )
    if skipped:
        debug.log(
            "workspace.promote.skipped",
            f"skipped {len(skipped)} deliverable(s): {skipped[:5]}"
            + (f" ... ({len(skipped)-5} more)" if len(skipped) > 5 else ""),
        )

    # 2026-05-10 Patch 58: 累积 main_name → helper_name(declared)映射,推送时
    # 用这个 mapping 去掉内部前缀,不暴露智能体内部命名结构。
    if name_remap:
        _displayed: dict[str, str] = {}
        for declared, actual in name_remap.items():
            if declared and actual and declared != actual:
                _displayed[actual] = declared  # 反向:main_name → helper-declared
        if _displayed:
            try:
                update_displayed_name_remap(main_ws, _displayed)
            except Exception:
                pass  # 写失败不阻塞

    return promoted, skipped, name_remap


# ─── 2026-05-10 Patch 58: helpers 显示名映射(用户友好命名) ───
# 用户反馈:"推送时有时候会带有内部维护的文件名,实际不需要,会暴露智能体内部结构"
# 设计:
#   - 主区文件保持带 task_id 前缀(内部交互稳定,避免 helper 间命名冲突)
#   - 推送给用户(QQ 群)的 displayed name **去掉内部前缀**(用户友好,不暴露结构)
#   - URL 仍指向带前缀的主区文件(下载链路精确)
# 实施:helper promote 完成时,把 main_name → helper_name 反向 mapping 累积写到
# main_ws 的 metadata 文件 `.helpers_displayed_name.json`。orchestrator 推送层
# 读取这个文件构造 displayed name。
_DISPLAYED_NAME_REMAP_FILE = ".helpers_displayed_name.json"

# 2026-05-10 Patch 74: metadata 文件原子写入(防 race 半文件 + 多 helper 并发覆盖)
#
# 病因(防御性):多个 helper 并行完成时同时调 update_displayed_name_remap 写同一
# `.helpers_displayed_name.json`。read-modify-write 路径下:
#   helper A: read existing={}
#   helper B: read existing={}
#   helper A: write {x: y}
#   helper B: write {z: w}    ← 覆盖 A
# 结果 A 的 mapping 丢,推送时 x 仍带前缀给用户。
#
# 完整防 race 需 file lock(跨平台复杂),这里用 atomic write 处理:
#   - tmp + os.replace 保证文件永远是完整 JSON(中间崩溃无半文件)
#   - read-modify-write race 仍可能丢条目,但**频率低**(helper 完成时间通常错开)
#   - 文件不损坏(损失映射 vs 文件不可读 — 前者影响小得多)



# 2026-05-10 Patch 57 v3: 用户主动 fetch 的群文件 / 用户上传的文件 basenames
# 病因(trace f973df3770544567):主线程调 fetch_group_file 把
# `update_paper_paper_database_index_paper.docx` 拉到主区。下游 helper(embed_charts)
# fork 时,P57 的 _looks_like_helper_artifact 把这个 docx 名识别为 helper 链产物
# (因为含多段下划线分隔),跳过不复制 → embed_charts inspect_file 失败。
# P57 之前的 allowed_prefixes 只覆盖当前会话的 helper task_id,没覆盖"用户拉来的文件"。
# 修法:fetch_group_file 完成时把 basename 写入 metadata,fork 时读取作为完整名白名单。
_USER_FETCHED_FILES_RECORD = ".user_fetched_files.json"


def add_user_fetched_basename(main_ws: str, basename: str) -> None:
    """记录一个用户 fetch / 上传的文件 basename,fork 时无条件放行。"""
    if not main_ws or not basename:
        return
    path = os.path.join(main_ws, _USER_FETCHED_FILES_RECORD)
    existing: list[str] = []
    try:
        if os.path.isfile(path):
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, list):
                    existing = data
    except (OSError, json.JSONDecodeError):
        existing = []
    if basename not in existing:
        existing.append(basename)
        # P74: atomic write
        _atomic_write_json(path, existing)


def load_user_fetched_basenames(main_ws: str) -> frozenset[str]:
    if not main_ws:
        return frozenset()
    path = os.path.join(main_ws, _USER_FETCHED_FILES_RECORD)
    try:
        if os.path.isfile(path):
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, list):
                    return frozenset(str(x) for x in data if x)
    except (OSError, json.JSONDecodeError):
        pass
    return frozenset()


def update_displayed_name_remap(main_ws: str, mapping: dict[str, str]) -> None:
    """累积写入 main_name → user_friendly_name 映射到 main_ws metadata 文件。

    每次 helper 完成 promote 后调用一次,合并到现有 mapping。
    `.` 开头文件名会被 _looks_like_helper_artifact 跳过(不会复制给下游 helper),
    也被 list_generated_files 跳过(不会推送给用户)。
    """
    if not main_ws or not mapping:
        return
    path = os.path.join(main_ws, _DISPLAYED_NAME_REMAP_FILE)
    existing: dict[str, str] = {}
    try:
        if os.path.isfile(path):
            with open(path, "r", encoding="utf-8") as f:
                existing = json.load(f) or {}
                if not isinstance(existing, dict):
                    existing = {}
    except (OSError, json.JSONDecodeError):
        existing = {}
    existing.update(mapping)
    # P74: atomic write 防 race 半文件
    _atomic_write_json(path, existing)


def load_displayed_name_remap(main_ws: str) -> dict[str, str]:
    """读 main_ws 的 displayed name 映射(orchestrator 推送层用)。"""
    if not main_ws:
        return {}
    path = os.path.join(main_ws, _DISPLAYED_NAME_REMAP_FILE)
    try:
        if os.path.isfile(path):
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f) or {}
                return data if isinstance(data, dict) else {}
    except (OSError, json.JSONDecodeError):
        pass
    return {}


# 2026-05-09 Patch 34: 多文件自动打 zip
#
# 病因:用户场景中产物可能很多(如同时给 docx + pptx + 数据 + 图表脚本 + readme),
# 一个一个推 QQ 群文件刷屏体验差。> 3 个文件时打成 zip 让用户一次拿走更顺。
# 设计:
#   - 仅在 main_ws / archive_id / group_id 都齐全时调用
#   - zip 内文件名扁平化(只取 basename),避免目录穿越
#   - zip 自身放在 main_ws 根目录,命名 deliverables_{trace8}_{ts}.zip,可被
#     /v1/chat/files/.../<zipname> 端点直接下载
#   - 失败返回 None,调用方降级到原始文件清单
def create_zip_archive(
    main_ws: str,
    file_basenames: list[str],
    *,
    trace_id: str = "",
    archive_id: str = "",
    group_id: str = "",
) -> tuple[str, str, str] | None:
    """把 main_ws 下的多个文件打成单个 zip,返回 (zip_basename, url, full_path)。

    Args:
        main_ws: 主工作区绝对路径
        file_basenames: 要打包的文件 basename 列表(必须已 promote 到 main_ws 根目录)
        trace_id: 用于命名 zip(取前 8 字符)
        archive_id, group_id: 用于构造下载 URL,留空则 url=""

    Returns:
        (zip_basename, url, full_path) 或 None(失败/无可打包文件)
    """
    import zipfile
    import time

    if not main_ws or not file_basenames:
        return None
    if not os.path.isdir(main_ws):
        log.warning("create_zip_archive: main_ws 不存在 %s", main_ws)
        return None

    # 过滤实际存在的文件
    _existing: list[tuple[str, str]] = []  # (basename, full_path)
    for _bn in file_basenames:
        # 防御:basename 必须不含路径分隔符,且文件确实在 main_ws 根
        _bn_clean = os.path.basename(_bn)
        if _bn_clean != _bn or not _bn_clean:
            log.warning("create_zip_archive: 跳过非纯 basename: %r", _bn)
            continue
        _full = os.path.join(main_ws, _bn_clean)
        if not os.path.isfile(_full):
            log.warning("create_zip_archive: 文件不存在跳过: %s", _full)
            continue
        _existing.append((_bn_clean, _full))

    if not _existing:
        return None

    # zip 命名:deliverables_{trace8}_{HHMMSS}.zip
    _trace8 = (trace_id or "x" * 8)[:8]
    _ts = time.strftime("%H%M%S")
    _zip_name = f"deliverables_{_trace8}_{_ts}.zip"
    _zip_full = os.path.join(main_ws, _zip_name)

    try:
        with zipfile.ZipFile(_zip_full, "w", zipfile.ZIP_DEFLATED, compresslevel=6) as zf:
            for _bn, _full in _existing:
                zf.write(_full, arcname=_bn)
    except (OSError, zipfile.BadZipFile) as e:
        log.warning("create_zip_archive: 打包失败 %s: %s", _zip_name, e)
        try:
            if os.path.isfile(_zip_full):
                os.unlink(_zip_full)
        except OSError:
            pass
        return None

    # URL(若 archive_id/group_id 给了)
    _url = ""
    if archive_id and group_id:
        _url = f"/v1/chat/files/{archive_id}/{group_id}/{_zip_name}"

    debug.log(
        "workspace.zip.created",
        f"打包 {len(_existing)} 个 deliverable 到 {_zip_name} "
        f"({os.path.getsize(_zip_full):,} bytes)",
        {"files": [bn for bn, _ in _existing], "zip": _zip_name},
    )

    # 2026-05-10 Patch 58 v3: zip 加 displayed name mapping
    # 主区名 deliverables_{trace8}_{HHMMSS}.zip 暴露 trace_id 内部结构,
    # 推送给用户应显示友好名"打包文件_HHMMSS.zip"(去 trace_id,保留时间区分)。
    try:
        update_displayed_name_remap(
            main_ws,
            {_zip_name: f"打包文件_{_ts}.zip"},
        )
    except Exception:
        pass  # 写失败不阻断 zip 返回

    return (_zip_name, _url, _zip_full)


def register_workspace(group_key: str, ws_dir: str) -> None:
    if not group_key or not ws_dir:
        return
    entries = [p for p in _workspace_registry.get(group_key, []) if p != ws_dir]
    entries.append(ws_dir)
    _workspace_registry[group_key] = entries


def unregister_workspace(group_key: str, ws_dir: str | None = None) -> None:
    if not group_key:
        return
    if ws_dir is None:
        _workspace_registry.pop(group_key, None)
        return
    entries = [p for p in _workspace_registry.get(group_key, []) if p != ws_dir]
    if entries:
        _workspace_registry[group_key] = entries
    else:
        _workspace_registry.pop(group_key, None)


def get_workspace(group_key: str) -> str | None:
    entries = _workspace_registry.get(group_key) or []
    return entries[-1] if entries else None


def get_registered_workspaces(group_key: str) -> list[str]:
    entries = _workspace_registry.get(group_key) or []
    return list(reversed(entries))


def workspace_token(ws_dir: str) -> str:
    if not ws_dir:
        return ""
    norm = os.path.normcase(os.path.abspath(ws_dir))
    return hashlib.sha1(norm.encode("utf-8", errors="ignore")).hexdigest()[:16]


def token_matches_workspace(ws_dir: str, token: str) -> bool:
    token = str(token or "").strip()
    return bool(token and ws_dir and workspace_token(ws_dir) == token)


# 已知框架/工具自动生成的中间产物（非用户可见文件）
# 2026-05-10 Patch 58: .helpers_displayed_name.json 是 P58 推送层 metadata,
# 不能被 list_generated_files 列出(否则会被推给用户)
_ARTIFACT_PREFIXES = ("fontlist-", ".matplotlib")
_METADATA_BASENAMES = frozenset({
    ".helpers_displayed_name.json",
    ".helper_completions.json",
    ".read_history.json", ".todos.json", ".edit_history.json",
    ".session_tag",
})
_NON_DELIVERABLE_SUFFIXES = (
    "_call_count.json", "_history.json", "_count.json", "_rewrite_count.json",
    ".todos_call_count.json",
)
_NON_DELIVERABLE_EXTS = {".o", ".obj", ".pyc", ".pyo", ".lib", ".a", ".so"}
_NON_DELIVERABLE_BASENAMES = {
    ".session_tag",
    "bench_out.txt",
    "build_output.txt",
    "compile_log.txt",
}


def _suggest_similar_files(
    ws_dir: str, target_path: str, *, limit: int = 5, max_scan: int = 200,
) -> list[dict]:
    """工作区内文件 fuzzy 匹配,返回 top-N 候选 + 评分。

    2026-05-11 新增:helper / 主线程在 read_file / fetch_to_temp / bash 引用
    路径找不到文件时,主动返回相似名清单,避免 helper 凭印象瞎猜。

    匹配策略(三层叠加,分数从高到低):
      1. basename 精确匹配 (path='_shared/foo.h' 找不到 → 找 basename='foo.h' 的文件)
      2. basename 子串包含
      3. token-overlap (按 `_`、`-`、`.` 切 token,计算 Jaccard)
    """
    if not ws_dir or not os.path.isdir(ws_dir):
        return []
    target_basename = os.path.basename(target_path) or target_path
    target_tokens = set(_tokenize_path(target_basename))

    # 扫描工作区(包含 _shared/_helpers_shared 一级子目录,不深入 _delegate_*)
    candidates: list[tuple[str, int, str]] = []  # (rel_path, score, reason)
    try:
        for root, dirs, files in os.walk(ws_dir):
            # 不深入隐藏 / helper 沙箱 / .temp 内的 _delegate_*
            dirs[:] = [
                d for d in dirs
                if not d.startswith((".", "_delegate_")) or d in ("_shared", "_helpers_shared")
            ]
            for name in files:
                if name.startswith("."):
                    continue
                full = os.path.join(root, name)
                try:
                    rel = os.path.relpath(full, ws_dir).replace("\\", "/")
                except ValueError:
                    continue
                score = 0
                reason = ""
                # 完整路径精确匹配优先(_shared/common.h 而非根 common.h)
                if rel == target_path or rel == target_path.replace("\\", "/"):
                    score = 110
                    reason = "exact path match"
                elif name == target_basename:
                    score = 100
                    reason = "basename exact match"
                elif name.endswith(target_basename) and "/" in target_basename:
                    score = 90
                    reason = f"endswith {target_basename!r}"
                elif target_basename in name:
                    score = 70
                    reason = f"basename {target_basename!r} contained"
                elif name in target_basename:
                    score = 60
                    reason = f"file basename contained in target"
                else:
                    # token Jaccard
                    file_tokens = set(_tokenize_path(name))
                    if file_tokens and target_tokens:
                        inter = len(file_tokens & target_tokens)
                        union = len(file_tokens | target_tokens)
                        jac = inter / union if union else 0.0
                        if jac >= 0.5:
                            score = int(50 * jac)
                            reason = f"token overlap {inter}/{union}"
                if score > 0:
                    candidates.append((rel, score, reason))
                # max_scan 限制(超大工作区不全扫)
                if len(candidates) > max_scan:
                    break
            if len(candidates) > max_scan:
                break
    except OSError:
        return []
    # 按分数倒序,取 top
    candidates.sort(key=lambda x: -x[1])
    return [
        {"path": rel, "score": score, "reason": reason}
        for rel, score, reason in candidates[:limit]
    ]


# 代码大纲/文本读取工具已抽离到 code_outline.py(2026-05-20 重构);re-export 兼容。
from app.llm.tools.code_outline import (  # noqa: E402,F401
    _smart_decode,
    _iter_text_lines,
    _truncate_head_tail,
    _tokenize_path,
    _find_function_end,
    _find_callees_in_file,
    _extract_c_outline,
    _extract_generic_outline,
)


# 工作区结果/提示文本 helper 已抽离到 workspace_text.py(2026-05-20 重构);re-export 兼容。
from app.llm.tools.workspace_text import (  # noqa: E402,F401
    _same_pattern,
    _extract_test_summary,
    _helper_missing_file_fetch_hint,
    _structured_read_file_rejection,
)


def _file_not_found_response(
    ws_dir: str, path: str, *, action_hint: str = "read_file",
) -> dict:
    """统一构造 file_not_found 响应,带 fuzzy 建议。

    替代散落各处的 `return {"ok": False, "error": f"file not found: {path}"}`。
    LLM 看到这个响应可立即决策(用建议的实际名重试,或停手报告无此文件)。

    2026-05-12 P42: 当 fuzzy 找到 score >= 95 的极强匹配时, 返回额外字段
    `_auto_redirect_path` 指示调用方应该自动重定向到该路径。
    病因(实测 21:05 trace): 主线程 read_file 用 `_helpers_shared/helper_X/Y`
    错路径,系统已给提示但主线程反复试错(31 次 fail). 修法: 极强匹配时直接
    redirect, 避免 LLM 反复试错浪费 token + 时间。
    """
    from app.core.filesystem import PathZone, classify_path

    classification = classify_path(path)
    if classification.zone == PathZone.STAGED_ROOT:
        return {
            "ok": False,
            "error": "path_is_directory_or_missing_staged_root",
            "path": path,
            "path_zone": str(classification.zone),
            "message": (
                "`_env/` is a staged directory area, not a readable file. Choose a concrete file under `_env/...`, "
                "or use environment directory tools for the real project. If `_env/` has not been created yet, first "
                "fetch a project file or ask a helper to stage explicit expected outputs.\n\n"
                "`_env/` 是暂存目录，不是文件；先定位具体文件或使用 env_* 项目工具。"
            ),
            "suggested_tools": ["env_list_tree", "env_search", "env_read", "workspace locate"],
            "_next_action_instruction": (
                "Use env_list_tree/env_search/env_read for project directories, "
                "or locate/read a concrete workspace file such as `_env/project_inventory.md` when it exists.\n\n"
                "项目目录用 env_*；工作区读取需先定位到具体文件。"
            ),
            "_suggestions": [],
        }
    if classification.is_directory_hint:
        return {
            "ok": False,
            "error": "path_is_directory",
            "path": path,
            "path_zone": str(classification.zone),
            "message": (
                "The supplied path is shaped like a directory, but this tool reads one concrete file. Use a listing, "
                "search, or environment inventory tool to choose exact files before reading.\n\n"
                "该路径是目录形态；先列目录或搜索，再读取具体文件。"
            ),
            "suggested_tools": ["workspace locate", "search_files", "search_across_files", "env_list_tree", "env_search"],
            "_next_action_instruction": (
                "First identify concrete file paths in the correct zone before calling read_file.\n\n"
                "先在正确区域定位具体文件路径，再读取。"
            ),
            "_suggestions": [],
        }
    if classification.zone == PathZone.STAGED_FILE:
        try:
            from app.core.runtime_mode import current_environment, is_environment_mode

            if is_environment_mode():
                env = current_environment()
                project_path = classification.project_path
                project_exists = False
                if env is not None and project_path:
                    try:
                        project_root = Path(env.root_dir).resolve()
                        candidate = (project_root / project_path).resolve()
                        candidate.relative_to(project_root)
                        project_exists = candidate.exists()
                    except Exception:
                        project_exists = False
                if project_exists:
                    return {
                        "ok": True,
                        "action": action_hint,
                        "path": path,
                        "content": "",
                        "content_omitted": True,
                        "content_compacted": True,
                        "content_omitted_reason": "staged_copy_missing_project_path_exists",
                        "path_zone": str(classification.zone),
                        "staged_path": classification.workspace_path,
                        "staged_path_exists": False,
                        "project_path": project_path,
                        "project_path_exists": True,
                        "staged_copy_handoff_fact": (
                            f"Fact: `{path}` is a staged workspace copy path, but that staged copy does not exist. "
                            f"The corresponding real project path `{project_path}` exists. No file body was read. "
                            "Use env_read/env_search/env_list_tree/env_run for compact project facts, or env_fetch "
                            "when a helper or workspace tool needs a staged copy."
                        ),
                        "事实": (
                            f"`{path}` 暂存副本不存在，但真实项目路径 `{project_path}` 存在；本次未读取正文。"
                            "需要项目事实用 env_*，需要工作区副本再 env_fetch。"
                        ),
                        "suggested_tools": ["env_read", "env_search", "env_list_tree", "env_run", "env_fetch"],
                        "_next_action_instruction": (
                            "Treat this as path evidence, not a failed read. Decide whether the task needs compact "
                            f"project facts from `{project_path}`, a staged copy via env_fetch, or helper-owned reading. "
                            "Do not full-read source/test bodies in the main thread merely to verify helper output.\n\n"
                            "这是路径事实而非读取失败；由模型决定使用 env_*、env_fetch 或 helper 读取。"
                        ),
                        "_suggestions": [],
                    }
                return {
                    "ok": False,
                    "error": (
                        f"file not found: {path}\n"
                        f"`{path}` is a chat-workspace staged copy path. The corresponding project path is "
                        f"`{project_path}`. The staged copy does not currently exist in the chat workspace"
                        + ("; the project path exists." if project_exists else "; the project path was not verified as existing.")
                        + "\nUse env_read/env_search/env_list_tree/env_run with the project path for read-only project facts, "
                        "or call env_fetch before using read_file/search_in_file/code_index on the `_env/...` staged copy. "
                        "If a helper needs this file, pass the exact project_path from inventory/resource manifest or request it explicitly.\n"
                        "该路径是暂存副本路径；当前副本不存在。只读项目事实用项目相对路径调用 env_*；需要工作区副本时先 env_fetch。"
                    ),
                    "path": path,
                    "path_zone": str(classification.zone),
                    "staged_path": classification.workspace_path,
                    "staged_path_exists": False,
                    "project_path": project_path,
                    "project_path_exists": project_exists,
                    "suggested_tools": ["env_read", "env_search", "env_list_tree", "env_run", "env_fetch"],
                    "_next_action_instruction": (
                        "This result is path evidence only. Decide whether the next step needs direct project evidence "
                        f"(`{project_path}` with env_* tools) or a staged workspace copy (env_fetch, then use `{path}`).\n\n"
                        "本结果只陈述路径事实；下一步由模型根据任务决定直接读项目路径还是先暂存。"
                    ),
                    "_suggestions": [],
                }
        except Exception:
            pass

    suggestions = _suggest_similar_files(ws_dir, path)
    error = f"file not found: {path}"
    auto_redirect = None
    if suggestions:
        top = suggestions[0]
        if top["score"] >= 95:
            # P42: 极强匹配, 直接重定向 (95+ 表示几乎确定就是这个)
            auto_redirect = top["path"]
            error += (
                f"\nStrong workspace-name match found (score={top['score']}): {top['path']}. "
                f"Retry {action_hint} with that exact path.\n"
                f"发现高置信相似路径；请直接改用它。"
            )
        elif top["score"] >= 90:
            error += (
                f"\nLikely intended path: {top['path']} ({top['reason']}). "
                f"Retry {action_hint} with the exact path.\n"
                f"可能是该路径；请用精确路径重试。"
            )
        else:
            cand_str = ", ".join(f"{s['path']}" for s in suggestions[:3])
            error += (
                f"\nSimilar workspace paths: {cand_str}. Choose one exact path before retrying.\n"
                f"存在相似路径；先确认具体文件。"
            )
    else:
        error += (
            "\nNo similar file exists in the chat workspace. Check whether the path is a real project path, "
            "whether a helper actually produced the file, or whether it is a stale path from a previous session. "
            "Use locate/search before retrying.\n"
            "工作区无相似文件；先定位文件或确认产物是否存在。"
        )
    try:
        from app.core.runtime_mode import is_environment_mode
        path_norm = str(path).replace("\\", "/").lstrip("./")
        if is_environment_mode() and not path_norm.startswith("_env/"):
            staged_candidate = f"_env/{path_norm}"
            staged_exists = False
            try:
                staged_abs = _safe_resolve(ws_dir, staged_candidate)
                staged_exists = os.path.exists(staged_abs)
            except Exception:
                staged_exists = False
            error += (
                "\nCurrent mode has a separate real project directory. read_file sees only chat-workspace files "
                "and fetched `_env/...` staged copies; use env_read/env_search/env_list_tree/env_run for real "
                "project paths.\n"
                "项目文件请用 env_*；read_file 只读工作区或已暂存副本。"
            )
            action_use = (
                "workspace editing"
                if action_hint in {"edit_file", "multi_edit", "insert_in_file"}
                else "workspace reading"
            )
            if staged_exists:
                error += (
                    f" A staged copy exists at `{staged_candidate}`; use that exact path if you need {action_use}.\n"
                    f"已存在暂存副本 `{staged_candidate}`；请直接改用它。"
                )
            else:
                error += (
                    f" No staged copy exists at `{staged_candidate}`. Use env tools directly, or fetch/stage a concrete "
                    "project file before using read_file. If a helper needs this file, pass the exact project_path "
                    f"`{path_norm}` from `_env/.resource_manifest.json` or use fetch_to_temp(source='main', "
                    f"paths=['{path_norm}']) inside that helper after the main process stages it; otherwise call "
                    "request_resource with the exact project_path."
                )
    except Exception:
        pass
    error += _helper_missing_file_fetch_hint(ws_dir, path)
    response = {
        "ok": False,
        "error": error,
        "_suggestions": suggestions,
    }
    if "staged_candidate" in locals():
        response["staged_candidate"] = staged_candidate
        response["staged_candidate_exists"] = staged_exists
        response["path_zone_fact"] = (
            "In environment mode, non-`_env/` paths are project-relative names; workspace file/edit tools "
            "operate on chat-workspace files or existing staged `_env/...` copies."
        )
    if auto_redirect:
        response["_auto_redirect_path"] = auto_redirect
    return response


def list_generated_files(ws_dir: str) -> list[str]:
    """列出工作区中可被交付的文件（递归子目录）。

    历史教训:这里曾按"源码扩展名"(.py/.js/.ts/...)过滤,导致 .py 永远不可交付。
    现在改为只过滤明确的"框架中间产物"(matplotlib font cache 等)和 _delegate_*
    helper 子目录的内部文件——后者是 helper 工作区,模型若想交付 helper 产物,
    delegate.py 的 _copy_results_to_main 已经把非源码产物复制回主工作区。
    """
    if not ws_dir or not os.path.isdir(ws_dir):
        return []
    ws_path = Path(ws_dir).resolve()
    files = []
    for entry in ws_path.rglob("*"):
        if not entry.is_file():
            continue
        rel = str(entry.relative_to(ws_path)).replace("\\", "/")
        top = rel.split("/", 1)[0]
        if top in {".temp", ".prev", "archive"}:
            continue
        if ws_path.name == ".temp" and top == "_sessions":
            continue
        name = entry.name
        if name.startswith(_ARTIFACT_PREFIXES):
            continue
        if name in _METADATA_BASENAMES:
            continue
        if name in _NON_DELIVERABLE_BASENAMES:
            continue
        if name.endswith(_NON_DELIVERABLE_SUFFIXES):
            continue
        if entry.suffix.lower() in _NON_DELIVERABLE_EXTS:
            continue
        if rel.startswith("_delegate_"):
            continue
        if rel.startswith("_env/"):
            continue
        if "/" not in rel and name.startswith("helper_"):
            continue
        if _looks_like_helper_artifact(name):
            continue
        files.append(rel)
    return sorted(files)


def workspace_disk_usage(ws_dir: str) -> dict:
    """Return recursive workspace usage, skipping filesystem entries that vanish mid-scan."""
    total = 0
    files = 0
    dirs = 0
    if not ws_dir or not os.path.isdir(ws_dir):
        return {"bytes": 0, "files": 0, "dirs": 0}
    try:
        for root, dirnames, filenames in os.walk(ws_dir):
            dirnames[:] = [d for d in dirnames if d not in {"__pycache__", ".git"}]
            dirs += len(dirnames)
            for name in filenames:
                path = os.path.join(root, name)
                try:
                    if os.path.islink(path):
                        continue
                    total += os.path.getsize(path)
                    files += 1
                except OSError:
                    continue
    except OSError:
        pass
    return {"bytes": total, "files": files, "dirs": dirs}






def enforce_workspace_capacity(
    ws_dir: str,
    *,
    max_bytes: int | None = None,
    label: str = "workspace",
) -> dict:
    """Keep one agent workspace under max_bytes, preferring temp/cache deletion first.

    Cleanup order: rebuildable temp files/dirs, duplicate deliverables, then oldest
    non-source deliverables. This can delete generated artifacts so an oversized
    workspace can still start, while preserving source/config/text files last.
    """
    limit = int(max_bytes if max_bytes is not None else settings.workspace_agent_max_bytes)
    before = workspace_disk_usage(ws_dir)
    result = {
        "ok": before["bytes"] <= limit,
        "before_bytes": before["bytes"],
        "after_bytes": before["bytes"],
        "max_bytes": limit,
        "removed_files": 0,
        "removed_dirs": 0,
        "removed_bytes": 0,
        "removed_duplicates": 0,
        "removed_deliverables": 0,
        "errors": [],
    }
    if not ws_dir or not os.path.isdir(ws_dir) or limit <= 0 or before["bytes"] <= limit:
        return result

    removed_files = 0
    removed_dirs = 0
    removed_bytes = 0
    removed_duplicates = 0
    removed_deliverables = 0
    errors: list[str] = []

    def _current_bytes() -> int:
        return workspace_disk_usage(ws_dir)["bytes"]

    def _remove_file(path: str) -> int:
        nonlocal removed_files, removed_bytes
        try:
            size = os.path.getsize(path)
        except OSError:
            size = 0
        try:
            os.remove(path)
            removed_files += 1
            removed_bytes += size
            return size
        except OSError as e:
            errors.append(f"{os.path.basename(path)}: {e}")
            return 0

    def _remove_dir(path: str) -> None:
        nonlocal removed_dirs, removed_bytes
        try:
            size = workspace_disk_usage(path)["bytes"]
            shutil.rmtree(path, ignore_errors=True)
            if not os.path.exists(path):
                removed_dirs += 1
                removed_bytes += size
        except OSError as e:
            errors.append(f"{os.path.basename(path)}: {e}")

    for root, dirnames, filenames in os.walk(ws_dir, topdown=True):
        for dirname in list(dirnames):
            if not _is_known_reclaimable_workspace_dir(dirname):
                continue
            _remove_dir(os.path.join(root, dirname))
            dirnames.remove(dirname)
        for filename in filenames:
            if _is_known_reclaimable_workspace_file(filename):
                _remove_file(os.path.join(root, filename))

    if _current_bytes() > limit:
        by_canonical: dict[str, list[tuple[float, int, str]]] = {}
        for root, dirnames, filenames in os.walk(ws_dir):
            dirnames[:] = [
                d for d in dirnames
                if d not in {"archive"} and not d.startswith("_delegate_")
            ]
            for filename in filenames:
                if not _is_deletable_deliverable(filename):
                    continue
                path = os.path.join(root, filename)
                try:
                    st = os.stat(path)
                except OSError:
                    continue
                key = _duplicate_canonical_name(filename)
                by_canonical.setdefault(key, []).append((st.st_mtime, st.st_size, path))
        duplicate_candidates: list[tuple[float, int, str]] = []
        for items in by_canonical.values():
            if len(items) <= 1:
                continue
            items.sort(reverse=True)
            duplicate_candidates.extend(items[1:])
        duplicate_candidates.sort(key=lambda x: (x[0], -x[1]))
        for _mtime, _size, path in duplicate_candidates:
            if _current_bytes() <= limit:
                break
            if _remove_file(path):
                removed_duplicates += 1

    if _current_bytes() > limit:
        deliverable_candidates: list[tuple[float, int, str]] = []
        for root, dirnames, filenames in os.walk(ws_dir):
            dirnames[:] = [
                d for d in dirnames
                if d not in {"archive"} and not d.startswith("_delegate_")
            ]
            for filename in filenames:
                if not _is_deletable_deliverable(filename):
                    continue
                path = os.path.join(root, filename)
                try:
                    st = os.stat(path)
                except OSError:
                    continue
                deliverable_candidates.append((st.st_mtime, st.st_size, path))
        deliverable_candidates.sort(key=lambda x: (x[0], -x[1]))
        for _mtime, _size, path in deliverable_candidates:
            if _current_bytes() <= limit:
                break
            if _remove_file(path):
                removed_deliverables += 1

    after = workspace_disk_usage(ws_dir)
    result.update({
        "ok": after["bytes"] <= limit,
        "after_bytes": after["bytes"],
        "removed_files": removed_files,
        "removed_dirs": removed_dirs,
        "removed_bytes": removed_bytes,
        "removed_duplicates": removed_duplicates,
        "removed_deliverables": removed_deliverables,
        "errors": errors[:10],
    })
    debug.log(
        "workspace.capacity",
        f"{label}: {before['bytes'] // 1024 // 1024}MB → {after['bytes'] // 1024 // 1024}MB "
        f"(limit={limit // 1024 // 1024}MB, removed_files={removed_files}, "
        f"removed_dirs={removed_dirs}, duplicates={removed_duplicates}, deliverables={removed_deliverables})",
        result,
    )
    return result


_DUP_SUFFIX_RE = re.compile(r"^(?P<stem>.+?)(?:_from_group|[_ -]?copy|\s*\(\d+\)|_\d+)$", re.IGNORECASE)


def _duplicate_canonical_name(name: str) -> str:
    stem, ext = os.path.splitext(name.lower())
    while True:
        m = _DUP_SUFFIX_RE.match(stem)
        if not m:
            break
        stem = m.group("stem").rstrip(" _-")
    return stem + ext


def _is_deletable_deliverable(name: str) -> bool:
    lower = name.lower()
    if lower.startswith("."):
        return False
    if lower in {"_session_manifest.json"}:
        return False
    ext = os.path.splitext(lower)[1]
    return ext not in _PROTECTED_EXTS


# ── Bug 6 修:helper 工作区目录清理 ────────────────────────────
def cleanup_delegate_dirs(ws_dir: str) -> int:
    """清理工作区中所有 _delegate_* 子目录(上一轮 helper 残留)。
    返回清理掉的目录数。

    设计意图:helper 工作区在 _run_one_helper 开头会调 clean_workspace_dir 重置内部,
    但目录本身永久存活——下次同名 task_id 来时会复用,不同名时新增。多次会话后
    workspace 里堆满 _delegate_avl_tree / _delegate_dqn_agent / ... 互相不相关,
    新任务的 helper 会困惑("怎么有 avl_tree.c?这是上次包涵的代码")。
    在 orchestrator 创建工作区后立即清理,保证每轮 round2 启动时是干净状态。
    """
    if not ws_dir or not os.path.isdir(ws_dir):
        return 0
    cleaned = 0
    for name in os.listdir(ws_dir):
        if name.startswith("_delegate_"):
            target = os.path.join(ws_dir, name)
            if os.path.isdir(target):
                try:
                    shutil.rmtree(target, ignore_errors=True)
                    cleaned += 1
                except OSError:
                    pass
    if cleaned:
        debug.log("workspace.cleanup_delegate", f"removed {cleaned} stale _delegate_* dirs")
    return cleaned


# ── Opt 3: 主工作区增长控制 ─────────────────────────────────────
# 制品文件类型(非源码,随时间累积)
_ARCHIVABLE_EXTS: tuple[str, ...] = (
    ".png", ".jpg", ".jpeg", ".gif", ".svg", ".webp", ".bmp",
    ".docx", ".pptx", ".xlsx", ".pdf", ".zip", ".csv",
    ".mp4", ".avi", ".mov", ".webm",
)

# 保护性扩展名(源码/配置,绝不归档)
_PROTECTED_EXTS: tuple[str, ...] = (
    ".py", ".c", ".cpp", ".h", ".hpp", ".js", ".ts", ".jsx", ".tsx",
    ".rs", ".go", ".java", ".rb", ".sh", ".bat", ".ps1", ".toml",
    ".json", ".yaml", ".yml", ".ini", ".cfg", ".env", ".md", ".txt",
    ".sql", ".html", ".css",
)


def archive_stale_artifacts(main_ws: str, max_age_days: int = 14) -> dict:
    """将主工作区超过 max_age_days 的制品文件归档到 archive/YYYY-MM/ 子目录。

    只归档制品类型(图片/文档/视频),源码和配置文件永远留在原地。
    archive/ 目录在 .temp/ sync 时会被跳过,不会污染 helper 工作区。

    Returns:
        {archived: count, archive_dir: path, bytes_freed: int, errors: [str]}
    """
    if not main_ws or not os.path.isdir(main_ws):
        return {"archived": 0, "archive_dir": "", "bytes_freed": 0, "errors": ["invalid dir"]}

    import time as _time_module
    now = _time_module.time()
    cutoff = now - (max_age_days * 86400)
    archive_root = os.path.join(main_ws, "archive")
    archived = 0
    bytes_freed = 0
    errors: list[str] = []

    for entry in os.listdir(main_ws):
        if entry.startswith(".") or entry.startswith("_") or entry == "archive":
            continue
        full = os.path.join(main_ws, entry)
        if os.path.isdir(full):
            continue  # 子目录暂不处理(可能是用户组织的)

        ext = os.path.splitext(entry)[1].lower()
        if ext not in _ARCHIVABLE_EXTS:
            continue

        try:
            mtime = os.path.getmtime(full)
        except OSError:
            continue
        if mtime > cutoff:
            continue  # 还不够"老"

        # 移动到 archive/YYYY-MM/
        month_dir = _time_module.strftime("%Y-%m", _time_module.localtime(mtime))
        target_dir = os.path.join(archive_root, month_dir)
        target_path = os.path.join(target_dir, entry)

        try:
            os.makedirs(target_dir, exist_ok=True)
            size = os.path.getsize(full)
            # 目标文件已存在则跳过(保留旧版)
            if os.path.exists(target_path):
                os.remove(full)  # 删源(归档版还在)
            else:
                shutil.move(full, target_path)
            archived += 1
            bytes_freed += size
        except OSError as e:
            errors.append(f"{entry}: {e}")

    if archived:
        debug.log(
            "workspace.archive",
            f"archived {archived} stale artifacts ({bytes_freed // 1024 // 1024}MB) "
            f"to archive/ (max_age>{max_age_days}d)",
        )
    return {
        "archived": archived,
        "archive_dir": archive_root,
        "bytes_freed": bytes_freed,
        "errors": errors,
    }




def _is_shared_readonly_path(rel_path: str) -> bool:
    """判断该路径是否落在 _shared/ 区域。仅对 helper 生效 — 主线程不受限。

    ── 2026-05-04 Razor 教训修复 ──
    _shared/ 是主线程预先准备的测试数据/timing 框架/输出格式 — helper 应该
    `read_file` 它,**不应该 write/edit/multi_edit/insert**。trace razor 实测:
    helper 改了 _shared/ 里的测试驱动,把 bench_run_all 替换成自己写的 main 但
    路径解析错,导致 12 个测试文件 SKIP 报"3/3 passed"虚假成绩。

    helper 想扩共享内容用 `_helpers_shared/`(允许写,也会反向合并到主区)。

    ── 2026-05-05 修复:主线程误拦截 ──
    原实现不区分调用方,连主线程写自己的 _shared/ 也拦截。主线程 owner 以
    "main:" 开头,helper 以 "helper:" 开头。现在只对 helper 生效。
    """
    if not rel_path:
        return False
    # 主线程不受 _shared/ 写保护
    if _is_main_thread():
        return False
    normalized = rel_path.lstrip("/").lstrip("\\").replace("\\", "/")
    return normalized == "_shared" or normalized.startswith("_shared/")


def _is_main_thread() -> bool:
    """当前调用方是否为主线程(非 helper)。"""
    try:
        owner = current_owner()
        return owner.startswith("main:")
    except Exception:
        return True  # 保守:取不到 owner 就当主线程,不拦截


_set_workspace_run_main_thread_provider(_is_main_thread)

_SHARED_READONLY_ERROR_MSG = (
    "`_shared/` is read-only scaffold supplied by the main process, such as test data, timing "
    "frameworks, or output-format contracts. Do not modify it with write/edit/insert tools. "
    "To extend shared content, write under `_helpers_shared/`; sibling helpers can see that area "
    "and it can be merged back to the main workspace. For your own test driver or task-local "
    "notes, write in the helper sandbox root or a task-specific subdirectory.\n\n"
    "`_shared/` 是主流程提供的只读脚手架；扩展共享内容写 `_helpers_shared/`，自己的测试驱动写 helper 沙箱。"
)


# ── 文件操作 handlers ─────────────────────────────────────────
async def handle_mkdir(ws_dir: str, path: str) -> dict:
    try:
        target = _safe_resolve(ws_dir, path)
    except ValueError as e:
        return {"ok": False, "error": str(e)}
    os.makedirs(target, exist_ok=True)
    return {"ok": True, "action": "mkdir", "path": path}


from app.llm.tools.workspace_file_ops import (  # noqa: E402,F401
    handle_write,
    handle_append,
    _extract_python_outline,
    _build_file_outline,
    _detect_file_type,
    _check_file_readable,
    handle_inspect_file,
    _track_edit_count,
    _read_text_safely,
    handle_read_file,
    handle_edit_file,
    handle_multi_edit,
    handle_insert_in_file,
    handle_search_in_file,
)



# ── 局部读写 handlers ─────────────────────────────────────────
# 设计目的:替代"全文 read + 全文 write"的低效模式。让模型可以:
#   1. read_file 看局部 + 行号
#   2. search_in_file 定位
#   3. edit_file 精准 str_replace
#   4. insert_in_file 在指定行后插入
# 所有操作复用 _safe_resolve 路径防护;helper 只能 read/search,不能 edit/insert。

# 二进制检测阈值:前 1KB 含 NULL 字节即视为二进制
_BINARY_DETECT_BYTES = 1024
# 单次 read 最大字符数(防止 prompt 爆炸)
# 2026-06-05: 用户要求任何时候单次操作不超过 30KB, 正常默认 6KB。
# 一次 read_file 应只返回合理大小片段。超限时截断并告知 LLM 用 search_in_file +
# start_line/end_line 细读, 而不是续读下一块(续读同样会超限)。
# 旧值从 500KB 降下来(实测 trace 990126 主线程全文读 4 份证据 ~22KB 后
# 因上下文膨胀而张冠李戴把 2FSK 写成 16QAM)。
_READ_MAX_CHARS_DEFAULT = 6_000   # 默认 6KB(鼓励按范围细读)
_READ_MAX_CHARS_HARD_CAP = 30_000  # HARD_CAP 30KB(任何操作不可超过)
# search 单次最大返回结果数 + 超时秒数
_SEARCH_MAX_RESULTS_DEFAULT = 50
_SEARCH_TIMEOUT_S = 2.0
# edit_file old_str 最小长度(防止误改 ; } 等危险短串)
_EDIT_OLD_STR_MIN_LEN = 5
# 文件大小上限分两档:
# - read/edit/insert 全文加载: 2MB(part22 微调:500KB→2MB,跟 _READ_MAX_CHARS_HARD_CAP 对齐)
#   原 500KB 是与 _READ_MAX_CHARS_HARD_CAP 不一致的死路径 — 模型传 max_chars=2M 也读不到
# - search 流式扫描:        50MB(再大就真的太巨型,引导 python)
_FILE_SIZE_CAP = 2_000_000
_SEARCH_FILE_SIZE_CAP = 50_000_000

# 2026-05-03 优化 #5:同文件累计 edit 次数到达此阈值,提示重写
# (trace 09ba132f bwt_fix:同函数 multi_edit×3 + edit_file 都没修通,
#  依然 multi_edit 改第 4 次。铁律 #2 说"≥3 次重写",但 multi_edit 不算 edit_file
#  独立调用。这里统一在 edit/multi_edit/insert handler 末尾计数。)
_EDIT_REWRITE_THRESHOLD = 3
# 2026-05-15 P69: 同文件 ≥10 次 edit → 硬阻断, 不再只是软提示。
# 病因(实测 05-15 comp_bench): 旧版只在 _rewrite_suggestion 字段塞文本警告,
# helper LLM 直接忽略 → 单 helper 对 bench/benchmark.c 编辑 50 次, bench/compress.h
# 编辑 55 次, 反复修不通同一组接口耦合的 .c/.h 死循环。
# 修法: 阈值 ≥10 时返 ok=false + error_kind="edit_thrashing_exceeded", 拒绝执行 edit,
# 强制 helper 切换策略 (workspace.write 整文件重写, 或承认失败 spawn 新 task)。
#
# 2026-06-05 调整: 实测 trace 394304 14:55:22 / 14:57:00 / 15:03:26 — helper 在
# 实现复杂数据结构 (acb_tree.py / existing_algos.py) 时 10 次 edit 仍属正常增量
# 开发, 命中 hard block 后被迫整体重写反而丢失增量进度。提高到 20, 配合 stuck
# detector 的 same_file_edit_fail (≥4 次 edit + run_fail 才软提示) 已经能在真正的
# 死循环里给出引导,不需要这条硬线。
_EDIT_HARD_BLOCK_THRESHOLD = 20

# ── 进程内 read tracker(2026-05-03 trace b78b242533a24a46 教训)──
# 磁盘 .read_history.json 在 helper sandbox 里偶尔写不进去(实测 lz78 helper 27 分钟里
# read 同一文件 17 次,任何一次都没收到 _already_read_full 警告 → 经查 sandbox 里
# .read_history.json 根本没创建)。原因可能是异常被 except OSError pass 吃了,或
# helper resume 路径下文件权限/路径解析异常。
#
# 兜底方案:除了写磁盘 history 外,**进程内也维护一份 read 计数 + 全读标记**。
# key = (ws_dir, path) 元组,value = {"full_read": bool, "read_count": int, "fragments": list}
# 即使磁盘写失败,内存仍能识别"重读"。也用计数而不仅是 boolean,这样:
#   - 第 1 次全读:正常
#   - 第 2 次重读(全文/片段):返回 _already_read_full=true 警告
#   - 第 3 次仍重读:**升级为 ERROR 拒绝**,强制 helper 改用已有上下文
#
# 这是回答"为什么 helper 27 分钟修不对"的关键修复 — lz78 helper 反复 read 同
# 一个 algo_lz78.c 17 次,如果第 3 次起就被强制拒绝,会节省至少 6-8 分钟。
_read_tracker: dict[tuple[str, str], dict] = {}
_READ_REPEAT_WARN_THRESHOLD = 2   # 第 2 次起警告
_READ_REPEAT_BLOCK_THRESHOLD = 4  # 第 4 次起拒绝


# ── 2026-05-11 A2: 文件结构摘要提取(用于重读拒绝的软退化) ────
# 旧版重读拒绝返回 ERROR,helper 拿不到任何信息只能凭印象 edit → old_str not found
# → 死循环。新版改成"拒绝重读全文,但返回结构化摘要"(函数列表/导入/常量/行数),
# helper 仍能继续工作,不必盲操作。
# 摘要是确定性提取(正则/AST),零 LLM 成本零延迟。

import ast as _ast_mod  # noqa










# ── 错误信息中的引导模板(让模型知道"被拒后该怎么办")──
_GUIDE_BINARY = (
    "binary file detected — local read/edit tools work on text only. "
    "Use `workspace run` with a Python script to convert it to text first. "
    "Common conversions:\n"
    "  - PDF       → pypdf / pdfplumber:  text = pypdf.PdfReader(p).pages[0].extract_text()\n"
    "  - Excel     → openpyxl / pandas:   pandas.read_excel(p).to_csv('out.csv')\n"
    "  - Word      → python-docx:         '\\n'.join(p.text for p in docx.Document(p).paragraphs)\n"
    "  - Image     → PIL (size/info) or pytesseract (OCR):  Image.open(p).size\n"
    "  - JSON/XML  → json/xml.etree:       extract subset to .txt\n"
    "  - zip/tar   → zipfile/tarfile:      extract then process\n"
    "Workflow: workspace write convert.py '...' → workspace run 'python convert.py' → "
    "read_file output.txt → (edit) → workspace run python (convert back if needed)"
)

_GUIDE_TOO_LARGE = (
    "file too large for full load. Two paths:\n"
    "  1. If you only need a slice: search_in_file(path, pattern) to locate, "
    "then read_file(path, start, end) for that range — both work on big files.\n"
    "  2. If you need to transform: workspace write proc.py "
    "'with open(\"{path}\") as f, open(\"out.txt\", \"w\") as o: ...processing... ' → "
    "workspace run 'python proc.py' → read_file out.txt → edit there → "
    "workspace run another script to write changes back."
)


# ── 精确的文件类型 → Python 处理方案对照表 ──
# 用于 inspect_file 工具和 _check_file_readable 的二进制错误升级。
# (category, friendly_name, extract_to_text_code, save_back_code)
# extract/save 中 PATH 是占位符,inspect_file 时会替换成真实路径。
_FILE_TYPE_TABLE = {
    # text — 直接 read_file
    ".txt": ("text", "plain text", None, None),
    ".md": ("text", "markdown", None, None),
    ".log": ("text", "log file", None, None),
    ".html": ("text", "HTML", None, None),
    ".py": ("text", "Python source", None, None),
    ".c": ("text", "C source", None, None),
    ".cpp": ("text", "C++ source", None, None),
    ".h": ("text", "C/C++ header", None, None),
    ".js": ("text", "JavaScript", None, None),
    ".ts": ("text", "TypeScript", None, None),
    ".java": ("text", "Java", None, None),
    ".go": ("text", "Go", None, None),
    ".rs": ("text", "Rust", None, None),
    ".sh": ("text", "shell script", None, None),
    # text-structured — 直接 read_file 或用 python 库结构化读
    ".csv": ("text-structured", "CSV table",
             "import pandas as pd; df = pd.read_csv('PATH')",
             "df.to_csv('PATH', index=False)"),
    ".tsv": ("text-structured", "TSV table",
             "import pandas as pd; df = pd.read_csv('PATH', sep='\\t')",
             "df.to_csv('PATH', sep='\\t', index=False)"),
    ".json": ("text-structured", "JSON",
              "import json; data = json.load(open('PATH', encoding='utf-8'))",
              "json.dump(data, open('PATH', 'w', encoding='utf-8'), ensure_ascii=False, indent=2)"),
    ".yaml": ("text-structured", "YAML",
              "import yaml; data = yaml.safe_load(open('PATH'))",
              "yaml.safe_dump(data, open('PATH', 'w'), allow_unicode=True)"),
    ".yml": ("text-structured", "YAML",
             "import yaml; data = yaml.safe_load(open('PATH'))",
             "yaml.safe_dump(data, open('PATH', 'w'), allow_unicode=True)"),
    ".xml": ("text-structured", "XML",
             "from xml.etree import ElementTree as ET; tree = ET.parse('PATH'); root = tree.getroot()",
             "tree.write('PATH', encoding='utf-8', xml_declaration=True)"),
    ".toml": ("text-structured", "TOML",
              "import tomllib; data = tomllib.load(open('PATH', 'rb'))",
              "# 写回需 tomli_w: tomli_w.dump(data, open('PATH', 'wb'))"),
    # document — normal route is the structured `office` tool. Python snippets
    # are fallback extraction probes only when the office tool lacks a needed
    # operation or exact unsupported inspection is required.
    ".docx": ("document", "Word document",
              "# Normal route: office(action='read', path='PATH') returns document structure.\n"
              "# Fallback probe only when a named unsupported detail is needed:\n"
              "from docx import Document\\n"
              "doc = Document('PATH')\\n"
              "text = '\\n'.join(p.text for p in doc.paragraphs)",
              "# Normal route: office(action='write'/'append'/'replace_block', path='PATH', blocks=[...])\n"
              "# blocks: heading/paragraph/list/table/image/page_break.\n"
              "# Use python-docx only for a named unsupported formatting/inspection need, not for ordinary assembly."),
    ".doc": ("document", "Word document (legacy .doc)",
             "# .doc 老格式 office 工具不直接支持,先转 .docx:\\n"
             "# workspace run 'libreoffice --headless --convert-to docx PATH'\\n"
             "# 然后 office(action='read', path='PATH.docx') 处理生成的 .docx",
             None),
    ".pdf": ("document", "PDF",
             "# === 读 PDF(两种方式互补) ===\\n"
             "# 方式1 — pypdf 提取文本层(结构化文字,快速):\\n"
             "import pypdf\\n"
             "reader = pypdf.PdfReader('PATH')\\n"
             "text = '\\n\\n'.join(page.extract_text() for page in reader.pages)\\n"
             "open('extracted.txt', 'w', encoding='utf-8').write(text)\\n"
             "# 方式2 — OCR 识别扫描件/图片中的文字(不可提取文本):\\n"
             "#   先用 pymupdf/fitz 将页面渲染为 PNG:\\n"
             "import fitz\\n"
             "doc = fitz.open('PATH')\\n"
             "page = doc[0].get_pixmap(dpi=200)\\n"
             "page.save('page_0.png')\\n"
             "#   再调 ocr(image_path='page_0.png') 获取图片中的文字\\n"
             "# === 写 PDF ===\\n"
             "# PDF 难以原地修改文字。两个选项:\\n"
             "# 1. 用 reportlab 生成全新 PDF\\n"
             "# 2. 转成 docx 编辑后再转回 PDF (LibreOffice 命令行)"),
    ".rtf": ("document", "Rich Text Format",
             "from striprtf.striprtf import rtf_to_text\\n"
             "text = rtf_to_text(open('PATH', encoding='utf-8').read())",
             "# RTF 写回较麻烦,通常做法是输出 .docx 或 .txt"),
    # presentation — PowerPoint
    ".pptx": ("presentation", "PowerPoint presentation",
              "# Preferred: office(action='read', path='PATH') returns slides[].title + body_texts JSON.\n"
              "# 代码方式:from pptx import Presentation; prs = Presentation('PATH')\n"
              "# for s in prs.slides: print(s.shapes.title.text if s.shapes.title else '')",
              "# Preferred: office(action='write', path='PATH', slides=[{title:'...', body:[...]}])\n"
              "# body 元素: text / bullets / image / table\n"
              "# 复杂幻灯片(动画/转场/精确版式)再考虑手写 python-pptx 代码"),
    ".ppt": ("presentation", "PowerPoint (legacy .ppt)",
             "# .ppt 老格式 office 工具不直接支持,先转 .pptx:\\n"
             "# workspace run 'libreoffice --headless --convert-to pptx PATH'",
             None),
    # spreadsheet — 优先用 `office` 工具
    ".xlsx": ("spreadsheet", "Excel workbook",
              "# Preferred: office(action='read', path='PATH') returns sheets[].rows JSON.\n"
              "# 代码方式(读样式/合并/条件格式时):import openpyxl\\n"
              "wb = openpyxl.load_workbook('PATH', data_only=True)\\n"
              "for sn in wb.sheetnames: ws = wb[sn]; print(sn, list(ws.iter_rows(values_only=True))[:3])",
              "# Preferred: office(action='write', path='PATH', sheets=[{name:'...', rows:[[...]], header:true}])\n"
              "# 公式以 '=' 开头: rows=[['A','B'],[1,2],['Total','=SUM(B2:B3)']]\n"
              "# 复杂场景(图表/合并单元格/条件格式)用 openpyxl 代码"),
    ".xls": ("spreadsheet", "Excel (legacy .xls)",
             "import pandas as pd; df = pd.read_excel('PATH')",
             "df.to_excel('PATH', index=False)  # 写回会变成 .xlsx 格式"),
    ".ods": ("spreadsheet", "OpenDocument Spreadsheet",
             "import pandas as pd; df = pd.read_excel('PATH', engine='odf')",
             "# .ods 写回需 odfpy"),
    # image — 一般不"转文本",但可以 PIL 处理
    ".png": ("image", "PNG image",
             "from PIL import Image\\n"
             "img = Image.open('PATH')\\n"
             "print('size:', img.size, 'mode:', img.mode)\\n"
             "# OCR(需要 pytesseract): import pytesseract; text = pytesseract.image_to_string(img)",
             "img.save('PATH')  # 或 'output.png'"),
    ".jpg": ("image", "JPEG image",
             "from PIL import Image; img = Image.open('PATH'); print(img.size, img.mode)",
             "img.save('PATH')"),
    ".jpeg": ("image", "JPEG image",
              "from PIL import Image; img = Image.open('PATH'); print(img.size, img.mode)",
              "img.save('PATH')"),
    ".gif": ("image", "GIF image",
             "from PIL import Image; img = Image.open('PATH')",
             "img.save('PATH')"),
    ".bmp": ("image", "Bitmap image",
             "from PIL import Image; img = Image.open('PATH')",
             "img.save('PATH')"),
    ".webp": ("image", "WebP image",
              "from PIL import Image; img = Image.open('PATH')",
              "img.save('PATH', 'WEBP')"),
    # archive — zipfile/tarfile
    ".zip": ("archive", "ZIP archive",
             "import zipfile\\n"
             "with zipfile.ZipFile('PATH') as z:\\n"
             "    print(z.namelist())\\n"
             "    z.extractall('extracted/')",
             "import zipfile\\n"
             "with zipfile.ZipFile('PATH', 'w') as z:\\n"
             "    z.write('file_to_add.txt')"),
    ".tar": ("archive", "TAR archive",
             "import tarfile\\n"
             "with tarfile.open('PATH') as t:\\n"
             "    print(t.getnames())\\n"
             "    t.extractall('extracted/')",
             "with tarfile.open('PATH', 'w') as t: t.add('file.txt')"),
    ".gz": ("archive", "GZip compressed",
            "import gzip; data = gzip.open('PATH', 'rb').read()",
            "import gzip; gzip.open('PATH', 'wb').write(data)"),
    # media — LLM 一般不直接处理
    ".mp3": ("media", "MP3 audio",
             "# 音频文件 — LLM 不直接处理。\\n"
             "# 选项:用 openai-whisper 做语音识别;或 librosa/pydub 做信号分析",
             None),
    ".mp4": ("media", "MP4 video",
             "# 视频 — LLM 不直接处理。\\n"
             "# 选项:cv2/moviepy 抽帧,ffmpeg 提取音轨",
             None),
    ".wav": ("media", "WAV audio",
             "# 音频 — 用 wave/scipy.io.wavfile/librosa 读样本",
             None),
    # binary — 不建议处理
    ".exe": ("binary", "Windows executable",
             "# 可执行文件,不建议直接处理。如需分析,用 pefile/lief",
             None),
    ".so": ("binary", "Linux shared library",
            "# 共享库二进制,不建议直接处理",
            None),
    ".dll": ("binary", "Windows DLL",
             "# DLL,不建议直接处理",
             None),
}






















# ── 2026-05-02 part20:multi_edit(Claude Code 风格 MultiEdit)──
# 一次原子改同文件多处。比连续 N 次 edit_file 强:
# 1. 减少 N-1 次 round-trip (LLM ↔ 服务端)
# 2. 一次看到全部 diff,避免改了 1 处忘了相关 2 处
# 3. 原子性:全部成功才提交,任一失败回滚 — 不会留半残状态
# 4. 模型一次"想清楚整组改动"比"一处一处试探"质量高得多






# ── 2026-05-02 part14:code_index 工具 ─────────────────────────
# 教训(trace 74b1295b):helper rdh_v2 read 同一 rdh.c 60 次共 26 次落在 huff_decode
# 区域。它没有"鸟瞰一眼代码结构"的工具,只能反复 read 不同行段试图建立 mental model。
# 1000 行代码 read 全文进 context 是大开销,但它需要的只是"哪些函数在哪行"的地图。
#
# 方法论灵感来自 ctags/cscope/tree-sitter:把符号位置压成紧凑索引(~30 行替代 ~1000 行)。
# 实现简单 — 正则匹配 .c/.h/.py/.js/.ts/.go/.rs 几种语言的 def/class/include。
# 不需要 100% 准确,准确度 90%+ 已经能让模型用 1 次调用代替 5-10 次 read。

# 每语言的 symbol 提取规则:(language_id, file_extensions, regex_patterns)
# 模式列表里每条:(re.compile, label),按行匹配,提取首个捕获组作为符号名。
_CODE_INDEX_RULES = {
    "c": {
        "exts": (".c", ".cc", ".cpp", ".cxx", ".h", ".hpp", ".hh"),
        "patterns": [
            # 函数定义: int foo(...) { 或 static int foo(...)\n{ 或 单行 int foo(...) { ... }
            # 末尾约束放宽:`)` 后任意字符(支持单行完整函数 / 多行签名 / { 在下一行)
            (re.compile(r'^[\w*\s]*?(\w+)\s*\([^)]*\)\s*(\{.*|\{?\s*)$'), "fn"),
            # struct / typedef / enum 定义
            (re.compile(r'^typedef\s+\w+(?:\s+\w+)?\s+(\w+)\s*;?\s*$'), "typedef"),
            (re.compile(r'^(?:typedef\s+)?struct\s+(\w+)\s*\{?\s*$'), "struct"),
            (re.compile(r'^(?:typedef\s+)?enum\s+(\w+)?\s*\{?\s*$'), "enum"),
            # #define MACRO
            (re.compile(r'^#define\s+([A-Z_][A-Z0-9_]+)'), "define"),
            # #include
            (re.compile(r'^#include\s+([<"][^>"]+[>"])'), "include"),
        ],
    },
    "py": {
        "exts": (".py",),
        "patterns": [
            (re.compile(r'^(?:async\s+)?def\s+(\w+)\s*\('), "def"),
            (re.compile(r'^class\s+(\w+)'), "class"),
            (re.compile(r'^(?:from\s+\S+\s+)?import\s+(.+)$'), "import"),
        ],
    },
    "js": {
        "exts": (".js", ".jsx", ".ts", ".tsx", ".mjs"),
        "patterns": [
            (re.compile(r'^(?:export\s+)?(?:async\s+)?function\s+(\w+)\s*\('), "fn"),
            (re.compile(r'^(?:export\s+)?(?:const|let|var)\s+(\w+)\s*=\s*(?:async\s*)?\('), "fn"),
            (re.compile(r'^(?:export\s+)?class\s+(\w+)'), "class"),
            (re.compile(r'^import\s+.+\s+from\s+[\'"]([^\'"]+)[\'"]'), "import"),
        ],
    },
    "go": {
        "exts": (".go",),
        "patterns": [
            (re.compile(r'^func\s+(?:\([^)]+\)\s+)?(\w+)\s*\('), "fn"),
            (re.compile(r'^type\s+(\w+)\s+(?:struct|interface)'), "type"),
        ],
    },
    "rs": {
        "exts": (".rs",),
        "patterns": [
            (re.compile(r'^(?:pub\s+)?(?:async\s+)?fn\s+(\w+)\s*[<(]'), "fn"),
            (re.compile(r'^(?:pub\s+)?struct\s+(\w+)'), "struct"),
            (re.compile(r'^(?:pub\s+)?enum\s+(\w+)'), "enum"),
            (re.compile(r'^(?:pub\s+)?trait\s+(\w+)'), "trait"),
            (re.compile(r'^impl\s+.*\s+for\s+(\w+)'), "impl"),
        ],
    },
}


def _detect_lang(path: str) -> str | None:
    p = path.lower()
    for lang, rules in _CODE_INDEX_RULES.items():
        if any(p.endswith(e) for e in rules["exts"]):
            return lang
    return None


# ── 2026-05-02 part16:_handle_code_index_glob — 多文件批量索引 ──
# 1M context 利用方式:helper 看陌生项目时一次调用拿"项目骨架",而不是逐文件 index。
# 实测场景:5 个 .c 文件总 4000 行 → glob 一次返回所有函数列表(~150 行索引),
# 比 5 次单文件调用节省 round-trip + token。

async def _handle_code_index_glob(
    ws_dir: str, glob_pattern: str, *,
    include_includes: bool = True,
    name_filter: str | None = None,
    kinds: list[str] | None = None,
    max_files: int = 30,
) -> dict:
    """code_index 的 glob 多文件实现(被 handle_code_index 检测到 *? 后调用)。"""
    import fnmatch as _fnm
    if not os.path.isdir(ws_dir):
        return {"ok": False, "error": "workspace dir not found"}

    # 收集匹配文件:走 os.walk 跳过常见污染目录
    candidate_files = []
    for root, dirs, files in os.walk(ws_dir):
        dirs[:] = [d for d in dirs
                   if not d.startswith(".")
                   and d not in {"__pycache__", "node_modules", ".git",
                                 ".pytest_cache"}
                   and not d.startswith("_delegate_")]
        for fn in files:
            if fn.startswith("."):
                continue
            full = os.path.join(root, fn)
            rel = os.path.relpath(full, ws_dir).replace(os.sep, "/")
            # 匹配 glob_pattern (相对路径或文件名)
            if _fnm.fnmatch(rel, glob_pattern) or _fnm.fnmatch(fn, glob_pattern):
                # 必须是支持的源代码类型
                if _detect_lang(rel) is not None:
                    candidate_files.append(rel)

    if not candidate_files:
        return {
            "ok": False,
            "error": (
                f"no source files match glob '{glob_pattern}'. "
                f"supported extensions: .c .cpp .h .py .js .ts .go .rs .java"
            ),
        }

    files_truncated = False
    if len(candidate_files) > max_files:
        candidate_files = sorted(candidate_files)[:max_files]
        files_truncated = True
    else:
        candidate_files = sorted(candidate_files)

    files_info = []
    all_symbols = []
    summary_parts = []
    total_lines_overall = 0

    for rel in candidate_files:
        sub = await handle_code_index(
            ws_dir, rel,
            include_includes=include_includes,
            name_filter=name_filter,
            kinds=kinds,
        )
        if not sub.get("ok"):
            continue
        files_info.append({
            "path": rel,
            "lang": sub["lang"],
            "total_lines": sub["total_lines"],
            "symbol_count": sub["symbol_count"],
        })
        for s in sub.get("symbols", []):
            s2 = dict(s)
            s2["path"] = rel
            all_symbols.append(s2)
        summary_parts.append(sub["summary"])
        total_lines_overall += sub["total_lines"]

    summary = (
        f"=== {len(files_info)} files indexed "
        f"({total_lines_overall} lines total) ===\n\n"
        + "\n\n".join(summary_parts)
    )

    result = {
        "ok": True,
        "action": "code_index",
        "path": glob_pattern,
        "files_indexed": len(files_info),
        "files": files_info,
        "symbols": all_symbols,
        "symbol_count": len(all_symbols),
        "summary": summary,
        "total_lines": total_lines_overall,
    }
    if files_truncated:
        result["files_truncated"] = True
        result["files_truncated_note"] = (
            f"达 max_files={max_files} 上限,只索引前 {max_files} 个。"
            f"用更窄的 glob(如 'src/*.c' 替代 '*')缩小范围。"
        )
    if name_filter or kinds:
        result["filter_applied"] = {
            "name_filter": name_filter, "kinds": kinds,
        }
    return result


async def handle_code_index(
    ws_dir: str, path: str, *,
    include_includes: bool = True,
    name_filter: str | None = None,
    kinds: list[str] | None = None,
) -> dict:
    """提取源文件的符号索引(函数 / 类 / struct / include 等)。

    设计目标:helper 看陌生代码时**用 1 次调用拿到全文骨架**,代替 5-10 次盲目 read。
    返回紧凑的 ASCII 表(每行约 50 chars),1000 行代码典型输出 ~30 行。

    Args:
        path: 工作区相对路径 OR glob 通配符(如 '*.c' / 'src/**/*.py')
            **2026-05-02 part16 增强**:含 * 或 ? 时 glob 多文件,返回汇总骨架。
            实测项目骨架场景:`code_index('*.c')` 一次拿所有 C 文件结构。
        include_includes: 是否包含 include/import 列表(默认 True;只想要函数列表设 False)
        name_filter: glob 或 regex 过滤名字(如 'huff_*' 只列 huff 开头)
        kinds: 只列指定类型 ([fn, struct, typedef, enum, define, class, def, import, include])

    Returns(单文件):
        {
          "ok": True, "path": "rdh.c", "lang": "c", "total_lines": 762,
          "symbols": [...],
          "summary": "rdh.c (762 lines, c)\n  L120 fn  huff_build_tree\n  ..."
        }
    Returns(多文件 glob):
        {
          "ok": True, "path": "*.c",
          "files_indexed": 5,
          "files": [{"path": "rdh.c", "lang": "c", "total_lines": 762, "symbol_count": 23}, ...],
          "symbols": [...]                            # 所有文件符号合并(每个有 path 字段)
          "summary": "=== 5 files indexed (4321 lines total) ===\n\nrdh.c (762 lines, c)\n  ...\n\nhuff.c ..."
        }
    """
    if not path:
        return {"ok": False, "error": "path is required"}

    # 2026-05-02 part16:检测 glob 模式,多文件批量索引
    is_glob = "*" in path or "?" in path
    if is_glob:
        return await _handle_code_index_glob(
            ws_dir, path,
            include_includes=include_includes,
            name_filter=name_filter,
            kinds=kinds,
        )

    try:
        target = _safe_resolve(ws_dir, path)
    except ValueError as e:
        return {"ok": False, "error": str(e)}
    if not os.path.isfile(target):
        return _file_not_found_response(ws_dir, path)

    lang = _detect_lang(path)
    if lang is None:
        return {
            "ok": False,
            "error": (
                f"unsupported file type for code_index: {path}. "
                f"supported: .c .cc .cpp .h .hpp .py .js .jsx .ts .tsx .go .rs. "
                f"for plain text use read_file or search_in_file instead."
            ),
        }

    content, err = _read_text_safely(target)
    if err is not None:
        if isinstance(err.get("error"), str):
            err["error"] = err["error"].replace("<path>", path)
        return err

    lines = content.split("\n")
    total_lines = len(lines)
    rules = _CODE_INDEX_RULES[lang]["patterns"]
    symbols = []

    # ── C 语言:预处理多行函数签名 ──
    # 例: `size_t huff_encode(const uint8_t *input,\n   size_t len,\n   uint8_t *out)\n{`
    # 单行正则匹配不到。把跨行签名"折叠"成虚拟单行(line 仍记录起始行号)。
    # 检测启发:行末是 `,` 或 `(` 且下一行不是 `{` → 合并。
    #
    # 同时检测 `} typedef_name;` 形式(anonymous struct typedef 末尾命名)。
    if lang == "c":
        folded_lines = []  # [(orig_line_no, folded_line_text)]
        i = 0
        while i < total_lines:
            line = lines[i]
            stripped = line.rstrip()
            # 检测:行非空,行末以 `,` 或 `(` 结尾,本行有 `(`,且不在注释中
            if (stripped and (stripped.endswith(",") or stripped.endswith("("))
                    and "(" in stripped and not stripped.lstrip().startswith("//")
                    and not stripped.lstrip().startswith("*")):
                # 收集后续行直到看到 `)` 或 `;` 或 `{` 或 5 行上限
                merged = stripped
                j = i + 1
                merge_count = 0
                while j < total_lines and merge_count < 5:
                    next_l = lines[j].rstrip()
                    if not next_l:
                        j += 1; continue
                    merged += " " + next_l.lstrip()
                    merge_count += 1
                    if ")" in next_l or ";" in next_l or "{" in next_l:
                        break
                    j += 1
                folded_lines.append((i + 1, merged))
                i = j + 1
            else:
                folded_lines.append((i + 1, stripped))
                i += 1

        # 在折叠后的行上跑正则匹配
        for line_no, line_text in folded_lines:
            for pat, kind in rules:
                if not include_includes and kind in ("include", "import"):
                    continue
                m = pat.match(line_text)
                if m:
                    name = m.group(1)
                    if kind == "fn":
                        # 排除控制流关键字(行首恰好是这些)
                        if name in ("if", "while", "for", "switch", "return",
                                    "sizeof", "case", "do", "else"):
                            break
                        # 必须是定义不是 prototype:折叠后行末有 `){` 或 line_no 之后 5 行内有 `{`
                        has_brace = bool(re.search(r'\)\s*\{', line_text))
                        if not has_brace:
                            for k in range(line_no, min(total_lines, line_no + 6)):
                                if "{" in lines[k]:
                                    has_brace = True
                                    break
                                if ";" in lines[k]:
                                    break  # 是 prototype,不算
                            if not has_brace:
                                break
                    symbols.append({"line": line_no, "kind": kind, "name": name})
                    break

        # 额外:检测 anonymous typedef end-naming (`} mytype_t;`)
        for i, line in enumerate(lines, start=1):
            m = re.match(r'^\s*\}\s*(\w+)\s*;', line)
            if m:
                # 反查最近的 `typedef struct {` 起始行作为 line
                for back in range(i - 1, max(0, i - 50), -1):
                    if re.match(r'^\s*typedef\s+(?:struct|enum|union)\s*\{?\s*$', lines[back - 1]):
                        symbols.append({"line": back, "kind": "typedef", "name": m.group(1)})
                        break
                else:
                    symbols.append({"line": i, "kind": "typedef", "name": m.group(1)})
    else:
        # 非 C 语言:直接单行匹配
        for i, line in enumerate(lines, start=1):
            stripped = line.rstrip()
            if not stripped:
                continue
            for pat, kind in rules:
                if not include_includes and kind in ("include", "import"):
                    continue
                m = pat.match(stripped)
                if m:
                    name = m.group(1)
                    symbols.append({"line": i, "kind": kind, "name": name})
                    break

    # 排序 + 去重(C 路径有可能 anonymous typedef + 同行多规则匹配)
    seen = set()
    deduped = []
    for s in sorted(symbols, key=lambda x: (x["line"], x["kind"], x["name"])):
        key = (s["line"], s["name"])
        if key in seen:
            continue
        seen.add(key)
        deduped.append(s)
    symbols = deduped

    # ── 2026-05-02 part16:filter 应用 ──
    n_total_before_filter = len(symbols)
    if kinds:
        kinds_set = set(kinds)
        symbols = [s for s in symbols if s["kind"] in kinds_set]
    if name_filter:
        # 模糊匹配:把 * 当 .* 处理(类似 glob);其他作为 regex
        try:
            import fnmatch as _fnm
            if "*" in name_filter or "?" in name_filter:
                # glob 模式
                symbols = [s for s in symbols if _fnm.fnmatch(s["name"], name_filter)]
            else:
                # regex
                _pat = re.compile(name_filter)
                symbols = [s for s in symbols if _pat.search(s["name"])]
        except re.error:
            return {
                "ok": False,
                "error": f"invalid name_filter regex: {name_filter!r}. "
                         f"use glob (huff_*) or valid regex.",
            }

    # 生成紧凑摘要
    summary_lines = [f"{path} ({total_lines} lines, {lang})"]
    # kind 排序:include/import 在前,然后按行号
    includes = [s for s in symbols if s["kind"] in ("include", "import")]
    others = [s for s in symbols if s["kind"] not in ("include", "import")]
    if includes:
        # 多 include 折叠成一行,空间紧凑
        if len(includes) <= 3:
            for s in includes:
                summary_lines.append(f"  L{s['line']:>4d} {s['kind']:8s} {s['name']}")
        else:
            inc_summary = ", ".join(s["name"] for s in includes[:8])
            if len(includes) > 8:
                inc_summary += f" ... +{len(includes)-8} more"
            summary_lines.append(f"  includes: {inc_summary}")
    for s in others:
        summary_lines.append(f"  L{s['line']:>4d} {s['kind']:8s} {s['name']}")

    summary = "\n".join(summary_lines)

    result = {
        "ok": True,
        "action": "code_index",
        "path": path,
        "lang": lang,
        "total_lines": total_lines,
        "symbols": symbols,
        "symbol_count": len(symbols),
        "summary": summary,
    }
    if name_filter or kinds:
        result["filter_applied"] = {
            "name_filter": name_filter,
            "kinds": kinds,
            "before_filter": n_total_before_filter,
            "after_filter": len(symbols),
        }

    # 2026-05-15 P72: 重复 code_index 调用追踪 (类似 read_file 的 _redundant_read_warning)。
    # 病因(实测 comp_bench 内 881 次 code_index, comp_bench_simple 内 722 次):
    #   helper 反复 code_index 同样的几个 .c/.h 文件, 每次 1-3ms 但 helper LLM 多算一轮决策。
    #   code_index 输出基于文件内容 — 文件未变时结果不变, 重读浪费 helper iter。
    # 修法: 用 .code_index_history.json 追踪 (path, mtime, count); count >= 3 时附 _redundant
    #   warning + 建议用 read_function (精准函数级) 或读上次缓存内容
    try:
        ci_history_path = os.path.join(ws_dir, ".code_index_history.json")
        ci_history: dict = {}
        try:
            if os.path.isfile(ci_history_path):
                with open(ci_history_path, "r", encoding="utf-8") as f:
                    ci_history = json.load(f) or {}
        except (OSError, json.JSONDecodeError):
            ci_history = {}
        cur_mtime = 0.0
        try:
            cur_mtime = os.path.getmtime(target)
        except OSError:
            pass
        entry = ci_history.get(path)
        if isinstance(entry, dict):
            prev_mtime = float(entry.get("mtime", 0))
            count = int(entry.get("count", 0))
        else:
            prev_mtime, count = 0.0, 0
        # 文件未变 (mtime 相同) → 增计数; 变了 → 计数清零
        if abs(cur_mtime - prev_mtime) < 1.0 and prev_mtime > 0:
            count += 1
        else:
            count = 1
        ci_history[path] = {"mtime": cur_mtime, "count": count}
        try:
            with open(ci_history_path, "w", encoding="utf-8") as f:
                json.dump(ci_history, f, ensure_ascii=False)
        except OSError:
            pass
        if count >= 3:
            result["_redundant_code_index_warning"] = (
                f"`code_index` has been called for `{path}` {count} times without a file mtime change. "
                "The index result is stable while the file is unchanged. "
                f"For a specific function body, use read_function('{path}', '<fn>'); "
                "for file structure, reuse the previous index in context; if the file changed, this latest result is enough.\n\n"
                "文件未变时 code_index 结果稳定；复用已有结构或改用 read_function。"
            )
    except Exception:
        pass  # 追踪失败不影响主功能

    return result


# ── 2026-05-02 part16:read_function — 函数级精准读 ──
# 教训(trace 74b1295b):helper 想看 huff_decode 时反复 read_file(rdh.c, 600, 720)
# 但 huff_decode 实际是 600-680 行 — 它不知道精确边界,反复猜 + 反复 read。
# 此工具用 code_index 找起始行,然后扫到第一个完整闭合的 } 拿到函数 body。

async def handle_read_function(
    ws_dir: str, path: str, function_name: str, *,
    include_xref: bool = True,
    xref_scope: str = "file",
) -> dict:
    """精准读一个函数的完整 body — 自动定位边界。

    比 read_file(path, start, end) 优秀的地方:
    1. **不需要先知道边界**(read_file 让模型猜行号经常猜错,反复 read)
    2. **自动找 body 闭合**(扫到匹配的 `}` 或 dedent,函数完整无截断)
    3. **附带 callers/callees**(include_xref=True 时,扫**同文件**内谁调它 + 它调谁)

    Args:
        path: 工作区相对路径
        function_name: 函数名(必须精确匹配)
        include_xref: 是否扫同文件内的 caller/callee(默认 True)
        xref_scope: 2026-05-02 part16 加。
            "file"(默认):callers 仅扫同文件
            "workspace":callers 扫整个工作区所有源文件 — 跨文件调用关系全图
            workspace 模式适合"找 huff_decode 在整个项目哪里被调用"

    Returns:
        {
          "ok": True,
          "function_name": "huff_decode",
          "lines": [600, 678],              # 起止行号
          "body": "size_t huff_decode(...)\n{\n    ...\n}",
          "called_by": ["L720 in rdh_decompress: ...", ...],
                                              # xref_scope=workspace 时含 path 前缀
          "calls": ["br_init @ L605", ...],  # 此函数调用的其他符号
          "n_lines": 79,
          "xref_scope": "file" | "workspace",
        }

    错误返回 {"ok": False, "error": ...}。
    函数有多个同名(C 静态本地)→ 返回所有匹配项。
    """
    if not path or not function_name:
        return {"ok": False, "error": "path and function_name are required"}
    if xref_scope not in ("file", "workspace"):
        return {"ok": False, "error": f"xref_scope must be 'file' or 'workspace', got {xref_scope!r}"}
    try:
        target = _safe_resolve(ws_dir, path)
    except ValueError as e:
        return {"ok": False, "error": str(e)}
    if not os.path.isfile(target):
        return _file_not_found_response(ws_dir, path)

    lang = _detect_lang(path)
    if lang is None:
        return {
            "ok": False,
            "error": f"unsupported file type for read_function: {path}",
            "next_action_instruction": (
                "read_function only supports recognized source-code function bodies. For HTML, documents, or unknown "
                "text formats, use search_in_file to locate a function name or keyword, then read_file with "
                "start_line/end_line for the smallest relevant fragment. Avoid repeating read_function on the same "
                "unsupported file.\n"
                "read_function 只适合受支持源码；未知文本先搜索定位，再分段读取。"
            ),
            "suggested_tools": ["search_in_file", "read_file"],
        }

    # 用 code_index 找符号位置
    idx = await handle_code_index(ws_dir, path, include_includes=False)
    if not idx.get("ok"):
        return idx

    # 找匹配的函数(可能有多个,如 C 文件不同 .c 重复;同一文件内 static 函数也可能重名)
    # 兼容多语言的 kind 名:C/JS/Go/Rust 用 "fn",Python 用 "def"
    _FN_KINDS = {"fn", "def"}
    candidates = [s for s in idx["symbols"]
                  if s["kind"] in _FN_KINDS and s["name"] == function_name]

    if not candidates:
        # 给个提示:列出近似匹配
        all_fns = [s["name"] for s in idx["symbols"] if s["kind"] in _FN_KINDS]
        from difflib import get_close_matches
        suggestions = get_close_matches(function_name, all_fns, n=5, cutoff=0.6)
        sug_str = f" did you mean: {suggestions}?" if suggestions else ""
        return {
            "ok": False,
            "error": f"function '{function_name}' not found in {path}.{sug_str}",
            "available_functions": all_fns[:30],
        }

    content, err = _read_text_safely(target)
    if err is not None:
        return err
    lines = content.split("\n")
    total_lines = len(lines)

    bodies = []
    for cand in candidates[:5]:  # 上限 5 个同名匹配,够用了
        start = cand["line"]
        end = _find_function_end(lines, start, lang)
        if end is None:
            continue  # 找不到边界,跳过
        body = "\n".join(lines[start - 1:end])
        body_with_lineno = "\n".join(
            f"{i+start}: {lines[start - 1 + i]}"
            for i in range(end - start + 1)
        )

        result_item = {
            "function_name": function_name,
            "lines": [start, end],
            "n_lines": end - start + 1,
            "body": body_with_lineno,
        }

        # xref 扫描(同文件内 OR 全工作区)
        if include_xref:
            # 准备 fn_ranges:每个 fn 的 [start, end, name] — 让 caller 输出带函数名
            fn_ranges = []
            for s in idx["symbols"]:
                if s["kind"] == "fn":
                    s_start = s["line"]
                    s_end = _find_function_end(lines, s_start, lang)
                    if s_end:
                        fn_ranges.append((s_start, s_end, s["name"]))
            called_by = _find_callers_in_file(
                lines, function_name, exclude_range=(start, end),
                fn_ranges=fn_ranges,
            )
            calls = _find_callees_in_file(lines, start, end)

            # 2026-05-02 part16:workspace scope — 扫所有源文件找 callers
            if xref_scope == "workspace":
                # 用 search_across_files 扫整个工作区,然后**排除**当前文件
                # (那部分已经被 _find_callers_in_file 覆盖)
                xfile_pattern = rf'(?<![\w]){re.escape(function_name)}\s*\('
                cross_file = await handle_search_across_files(
                    ws_dir, xfile_pattern, is_regex=True,
                    max_results_per_file=8, max_files=50,
                )
                for fp, hits in (cross_file.get("results") or {}).items():
                    if fp == path:
                        continue  # 当前文件已经处理
                    for h in hits:
                        line_text = h["text"][:80]
                        called_by.append(f"{fp}:L{h['line']}: {line_text}")

            result_item["called_by"] = called_by[:15]  # workspace 模式上限放宽到 15
            result_item["calls"] = calls[:15]

        bodies.append(result_item)

    if not bodies:
        return {
            "ok": False,
            "error": f"function '{function_name}' located but body extraction failed (parser error)",
        }

    if len(bodies) == 1:
        # 单个匹配:扁平返回
        result = {"ok": True, "action": "read_function", "path": path,
                  "xref_scope": xref_scope if include_xref else "off"}
        result.update(bodies[0])
        return result

    # 多个匹配:返回数组
    return {
        "ok": True,
        "action": "read_function",
        "path": path,
        "function_name": function_name,
        "matches": bodies,
        "match_count": len(bodies),
        "xref_scope": xref_scope if include_xref else "off",
        "note": f"multiple functions named '{function_name}' found ({len(bodies)} matches)",
    }




# 函数调用扫描:`name(` 开头,前面不是字母/数字/_(避免 prefix 误匹配)
_CALLEE_PATTERN_CACHE: dict[str, "re.Pattern"] = {}


def _find_callers_in_file(
    lines: list[str], func_name: str, exclude_range: tuple,
    *, fn_ranges: list[tuple] | None = None,
) -> list[str]:
    """在文件里找哪些行调用了 func_name (排除 exclude_range 内的行,即函数自身体内)。

    2026-05-02 part16 加 fn_ranges 参数:[(start, end, name), ...] 函数边界列表。
    传入时 caller 输出会附带"调用方所在函数",格式 "L34 in rdh_decompress: ..."。
    这样模型一眼看到"谁调用我",而不只是"哪行调用我"。
    """
    pat = _CALLEE_PATTERN_CACHE.get(func_name)
    if pat is None:
        pat = re.compile(rf'(?<![\w]){re.escape(func_name)}\s*\(')
        _CALLEE_PATTERN_CACHE[func_name] = pat
    callers = []
    excl_start, excl_end = exclude_range
    for i, line in enumerate(lines, start=1):
        if excl_start <= i <= excl_end:
            continue
        if pat.search(line):
            # 找出 i 所属的函数名
            in_fn = ""
            if fn_ranges:
                for s, e, n in fn_ranges:
                    if s <= i <= e:
                        in_fn = f" in {n}"
                        break
            callers.append(f"L{i}{in_fn}: {line.strip()[:80]}")
    return callers




# ── 2026-05-02 part16:search_across_files — 跨文件 grep ──
# 教训(trace 74b1295b):helper 想找 "huff_build_canonical 在哪些地方被调用",
# 现有 search_in_file 只单文件,要分 N 次扫 N 个文件。跨文件能力是必需的。

async def handle_search_across_files(
    ws_dir: str, pattern: str, *,
    file_glob: str = "*",
    is_regex: bool = False,
    max_results_per_file: int = 5,
    max_files: int = 30,
    max_total_results: int = 100,
) -> dict:
    """跨工作区文件搜索 pattern。

    Args:
        pattern: 文本或 regex 模式
        file_glob: glob 通配符(*.c / *.py / **/*.h 等),默认 "*"  匹配所有文件
        is_regex: 默认 False(普通文本)
        max_results_per_file: 单文件最多返回行数(默认 5)
        max_files: 最多扫多少个匹配文件(默认 30,防扫超大项目)
        max_total_results: 全局结果上限(默认 100)

    Returns:
        {
          "ok": True,
          "pattern": "huff_build_canonical",
          "files_scanned": 8,
          "files_with_matches": 3,
          "total_matches": 12,
          "results": {
            "rdh.c": [
              {"line": 530, "text": "    huff_build_canonical(symbol_count, code_lens, codes);"},
              ...
            ],
            "rdh.h": [...]
          }
        }
    """
    if not pattern:
        return {"ok": False, "error": "pattern is required"}
    try:
        compiled = re.compile(pattern) if is_regex else re.compile(re.escape(pattern))
    except re.error as e:
        return {"ok": False, "error": f"invalid regex: {e}"}

    # 收集匹配的文件(用 glob)
    import fnmatch
    if not os.path.isdir(ws_dir):
        return {"ok": False, "error": "workspace dir not found"}

    candidate_files = []
    for root, dirs, files in os.walk(ws_dir):
        # 跳过隐藏目录 + node_modules / __pycache__ 等
        dirs[:] = [d for d in dirs
                   if not d.startswith(".")
                   and d not in {"__pycache__", "node_modules", ".git",
                                 ".pytest_cache", "_delegate_"}
                   and not d.startswith("_delegate_")]
        for fn in files:
            if fn.startswith("."):
                continue
            full_path = os.path.join(root, fn)
            rel_path = os.path.relpath(full_path, ws_dir)
            # glob 匹配相对路径或文件名
            if file_glob == "*" or fnmatch.fnmatch(rel_path, file_glob) or fnmatch.fnmatch(fn, file_glob):
                # 跳过二进制扩展名
                ext = os.path.splitext(fn)[1].lower()
                if ext in {".o", ".obj", ".exe", ".pyc", ".pyo", ".so", ".dll",
                           ".lib", ".a", ".bin", ".dat", ".png", ".jpg", ".jpeg",
                           ".gif", ".pdf", ".zip", ".tar", ".gz"}:
                    continue
                # 跳过过大文件
                try:
                    if os.path.getsize(full_path) > _SEARCH_FILE_SIZE_CAP:
                        continue
                except OSError:
                    continue
                candidate_files.append((rel_path, full_path))

    files_scanned = 0
    results: dict[str, list] = {}
    total_matches = 0
    truncated = False
    files_truncated = False

    for rel_path, full_path in sorted(candidate_files):
        if files_scanned >= max_files:
            files_truncated = True
            break
        if total_matches >= max_total_results:
            truncated = True
            break
        files_scanned += 1
        try:
            with open(full_path, "r", encoding="utf-8", errors="replace") as f:
                file_matches = []
                for line_no, line in enumerate(f, start=1):
                    if compiled.search(line):
                        file_matches.append({
                            "line": line_no,
                            "text": line.rstrip("\n").rstrip("\r")[:200],
                        })
                        if len(file_matches) >= max_results_per_file:
                            break
                if file_matches:
                    results[rel_path] = file_matches
                    total_matches += len(file_matches)
        except OSError:
            continue

    return {
        "ok": True,
        "action": "search_across_files",
        "pattern": pattern,
        "is_regex": is_regex,
        "file_glob": file_glob,
        "files_scanned": files_scanned,
        "files_with_matches": len(results),
        "total_matches": total_matches,
        "results": results,
        **({"files_truncated": True, "files_truncated_note":
            f"扫文件数达上限 {max_files},还有 {len(candidate_files) - files_scanned} 个未扫"}
           if files_truncated else {}),
        **({"results_truncated": True} if truncated else {}),
    }


# ── 命令执行 + 安全检查 ──────────────────────────────────────


def _translate_windows_command(cmd: str, ws_dir: str) -> str:
    """在 cmd.exe 上把 LLM 写的 Unix 风格命令翻译成 Windows 等价形式。

    返回翻译后的命令；无法翻译时返回原命令。
    """
    import re

    def _decode_python_c_line_escapes_outside_strings(src: str) -> str:
        """Decode shell-escaped line separators without changing Python string literals."""
        out: list[str] = []
        i = 0
        in_string = False
        quote = ""
        triple = False
        escaped = False
        while i < len(src):
            ch = src[i]
            if not in_string:
                if src.startswith("\\r\\n", i):
                    out.append("\n")
                    i += 4
                    continue
                if src.startswith("\\n", i):
                    out.append("\n")
                    i += 2
                    continue
                if src.startswith("\\r", i):
                    out.append("\n")
                    i += 2
                    continue
                if ch in {"'", '"'}:
                    quote = ch
                    triple = src.startswith(ch * 3, i)
                    in_string = True
                    escaped = False
                    if triple:
                        out.append(ch * 3)
                        i += 3
                    else:
                        out.append(ch)
                        i += 1
                    continue
                out.append(ch)
                i += 1
                continue

            out.append(ch)
            if triple:
                if src.startswith(quote * 3, i):
                    out.extend([quote, quote])
                    i += 3
                    in_string = False
                    quote = ""
                    triple = False
                else:
                    i += 1
                continue

            if escaped:
                escaped = False
                i += 1
                continue
            if ch == "\\":
                escaped = True
                i += 1
                continue
            if ch == quote:
                in_string = False
                quote = ""
            i += 1
        return "".join(out)

    def _normalize_python_c_source(src: str) -> str:
        normalized = src.replace(r'\"', '"').replace(r"\'", "'")
        normalized = _decode_python_c_line_escapes_outside_strings(normalized)
        normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
        return normalized

    def _strip_outer_command_quotes(src: str) -> str:
        src = src.strip()
        if len(src) >= 2 and src[0] == src[-1] and src[0] in ('"', "'"):
            return src[1:-1].strip()
        return src

    def _format_cmd_set_assignment(assign: str) -> str | None:
        """Return a cmd.exe-safe `set "NAME=value"` assignment.

        Plain `set NAME=value && ...` keeps the space before `&&` in the value
        on cmd.exe. Use `set NAME=value&& ...` so explicit command-local
        environment variables survive even when the sandbox removes inherited
        service variables such as PYTHONPATH.

        cmd.exe 的裸 set 会把 && 前空格写进变量值；用无空格 && 保留显式环境变量。
        """
        assign = (assign or "").strip()
        if len(assign) >= 2 and assign[0] == assign[-1] == '"':
            assign = assign[1:-1]
        if "=" not in assign:
            return None
        name, value = assign.split("=", 1)
        name = name.strip()
        value = value.strip()
        if not name or any(ch in name for ch in '"&|<>'):
            return None
        value = value.replace('"', "")
        return f"set {name}={value}"

    def _consume_leading_set_assignments(src: str) -> tuple[list[str], str] | None:
        assignments: list[str] = []
        pos = 0
        one = re.compile(
            r'\s*set\s+(?P<assign>"[^"]+"|[^&]+?)\s*&&\s*',
            re.IGNORECASE | re.DOTALL,
        )
        while True:
            m = one.match(src, pos)
            if not m:
                break
            formatted = _format_cmd_set_assignment(m.group("assign"))
            if not formatted:
                return None
            assignments.append(formatted)
            pos = m.end()
        if not assignments:
            return None
        rest = src[pos:].strip()
        if not rest:
            return None
        return assignments, rest

    stripped_cmd = cmd.strip()
    _CMD_WRAP_PAT = re.compile(r'^(?P<exe>cmd(?:\.exe)?)\s+/(?P<flag>[ck])\s+(?P<rest>.*)$', re.IGNORECASE | re.DOTALL)
    wrapped = _CMD_WRAP_PAT.match(stripped_cmd)
    if wrapped:
        rest = _strip_outer_command_quotes(wrapped.group("rest"))
        translated_rest = _translate_windows_command(rest, ws_dir)
        if translated_rest != rest:
            return f"{wrapped.group('exe')} /{wrapped.group('flag')} {translated_rest}"
        return cmd

    def _translate_windows_null_device(src: str) -> str:
        """Translate Unix null-device targets to cmd.exe's NUL device.

        Only rewrite when `/dev/null` is used as a redirection target or a
        command output target such as `curl -o /dev/null`; path arguments and
        URLs are left alone.
        """
        def repl(match: re.Match) -> str:
            return f"{match.group('prefix')}{match.group('quote') or ''}NUL{match.group('quote') or ''}"

        return re.sub(
            r'(?P<prefix>(?:\d?\s*>\s*|--output\s+|-o\s+))(?P<quote>["\']?)/dev/null(?P=quote)',
            repl,
            src,
            flags=re.IGNORECASE,
        )

    null_translated = _translate_windows_null_device(stripped_cmd)
    if null_translated != stripped_cmd:
        debug.log(
            "workspace.translate.null_device",
            f"/dev/null output target → NUL: {null_translated[:120]}",
        )
        stripped_cmd = null_translated
        cmd = stripped_cmd

    # set VAR=value && python -c "..." → preserve the env prefix safely, but still move -c code to a temp file.
    consumed_set = _consume_leading_set_assignments(stripped_cmd)
    if consumed_set:
        set_prefixes, rest = consumed_set
        translated_rest = _translate_windows_command(rest, ws_dir)
        translated = f"{'&& '.join(set_prefixes)}&& {translated_rest}"
        if translated != stripped_cmd:
            debug.log(
                "workspace.translate.set_and",
                f"set VAR=... && command → {translated[:120]}",
            )
            return translated

    # Unix env prefix on Windows: PYTHONPATH=x python -c "..." → set PYTHONPATH=x && python ...
    _PY_LAUNCHER_RE = r'(?:python3?|py)(?:\.(?:exe|cmd|bat))?'

    _ENV_PREFIX_PAT = re.compile(
        r'^(?P<assigns>(?:[A-Za-z_][A-Za-z0-9_]*=[^\s&|<>]+\s+)+)(?P<rest>'
        + _PY_LAUNCHER_RE + r'\s+.+)$',
        re.IGNORECASE | re.DOTALL,
    )
    m = _ENV_PREFIX_PAT.match(stripped_cmd)
    if m:
        assigns = [x for x in m.group("assigns").strip().split() if "=" in x]
        translated_rest = _translate_windows_command(m.group("rest").strip(), ws_dir)
        formatted = [_format_cmd_set_assignment(a) for a in assigns]
        if any(x is None for x in formatted):
            return cmd
        prefix = "&& ".join(x for x in formatted if x)
        translated = f"{prefix}&& {translated_rest}"
        debug.log(
            "workspace.translate.env_prefix",
            f"VAR=... python ... → {translated[:120]}",
        )
        return translated

    def _translate_simple_cp_command(src: str) -> str | None:
        try:
            import shlex
            parts = shlex.split(src, posix=False)
        except ValueError:
            return None
        if not parts or parts[0].lower() != "cp":
            return None
        opts = []
        paths = []
        for part in parts[1:]:
            clean = part.strip()
            if clean in {"-f", "--force"} and not paths:
                opts.append(clean)
                continue
            if clean.startswith("-"):
                return None
            paths.append(clean.strip("\"'"))
        if len(paths) != 2:
            return None
        workspace = Path(ws_dir).resolve()

        def _inside_workspace(raw: str) -> Path | None:
            if not raw or any(ch in raw for ch in "<>|&"):
                return None
            candidate = Path(raw)
            resolved = candidate.resolve() if candidate.is_absolute() else (workspace / candidate).resolve()
            try:
                resolved.relative_to(workspace)
            except ValueError:
                return None
            return resolved

        src_path = _inside_workspace(paths[0])
        dst_path = _inside_workspace(paths[1])
        if src_path is None or dst_path is None:
            return None
        try:
            import uuid
            tmp_name = f"_py_cmd_{uuid.uuid4().hex[:8]}.py"
            tmp_path = workspace / tmp_name
            script = (
                "from pathlib import Path\n"
                "import shutil\n"
                f"workspace = Path({str(workspace)!r}).resolve()\n"
                f"src = Path({str(src_path)!r}).resolve()\n"
                f"dst = Path({str(dst_path)!r}).resolve()\n"
                "for p in (src, dst):\n"
                "    p.relative_to(workspace)\n"
                "if src.is_dir():\n"
                "    shutil.copytree(src, dst, dirs_exist_ok=True)\n"
                "else:\n"
                "    dst.parent.mkdir(parents=True, exist_ok=True)\n"
                "    shutil.copy2(src, dst)\n"
            )
            tmp_path.write_text(script, encoding="utf-8")
            debug.log(
                "workspace.translate.cp",
                f"cp {paths[0]} {paths[1]} -> python {tmp_name}",
            )
            return f"python {tmp_name}"
        except OSError:
            return None

    translated_cp = _translate_simple_cp_command(cmd.strip())
    if translated_cp is not None:
        return translated_cp

    def _collapse_python_runner_forwarding(src: str) -> str | None:
        """python python3.cmd -c "..." is a runner-forwarding shape, not a Python script."""
        _PY_FORWARD_PAT = re.compile(
            r'^(?P<outer>' + _PY_LAUNCHER_RE + r')'
            r'(?P<opts>(?:\s+-(?!c(?:\s|$))[A-Za-z0-9]+(?:\s+[A-Za-z0-9_.:/\\-]+)?)*)'
            r'\s+(?P<shim>(?:\.?[\\/])?(?:python3?|py)\.(?:cmd|bat|exe))\s+'
            r'(?P<rest>-c\s+.*)$',
            re.IGNORECASE | re.DOTALL,
        )
        m = _PY_FORWARD_PAT.match(src.strip())
        if not m:
            return None
        shim = m.group("shim").replace("\\", "/")
        shim_name = os.path.basename(shim)
        rest = m.group("rest").strip()
        collapsed = f"{shim_name} {rest}"
        debug.log(
            "workspace.translate.python_runner_forward",
            f"{m.group('outer')} {m.group('shim')} -c ... → {collapsed[:120]}",
        )
        return collapsed

    def _translate_python_c_command(src: str, script_dir: str = "", display_prefix: str = "") -> str | None:
        # python [-B/-u/-X utf8/...] -c "..." → write a temporary .py file.
        # Keep interpreter flags before -c; they affect script execution too and
        # helpers often add `-B` for smoke checks.
        _PY_OPTS = r'(?P<opts>(?:\s+-(?!c(?:\s|$))[A-Za-z0-9]+(?:\s+[A-Za-z0-9_.:/\\-]+)?)*)'
        _PY_C_PAT = re.compile(
            r'^(?P<exe>' + _PY_LAUNCHER_RE + r')' + _PY_OPTS + r'\s+-c\s+'
            r'(?P<quote>["\'])(?P<code>.*)(?P=quote)\s*$',
            re.DOTALL,
        )
        _PY_C_WITH_SUFFIX_PAT = re.compile(
            r'^(?P<exe>' + _PY_LAUNCHER_RE + r')' + _PY_OPTS + r'\s+-c\s+'
            r'(?P<quote>["\'])(?P<code>.*)(?P=quote)\s+'
            r'(?P<suffix>(?:\d?>\S+|[<>|&].*)\s*)$',
            re.DOTALL,
        )
        m = _PY_C_PAT.match(src.strip())
        suffix = ""
        if not m:
            m = _PY_C_WITH_SUFFIX_PAT.match(src.strip())
            if m:
                suffix = " " + m.group("suffix").strip()
        if not m:
            return None
        inner_code = _normalize_python_c_source(m.group("code"))
        import uuid
        tmp_name = f"_py_cmd_{uuid.uuid4().hex[:8]}.py"
        tmp_rel = f"{display_prefix}{tmp_name}" if display_prefix else tmp_name
        tmp_path = os.path.join(ws_dir, script_dir, tmp_name) if script_dir else os.path.join(ws_dir, tmp_name)
        try:
            os.makedirs(os.path.dirname(tmp_path) or ws_dir, exist_ok=True)
            with open(tmp_path, "w", encoding="utf-8") as f:
                f.write(inner_code)
            debug.log(
                "workspace.translate.py_c",
                f"python ... -c ... → python {tmp_rel} (code written to file)",
            )
            return f"{m.group('exe')}{m.group('opts') or ''} {tmp_name if script_dir else tmp_rel}{suffix}"
        except OSError:
            return None

    collapsed_forwarding = _collapse_python_runner_forwarding(cmd.strip())
    if collapsed_forwarding is not None:
        translated_forwarding = _translate_python_c_command(collapsed_forwarding)
        if translated_forwarding is not None:
            return translated_forwarding
        return collapsed_forwarding

    translated_py_c = _translate_python_c_command(cmd.strip())
    if translated_py_c is not None:
        return translated_py_c

    # cd _helpers_shared/... && python foo.py → python _helpers_shared/.../foo.py
    # 病因(trace 85cfe2): helper 先 cd 进 _helpers_shared 子目录,脚本内部又把产物写到
    # _helpers_shared/... 相对路径,会重复拼前缀导致 FileNotFoundError。
    # 仅对共享目录里的简单 python 脚本执行做重写,避免改变一般 shell 语义。
    _CD_SHARED_PY_PAT = re.compile(
        r'^cd\s+(?P<subdir>(?:_helpers_shared|_shared)[^\s]*)\s+&&\s+'
        r'(?P<exe>' + _PY_LAUNCHER_RE + r')\s+(?P<script>[^\s"\']+\.py)(?P<args>(?:\s+.*)?)$',
        re.IGNORECASE,
    )
    m = _CD_SHARED_PY_PAT.match(cmd.strip())
    if m:
        subdir = m.group("subdir")
        script = m.group("script")
        args = m.group("args") or ""
        if not any(sep in script for sep in ("/", "\\")):
            translated = f"{m.group('exe')} {subdir}/{script}{args}"
            debug.log(
                "workspace.translate.cd_shared_python",
                f"cd {subdir} && python {script} → {translated[:120]}",
            )
            return translated

    # cd X && cmd → pushd X && cmd (或 cd /d X && cmd)
    _CD_AND_PAT = re.compile(r'^cd\s+(\S+)\s+&&\s+(.+)$', re.IGNORECASE)
    m = _CD_AND_PAT.match(cmd)
    if m:
        subdir, rest = m.group(1), m.group(2)
        py_c_in_subdir = _translate_python_c_command(rest, script_dir=subdir.replace("\\", "/"))
        translated_rest = py_c_in_subdir if py_c_in_subdir is not None else _translate_windows_command(rest, ws_dir)
        translated = f"cd /d {subdir} && {rest}"
        if translated_rest != rest:
            translated = f"cd /d {subdir} && {translated_rest}"
        debug.log(
            "workspace.translate.cd_and",
            f"cd X && Y → {translated[:120]}",
        )
        return translated

    return cmd




from app.llm.tools.workspace_run import handle_run  # noqa: E402,F401







# ── 2026-05-02 part14:stdout 智能摘要工具函数 ──
# 教训(trace 74b1295b iter 89):helper 自加 fprintf 把 stdout 从 700→13K chars,
# PASS/FAIL 总结被淹没。这两个函数把"先看 head/tail + grep 关键词"的人类阅读策略
# 内置到工具层。









# ═══════════════════════════════════════════════════════════════
# 安全检查核心
# ═══════════════════════════════════════════════════════════════


# ── CMD 安全检查 ──────────────────────────────────────────────




# ── 路径工具 ──────────────────────────────────────────────────








# ─── 2026-05-02 part21:TodoWrite 工具(参考 Claude Code 设计) ─────────
# 核心洞察:LLM 在做长链任务时容易"忘事"。Claude Code 的 TodoWrite 让模型把
# 任务分解外化到工作区,每完成一项 mark complete。这个简单机制让 Claude Code
# 在长任务上比"靠 prompt 记住" 强得多。
#
# 实现:.todos.json 存储,显示用 ☐/▶/✓ 三态可视化。
# 强制约束:同时只能 1 个 in_progress；多余项自动降为 pending,避免模型卡在重复纠错循环

# todo 工具处理器已抽离到 todo_handlers.py(2026-05-20 重构);re-export 兼容
# (registry 经 ws_tool.handle_todo_* 访问、测试经 from workspace import,均由 re-export 覆盖)。
from app.llm.tools.todo_handlers import (  # noqa: E402,F401
    handle_todo_write,
    handle_todo_read,
)




# ── L4-1 (2026-05-09): workspace.locate —— 全工作区模糊文件搜索 ──

_LOCATE_MAX_MATCHES = 80
_LOCATE_SKIP_PREFIXES = ("_downloaded_media", ".git", "__pycache__", ".claude")
# 2026-05-10 Patch 80b: 图片/媒体扩展名 — pattern 命中时自动包含 _downloaded_media/
# 病因(trace 14:58):用户发图后主线程 workspace.locate "*.jpg" 找不到图,
# 因为 _LOCATE_SKIP_PREFIXES 默认跳过 _downloaded_media/。主线程只能用 dir 命令绕过,
# 多浪费 1-2 个工具调用。修法:locate 检测 pattern 是图片/媒体扩展名时,自动允许扫
# _downloaded_media/(用户上传的媒体本就该被找到)。
_MEDIA_EXTENSIONS = (
    ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".tiff",  # 图片
    ".mp4", ".mov", ".avi", ".webm",  # 视频
    ".mp3", ".wav", ".ogg", ".m4a", ".flac",  # 音频
)


def _pattern_is_media(pattern: str) -> bool:
    """检测 glob pattern 是否在找媒体文件(图片/视频/音频)。
    例:'*.jpg' / '*.png' / 'img_*' / '*media*' → True
    """
    p = pattern.lower()
    if any(p.endswith(ext) or p.endswith(ext + "*") for ext in _MEDIA_EXTENSIONS):
        return True
    if "img_" in p or "media" in p or "image" in p:
        return True
    return False


def _is_internal_metadata_file(path: str) -> bool:
    name = os.path.basename(path)
    ext = os.path.splitext(name)[1].lower()
    return (
        name in {".read_history.json", ".todos.json", ".edit_history.json", ".session_tag"}
        or name in _NON_DELIVERABLE_BASENAMES
        or name.endswith(_NON_DELIVERABLE_SUFFIXES)
        or ext in _NON_DELIVERABLE_EXTS
        or name.startswith(".helper_")
        or name.startswith(_ARTIFACT_PREFIXES)
    )


async def handle_locate(ws_dir: str, pattern: str) -> dict:
    """扫描整个工作区(含 _delegate_* / _helpers_shared),按 glob 模式匹配文件。

    跳过 _downloaded_media、.git、__pycache__ 等非生产目录。
    **2026-05-10 P80b**: 找媒体文件时(*.jpg / img_* 等)自动包含 _downloaded_media/。
    pattern 不含通配符(*、?、[)时自动包裹为 *substring*。
    结果按层级(main/temp/delegate/shared)分类,最多 _LOCATE_MAX_MATCHES 条。
    """
    import fnmatch

    if not pattern or not pattern.strip():
        return {"ok": False, "error": "Missing pattern. For example: pattern='*.docx' or pattern='chart'.\n缺少 pattern 字段。"}
    pattern = pattern.strip()
    if not any(ch in pattern for ch in ("*", "?", "[")):
        pattern = f"*{pattern}*"
    basename_pattern = pattern
    rel_pattern = pattern
    if "/" not in pattern and "\\" not in pattern:
        rel_pattern = f"*/{pattern}"
        basename_pattern = f"*{pattern}"

    matches: list[dict] = []
    total = 0
    ws_root = os.path.normpath(ws_dir)
    skip_prefixes = _LOCATE_SKIP_PREFIXES
    if _pattern_is_media(pattern):
        skip_prefixes = tuple(p for p in skip_prefixes if p != "_downloaded_media")

    for dirpath, dirnames, filenames in os.walk(ws_root):
        dirnames[:] = [d for d in dirnames if not d.startswith(skip_prefixes)]
        rel_dir = os.path.relpath(dirpath, ws_root)
        if rel_dir == ".":
            layer = "main"
        elif os.path.normpath(rel_dir).startswith(os.path.normpath(".temp")):
            layer = "temp"
        else:
            layer = "main"

        for fname in filenames:
            full = os.path.join(dirpath, fname)
            rel = os.path.relpath(full, ws_root)
            if _is_internal_metadata_file(rel):
                continue
            if fnmatch.fnmatch(fname, basename_pattern) or fnmatch.fnmatch(rel.replace("\\", "/"), rel_pattern):
                total += 1
                if len(matches) >= _LOCATE_MAX_MATCHES:
                    continue
                # 细化层级
                if "_delegate_" in rel:
                    layer = "delegate"
                elif "_helpers_shared" in rel:
                    layer = "shared"
                elif "_downloaded_media" in rel:
                    layer = "media"  # P80b: 标记媒体层
                matches.append({
                    "path": rel.replace("\\", "/"),
                    "layer": layer,
                    "size": os.path.getsize(full),
                })

    truncated = total > len(matches)
    return {
        "ok": True,
        "pattern": pattern,
        "total": 80 if total >= 80 and _LOCATE_MAX_MATCHES == 80 else total,
        "total_matches": total,
        "returned": len(matches),
        "truncated": truncated,
        "matches": matches,
        "hint": (
            None if matches
            else f"没有匹配 '{pattern}' 的文件。试试更宽的模式(如 '*.docx')或检查文件名拼写。"
        ),
    }


# ── 2026-05-09 Patch 22: inspect_file —— narrow structured metadata inspection ──
#
# 病因(trace 963c236c5a6f4f27):主线程编码能力弱于 helper(对长串迭代写代码
# 效率低正确率低),所以默认策略是把代码工作 delegate 给 helper。但 helper 经常
# 产出 docx/pptx/xlsx/png 这类二进制,helper 的 report 可能缺结构事实或有矛盾。
# 主线程需要窄元数据工具处理这种缺口,但又不该为每个简单结构事实写脚本(就算 < 150 行
# 自己写没问题,反复来也烦)。
#
# 解决:inspect_file 是预制工具,纯 ZIP/header 解析,**不需要主线程写代码**。
# 给主线程结构化元数据:页数/表格数/图片数/宽高/首行预览等,让它一调即得。
# 常见结构事实靠 inspect_file 一行搞定;字段不够再写小 .py(Patch 28 允许);
# 真正复杂校验交给 verify/helper。三层阶梯见 Round 2 prompt §5.7。

_INSPECT_TEXT_EXTS = {
    ".txt", ".md", ".json", ".csv", ".tsv", ".yaml", ".yml", ".xml",
    ".log", ".ini", ".toml", ".cfg",
    ".py", ".c", ".h", ".cpp", ".hpp", ".cc", ".cxx",
    ".js", ".ts", ".jsx", ".tsx",
    ".go", ".rs", ".java", ".kt", ".rb", ".sh", ".bat", ".ps1",
    ".html", ".css", ".sql",
}
_INSPECT_MAX_PREVIEW_LINES = 10
_INSPECT_MAX_PREVIEW_LINE_LEN = 200


def _inspect_text(abs_path: str) -> dict:
    """文本类:行数 + 头 N 行预览。"""
    try:
        with open(abs_path, "r", encoding="utf-8", errors="replace") as f:
            # 读全文做行数;限制 1MB 防 OOM
            data = f.read(1024 * 1024)
            truncated = len(data) >= 1024 * 1024
        lines = data.split("\n")
        preview = []
        for ln in lines[:_INSPECT_MAX_PREVIEW_LINES]:
            if len(ln) > _INSPECT_MAX_PREVIEW_LINE_LEN:
                ln = ln[:_INSPECT_MAX_PREVIEW_LINE_LEN] + "...[行截]"
            preview.append(ln)
        return {
            "line_count": len(lines) if not truncated else f">{len(lines)} (1MB 截)",
            "char_count": len(data),
            "preview_first_lines": preview,
        }
    except OSError as e:
        return {"error": f"text inspection failed: {e}"}


# 文件类型探查族已抽离到 file_inspect.py(2026-05-20 重构);re-export 兼容。
from app.llm.tools.file_inspect import (  # noqa: E402,F401
    _inspect_docx,
    _inspect_pptx,
    _inspect_xlsx,
    _inspect_image,
    _inspect_pdf,
    _inspect_wav,
    _inspect_warnings,
)














_INSPECT_EXT_TYPE_MAP = {
    ".docx": "docx", ".pptx": "pptx", ".xlsx": "xlsx",
    ".pdf": "pdf", ".wav": "wav",
    ".png": "image", ".jpg": "image", ".jpeg": "image",
    ".gif": "image", ".bmp": "image", ".webp": "image",
}


def inspect_file(workspace_dir: str, path: str) -> dict:
    """Narrow structured metadata tool entry point.

    返回结构:
      {
        "ok": True/False,
        "path": <relative path>,
        "size_bytes": int,
        "mtime": float,
        "type": "docx"|"pptx"|"xlsx"|"pdf"|"wav"|"image"|"text"|"binary"|"unknown",
        "metadata": { type-specific },
        "warnings": [可选,structural 异常],
        "error": 仅在 ok=False 时
      }

    设计原则:
      - 不执行任何 user 代码 / Python — 这是预制只读校验,主线程一调即得
      - 纯 ZIP / header 解析,失败时返回 error 字段而不抛异常
      - 给出主线程能据此决策的具体数字(段落/表格/图片/页/行/尺寸)
      - 是验证三层阶梯的第一层(Round 2 §5.7);字段不够时主线程可写 ≤5KB .py 自验
    """
    try:
        abs_path = _safe_resolve(workspace_dir, path)
    except ValueError as e:
        return {"ok": False, "error": f"Invalid path: {e}.\n路径无效。"}

    if os.path.isdir(abs_path):
        display_path = path or "."
        entries: list[dict] = []
        truncated = False
        try:
            names = sorted(os.listdir(abs_path))
            truncated = len(names) > 80
            for name in names[:80]:
                full = os.path.join(abs_path, name)
                entries.append({
                    "name": name,
                    "type": "directory" if os.path.isdir(full) else "file",
                    **({"size_bytes": os.path.getsize(full)} if os.path.isfile(full) else {}),
                })
        except OSError:
            entries = []
        return {
            "ok": False,
            "path": display_path,
            "type": "directory",
            "error": (
                f"Path is a directory, not a file: {display_path}.\n"
                "Use workspace(action='list' or 'locate'), search_files, env_read/env_search, "
                "or inspect a concrete file path. In project mode, `_env/...` is a sparse staged "
                "directory tree, not a single readable file.\n"
                "路径是目录而不是文件；请列目录、搜索或检查具体文件。项目模式下 _env 是稀疏暂存目录。"
            ),
            "directory_entries": entries,
            "entries_truncated": truncated,
            "suggested_tools": ["workspace list", "workspace locate", "search_files", "env_read"],
        }

    if not os.path.isfile(abs_path):
        return {"ok": False, "error": f"File does not exist: {path}.\n文件不存在。"}

    try:
        size = os.path.getsize(abs_path)
        mtime = os.path.getmtime(abs_path)
    except OSError as e:
        return {"ok": False, "error": f"Failed to read file metadata: {e}.\n读元信息失败。"}

    ext = os.path.splitext(path)[1].lower()
    file_type = _INSPECT_EXT_TYPE_MAP.get(ext)
    if file_type is None and ext in _INSPECT_TEXT_EXTS:
        file_type = "text"
    if file_type is None:
        file_type = "binary"

    base = {
        "ok": True,
        "path": path,
        "size_bytes": size,
        "mtime": round(mtime, 1),
        "type": file_type,
    }

    if file_type == "docx":
        base["metadata"] = _inspect_docx(abs_path)
    elif file_type == "pptx":
        base["metadata"] = _inspect_pptx(abs_path)
    elif file_type == "xlsx":
        base["metadata"] = _inspect_xlsx(abs_path)
    elif file_type == "image":
        base["metadata"] = _inspect_image(abs_path, ext)
    elif file_type == "pdf":
        base["metadata"] = _inspect_pdf(abs_path)
    elif file_type == "wav":
        base["metadata"] = _inspect_wav(abs_path)
    elif file_type == "text":
        base["metadata"] = _inspect_text(abs_path)
    else:
        # binary unknown:只给 size + 头部 16 字节 hex
        try:
            with open(abs_path, "rb") as f:
                head = f.read(16)
            base["metadata"] = {
                "head_hex": head.hex(),
                "note": f"Unrecognized format ({ext or 'no extension'}); only the file header is shown.\n未识别格式，仅返回文件头。",
            }
        except OSError as e:
            base["metadata"] = {"error": f"binary read failed: {e}"}

    warnings = _inspect_warnings(base)
    if warnings:
        base["warnings"] = warnings

    return base
