"""docx 文档处理:正文枚举、孤立媒体清理、文本清洗、断行 run、默认样式、占位符查找、
通用表格抽取、part 内图片枚举、读取、块删除 + 返回上限常量。

2026-05-20 重构: 从 llm/tools/office.py 原样抽出。closure 自包含(13 符号, 0 unsafe);
_err 复用 office_common;模块自建 logger;python-docx 在函数内惰性 import。office.py re-export 兼容。
"""
from __future__ import annotations

import asyncio
import json
import logging
import os
import re
import shutil
import tempfile
import zipfile
from typing import Any

from app.llm.tools.office_common import _err

log = logging.getLogger(__name__)


def _docx_open_error(e: Exception, rel_path: str) -> str:
    return (
        f"open failed: {type(e).__name__}: {e}. The DOCX package may be damaged or half-written: {rel_path}. "
        "If a sibling .bak file exists, copy it to a new clean filename and continue from that backup; otherwise "
        "rebuild the document from verified source/evidence instead of repeatedly editing this damaged file.\n"
        "DOCX 可能已损坏或半写入；优先用 .bak 复制成干净文件继续，没有备份则从证据重建。"
    )


def _save_docx_atomic(doc, target: str) -> tuple[int, str | None]:
    directory = os.path.dirname(target) or "."
    os.makedirs(directory, exist_ok=True)
    fd, tmp_path = tempfile.mkstemp(prefix=".docx_tmp_", suffix=".docx", dir=directory)
    os.close(fd)
    bak_path = target + ".bak"
    try:
        doc.save(tmp_path)
        with zipfile.ZipFile(tmp_path, "r") as zf:
            bad_member = zf.testzip()
        if bad_member:
            return 0, f"save verification failed: bad zip member {bad_member}"
        if os.path.exists(target):
            try:
                shutil.copy2(target, bak_path)
            except OSError:
                pass
        os.replace(tmp_path, target)
        return os.path.getsize(target) if os.path.exists(target) else 0, None
    except Exception as e:
        return 0, f"save failed: {type(e).__name__}: {e}"
    finally:
        try:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
        except OSError:
            pass


# 各种保护性上限
_MAX_PARAGRAPHS_RETURNED = 2000


_MAX_TABLE_ROWS_RETURNED = 200


_MAX_TABLE_COLS_RETURNED = 30


async def _docx_read(workspace_dir: str, target: str, rel_path: str, args: dict) -> str:
    if not os.path.isfile(target):
        return _err(f"file not found: {rel_path}")

    from docx import Document

    def _do_read():
        try:
            doc = Document(target)
        except Exception as e:
            return None, _docx_open_error(e, rel_path)

        # 2026-05-04 v19.1: 用 _docx_enumerate_body 同时输出 block_index
        # 让 LLM 后续可以直接 replace_block / delete_block / insert_block
        items = _docx_enumerate_body(doc)
        try:
            start_block = int(args.get("start_block", 0) or 0)
        except (TypeError, ValueError):
            start_block = 0
        try:
            end_block_raw = args.get("end_block", None)
            end_block = int(end_block_raw) if end_block_raw is not None else -1
        except (TypeError, ValueError):
            end_block = -1
        try:
            max_blocks = int(args.get("max_blocks", 0) or 0)
        except (TypeError, ValueError):
            max_blocks = 0
        start_block = max(0, start_block)
        if end_block >= 0:
            selected_items = [it for it in items if start_block <= int(it.get("index", -1)) <= end_block]
        else:
            selected_items = [it for it in items if int(it.get("index", -1)) >= start_block]
        if max_blocks > 0:
            selected_items = selected_items[:max_blocks]
        paragraphs: list[dict] = []
        headings: list[dict] = []
        truncated_paragraphs = False
        para_count = 0
        for it in selected_items:
            if it["kind"] == "table" or it["kind"] == "other":
                continue
            if para_count >= _MAX_PARAGRAPHS_RETURNED:
                truncated_paragraphs = True
                break
            text = it["text"] or ""
            if len(text) > 1500:
                text = text[:1500] + "…[段落截断]"
            entry = {
                "block_index": it["index"],   # ← LLM 可用于 replace_block 等
                "kind": it["kind"],
                "text": text,
            }
            if it["kind"] == "heading":
                entry["level"] = it.get("level", 1)
                headings.append({
                    "block_index": it["index"],
                    "level": it.get("level", 1),
                    "text": text,
                })
            paragraphs.append(entry)
            para_count += 1

        tables = _extract_tables_generic(doc.tables)
        images = _list_images_in_part(doc.part)

        # 2026-05-17 P153: 一致性自检 — 找正文里提到的"图 N" vs 实际嵌入的 image
        # 帮 helper 在写完文档后看出"我说有图但实际没嵌入"的情况。
        figure_refs: list[str] = []
        for it in selected_items:
            if it["kind"] == "table" or it["kind"] == "other":
                continue
            text = it["text"] or ""
            # 匹配中文 "图 1"/"图1"/英文 "Figure 1"/"Fig. 1"
            for m in re.finditer(r"图\s*(\d+)|[Ff]ig(?:ure)?\.?\s*(\d+)", text):
                num = m.group(1) or m.group(2)
                if num and num not in figure_refs:
                    figure_refs.append(num)

        # 嵌入的 image 数量与正文提到的图编号数量比对
        # 简单启发: 如果正文提到 N 张图但只嵌入了 M 张 (M < N),flag 出来
        figure_consistency = None
        if figure_refs:
            n_referenced = len(figure_refs)
            n_embedded = len(images)
            if n_embedded < n_referenced:
                figure_consistency = {
                    "warning": "figure_count_mismatch",
                    "referenced_in_text": sorted(figure_refs, key=lambda x: int(x)),
                    "referenced_count": n_referenced,
                    "embedded_count": n_embedded,
                    "hint": (
                        f"The document text references {n_referenced} figure(s) ({figure_refs}), but only "
                        f"{n_embedded} image block(s) are embedded. Check whether each figure reference has a "
                        f"matching {{type:'image', path:'...'}} block.\n"
                        "正文提到的图数量多于实际嵌图数量。"
                    ),
                }

        # 孤儿 image rels 检测 (有 rId 引用但 image part 也存在,但反之亦然)
        # 这里只能看到 main part 的 rels — 不深扫 package
        try:
            from docx.oxml.ns import qn
            R_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
            used_rids: set[str] = set()
            for blip in doc.element.iter(qn("a:blip")):
                rid = blip.get(R_NS + "embed")
                if rid:
                    used_rids.add(rid)
            orphan_rels: list[str] = []
            for rid, rel in doc.part.rels.items():
                if "image" in (getattr(rel, "reltype", "") or "").lower():
                    if rid not in used_rids:
                        orphan_rels.append(getattr(rel, "target_ref", "") or rid)
        except Exception:
            orphan_rels = []

        return {
            "paragraphs": paragraphs,
            "headings": headings,
            "truncated_paragraphs": truncated_paragraphs,
            "paragraph_count": len(doc.paragraphs),
            "block_count": len(items),
            "returned_block_start": start_block,
            "returned_block_end": selected_items[-1]["index"] if selected_items else None,
            "returned_block_count": len(selected_items),
            "has_more_blocks": bool(selected_items and selected_items[-1]["index"] < items[-1]["index"]),
            "next_start_block": (selected_items[-1]["index"] + 1) if selected_items and selected_items[-1]["index"] < items[-1]["index"] else None,
            "table_count": len(doc.tables),
            "tables": tables,
            "images": images,
            "figure_consistency": figure_consistency,
            "orphan_image_rels": orphan_rels,
        }, None

    result, err = await asyncio.to_thread(_do_read)
    if err:
        return _err(err)

    out = {
        "ok": True, "action": "read", "format": "docx", "path": rel_path,
        "paragraph_count": result["paragraph_count"],
        "block_count": result["block_count"],
        "returned_block_start": result["returned_block_start"],
        "returned_block_end": result["returned_block_end"],
        "returned_block_count": result["returned_block_count"],
        "has_more_blocks": result["has_more_blocks"],
        **({"next_start_block": result["next_start_block"]} if result["next_start_block"] is not None else {}),
        "table_count": result["table_count"],
        "image_count": len(result["images"]),
        "headings": result["headings"],   # 章节大纲(供 replace_section 用)
        "paragraphs": result["paragraphs"],
        "tables": result["tables"],
        "images": result["images"],
        **({"truncated_paragraphs": True} if result["truncated_paragraphs"] else {}),
        "hint": (
            "Image metadata is available in the `images` field, including width, height, and filename. "
            "Use office(action='ocr_images') only when text inside images is needed. For image count or render-state "
            "checks, use image_count without OCR. Use extract_images to release embedded images as files. "
            "For large documents, read body text with start_block/end_block/max_blocks; for image text, batch "
            "ocr_images with image_offset/max_images/save_to and then read the saved text in chunks. "
            "For incremental edits, use replace_block, replace_blocks, or replace_section with heading_text.\n"
            "图片元数据在 images 字段；大文档和图片文字都应分段读取。"
        ),
    }
    # 2026-05-17 P153: figure_consistency 警告 (只在有问题时附加)
    if result.get("figure_consistency"):
        out["figure_consistency"] = result["figure_consistency"]
    if result.get("orphan_image_rels"):
        out["orphan_image_rels"] = result["orphan_image_rels"]
        out["orphan_image_hint"] = (
            f"The DOCX has {len(result['orphan_image_rels'])} orphan image relationship(s): image rels exist but are not referenced by the body. "
            "The next replace_section, replace_block, or delete_block call will clean them automatically. "
            "You may also trigger cleanup with an equivalent replace_block call if no content change is needed.\n\n"
            "文档存在孤儿图片引用；后续块级编辑会自动清理。"
        )
    save_to = str(args.get("save_to", "") or "").strip()
    if save_to:
        try:
            from app.llm.tools import workspace as ws_tool

            save_path = ws_tool._safe_resolve(workspace_dir, save_to)
            os.makedirs(os.path.dirname(save_path) or ".", exist_ok=True)
            lines: list[str] = []
            for block in result["paragraphs"]:
                text = str(block.get("text") or "")
                if text:
                    lines.append(text)
            for table_idx, table in enumerate(result["tables"], 1):
                lines.append(f"[table {table_idx}]")
                for row in table.get("rows", []) or []:
                    lines.append("\t".join(str(cell) for cell in row))
            body = "\n".join(lines)
            with open(save_path, "w", encoding="utf-8") as fh:
                fh.write(body)
            out["saved_to"] = save_to
            out["saved_text_lines"] = body.count("\n") + (1 if body else 0)
            out["saved_text_bytes"] = len(body.encode("utf-8"))
            out["hint"] += (
                f"\nRead body text was also saved to {save_to}; use read_file with line ranges if you need "
                "paged text evidence. office(read) still returns structured blocks and tables directly.\n"
                "正文已保存到 save_to，可用 read_file 分段读取；office(read) 本身也已返回结构化内容。"
            )
        except (OSError, ValueError) as e:
            out["save_to_error"] = f"{type(e).__name__}: {e}"
    return json.dumps(out, ensure_ascii=False)


def _docx_enumerate_body(doc):
    """枚举 docx body 的所有顶层元素,按文档顺序返回 (index, kind, element, text_preview)。

    kind ∈ {"heading", "paragraph", "table", "other"}
    text_preview: heading/paragraph 是纯文本,table 是首行 join,other 是空字符串
    """
    from docx.oxml.ns import qn

    body = doc.element.body
    items = []
    para_iter = iter(doc.paragraphs)
    table_iter = iter(doc.tables)

    for child in body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            try:
                p = next(para_iter)
            except StopIteration:
                continue
            style = (p.style.name or "").strip() if p.style else ""
            if style.lower().startswith("heading"):
                # 提取 heading level
                try:
                    level = int(style.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                items.append({
                    "index": len(items),
                    "kind": "heading",
                    "level": level,
                    "text": p.text,
                    "_element": p._element,
                    "_para": p,
                })
            else:
                items.append({
                    "index": len(items),
                    "kind": "paragraph",
                    "text": p.text,
                    "_element": p._element,
                    "_para": p,
                })
        elif tag == qn("w:tbl"):
            try:
                t = next(table_iter)
            except StopIteration:
                continue
            preview = ""
            if t.rows:
                preview = " | ".join(
                    (c.text or "").strip() for c in t.rows[0].cells
                )[:120]
            items.append({
                "index": len(items),
                "kind": "table",
                "text": preview,
                "_element": t._element,
                "_table": t,
            })
        elif tag == qn("w:sectPr"):
            # section properties 不算 block,跳过
            continue
        else:
            items.append({
                "index": len(items),
                "kind": "other",
                "text": "",
                "_element": child,
            })
    return items


# 2026-05-17 P152: 清理孤儿 image rels
# 病因: replace_section / replace_block / delete_block 在 python-docx 里只删
# document.xml 的 <w:p> 元素, 但不动 word/_rels/document.xml.rels 里的 rId 引用,
# 也不删 word/media/image*.png 文件本身。结果是 docx 内部留着死引用 + 死图片,
# 文档体积虚高 (一次 trace 里 10 张图只用 4 张, 6 张孤儿)。
def _drop_orphan_media(doc) -> dict:
    """删除 document.xml 不再引用的 image rels 及 media 文件。

    返回 {"dropped_rels": N, "dropped_filenames": [...]}

    2026-05-18 P168: 修复 KeyError 'v' — qn('v:imagedata') 在 python-docx 默认 nsmap
    中找不到 'v' 前缀, 直接 raise KeyError。改用原生命名空间 URI 字符串避开
    qn() 解析。修复前: replace_block / replace_section / delete_block 全部失败。
    """
    from docx.oxml.ns import qn
    R_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
    # VML namespace URI (不走 qn() 因为 'v' 不在 docx 默认 nsmap 里)
    V_NS = "{urn:schemas-microsoft-com:vml}"

    # 1. 扫描 body 找所有还在用的 rIds (a:blip r:embed)
    used_rids: set[str] = set()
    for blip in doc.element.iter(qn("a:blip")):
        rid = blip.get(R_NS + "embed") or blip.get(R_NS + "link")
        if rid:
            used_rids.add(rid)
    # v:imagedata (老式 VML 图片) — 直接用 namespace URI 避开 qn()
    for img in doc.element.iter(V_NS + "imagedata"):
        rid = img.get(R_NS + "id") or img.get(R_NS + "pict")
        if rid:
            used_rids.add(rid)

    # 2. headers / footers 也会引用 image,但 main doc 的 rels 通常不被它们用
    # (header/footer 自己有 rels)。保守起见,只清理 main part 的 rels。

    dropped_filenames: list[str] = []
    dropped_count = 0
    # 复制 rels keys 避免迭代中修改
    rel_ids = list(doc.part.rels.keys())
    for rid in rel_ids:
        try:
            rel = doc.part.rels[rid]
        except KeyError:
            continue
        # 只处理 image 类型
        rel_type = getattr(rel, "reltype", "") or ""
        if "image" not in rel_type.lower():
            continue
        if rid in used_rids:
            continue
        # 孤儿 — 删 rel
        target_ref = getattr(rel, "target_ref", "") or ""
        try:
            doc.part.drop_rel(rid)
            dropped_count += 1
            if target_ref:
                dropped_filenames.append(target_ref)
        except Exception:
            # drop_rel 失败也继续 — 孤儿 rel 不致命
            pass

    # 3. 清理 word/media/ 里没人引用的 image part (Package 级别)
    # python-docx 不自动 GC image parts。要手动从 package 里移除。
    try:
        package = doc.part.package
        # 收集所有 part 的 target_ref → part 的反向映射 + 用到的 image partname
        used_partnames: set[str] = set()
        # 重新扫一遍所有 rels (包括 headers/footers 的)
        for part_uri, part in package.iter_parts():
            for rel in getattr(part, "rels", {}).values():
                if "image" in (getattr(rel, "reltype", "") or "").lower():
                    tgt_part = getattr(rel, "target_part", None)
                    if tgt_part is not None:
                        used_partnames.add(str(tgt_part.partname))

        # 找出 word/media/ 下没人引用的 image part
        orphan_parts = []
        for part_uri, part in list(package.iter_parts()):
            pn = str(part.partname)
            if pn.startswith("/word/media/") and pn not in used_partnames:
                orphan_parts.append(part)

        for op in orphan_parts:
            try:
                # python-docx Package 没有公开 drop API
                # 但内部 parts dict 可以删
                pn = op.partname
                if hasattr(package, "_parts") and pn in package._parts:
                    del package._parts[pn]
                    fn = str(pn).rsplit("/", 1)[-1]
                    if fn not in dropped_filenames:
                        dropped_filenames.append(fn)
            except Exception:
                pass
    except Exception:
        # 兼容不同 python-docx 版本 — 失败也不致命
        pass

    return {"dropped_rels": dropped_count, "dropped_filenames": dropped_filenames}


async def _docx_delete_block(
    workspace_dir: str, target: str, rel_path: str, args: dict,
) -> str:
    """删除指定 index 的 block。

    args:
      index: 要删除的 block 序号(从 0 开始)
      count: 连续删除 count 个(默认 1)
    """
    from docx import Document

    if not os.path.isfile(target):
        return _err(f"file not found: {rel_path}", action="delete_block")
    try:
        index = int(args.get("index", -1))
        count = int(args.get("count", 1))
    except (TypeError, ValueError):
        return _err("index/count must be integers", action="delete_block", hint="Provide numeric values, e.g. index:3, count:1")
    if index < 0 or count < 1:
        return _err("index >= 0 and count >= 1 required", action="delete_block", hint="index starts at 0, count must be at least 1")

    def _do():
        try:
            doc = Document(target)
        except Exception as e:
            return None, _docx_open_error(e, rel_path)
        items = _docx_enumerate_body(doc)
        if index >= len(items):
            return None, f"index {index} out of range (have {len(items)} blocks)"
        end = min(index + count, len(items))
        elements_to_delete = [items[i]["_element"] for i in range(index, end)]
        deleted_previews = [items[i]["text"][:60] for i in range(index, end)]
        for el in reversed(elements_to_delete):
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
        # 2026-05-17 P152: 清理孤儿 image
        orphan_info = _drop_orphan_media(doc)
        _, save_err = _save_docx_atomic(doc, target)
        if save_err:
            return None, save_err
        return {
            "deleted_count": end - index,
            "start_index": index,
            "deleted_previews": deleted_previews,
            "orphan_cleanup": orphan_info,
        }, None

    info, err = await asyncio.to_thread(_do)
    if err:
        return _err(err)
    return json.dumps({
        "ok": True, "action": "delete_block", "format": "docx",
        "path": rel_path,
        **info,
    }, ensure_ascii=False)


def _clean_docx_text(text: Any) -> str:
    s = str(text) if text is not None else ""
    s = re.sub(r"```[a-zA-Z0-9_-]*", "", s).replace("```", "")
    s = re.sub(r"(?m)^\s{0,3}#{1,6}\s+", "", s)
    s = re.sub(r"(?m)^\s{0,3}[-*_]{3,}\s*$", "", s)
    s = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)
    s = re.sub(r"__([^_]+)__", r"\1", s)
    s = re.sub(r"(?<!\*)\*([^*\n]+)\*(?!\*)", r"\1", s)
    # 2026-05-18 P189: 原 regex `(?<!_)_([^_\n]+)_(?!_)` 把单下划线 markdown emphasis 剥离,
    # 但 **数学表达式的下标也会被破坏**:
    #   `x_0 + y_1` → `x0 + y1` (下标完全消失)
    #   `a_max - a_min` → `amax - amin` (变量名残缺)
    #   `Σ_i=1^n a_i` → `Σi=1^n ai`
    # 修复: 严格的 markdown emphasis 形式必须 **下划线两侧是空白或字符串边界**;
    # 当下划线**两侧有字母数字字符**时, 它是数学下标 (`x_0`, `a_max`), 不该剥离。
    s = re.sub(
        r"(?<![A-Za-z0-9_])"   # 左侧非字母数字 (空白/标点/边界)
        r"_"
        r"([^_\n]+?)"
        r"_"
        r"(?![A-Za-z0-9_])",   # 右侧非字母数字
        r"\1",
        s,
    )
    s = re.sub(r"`([^`]+)`", r"\1", s)
    s = re.sub(r"(?m)^\s{0,3}[-*+]\s+", "", s)
    return s.strip()


def _add_run_with_breaks(paragraph, text: str, *, bold: bool = False, italic: bool = False,
                         font_size_pt: float | None = None):
    """添加 run 到段落, 自动把 `\\n` 转为软换行 (<w:br/>)。

    2026-05-18 P169: helper 经常写多行 step 文本(如「步骤1: ...\\n步骤2: ...」)放进一个
    paragraph block, 但 Word 把 `<w:t>` 里的 `\\n` 当成空白折叠, 不是换行 — 用户看到
    「步骤1...步骤2」挤在一行没间隔。修正: 按 `\\n` 拆 segment, 每段加一个 run, 之间
    插 `<w:br/>` 软换行。
    """
    from docx.shared import Pt
    parts = (text or "").split("\n")
    for idx, segment in enumerate(parts):
        if idx > 0:
            # 插入软换行
            br_run = paragraph.add_run()
            br_run.add_break()
        run = paragraph.add_run(segment)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if isinstance(font_size_pt, (int, float)):
            run.font.size = Pt(font_size_pt)


def _setup_docx_defaults(doc) -> None:
    """初始化 docx 默认设置: A4 纸张 + 中英文字体 + 段落间距.

    2026-05-18 P196: 之前 docx 默认 Letter (8.5"×11") + 空 Normal style. 中文文档实际
    用 A4 (210mm×297mm = 11906×16838 twips), 视觉上应配宋体+Times New Roman.
    含公式 PNG 的段落需要 lineRule=auto 让行高自适应图片高度 (避免图片溢出/被截).

    设置:
      - 页面: A4 纸张, 边距 2.5cm
      - 默认字体: 西文 Times New Roman, 中文 宋体 (东亚语言典型组合)
      - 默认段落间距: 0pt 段前, 6pt 段后, line=1.15 (Word 默认接近)
      - lineRule="auto" — 关键! 让含图段落自动扩高, 不强制定行高把图压扁
    """
    from docx.shared import Cm, Pt, Twips
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # 1. 页面 A4
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 2. Default font on Normal style
    normal_style = doc.styles["Normal"]
    rpr = normal_style.element.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        normal_style.element.insert(0, rpr)
    # Remove old fonts
    for old_rf in rpr.findall(qn("w:rFonts")):
        rpr.remove(old_rf)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rpr.append(rFonts)
    # font size 11pt (val is half-point, so 22)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    rpr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "22")
    rpr.append(szCs)

    # 3. Default paragraph spacing — line=auto (含图段落自适应)
    ppr = normal_style.element.find(qn("w:pPr"))
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        # insert before rPr if exists
        rpr_existing = normal_style.element.find(qn("w:rPr"))
        if rpr_existing is not None:
            normal_style.element.insert(list(normal_style.element).index(rpr_existing), ppr)
        else:
            normal_style.element.append(ppr)
    # Remove old spacing
    for old_sp in ppr.findall(qn("w:spacing")):
        ppr.remove(old_sp)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "120")        # 6pt after
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:line"), "276")          # 1.15× of 240 = 276
    spacing.set(qn("w:lineRule"), "auto")     # 关键: 含图自适应
    ppr.append(spacing)


def _find_body_placeholder(slide):
    """找正文 placeholder(idx=1, 排除 title idx=0)。"""
    for shp in slide.placeholders:
        pf = getattr(shp, "placeholder_format", None)
        if pf and pf.idx == 1 and shp.has_text_frame:
            return shp
    return None


def _extract_tables_generic(tables) -> list[dict]:
    """python-docx tables 列表 → 简化 JSON。"""
    out: list[dict] = []
    for ti, table in enumerate(tables):
        rows = []
        truncated_rows = False
        truncated_cols = False
        for ri, row in enumerate(table.rows):
            if ri >= _MAX_TABLE_ROWS_RETURNED:
                truncated_rows = True
                break
            cells = []
            for ci, cell in enumerate(row.cells):
                if ci >= _MAX_TABLE_COLS_RETURNED:
                    truncated_cols = True
                    break
                t = cell.text or ""
                if len(t) > 500:
                    t = t[:500] + "…"
                cells.append(t)
            rows.append(cells)
        out.append({
            "index": ti,
            "rows": rows,
            **({"truncated_rows": True} if truncated_rows else {}),
            **({"truncated_cols": True} if truncated_cols else {}),
        })
    return out


def _list_images_in_part(part) -> list[dict]:
    """列出 docx/pptx Part 中所有图片关系(只元数据,不读 blob)。"""
    images: list[dict] = []
    try:
        rels = part.rels
        seen = set()
        for rel_id, rel in rels.items():
            if "image" not in rel.reltype:
                continue
            partname = str(getattr(rel.target_part, "partname", ""))
            if partname in seen:
                continue
            seen.add(partname)
            try:
                size = len(rel.target_part.blob)
            except Exception:
                size = -1
            ext = os.path.splitext(partname)[1].lower() or ".png"
            images.append({
                "rel_id": rel_id,
                "internal_name": os.path.basename(partname) or f"image{len(images)}{ext}",
                "size_bytes": size,
            })
    except Exception:
        log.exception("list_images_in_part failed")
    return images
