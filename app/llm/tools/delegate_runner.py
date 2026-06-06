"""Execution core for a single delegate helper."""
from __future__ import annotations

import os
import re

from app.core.filesystem import FileRegistry
from app.core.filesystem.models import FileKind, FileStatus, Visibility


def _sync_delegate_globals() -> None:
    from app.llm.tools import delegate as _delegate
    globals().update({
        name: value
        for name, value in vars(_delegate).items()
        if not name.startswith("__") and name != "_run_one_helper"
    })


def _is_environment_helper_workspace() -> bool:
    try:
        from app.core.runtime_mode import is_environment_mode
        return bool(is_environment_mode())
    except Exception:
        return False


def _registry_helper_outputs(main_workspace: str, task_id: str, kind: str) -> tuple[list[str], list[str]]:
    """Return registry-backed helper outputs visible to the main process."""
    if not main_workspace or not task_id:
        return [], []
    try:
        registry = FileRegistry.load(scope_id=f"workspace:{os.path.abspath(main_workspace)}", workspace_root=main_workspace)
    except Exception:
        return [], []
    ready_statuses = {
        FileStatus.READY,
        FileStatus.VERIFIED,
        FileStatus.PROMOTED,
        FileStatus.APPLIED,
        FileStatus.DELIVERED,
    }
    all_paths: list[str] = []
    visible_paths: list[str] = []
    for record in registry.list_records(kind=FileKind.HELPER_OUTPUT):
        if record.owner_task_id != task_id:
            continue
        if record.status not in ready_statuses and not record.verified:
            continue
        path = str(record.workspace_path or "").replace("\\", "/").strip()
        if not path:
            continue
        all_paths.append(path)
        if record.visibility in {Visibility.PROJECT, Visibility.DELIVERABLE, Visibility.EVIDENCE}:
            visible_paths.append(path)
    if kind in {"read", "ocr"}:
        visible_paths = []
    return sorted(set(all_paths)), sorted(set(visible_paths))


def _converge_terminal_state_for_complete_outputs(
    result: dict,
    *,
    expected_outputs: list[str] | None,
    kind: str,
    main_resource_request: dict | None = None,
) -> bool:
    """Normalize terminal state when system-verified outputs are complete.

    A helper may be cooperatively stopped after repeated tool failures while it
    has already delivered every declared output. In that case the authoritative
    acceptance fact is the system output check, not the stale stuck/interrupted
    lifecycle flag. This only converges when expected outputs exist, all are
    present, no blocking quality/resource issue exists, and the helper is not a
    read-helper evidence task.

    当系统已确认声明产物完整且无阻塞问题时，完成事实优先于旧的中断/卡住状态。
    """
    if not isinstance(result, dict):
        return False
    helper_kind = str(kind or "").strip().lower()
    if helper_kind in {"read", "ocr"} or not expected_outputs:
        return False
    outputs_check = result.get("outputs_check")
    if not isinstance(outputs_check, dict):
        return False
    if outputs_check.get("outputs_complete") is not True:
        return False
    if main_resource_request or result.get("resource_required"):
        return False
    quality_warnings = outputs_check.get("quality_warnings") or []
    if quality_warnings or outputs_check.get("quality_blocked") or result.get("quality_blocked"):
        return False
    if result.get("terminal_reason") in {
        "failed",
        "quality_blocked",
        "resource_required",
        "outputs_missing",
        "copyback_blocked",
        "crashed",
    }:
        return False

    previous_reason = result.get("terminal_reason")
    had_lifecycle_blocker = any(
        result.get(key)
        for key in (
            "interrupted",
            "stuck",
            "stuck_reason",
            "suggested_retry_kind",
            "suggested_retry_mode",
            "next_action",
            "retry_instruction",
        )
    )
    result["ok"] = True
    result["terminal_reason"] = "completed"
    for key in (
        "interrupted",
        "stuck",
        "stuck_reason",
        "suggested_retry_kind",
        "suggested_retry_mode",
        "next_action",
        "retry_instruction",
    ):
        result.pop(key, None)
    if had_lifecycle_blocker or previous_reason not in (None, "completed"):
        result["_terminal_converged_from"] = previous_reason or "lifecycle_blocker"
        result["post_helper_usage_hint"] = (
            str(result.get("post_helper_usage_hint") or "").strip()
            + ("\n" if result.get("post_helper_usage_hint") else "")
            + "System output checks verified every declared file; treat this helper as completed unless later artifact inspection finds a concrete defect.\n\n"
              "系统验收已确认声明产物完整；除非后续检查发现具体缺陷，否则按完成处理。"
        )
    return True


def _is_framework_future_output_reference(
    *,
    task_id: str,
    kind: str,
    prompt: str,
    filename: str,
    expected_outputs: list[str] | None,
) -> bool:
    """Return True when a filename is only a planned downstream output name."""
    role_text = f"{task_id} {kind} {' '.join(str(x) for x in (expected_outputs or []))} {prompt}".lower()
    if not any(
        marker in role_text
        for marker in (
            "framework", "contract", "output matrix", "file naming", "naming convention",
            "expected_outputs", "downstream", "later helper", "later slice",
            "框架", "契约", "输出矩阵", "文件命名", "命名规范", "后续 helper", "后续分片",
        )
    ):
        return False
    if any(
        marker in role_text
        for marker in (
            "read the file", "load the file", "consume the file", "use existing",
            "already produced", "previous helper output", "读取该文件", "读取文件",
            "使用已有", "前置产物", "已产出",
        )
    ):
        return False
    basename = os.path.basename(str(filename or "").replace("\\", "/")).lower()
    if not basename:
        return False
    stem, ext = os.path.splitext(basename)
    if ext not in {
        ".md", ".txt", ".json", ".csv", ".docx", ".xlsx", ".pptx",
        ".py", ".c", ".cpp", ".h", ".hpp", ".png", ".jpg", ".svg",
    }:
        return False
    future_name_patterns = (
        r"^(?:ds|analysis|bench|benchmark|complexity|experiment|paper|verify|verification|report|novel|new|final)[_-]",
        r"^(?:db|database|index|db_index|database_index)[_-]",
        r"^(?:rb|rbtree|redblack|skiplist|btree|bplus|b_tree|bplus_tree|novel_ds)$",
        r"^(?:framework|contract|outline|schema|manifest|matrix|check_framework)$",
    )
    return any(re.search(pattern, stem) for pattern in future_name_patterns)


def _needs_helper_report_format_self_repair(report: str, *, expected_outputs: list[str] | None, kind: str) -> bool:
    """Return true when a helper should get one local chance to repair its report."""
    helper_kind = str(kind or "").strip().lower()
    if helper_kind in {"read", "ocr"} or not expected_outputs:
        return False
    try:
        if _extract_reported_output_files(report):
            return False
        return bool(_has_malformed_output_files_attempt(report) or _extract_declared_files(report))
    except Exception:
        return False


def _build_helper_report_format_repair_prompt(
    *,
    original_report: str,
    expected_outputs: list[str] | None,
) -> str:
    outputs = "\n".join(f"- {path}" for path in (expected_outputs or []))
    excerpt = str(original_report or "")[:5000]
    return (
        "## Final Report Format Repair\n"
        "Your previous final report did not satisfy the required output-file declaration format. "
        "Resume the existing task and preserve existing artifacts unless you find a concrete missing file. "
        "Inspect only if needed, then produce a corrected final report.\n\n"
        "Required section:\n"
        "## Output files\n"
        "```json\n"
        "{\"files\": [\"_env/relative_output.ext\"]}\n"
        "```\n\n"
        "Expected output paths from the main task:\n"
        f"{outputs or '- <none>'}\n\n"
        "When those expected paths exist, copy them exactly into the `files` array, including `_env/` or "
        "`_helpers_shared/` prefixes. Keep full relative paths rather than bare basenames.\n\n"
        "If an expected file is missing, say that in `## Missing or warnings` and do not claim completion. "
        "If the files exist, list their exact paths in the JSON block and complete the report.\n\n"
        "Previous report excerpt:\n"
        f"```\n{excerpt}\n```\n\n"
        "Repair only the final report format. Prefer a compliant Output files JSON block; state missing files honestly.\n\n"
        "只修最终报告格式；优先逐字使用 expected_outputs 的完整相对路径，文件存在才声明完成，缺失则如实说明。"
    )


def _build_helper_user_prompt(
    *,
    prompt: str,
    dynamic_prompt_prefix_parts: list[str],
    prior_summary: str = "",
    resume: bool = False,
    resume_workspace_empty: bool = True,
    kind: str = "code",
) -> str:
    """Build the user-side helper prompt with stable framing before the task.

    The helper system prompt stays stable for cache reuse. Per-turn workspace
    listings, skill listings, preserved summaries, and the delegated task live
    in the user message. Keep reusable framing and shared dynamic snapshots
    before the latest main-process request so sibling helpers can share a longer
    provider-cache prefix and the freshest task stays near the end.
    """
    dynamic_prompt_prefix = "\n\n".join(
        str(part).strip()
        for part in dynamic_prompt_prefix_parts
        if str(part or "").strip()
    )
    helper_kind = str(kind or "").strip().lower()
    stable_kind_frame = ""
    if helper_kind in {"read", "ocr"}:
        stable_kind_frame = (
            "## Read Helper Operating Contract\n"
            "Use search, outlines, and targeted ranges to answer the delegated evidence question. "
            "When requested facts, line numbers, coverage notes, or uncertainty are supported by evidence, return the concise handoff report. "
            "Use overlapping reads only to close a named evidence gap. For broad material, save long extraction as text evidence and report coverage plus recommended ranges.\n\n"
            "读取类 helper 用搜索、结构和定向范围取证；证据足够即交接，长材料写入证据文件。"
            "\n\n"
        )
    elif helper_kind == "code":
        stable_kind_frame = (
            "## Code Helper Operating Contract\n"
            "Implement the delegated module, feature, or fix. Write working, tested code; "
            "declare all output files in the final report. Run the relevant test command before "
            "reporting done — if tests are not runnable, state the blocker explicitly. "
            "Prefer targeted edits over full rewrites; for new files, lay scaffold first then fill. "
            "Read only assigned material and consume evidence files provided by prior helpers.\n\n"
            "代码类 helper 实现功能、运行自测、在报告中声明产物；优先局部改动，先搭骨架再填充。"
            "\n\n"
        )
    elif helper_kind == "edit":
        stable_kind_frame = (
            "## Edit Helper Operating Contract\n"
            "Produce the delegated document, report, or prose artifact. "
            "Consume evidence files and coverage summaries rather than re-reading raw source material. "
            "Follow any provided template or outline exactly; fill all required sections before finishing. "
            "Declare the output file(s) in the final report; do not emit the full document body in the report itself.\n\n"
            "编辑类 helper 按模板生成文档产物，消费已有证据文件，在报告中声明产物路径，不在报告体内重复全文。"
            "\n\n"
        )
    elif helper_kind == "verify":
        stable_kind_frame = (
            "## Verify Helper Operating Contract\n"
            "Check the delegated artifact or claim against the stated acceptance criteria. "
            "Run tests, inspect files, compare against spec — do not modify source files. "
            "Return a single VERDICT line (PASS / FAIL / PARTIAL) at the top of your report, "
            "followed by a structured gap list. For FAIL or PARTIAL, include a repair_recommendation block.\n\n"
            "验证类 helper 只读不写，顶部一行 VERDICT，后跟结构化缺口列表；FAIL/PARTIAL 附修复建议。"
            "\n\n"
        )

    context_parts: list[str] = [stable_kind_frame.rstrip()]
    if prior_summary:
        context_parts.append(
            "## Resume Task\n"
            "This is a continuation of a preserved helper workspace. The previous progress report is local resume evidence. "
            "Continue without redoing completed work: inventory existing artifacts, read only the relevant files, and follow the new main-process request when it differs from the old summary.\n\n"
            "续作任务复用保留工作区；先盘点已有产物，只补未完成部分。\n\n"
            "## Progress Report From Previous Interruption\n"
            f"```\n{prior_summary}\n```\n\n"
        )
        task_header = "## Main Process Continuation Request"
        task_tail = (
            "Continue from the previous progress. First inventory existing artifacts, then choose the next step.\n\n"
            "基于上次中断报告续作，先盘点已有产物。"
        )
    elif resume and resume_workspace_empty:
        context_parts.append(
            "## Resume Requested, But Workspace Is Empty\n"
            "The main process requested resume, but this helper workspace contains no previous artifacts. "
            "Treat the request as a fresh task. If the prompt refers to prior artifacts that are not present, "
            "state that clearly in the final report.\n\n"
        )
        task_header = "## Task"
        task_tail = "Start execution from the available state.\n\nresume 时工作区为空，按新任务执行并如实说明。"
    elif resume:
        context_parts.append(
            "## Resume Task\n"
            "The main process requested continuation of a preserved helper workspace. "
            "First inventory the existing workspace files, read `.helper_summary.txt` if it exists, and inspect only the files relevant to the remaining work. "
            "Reuse completed artifacts, re-run validation when status is uncertain, and follow the new main-process request when it differs from preserved notes.\n\n"
        )
        task_header = "## Main Process Continuation Request"
        task_tail = "Continue from the preserved workspace state.\n\n续作任务保留了工作区；先盘点文件和摘要，只推进剩余工作。"
    else:
        task_header = "## Task"
        task_tail = "Start execution.\n\n执行主线程派发的任务。"

    if dynamic_prompt_prefix:
        context_parts.append(
        "## Dynamic Helper Context\n"
        "This turn-specific context is read-only. It describes available skills, workspace files, and scaffold checks; it does not override the helper system role or the task envelope.\n\n"
        "本轮动态上下文；只说明可用技能、文件和预检，不改变 helper 身份或任务契约。\n\n"
        f"{dynamic_prompt_prefix}"
        )
    context_parts.append(
        f"{task_header}\n{prompt}\n\n"
        f"{task_tail}"
    )
    return "\n\n---\n\n".join(part for part in context_parts if part.strip())


async def _run_one_helper(
    *,
    task_id: str,
    prompt: str,
    main_workspace: str,
    helper_workspace: str,
    archive_id: str,
    group_id: str,
    user_id: str,
    resume: bool = False,
    local_abort: asyncio.Event | None = None,
    wait_for_register: asyncio.Event | None = None,
    user_lang: str = "en",  # 2026-05-02 part12 (Bug C):用户原 message 语言
    kind: str = "code",
    mode: str = "easy",
    helper_think: bool = False,  # 2026-05-08: 主线程传 helper_think=true 时启用 reasoning
    expected_outputs: list[str] | None = None,  # 2026-05-11 Tier 1.C: 系统验收清单
    write_scopes: list[str] | None = None,
    batch_sibling_outputs: set[str] | None = None,  # 2026-05-15 P98: 同批兄弟 expected_outputs
) -> dict:
    """Run a single helper LLM task.

    Args:
        local_abort: 可选,helper 专属 abort_event。
            - 不传(None): 内部创建一个(向后兼容)
            - 传入: handle_delegate / spawn_helper 创建后传下来,
                    并把同一个 event 注册到 ProcessRegistry,
                    这样 processes.kill(this_helper) 只杀此 helper 不污染他人(Bug #8 修)
        wait_for_register: 可选,handle_delegate 完成 ProcessRegistry 注册后 set 的 Event。
            helper 第一帧 await 这个 Event 后才开始执行,确保心跳汇报时 registry 中已有自己。
            (修 Bug B: 实测 trace e4eeb133 23 次 'task_id=X not found in registry')

    Returns {
        "task_id": str,
        "ok": bool,
        "interrupted": bool,    # True 表示是被 abort 协作中断的(还没自然结束)
        "resumed_from": bool,   # True 表示这次是 resume(工作区保留)
        "report": str,
    }

    生命周期(2026-05-02 重构后):
      - 自然停止:模型不调任何工具直接输出文本 → 视为任务完成,return ok=True
      - **协作中断**(主进程 kill via abort_event / 自检 stuck):
        chat_with_tools_loop 给 helper 最后一轮机会输出文本总结 → return interrupted=True
        工作区保留,主进程可决定 resume=true 续作 or 用新 prompt 继续
      - **用户全局 abort**(`/v1/chat/abort`):
        主进程 chat_with_tools_loop 通过 racing 机制 cancel 当前 tool dispatch →
        handle_delegate 的 except CancelledError 块 cancel 所有 helper task →
        本函数任意 await 处抛 CancelledError → finally 清理 → task 进入 cancelled 状态
        **不会**经过这里的正常 return,result 是 CancelledError 不是 dict
      - **不再有时间硬墙** — helper 跑多久都行,主进程通过心跳监控决定是否 kill
    """
    _sync_delegate_globals()

    from app.llm.client import _parse_json_strict
    from app.llm.tools.registry import dispatch as tool_dispatch
    from app.core.locks import get_group_guard

    # ── 心跳 race 修(Bug B): 等 handle_delegate 写完 registry 再开干 ──
    # 不传 wait_for_register 的 caller(老 spawn_helper / 单元测试)直接跳过 = 旧行为。
    if wait_for_register is not None:
        try:
            await asyncio.wait_for(wait_for_register.wait(), timeout=2.0)
        except asyncio.TimeoutError:
            log.warning(
                "helper %s: wait_for_register timeout (2s) — proceeding without registry sync; "
                "heartbeat may be missing",
                task_id,
            )

    kind, mode = _normalize_helper_kind_mode(kind, mode)

    # Bug 10 修:helper 设 sub-trace_id,使 log 可读
    # 形如 "abc123.huffman_v2" — 前 6 位主 trace + helper task_id 前 12 位
    # 注意 ContextVar 在 asyncio.gather 起的子 task 里是隔离的(asyncio task 各自 copy context),
    # 这里 set 不会污染主 trace_id。
    # ── 2026-05-04 Bug #8 修复 ──
    # 旧版 task_id[:6] 把 "huffman_v2"/"huffman" 都截成 "huffma",
    # 把 "bwt_final"/"bwt_v2" 截成 "bwt_fi"/"bwt_v2" — grep 不友好且歧义。
    # 改成 [:12] 让常见 task_id 完整显示("huffman_v2" 长度 10 完整保留,
    # "compression_v3" 截到 "compression_" 仍可识别)。
    parent_trace = debug.current_trace_id()
    main_trace_id = parent_trace or ""
    sub_trace = f"{parent_trace[:6]}.{task_id[:12]}"
    debug.set_trace_id(sub_trace)

    # ── P0 race condition 防御(2026-05-03 修,trace b78b242533a24a46) ──
    # 主线程 wait_window_sec 超时返回后立刻调 delegate(resume=true) 时,旧 helper
    # 的 LLM stream 可能还没自然 finalize。如果不 cancel 旧 stream 直接开新 stream,
    # 会出现**同一 task_id 两个 LLM stream 并行跑**:
    #   - 旧 stream 自然结束 → 触发 delegate.X.done #1
    #   - 新 stream 不知道任务已 done,继续干活 → 又触发 delegate.X.done #2
    #     生成大量重复/无关文件(实测 razor 第 2 次 done 复制了 50 个垃圾文件,
    #     代码自己的 fork-copy pollution snowball 保护命中但已污染主区)。
    # 防御:resume=True 时先查同 task_id 是否还有活跃 helper,有则 set 其 abort_event
    # 等其优雅 finalize(最多 30s),再开新 stream。这保证全程"一个 task_id 只有一个流"。
    if resume:
        try:
            # ── 2026-05-04 Bug #19 修复:self-kill 防护 ──
            # 新 helper 在 wait_for_register 后自己已在 registry 里,
            # find_helper_by_task_id 会返回自己。如果不排除自己,resume_preempt
            # 会把自己当"旧 helper" abort + force-kill (trace 904c47ec61e74d85)。
            self_proc_id_for_resume = None
            try:
                _self_helper_owner = ProcessRegistry.make_helper_owner(
                    parent_trace, task_id,
                )
                _h_self = await proc_registry().find_helper_by_task_id(
                    task_id, same_trace_as=_self_helper_owner,
                )
                if _h_self is not None:
                    self_proc_id_for_resume = _h_self.proc_id
            except Exception:
                pass  # self-lookup 失败不影响主流程

            # ACL:只看自己 owner 创建的 helper(主线程发起的 resume,owner_id 已被
            # set_current_owner 设为 main_owner;helper 自己 spawn helper 的 resume 路径
            # 同理 owner 是父 helper 的)
            cur_owner = current_owner()
            existing = None
            if cur_owner is not None:
                existing = await proc_registry().find_helper_by_task_id(
                    task_id, owner=cur_owner,
                )
            if existing is None:
                # owner 严格匹配找不到,放宽到 trace 范围(防止 owner 换了 token 但还在跑)
                existing = await proc_registry().find_helper_by_task_id(
                    task_id, same_trace_as=cur_owner or "",
                )
            # ── self-kill 防护:排除自己 ──
            if (
                existing is not None
                and self_proc_id_for_resume is not None
                and existing.proc_id == self_proc_id_for_resume
            ):
                debug.log(
                    f"delegate.{task_id}.resume_self_skip",
                    f"find_helper_by_task_id returned self (proc_id={self_proc_id_for_resume}, "
                    f"iter={existing.last_iter}); skipping resume_preempt — main thread's "
                    f"race_protect already handled the real old helper",
                )
                existing = None
            if existing is not None and existing.abort_event is not None:
                # 命中 race:旧 helper 还在跑。set abort_event 让它协作退出。
                debug.log(
                    f"delegate.{task_id}.resume_preempt",
                    f"existing helper proc_id={existing.proc_id} still active "
                    f"(iter={existing.last_iter}); aborting before resume",
                )
                log.warning(
                    "delegate %s: resume requested but existing helper still active "
                    "(proc_id=%s, iter=%d); aborting old stream first",
                    task_id, existing.proc_id, existing.last_iter,
                )
                existing.abort_event.set()
                # 等其退出(用 ProcessRegistry 二次轮询;abort 后 helper 一般 1-3s 自然 finalize)
                # ── 2026-05-04 Bug #24 修复 ──
                # trace 85e330f3d6f04985: 30s → 60s。trace 904c47ec61e74d85: 60s → 180s。
                # helper 收到 abort 后需要: (1) 当前 LLM stream 完成 (2) forced_finalize
                # (reasoning=high 下 30-90s) (3) summary_persisted + 文件复制。
                # 60s 经常只够第 1 步,180s 给充分时间走完所有步骤。
                _RESUME_PREEMPT_TIMEOUT_SEC = 180.0
                _settled = False
                _poll_intervals = int(_RESUME_PREEMPT_TIMEOUT_SEC / 0.5)
                for _wait in range(_poll_intervals):
                    await asyncio.sleep(0.5)
                    # 重新 find,如果 helper 已 unregister 就找不到了
                    # 传 exclude_proc_id 防止在 finalize 时 find 到自己
                    h2 = await proc_registry().find_helper_by_task_id(
                        task_id, owner=cur_owner,
                        exclude_proc_id=self_proc_id_for_resume,
                    ) if cur_owner is not None else None
                    if h2 is None:
                        h2 = await proc_registry().find_helper_by_task_id(
                            task_id, same_trace_as=cur_owner or "",
                            exclude_proc_id=self_proc_id_for_resume,
                        )
                    if h2 is None:
                        _settled = True
                        break
                if _settled:
                    debug.log(
                        f"delegate.{task_id}.resume_preempt_settled",
                        f"old helper finalized; proceeding with resume",
                    )
                else:
                    # 180s 内旧 helper 没有 natural finalize — 但禁止 force-kill。
                    # 2026-05-05: resume_preempt 原来的 force=True kill 绕过了
                    # _should_kill_helper kill gate。正确做法:返回错误让主线程
                    # 通过 delegate(action="kill", reason=...) 明确 kill,
                    # 或换 task_id 派新 helper 让新旧并行竞速。
                    debug.error(
                        f"delegate.{task_id}.resume_preempt_blocked",
                        f"old helper (proc={existing.proc_id}) still alive after "
                        f"{_RESUME_PREEMPT_TIMEOUT_SEC:.0f}s abort wait; "
                        f"NOT force-killing (bypasses kill gate). "
                        f"Main thread should: delegate(action='kill', task_id='{task_id}', "
                        f"reason='content_deemed_useless') then resume with same task_id. "
                        f"Keep the same task_id for the same work boundary.",
                    )
                    return {
                        "task_id": task_id,
                        "ok": False,
                        "interrupted": True,
                        "terminal_reason": "resume_preempt_blocked",  # P2.1
                        "report": (
                            f"Resume is blocked: the previous helper (proc={existing.proc_id}) "
                            f"did not exit after {_RESUME_PREEMPT_TIMEOUT_SEC:.0f}s of cooperative abort wait.\n"
                            f"The kill gate prevents automatic force-kill. Call explicitly:\n"
                            f"  delegate(action='kill', task_id='{task_id}', "
                            f"reason='content_deemed_useless')\n"
                            f"After kill succeeds, resume the same task_id to keep the same work boundary.\n\n"
                            f"续作被旧 helper 阻塞；先显式协作中断，再用同一 task_id 续作。"
                        ),
                        "_resume_preempt_blocked": True,
                        "elapsed_sec": 0,
                    }
        except Exception:
            log.exception("delegate %s: resume_preempt check failed; proceeding", task_id)

    # 工作区策略:resume=True 保留上次文件,resume=False 全清重建
    # 2026-05-02 part9 #20:跟踪 resume 是否实际生效(可能 resume=True 但工作区空 → 降级 fresh)
    # 主线程靠 result["resumed_actually"] 字段知道是否真续作了 vs 等价于 fresh。
    resumed_actually = False
    resume_workspace_empty = bool(resume)
    if resume and os.path.isdir(helper_workspace):
        _all_files = os.listdir(helper_workspace)
        existing = [
            f for f in _all_files
            if not f.startswith(".") and not f.startswith("_")
        ]
        resume_workspace_empty = not existing
        log.info("delegate helper %s: RESUME (workspace preserved, %d files)",
                 task_id, len(existing))
        debug.log(f"delegate.{task_id}.resume",
                  f"workspace preserved, {len(existing)} top-level files",
                  existing[:30])
        # 2026-05-02 part22:根据 kind+mode 选择 prompt (2026-05-12 P20)
        sys_prompt = _select_helper_system(kind, mode)
        # Keep helper system prompts cache-friendly: stable capability text
        # stays in system; per-helper workspace listings move to the user turn.
        dynamic_prompt_prefix_parts: list[str] = [_helper_lang_hint(user_lang), _build_skills_listing()]
        # 2026-05-15 P115: resume 路径也注入工作区清单 - 让 helper 立刻知道上次产物
        try:
            _ws_listing = _list_helper_workspace_for_prompt(helper_workspace)
            if _ws_listing:
                dynamic_prompt_prefix_parts.append(_ws_listing)
        except Exception as _e_ws:
            debug.log(f"delegate.{task_id}.ws_listing_inject_failed", repr(_e_ws))
        resumed_actually = True  # 真 resume 上去了

        # #17: 读 .helper_summary.txt — 上次中断时持久化的状态摘要
        # 把它直接注入 user prompt,helper 看到上次写的内容,不需要主模型从 hot 重新转述
        prior_summary = ""
        summary_path = os.path.join(helper_workspace, ".helper_summary.txt")
        if os.path.isfile(summary_path):
            try:
                with open(summary_path, "r", encoding="utf-8") as f:
                    prior_summary = f.read()
                debug.log(f"delegate.{task_id}.summary_loaded",
                          f"loaded .helper_summary.txt ({len(prior_summary)} chars)")
            except OSError:
                log.exception("failed to read helper summary")
        # 2026-05-03 v18 修复(trace 85e330f3d6f04985 实测教训):
        # 旧逻辑 resume 路径下 fork_snapshot=None,全量复制路径触发。
        # 实际场景:resume helper 跑 30s 被 abort → 沙箱里 81 个 fork 时就在的文件
        # 全部当成 "helper 产出" → 触发 cap 50 → 整批拒绝。
        # 修复:resume 路径也拍 snapshot(此时沙箱 = 上次 helper 留下的状态),
        # helper 这一轮 abort 时 diff 出来的就是真实变化文件。
        fork_snapshot = take_workspace_snapshot(helper_workspace)
        debug.log(
            f"delegate.{task_id}.resume_snapshot",
            f"resume snapshot taken ({len(fork_snapshot)} files); "
            f"only files modified during this resume turn will be copied back",
        )

        # L6-2 (2026-05-09): 增量同步 — 仅同步 main→helper 的变更文件
        _sync_stats = _resume_incremental_sync(
            helper_workspace=helper_workspace,
            main_workspace=main_workspace,
        )
        if _sync_stats["copied"] or _sync_stats["skipped_unchanged"]:
            debug.log(
                f"delegate.{task_id}.resume_incr",
                f"resume sync: copied={_sync_stats['copied']} "
                f"skipped_unchanged={_sync_stats['skipped_unchanged']} "
                f"helper_kept={_sync_stats['helper_kept']} "
                f"(saved {_sync_stats['skipped_unchanged']} files vs full re-fork)",
            )
    else:
        if resume:
            log.info("delegate helper %s: resume=True 但工作区为空,降级为 fresh", task_id)
        clean_workspace_dir(helper_workspace)
        # 2026-05-10 Patch 57:从 ProcessRegistry 拿同 trace 的所有 helper task_id
        # 作为 fork 白名单,防止 _looks_like_helper_artifact 误把当前会话的
        # promote 产物(如 gen_charts_*.png)当成"历史污染"过滤掉。
        # trace b430c4f228eb40c7 暴露:gen_charts 产出 4 张 PNG,embed_charts fork 时
        # 4 张 PNG 全被 B8 过滤跳过 → embed_charts 沙箱里没图 → 任务失败 → 文件没推送。
        # 2026-05-12 P29 修复: P57 只从 ProcessRegistry 拿**活跃** helper, 已完成的退出 registry!
        # 实测 12:10 trace: paper fork 时 charts 已 done 130s, 不在 registry,
        # → allowed_prefixes 不含 'charts' → charts_chart_*.png 被当 artifact 跳过 →
        # paper 找不到图死循环 dir /b *.png → STUCK
        # 修复: 同时从 .helpers_displayed_name.json (P58 推送层 metadata) 加载, 
        # 它包含所有 push 过的 helper task_id (含已完成).
        _allowed_prefixes: set[str] = {task_id}  # 自己也加(虽然主要场景是兄弟 helper)
        try:
            _trace_id = debug.current_trace_id() or ""
            if _trace_id:
                _main_owner = ProcessRegistry.make_main_owner(_trace_id)
                _siblings = await proc_registry().list_owned_by(_main_owner)
                for h in _siblings:
                    _sib_tid = h.get("helper_task_id") or h.get("task_id") or ""
                    if _sib_tid:
                        _allowed_prefixes.add(_sib_tid)
        except Exception:
            log.exception("P57 sibling task_id collect failed (non-fatal)")
        # 2026-05-12 P29: 补充已完成 helper task_id (从 P58 推送层 metadata)
        try:
            from app.llm.tools.workspace import load_displayed_name_remap
            _displayed_remap = load_displayed_name_remap(main_workspace)
            if _displayed_remap:
                for _completed_tid in _displayed_remap.keys():
                    _allowed_prefixes.add(_completed_tid)
                debug.log(
                    f"delegate.{task_id}.fork_allowlist",
                    f"P29: 合并 P58 推送层 task_id 共 {len(_displayed_remap)} 个 "
                    f"(含已完成 helper, 防止其产物被 fork 误判 artifact)"
                )
        except Exception:
            log.exception("P29 displayed_name_remap collect failed (non-fatal)")
        # 2026-05-10 Patch 57 v3:加载用户 fetch 的群文件 basenames,fork 时无条件放行
        try:
            from app.llm.tools.workspace import load_user_fetched_basenames
            _allowed_basenames = load_user_fetched_basenames(main_workspace)
        except Exception:
            _allowed_basenames = frozenset()
        copied = copy_workspace_contents(
            main_workspace, helper_workspace,
            allowed_prefixes=frozenset(_allowed_prefixes),
            allowed_basenames=_allowed_basenames,
            include_downloaded_media=(kind in {"read", "ocr"}),
            # 2026-06-05: 在普通 chat 模式也传递 `_env/` 给下游 helper。
            # 病因(实测 trace 373640 17:07:48): design-act-algo / impl-baseline-algos
            # 把源码 merge 到 <main>/_env/，但 run-benchmarks helper 因 environment
            # mode 关而 _env 未递归 fork 进沙箱 → 对所有源文件 `read_file _env/...`
            # 全 not found。helper 浪费 90s 找文件最终 stuck。
            # 修法: 只要主区有 _env/，就传给下游。chat 模式下 _env 内容来自前序 helper，
            # 不会污染用户项目；environment 模式行为不变。
            include_environment_files=(
                _is_environment_helper_workspace()
                or os.path.isdir(os.path.join(main_workspace, "_env"))
            ),
        )
        _cap = enforce_workspace_capacity(
            helper_workspace,
            label=f"helper:{task_id}",
        )
        if not _cap.get("ok", True):
            raise RuntimeError(
                f"helper 工作区超过容量上限: {_cap.get('after_bytes', 0) // 1024 // 1024}MB "
                f"> {_cap.get('max_bytes', 0) // 1024 // 1024}MB"
            )
        log.info("delegate helper %s: FRESH (copied %d files from main)", task_id, copied)
        debug.log(f"delegate.{task_id}.fresh", f"workspace re-populated, copied={copied}")
        # ── B1 修复关键: fork 完成后拍快照,helper 跑完时 diff 出真正新增/修改文件 ──
        fork_snapshot = take_workspace_snapshot(helper_workspace)
        debug.log(
            f"delegate.{task_id}.snapshot",
            f"fork-snapshot taken ({len(fork_snapshot)} files); "
            f"only modified/new files will be copied back to main workspace",
        )
        # 2026-05-02 part22:根据 kind+mode 选择 prompt (2026-05-12 P20)
        sys_prompt = _select_helper_system(kind, mode)
        # 2026-05-11 P1.1 Skills 清单: moved to dynamic user prefix so
        # system prompts remain stable across sibling helpers.
        # 设计哲学(来自 Claude Code bundledSkills.ts): 主 prompt 只列清单,详细内容按需 read。
        # 每个 skill 名 + 简短描述 ~80 字符,5 个 skill 总占 ~400 字符,
        # 但节省了 _SHARED_WORKSPACE 详细教学的 ~1900 字符(净省 1500 字符 / 15% token)
        dynamic_prompt_prefix_parts: list[str] = [_helper_lang_hint(user_lang), _build_skills_listing()]
        # 2026-05-15 P115: 注入工作区清单 - helper 不用 workspace.list 探索就知道有啥
        # 病因(实测): helper 刚启动通常先 search_files / workspace.list 看工作区 (1-2 iter)
        # 主线程已知工作区状态, 直接给 helper 看省事
        # 修法: 列出沙箱前 30 个文件 + _shared/ 和 _helpers_shared/ 内容
        try:
            _ws_listing = _list_helper_workspace_for_prompt(helper_workspace)
            if _ws_listing:
                dynamic_prompt_prefix_parts.append(_ws_listing)
        except Exception as _e_ws:
            debug.log(f"delegate.{task_id}.ws_listing_inject_failed", repr(_e_ws))
        # 2026-05-05: _shared_ 脚手架预检,编译验证 C 源文件
        # 只对 code helper 做——edit/verify/hard-mode helper 不一定写 C,白浪费 subprocess
        if kind in ("code", "coding"):
            _shared_warn = _validate_shared_scaffold(helper_workspace)
            if _shared_warn:
                dynamic_prompt_prefix_parts.append(_shared_warn)
                log.warning(
                    "delegate helper %s: _shared_ scaffold validation warnings injected",
                    task_id,
                )
        prior_summary = ""

    # 2026-05-12 P34: 自动注入依赖路径清单(铁律 12 实操层)
    # 病因(实测 15:47 trace): 主线程派 bench_all 时 prompt 写
    # "你需要写 benchmark.c, 链接 ssl.o avl.o rb.o skip.o" 但没说在哪个目录。
    # bench_all 凭经验猜 _shared/skip.h → 不存在 → 3 次 STUCK。
    # 修法: 扫 prompt 引用的文件名(*.h/*.c/*.docx/*.csv...), 查 P58 推送层 metadata,
    # 看哪些是前置 helper 产的,自动在动态上下文注入路径清单。
    # helper LLM 看到清晰路径就不会瞎猜。
    if not resume:  # resume 时 user_prompt 已有进度报告,不重复注入
        try:
            from app.llm.tools.workspace import load_displayed_name_remap
            import re as _re_p34
            _file_pattern = _re_p34.compile(
                r'\b([a-zA-Z][\w\-]{1,40}\.(?:c|h|cpp|hpp|cc|py|js|ts|csv|tsv|json|xlsx|docx|pptx|png|jpg|svg|md|txt|html|yaml|yml))\b'
            )
            _referenced_files = set()
            _expected_output_basenames = {
                os.path.basename(str(x).replace("\\", "/")).lower()
                for x in (expected_outputs or [])
                if str(x).strip()
            }
            for _m in _file_pattern.finditer(prompt):
                _fname = _m.group(1)
                _fname_base_l = os.path.basename(_fname.replace("\\", "/")).lower()
                # 排除自身 expected_outputs (helper 自己要产的, 不是依赖)。
                # expected_outputs 可能带 _helpers_shared/task/ 前缀, prompt 通常只写 basename。
                if _fname_base_l in _expected_output_basenames:
                    continue
                # 2026-05-12 P40: 排除 C/C++/Python 标准库 + 常见 build artifact
                # 病因(实测 18:37 trace): P35 把 #include <stdint.h> 也警告"缺失依赖",
                # 误导 helper, 不应该警告系统库
                _STDLIB_HEADERS = {
                    "stdint.h", "stdbool.h", "stdio.h", "stdlib.h", "string.h",
                    "math.h", "time.h", "assert.h", "errno.h", "limits.h",
                    "ctype.h", "signal.h", "locale.h", "setjmp.h", "stdarg.h",
                    "stddef.h", "float.h", "wchar.h", "wctype.h", "iso646.h",
                    "pthread.h", "unistd.h", "fcntl.h", "sys/time.h", "sys/types.h",
                    "sys/stat.h", "sys/wait.h", "windows.h", "winsock2.h",
                    "iostream", "vector", "string", "map", "set", "algorithm",
                    "memory", "thread", "mutex", "atomic", "chrono",
                }
                if _fname.lower() in _STDLIB_HEADERS:
                    continue
                _referenced_files.add(_fname)

            if _referenced_files:
                # 2026-05-12 P40: 修 P34 schema bug — load_displayed_name_remap 返回
                # dict[str, str] (displayed_name → real_path), 不是 dict[str, dict].
                # 旧 P34 假设错 schema → 永远找不到匹配 → P34 触发 0 次 (18:37 trace 确认)
                _remap = load_displayed_name_remap(main_workspace)
                _dependency_paths: list[str] = []
                _fetchable_from_main: list[str] = []
                # _remap 形如: {"draw_charts_v2_chart1_overview.png": "draw_charts_v2/chart1_overview.png"}
                # 或 {"ssl_algo_ssl.c": "ssl_algo/ssl.c"}
                for _fname in sorted(_referenced_files):
                    _lower_fname = _fname.lower()
                    _fname_base_l = os.path.basename(_fname.replace("\\", "/")).lower()
                    _matched = False
                    if _fname_base_l in _expected_output_basenames:
                        _matched = True
                    if not _matched and helper_workspace:
                        _candidate_path = os.path.join(helper_workspace, _fname)
                        try:
                            if os.path.exists(_candidate_path):
                                _dependency_paths.append(
                                    f"  - `{_fname}` → `{_fname}` (已在当前 helper 工作区)"
                                )
                                _matched = True
                        except OSError:
                            pass
                    if not _matched and main_workspace:
                        _candidate_path = os.path.join(main_workspace, _fname)
                        try:
                            if os.path.exists(_candidate_path):
                                _fetchable_from_main.append(_fname)
                                _matched = True
                        except OSError:
                            pass
                    if _matched:
                        continue
                    for _displayed_name, _real_relpath in _remap.items():
                        if not isinstance(_displayed_name, str) or not isinstance(_real_relpath, str):
                            continue
                        _lower_disp = _displayed_name.lower()
                        _real_norm_for_match = _real_relpath.replace("\\", "/")
                        _real_base_l = os.path.basename(_real_norm_for_match).lower()
                        # 匹配: displayed_name endswith "_X.png" / 就是 X.png / real_relpath basename 是 X.png
                        if (_lower_disp.endswith("_" + _lower_fname)
                                or _lower_disp == _lower_fname
                                or _real_base_l == _lower_fname):
                            if "/" in _real_norm_for_match:
                                _real_path = "_helpers_shared/" + _real_norm_for_match
                            else:
                                _tid_part = _displayed_name[:-(len(_fname)+1)] if _lower_disp.endswith("_" + _lower_fname) else ""
                                if _tid_part:
                                    _real_path = f"_helpers_shared/{_tid_part}/{_fname}"
                                else:
                                    _real_path = _displayed_name
                            _dependency_paths.append(
                                f"  - `{_fname}` → `{_real_path}`"
                            )
                            _matched = True
                            break
                    # 2026-05-12 P41.B: fuzzy match 兜底 — 主线程用 `task_id.ext` 错命名时
                    # 病因(实测 18:37 trace): 主线程派 embed_charts prompt 写 `paper_text.docx`
                    # (用 task_id+ext 错命名), 实际真名 `paper_text_paper.docx`. 严格匹配失败.
                    # 修法: 如果 fname 形如 `X.ext` 且 X 完全等于某 helper 的 task_id,
                    # 找该 helper 的同扩展名产物。
                    # P41.B 加强版: stem 必须**完全等于** task_id (从 real_relpath 第一段提取),
                    # 避免 `ssl.h` 误匹配 `ssl_algo_ssl.h` (stem=ssl 不等于 tid=ssl_algo)。
                    if not _matched and "." in _fname:
                        _stem, _ext = _fname.rsplit(".", 1)
                        _ext_lower = _ext.lower()
                        _stem_lower = _stem.lower()
                        for _displayed_name, _real_relpath in _remap.items():
                            if not isinstance(_displayed_name, str) or not isinstance(_real_relpath, str):
                                continue
                            _lower_disp = _displayed_name.lower()
                            # 提取真实 task_id (real_relpath 的第一段)
                            _real_norm = _real_relpath.replace("\\", "/")
                            if "/" not in _real_norm:
                                continue
                            _real_tid = _real_norm.split("/", 1)[0].lower()
                            # 严格校验: stem 必须完全等于 task_id, ext 必须匹配
                            if _stem_lower == _real_tid and _lower_disp.endswith("." + _ext_lower):
                                _real_path = "_helpers_shared/" + _real_norm
                                _dependency_paths.append(
                                    f"  - `{_fname}` → `{_real_path}` "
                                    f"(display name is `{_displayed_name}`)"
                                )
                                _matched = True
                                break

                if _dependency_paths or _fetchable_from_main:
                    _dependency_hint_parts = [
                        "## Dependency File Paths (system-provided)\n"
                    ]
                    if _dependency_paths:
                        _dependency_hint_parts.append(
                            "The following files are already available in this helper workspace or shared area. Use these relative paths directly:\n"
                            + "\n".join(_dependency_paths) + "\n"
                        )
                    if _fetchable_from_main:
                        _fetch_list = ", ".join(repr(p) for p in sorted(_fetchable_from_main))
                        _dependency_hint_parts.append(
                            "The following files exist in the main workspace but are not yet in this helper sandbox. Before reading them, call:\n"
                            f"`fetch_to_temp(source='main', paths=[{_fetch_list}])`\n"
                            "After the copy succeeds, read them by the same relative names. Use workspace-relative paths.\n"
                        )
                    _dependency_hint_parts.append(
                        "Use the paths above as evidence. If fetch/read fails, report the actual error.\n\n"
                        "系统提供依赖文件路径，优先按给定路径读取。"
                    )
                    dynamic_prompt_prefix_parts.append("".join(_dependency_hint_parts))
                    debug.log(
                        f"delegate.{task_id}.p34_inject",
                        f"P34: 注入 {len(_dependency_paths)} 个直接依赖路径, "
                        f"{len(_fetchable_from_main)} 个 main fetch 依赖 "
                        f"(referenced={len(_referenced_files)})"
                    )

                # 2026-05-12 P35: 找不到的文件 → 警告 helper "依赖缺失" 触发诚实声明
                # 病因(实测 15:47 trace): 主线程派 gen_paper_v4 prompt 引用 6 张 PNG + 1 csv,
                # 但任务 1 时 charts/merge_csv helper 都没派 → 7 个文件全不存在。
                # gen_paper_v4 没查直接编 12 分钟 → 用户被误导。
                # 修法: 系统主动告诉 helper "你引用的文件不在前置 helper 产物里",
                # 触发 helper 用 P33 ## 不可行 / ## 缺失或警告 自停 + 诚实报告。
                #
                # 2026-05-15 P98 修复 P35 误报 (实测排序论文 trace):
                #   误报 1 (framework): prompt 含模板占位符 "sort_xxx.c" / "xxx.h",
                #     P35 把这些当依赖文件警告 — 但 xxx 明显是 placeholder。
                #   误报 2 (sort_paper): 同批 spawn 的 sort_charts 会产 sort_comparison_random.png,
                #     P35 没看同批 siblings 的 expected_outputs → 误报 PNG 不存在。
                #   修法:
                #     (a) 占位符检测: 名字含 xxx/your_/<...>/__/placeholder 等模式跳过
                #     (b) 同批 siblings 检查: 用 batch_sibling_outputs 参数也作为 found
                _existing_basenames = _workspace_has_basename(helper_workspace, main_workspace)

                # P98 (a): 占位符模式 (template placeholder)
                _PLACEHOLDER_PATTERNS = (
                    "xxx", "yyy", "zzz",
                    "your_", "_your", "<", ">",
                    "placeholder", "todo", "filename",
                    "name_here", "_xxx_", "_yyy_",
                )

                _missing_dep_files = []
                _found_fnames = set()
                for _line in _dependency_paths:
                    # 从 line 抠出 fname
                    _name_m = _re_p34.search(r'`([^`]+\.[a-z]+)`', _line)
                    if _name_m:
                        _found_fnames.add(os.path.basename(_name_m.group(1).replace("\\", "/")).lower())
                _found_fnames.update(os.path.basename(p.replace("\\", "/")).lower() for p in _fetchable_from_main)
                _found_fnames.update(_expected_output_basenames)
                # P98 (b): 加入同批 sibling 的 expected_outputs
                if batch_sibling_outputs:
                    _found_fnames.update(
                        os.path.basename(str(p).replace("\\", "/")).lower()
                        for p in batch_sibling_outputs
                    )
                for _fname in _referenced_files:
                    _fname_base_l = os.path.basename(_fname.replace("\\", "/")).lower()
                    if _fname_base_l in _existing_basenames:
                        continue
                    if _fname_base_l not in _found_fnames:
                        # 排除常见非依赖类提及 (skill 名 / 文档名)
                        if _fname_base_l in {"readme.md", "license.txt", "package.json"}:
                            continue
                        # P98 (a): 跳过占位符
                        if any(pat in _fname_base_l for pat in _PLACEHOLDER_PATTERNS):
                            continue
                        if _is_framework_future_output_reference(
                            task_id=task_id,
                            kind=kind,
                            prompt=prompt,
                            filename=_fname,
                            expected_outputs=expected_outputs,
                        ):
                            continue
                        _missing_dep_files.append(_fname)

                if _missing_dep_files:
                    dynamic_prompt_prefix_parts.append(
                        "## Referenced File Availability Facts\n"
                        "The task prompt mentions the following file names. They were not found in known previous helper outputs "
                        "or current workspace evidence at helper start:\n"
                        + "\n".join(f"  - `{f}`" for f in sorted(_missing_dep_files)) +
                        "\n\nUse this as availability evidence, not as a conclusion about the task. "
                        "If a name is an output target, create or update it as requested. If a name is required input evidence, "
                        "confirm with read/search/fetch before relying on it; if it remains unavailable, report the exact missing "
                        "input and the resource needed from the main process.\n\n"
                        "这是文件可用性事实，不代表任务结论；若是目标文件则按要求创建，若是输入依赖则先确认再报告缺失资源。"
                    )
                    debug.log(
                        f"delegate.{task_id}.p35_facts",
                        f"P35: referenced file availability facts for {len(_missing_dep_files)} file(s): "
                        f"{_missing_dep_files[:5]}"
                    )
        except Exception:
            log.exception("P34 dependency path injection failed (non-fatal)")

    if mode == "hard" and not prior_summary and not resume:
        dynamic_prompt_prefix_parts.append(
            "## Hard Mode Runtime Focus\n"
            "Hard mode is a same-kind higher-rigor pass. Start by reading the task, available files, prior evidence, "
            "and any failure signal. Diagnose whether the blocker is routing, resources, paths, dependencies, scope, "
            "or acceptance evidence. If the same helper kind remains correct, proceed in small verified steps and "
            "report completed evidence, remaining gaps, and exact artifact paths.\n\n"
            "hard 是同类高严谨流程：先读任务和失败证据，诊断路径/资源/依赖/范围/验收，再分步验证。"
        )

    # Keep the helper system prompt cacheable. Turn-specific workspace,
    # resume, dependency, and task facts are all assembled into the user message.
    user_prompt = _build_helper_user_prompt(
        prompt=prompt,
        dynamic_prompt_prefix_parts=dynamic_prompt_prefix_parts,
        prior_summary=prior_summary,
        resume=resume,
        resume_workspace_empty=resume_workspace_empty,
        kind=kind,
    )

    msgs = [
        {"role": "system", "content": sys_prompt},
        {"role": "user", "content": user_prompt},
    ]

    # ── P1A 已废弃(2026-05-03 v17 修复)──
    # 原 v12 设计:msgs 末尾追加 {"role":"assistant", "prefix":True} 引导 LLM 续写。
    # **致命问题**:DeepSeek API 不允许 prefix + tools 同时使用,返回
    #   400 "Function call should not be used with prefix"
    # 实测 trace 9b3a6a6cb1fd4cd8:5 个 helper(naive/kmp/bm/rk/hybrid)在 iter 1
    # 全部 LLM call 失败 → 报告为空 → 用户主任务白做。
    # 解决:**完全移除 prefix 注入**。开局引导改靠 user_prompt 的"开局必做"段(已存在)。
    # 配合 system prompt 的「任务类型与 6 步流程」段 + dispatch 软 hint,三层引导依然有效。

    # 协作中断信号:复用主进程的 (archive, group, user) abort_event。
    # 同 user 发新消息触发 abort 时,helper 的 chat_with_tools_loop 在下一轮 iter
    # 顶部检测到, break 进 forced finalize,由于 finalize_kind="text_summary",
    # helper 输出纯文本进度报告,正好是我们要的。
    # 2026-05-01: per-user 锁改造后 get_abort_event 签名加了 user_id,这里必须传——
    # 否则 TypeError 会让 delegate 工具直接失败。
    shared_abort = get_group_guard().get_abort_event(archive_id, group_id, user_id)

    # ── Bug #6 修(Phase 5 收尾): helper-local abort,不污染共享 ──
    # 之前 helper stuck → abort_event.set() 直接污染共享 event,
    # 主线程 _round2 末尾的 `if abort_event.is_set(): suppress upgrade flags`
    # 会把 macro_escalation 触发的 upgrade 也压制掉。
    # 修复: 用本地 event,仅在用户级 abort 时桥接到 shared_abort。
    #
    # ── Bug #8 修 (v3 全面审计): processes.kill(this_helper) 只杀本 helper ──
    # 优先用 caller(handle_delegate / spawn_helper)创建并传下来的 local_abort。
    # 这样 ProcessRegistry 注册的 abort_event 就是同一个,
    # processes.kill 调用 h.abort_event.set() 不会污染他人。
    # 没传(直接调用本函数测试场景)则回退到内部新建。
    if local_abort is None:
        local_abort = asyncio.Event()

    async def _bridge_shared_to_local() -> None:
        """监听共享 abort(用户主动 abort),触发时同步到本地。

        注:用户 abort 时实际有两条并发路径:
          A. 主线程 chat_with_tools_loop 通过 racing cancel 当前 tool dispatch
             → handle_delegate 的 except CancelledError → cancel 所有 helper task(硬杀)
          B. 本 bridge 把 group abort 转成 local abort
             → helper 的 chat_with_tools_loop 看到 → 走 forced finalize 出报告

        路径 A 几乎总是更快(asyncio cancel 几 ms,LLM 调用做 forced finalize 要几秒),
        所以用户 abort 时 bridge 实际上是冗余兜底——helper 通常已经被 cancel 了
        而没机会走到 forced finalize。这是符合"用户 abort = 直接抛弃"语义的设计。

        反方向不会传染:本 helper stuck → 只 set local_abort,主线程不受影响。
        """
        try:
            await shared_abort.wait()
            local_abort.set()
        except asyncio.CancelledError:
            pass

    bridge_task = asyncio.create_task(_bridge_shared_to_local())
    abort_event = local_abort  # 给下面所有代码用本地 event(stuck/forced finalize 都走它)

    # ── ProcessRegistry 注册 + owner + abort ContextVar 设置 ──
    # owner_id 让 dispatch 能识别当前 LLM 调用归属(主线程 / 哪个 helper)
    # ProcessRegistry 让主线程 / 自己 / 后台 cleanup 能管理这个 helper 的生命周期
    # abort_event 让 workspace.run 子进程能被中途 kill(Phase 5++ 修)
    helper_owner = ProcessRegistry.make_helper_owner(parent_trace, task_id)
    owner_token = set_current_owner(helper_owner)
    abort_token = set_current_abort_event(local_abort)
    # 2026-05-11 P9: 让工具层(workspace.write 等)能识别当前 helper 的 kind
    from app.core.core_processes import (
        set_current_helper_expected_outputs,
        set_current_helper_kind,
        set_current_helper_write_scopes,
        reset_current_helper_expected_outputs,
        reset_current_helper_kind,
        reset_current_helper_write_scopes,
    )
    kind_token = set_current_helper_kind(kind)
    expected_outputs_token = set_current_helper_expected_outputs(expected_outputs or [])
    write_scopes_token = set_current_helper_write_scopes(write_scopes or expected_outputs or [])

    # 2026-05-03:helper 也 set 自己的 thread context 让 recall_thread 工具可用。
    # helper 的"原始任务"= prompt(主线程派发时给的完整任务描述)。
    # plan 字段留空(helper 没有 round2 plan,只有任务描述)。
    # 2026-05-03 fix:旧版本这里用了 task_description / extra_instructions 两个
    # 未定义变量,helper 启动 0.1 秒就 NameError 崩溃(实测 trace be9eba71 两次
    # 全部 6 个 helper 全崩 → 主线程 fallback 单线程实现树结构,浪费 9 分钟)。
    # 函数签名只有 `prompt` — 直接用它即可。
    from app.core.core_processes import (
        ThreadContext as _HelperThreadContext,
        set_current_thread_context as _h_set_thread_ctx,
        reset_current_thread_context as _h_reset_thread_ctx,
    )
    _helper_user_msg = (prompt or "")[:600]
    _helper_thread_ctx = _HelperThreadContext(
        user_message=_helper_user_msg,
        plan_intent=f"helper task: {task_id}",
        plan_key_points=[],
        plan_deliverables=[],
        role_label=f"helper:{task_id}",
    )
    _helper_thread_token = _h_set_thread_ctx(_helper_thread_ctx)

    # 注意:这里的 task 引用是 None,因为 helper 还没被 wrap 成 task。
    # 真正的 register 在 handle_delegate / spawn_helper 那一层完成(它持有 task ref)。
    # 这里只设 owner ContextVar 让本 helper 内部的工具调用能识别归属。

    # ── helper 自找 proc_id(2026-05-02 重构:无 race 路径) ──
    # 现在 wait_for_register 保证了 registry 在 helper 第一帧前已写完,
    # 所以单次查找即可,不需要 retry 轮询。
    # 找到后 set ContextVar,helper 内部 chat_with_tools_loop 调
    # report_helper_progress 时就能更新自己的 ProcessHandle 心跳字段。
    #
    # ── Bug B 真修(2026-05-02 复盘):之前用 `owner=helper_owner` 永远找不到 ──
    # 注册路径(handle_delegate / spawn_helper)用的 owner 是**主线程 owner**(main_owner),
    # 不是 helper 自己的 owner。helper 内部不知道 main_owner 是什么,所以必须用
    # `same_trace_as=helper_owner` 走 trace 范围匹配 — registry 会扫所有同 trace 的
    # 注册项,找到 task_id 匹配的那个。
    # 实测 trace 30079526 / b00c015a 5+ 个 helper 全部 heartbeat_skipped,就是因为
    # 这一行用了 owner= 严格匹配。
    from app.core.core_processes import (
        set_current_helper_proc_id, reset_current_helper_proc_id,
    )
    helper_proc_id_token = None
    helper_proc_id = None
    h = await proc_registry().find_helper_by_task_id(
        task_id, same_trace_as=helper_owner,
    )
    if h is not None:
        helper_proc_id = h.proc_id
        helper_proc_id_token = set_current_helper_proc_id(helper_proc_id)
        debug.log(
            f"delegate.{task_id}.heartbeat_armed",
            f"helper proc_id={helper_proc_id} found in registry, "
            f"heartbeat reporting enabled",
        )
    else:
        # 单元测试场景或异常 — caller 没走 _register_helper_with_autoclean。
        # 心跳汇报会 no-op,helper 仍能正常跑。
        # 真生产环境(handle_delegate / handle_spawn_helper)有 wait_for_register 兜底,
        # 不应到这里;到了说明 caller 路径有 bug,日志要醒目。
        log.warning(
            "helper %s: NOT registered in registry (heartbeat disabled). "
            "Production callers should pass wait_for_register Event — "
            "this is OK in unit tests but a bug in production.",
            task_id,
        )
        debug.log(
            f"delegate.{task_id}.heartbeat_skipped",
            f"helper task_id={task_id} not found in registry; running without heartbeat "
            f"(unit test path or registration bug)",
        )

    # ── Stuck detector — 反复失败时自动触发协作中断 ──
    # 2026-05-15 P106: 传 mode 给 StuckDetector, hard 模式阈值放宽 3x ("资源无限"策略)
    # 2026-06-04: 传 kind 给 StuckDetector, 用于 read helper 早期循环检测
    stuck_detector = StuckDetector(task_id, mode=mode, kind=kind)
    main_resource_request: dict | None = None

    # 2026-05-04 Claude Code-ified: 移除 _flow_hint 引导 + 四阶段状态机 +
    # _final_reasoning_callback + 时间硬上限。hard-mode helper 全自主,无监管。
    # 模型自己决定工作流程与何时停手(Claude Code 终止机制:不调工具 = 完成)。

    async def _dispatch(name: str, args: dict) -> str:
        """Helper dispatcher: routes to helper workspace (isolated from main)."""

        if mode == "hard" and name == "spawn_helper":
            return json.dumps({
                "ok": False,
                "error": (
                    "Hard-mode helpers cannot call spawn_helper. Complete the assigned work within this helper, "
                    "or report the precise blocker/resource need so the main process can coordinate recovery.\n"
                    "hard helper 不自行派发子 helper；需要资源时报告给主进程协调。"
                ),
            }, ensure_ascii=False)

        # ── 空 args 早检测(常见于 max_tokens 截断,给 LLM 友好错误) ──
        # 2026-05-11 升级(实测 trace 15:34-15:43 abpt 同错重复 3 次 9 分钟才 stuck):
        # 引入"复发感知"——第 N 次同错时,error 文本本身递进升级,从"建议"到"强制",
        # 让 LLM 看到这次比上次更严厉的话,改变行为概率↑。
        if not args and name in ("workspace", "office", "edit_file", "insert_in_file",
                                  "search_in_file", "read_file", "delegate"):
            # 累计本 helper 的空 args 错误数(用 stuck_detector 的 _calls 历史)
            _empty_args_count = sum(
                1 for c in getattr(stuck_detector, "_calls", [])
                if not c[1] and (
                    "empty arguments" in (c[2] or "").lower()
                    or "空对象" in (c[2] or "")
                )
            ) + 1  # +1 是本次
            if _empty_args_count == 1:
                err_text = (
                    f"Tool `{name}` received empty arguments `{{}}`. "
                    "The usual cause is an oversized tool-call batch: several parallel calls plus large content "
                    "can make the API output hit max_tokens before the last `arguments` field is complete. "
                    "Change strategy now: keep each round to at most two parallel tool calls, keep each content "
                    "payload near 5KB or less, and split larger writes across multiple turns. If this is just "
                    "one large-file write, do one workspace.write call in a turn.\n\n"
                    "空参数通常是输出被截断；减少并行工具调用并拆小内容。"
                )
            elif _empty_args_count == 2:
                err_text = (
                    f"Second empty-arguments error for `{name}`. The previous retry kept the same shape, "
                    "so another identical retry is likely to fail again. Use a smaller recovery plan now:\n"
                    "  A. For a large file, split the write into two or more calls: write the skeleton or first half, "
                    "then append or edit the remaining sections. Keep each content payload under about 4KB.\n"
                    "  B. For source files, split a large implementation into a few coherent files or modules.\n"
                    "  C. For multiple files, make tool calls sequential and write one file per turn.\n"
                    "Do not retry the same oversized call shape.\n\n"
                    "第二次空参数说明需要换策略：拆文件、拆调用、降低单轮并行。"
                )
            else:  # ≥ 3
                err_text = (
                    f"This is the {_empty_args_count}th repeated empty-arguments error. "
                    "Stop attempting the same oversized workspace/tool call. Report to the main thread that "
                    "the helper needs task splitting: smaller helper slices, a framework/skeleton helper first, "
                    "or one file/section per call. Include what you were trying to write and what slice boundary "
                    "would let the next attempt succeed.\n\n"
                    "多次空参数后应停止同形重试，向主线程报告需要拆分任务。"
                )
            err = json.dumps({"ok": False, "error": err_text}, ensure_ascii=False)
            if mode != "hard":
                stuck_detector.record(name, err)
                if stuck_detector.stuck and not abort_event.is_set():
                    abort_event.set()
                    log.warning(
                        "helper %s: stuck detected (%s), setting abort_event",
                        task_id, stuck_detector.stuck_reason,
                    )
            return err

        result = await tool_dispatch(
            name, args,
            archive_id=archive_id,
            group_id=group_id,
            user_id=user_id,
            workspace_dir=helper_workspace,
            caller="helper",
        )
        nonlocal main_resource_request
        _resource_request = _parse_main_resource_request(result)
        if _resource_request and main_resource_request is None:
            main_resource_request = _resource_request
            try:
                from app.core import agent_state
                agent_state.register_helper_resource_request(
                    trace_id=main_trace_id,
                    task_id=task_id,
                    helper_kind=kind,
                    request=_resource_request,
                    report=str(_resource_request.get("blocked_reason") or ""),
                )
            except Exception as e:
                log.warning("agent state resource request write failed: %r", e)
            abort_event.set()
            debug.log(
                f"delegate.{task_id}.resource_required",
                (
                    f"helper requested main resource kind={_resource_request.get('resource_kind')} "
                    f"reason={_resource_request.get('blocked_reason')} "
                    f"path={_resource_request.get('blocked_path')}"
                ),
            )
        # 记录 + 检查 stuck(hard mode 跳过:让高难重试完整跑完)
        if mode != "hard":
            stuck_detector.record(name, result)
            if stuck_detector.stuck and not abort_event.is_set():
                abort_event.set()
                log.warning(
                    "helper %s: stuck detected (%s), setting abort_event for forced finalize",
                    task_id, stuck_detector.stuck_reason,
                )
                debug.log(
                    f"delegate.{task_id}.stuck",
                    stuck_detector.stuck_reason,
                )

        return result

    interrupted = False
    # 1.3: helper 独立 FIX_HINT 重复计数
    from app.llm.tools.workspace import reset_fix_hint_counts
    reset_fix_hint_counts()

    # ── long-run observer: 跑超过 30 分钟时记一条 warning 到日志,但**不强制中断** ──
    # 主进程靠 processes.list 看心跳决定是否 kill;observer 只是给运维 grep 用。
    helper_started_at = time.monotonic()

    # 2026-05-11 P5.2: SUBAGENT_START hook dispatch
    try:
        from app.core.hooks import dispatch_hook, HookEvent
        dispatch_hook(HookEvent.SUBAGENT_START, {
            "task_id": task_id, "kind": kind, "resume": resume,
            "prompt_len": len(prompt), "user_lang": user_lang,
            "helper_think": helper_think,
            "archive_id": archive_id, "group_id": group_id, "user_id": user_id,
        })
    except Exception as _hook_e:
        log.debug("subagent_start hook dispatch failed: %r", _hook_e)

    long_run_observer_done = asyncio.Event()

    async def _long_run_observer():
        # 2026-05-10 Patch 60: 移除 P40 自动 paired helper(原阶段 0)
        #
        # 旧设计(P40 阶段 0):helper kind=code 跑 15min 时自动派 paired helper,
        # paired helper 与父 helper 并行,"谁先出可用结果谁赢"。
        #
        # 删除原因(trace f973df3770544567 暴露):
        #   - paired helper 也跑长(woat_impl_auto_final 跑 1h+,父 woat_impl 也跑 45min)
        #     "并行加速"完全没生效 — 任务复杂时 paired helper 也救不了
        #   - paired helper 是 helper 派的 sub-helper,owner=helper:{trace_id}:{parent_task},
        #     不是 main_owner → P55 cancel 漏掉(P59 修但根因是 P40 设计错)
        #   - paired helper 跟父 helper 数据竞争:父还在跑,paired helper 用父的中间状态做 fork → 数据不一致
        #     主线程拿到 paired helper 中间结果可能用错(trace 中 embed_charts 用了 stale 数据)
        #   - paired helper 占用一份完整 LLM API + ProcessRegistry 槽位,资源浪费
        #   - 跟 P55 哲学冲突:"系统不替模型决策" — paired helper 是系统强加的并行 race
        #
        # 新设计:codehelper 在报告 ## 验证建议 段(P55 已加)里说明是否需要同 task_id resume 或 mode='hard' 续作。
        # 主线程综合判断 + 显式派 base kind + mode(owner 自然是 main_owner,P55 cancel 覆盖)。
        # 替代方案:主线程 wait_helper 看心跳,觉得卡住就 processes.kill + 同 task_id resume。
        #
        # 简化为 2 阶段:
        #   阶段 1: LONG_RUN_OBSERVE(30 min)→ long_run 警告(给运维 grep)
        #   阶段 2: HARD_KILL_THRESHOLD(45 min)→ P33 强制 abort 走 forced finalize

        # 阶段 1:等到 LONG_RUN_OBSERVE
        try:
            await asyncio.wait_for(
                long_run_observer_done.wait(),
                timeout=_HELPER_LONG_RUN_OBSERVE,
            )
            return  # helper 在阶段 1 前结束,正常退出
        except asyncio.TimeoutError:
            pass
        elapsed = time.monotonic() - helper_started_at
        log.warning(
            "helper %s long_run: %.0fs elapsed (no hard limit, observer only). "
            "main thread should poll heartbeat & decide if kill is warranted.",
            task_id, elapsed,
        )
        debug.log(
            f"delegate.{task_id}.long_run",
            f"helper 已运行 {elapsed:.0f}s 仍未完成 — 主进程请用 processes.list 看心跳决定下一步",
        )

        # 阶段 2:再等到 HARD_KILL_THRESHOLD 后强制 abort
        # (Patch 33: 病态 helper 死磕,主线程被 wait_window 锁住时兜底)
        # 二次审计修(Bug F): mode=hard 或 kind=verify 设计上跳过兜底 hard kill,
        # 时间换正确性是它们的核心设计。这里 skip 它们的 hard kill。
        # 父 helper(kind=code, mode=easy)正常被 hard kill,主线程再按报告决定
        # 同 task_id resume、mode='hard' 续作、拆小任务或放弃。
        # 如果 hard-mode 续作仍然长跑,主线程的 wait_window 会自动返回 still_running 心跳,让模型决定是否手动 kill。
        if mode == "hard" or kind == "verify":
            # 仍然等 helper 自然结束,避免 observer 提前退出导致后续清理被错过
            await long_run_observer_done.wait()
            return
        _hard_kill_remaining = max(
            0.0, _HELPER_HARD_KILL_THRESHOLD - _HELPER_LONG_RUN_OBSERVE,
        )
        if _hard_kill_remaining > 0:
            try:
                await asyncio.wait_for(
                    long_run_observer_done.wait(),
                    timeout=_hard_kill_remaining,
                )
                return
            except asyncio.TimeoutError:
                pass
            elapsed = time.monotonic() - helper_started_at
            log.warning(
                "helper %s hard kill: %.0fs >= %.0fs threshold, set local_abort",
                task_id, elapsed, _HELPER_HARD_KILL_THRESHOLD,
            )
            debug.log(
                f"delegate.{task_id}.hard_kill",
                f"helper 已运行 {elapsed:.0f}s,超过 {_HELPER_HARD_KILL_THRESHOLD}s 硬阈值,"
                f"observer 自动 set abort 让它走 forced finalize",
            )
            try:
                local_abort.set()
            except Exception:
                pass

    long_run_task = asyncio.create_task(_long_run_observer())

    try:
        from app.llm.model_pool import resolve_task

        use_lite = _helper_lite_var.get()
        # ── 读取上次失败计数(供 resume 升级用) ──
        _fail_count = 0
        _fail_path = os.path.join(helper_workspace, "_failure_count.txt")
        if os.path.isfile(_fail_path):
            try:
                with open(_fail_path, "r", encoding="utf-8") as _f:
                    _fail_count = int(_f.read().strip())
            except (ValueError, OSError):
                pass

        # hard/verify helpers never downgrade to lite — they need analytical rigor
        if mode == "hard":
            if kind in ("code", "coding"):
                _h_spec = resolve_task("helper_full_legacy_hard")
            elif kind == "verify":
                _h_spec = resolve_task("helper_full_verify_think")
            elif kind == "read":
                _h_spec = resolve_task("helper_full_read_hard")
            elif kind == "edit":
                _h_spec = resolve_task("helper_full_edit_hard")
            elif kind == "draw":
                _h_spec = resolve_task("helper_full_draw_hard")
            elif kind in {"project_map", "file_summary", "impact_review", "inventory", "summarize"}:
                _h_spec = resolve_task("helper_full_project_analysis_hard")
            elif kind == "tts":
                _h_spec = resolve_task("helper_full_tts_hard")
            else:
                _h_spec = resolve_task("helper_full_coding_think")
        elif use_lite and kind in ("code", "coding", "edit"):
            _h_spec = resolve_task("helper_lite")
        elif kind == "verify":
            _h_spec = resolve_task("helper_full_verify")
        elif kind in ("code", "coding"):
            _h_spec = resolve_task("helper_full_coding")
        elif kind == "edit":
            _h_spec = resolve_task("helper_full_edit")
        else:
            _h_spec = resolve_task("helper_full_coding")

        # ── Opt 5: 连续失败 ≥2 次自动升级 lite → full ──
        if _fail_count >= 2 and use_lite and kind in ("code", "coding", "edit"):
            _upgrade_kind = "coding" if kind == "coding" else kind
            _upgrade_key = f"helper_full_{_upgrade_kind}" if _upgrade_kind != "code" else "helper_full_coding"
            _h_spec = resolve_task(_upgrade_key)
            debug.log(
                f"delegate.{task_id}.model_upgrade",
                f"upgraded {kind} helper from lite→{_upgrade_key} "
                f"after {_fail_count} prior failures",
            )

        # ── 2026-05-08: helper_think 覆盖 ──
        # 主线程 delegate 时传 helper_think=true → coding/verify helper 启用 reasoning=low
        if helper_think:
            if kind in ("code", "coding"):
                _h_spec = resolve_task("helper_full_coding_think")
                debug.log(
                    f"delegate.{task_id}.helper_think",
                    f"enabled reasoning=low for {kind} helper (helper_think=true)",
                )
            elif kind == "verify":
                _h_spec = resolve_task("helper_full_verify_think")
                debug.log(
                    f"delegate.{task_id}.helper_think",
                    f"enabled reasoning=low for verify helper (helper_think=true)",
                )

        debug.log(
            "delegate.helper_route",
            (
                f"task_id={task_id} kind={kind} mode={mode} resume={bool(resume)} "
                f"resumed_actually={bool(resumed_actually)} helper_lite={bool(use_lite)} "
                f"helper_think={bool(helper_think)} fail_count={_fail_count} "
                f"model={_h_spec.model} reasoning={_h_spec.reasoning}"
            ),
        )

        # ── 2026-05-05: API 卡死检测(无流式 chunk 超阈值 → kill + 重试)──
        # 2026-05-09 Patch 16: 阈值从 60s 提到 90s。trace 96c47298 实测 bplus_tree/rb_tree
        # 在 69s 处被砍掉(刚过 60s),而它们其实是模型在编译错误诊断阶段思考时间长,
        # 不是真卡死。90s 给思考更宽容,但仍能在 1-2 分钟内识别真死锁。
        # 2026-05-09 Patch 17: 把 time.time() 换成 time.monotonic()。
        # NTP 时钟跳变(尤其向前)会让 _last_chunk_time 与 now 之差突增,触发虚假 stall。
        # monotonic 时钟保证单调递增,杜绝时钟跳坑。
        # 2026-05-11 B2 改: 阈值按上下文大小动态化。
        # 实测 trace 822f2aaa: skiplist iter 30 / fct iter 25 / verify_bench iter 13
        # 三次被 90s 误杀,这些 helper 在中后期上下文 100-300K tokens,DeepSeek V4
        # 在大上下文下 TTFT 显著拉长(模型并不卡,只是首 token 延迟)。
        # 按当前 messages 字节数估算 token,分档给阈值:
        _API_STALL_TIMEOUT = _stall_threshold_for(msgs)
        debug.log(
            f"delegate.{task_id}.stall_threshold",
            f"adaptive stall threshold = {_API_STALL_TIMEOUT:.0f}s "
            f"(initial msgs ≈ {_estimate_msgs_tokens(msgs)} tokens)",
        )
        _last_chunk_time = [time.monotonic()]  # mutable container for closure
        # 2026-05-18 P173: stream_open 状态门 — 仅在 stream 打开时跑 stall 监控。
        # 病因(trace 2026-05-18 124610): helper 调 office.ocr_images 16 张图,
        # 工具 serial 执行 60s+, 期间 LLM stream 已 close, 但 stall monitor 仍按
        # _last_chunk_time 老时间戳计时, 90s 后误判 API 卡死 → 杀 helper。
        # 修复: stream close 时停止计时, stream open 时重置, 工具执行不计入 stall。
        _stream_open = [False]

        def _on_stream_chunk():
            """每次收到流式 chunk 时调用,重置卡死计时器。"""
            _last_chunk_time[0] = time.monotonic()

        def _on_stream_event(event: str, reason: str | None = None):
            """stream open/close 状态切换。close 期间不计 stall (工具在跑)。"""
            if event == "open":
                _stream_open[0] = True
                _last_chunk_time[0] = time.monotonic()
            elif event == "close":
                _stream_open[0] = False
                # 不更新 _last_chunk_time, 下次 open 时会重置

        async def _api_stall_monitor():
            """后台监控:超过 _API_STALL_TIMEOUT 无 chunk → 判定 API 卡死。
            仅在 stream 打开期间计时, 工具派发 (stream close) 期间跳过。"""
            while not abort_event.is_set():
                await asyncio.sleep(10.0)
                # 2026-05-18 P173: stream 没开就不算 stall (工具在跑或回合间隙)
                if not _stream_open[0]:
                    continue
                since_last = time.monotonic() - _last_chunk_time[0]
                if since_last > _API_STALL_TIMEOUT:
                    log.warning(
                        "helper %s: API stall detected (%.0fs no chunk during open stream) — "
                        "triggering abort for kill+retry",
                        task_id, since_last,
                    )
                    debug.log(
                        f"delegate.{task_id}.api_stall",
                        f"API stall: {since_last:.0f}s since last chunk, "
                        f"setting abort_event",
                    )
                    abort_event.set()
                    break
                # 前 60s 每秒检查一次
                if since_last > 30.0:
                    await asyncio.sleep(5.0)
                else:
                    await asyncio.sleep(10.0)

        _stall_monitor_task = asyncio.create_task(_api_stall_monitor())

        # ── 2026-05-11 B4: helper auto-promote to think tier ──
        # 当 code/edit helper 连续失败 N 次(timeout / edit_fail / read_reject)时,
        # 自动切到 think 档(coding_think / verify_think),给后续 iter 开 reasoning=low。
        # 不再依赖主线程主动传 helper_think=true。
        #
        # 触发条件(满足任一即升档):
        #   - 累计 timeout >= 3 次
        #   - StuckDetector 在最近窗口报 stuck (反复同种错误)
        #   - iter > 30 + 未产出任何文件
        #
        # 一旦升档,本 helper 后续 iter 全用 think,不再来回切换。
        # final / verify_think / 已是 think 档的不重复升档。
        _auto_promote_state = {
            "promoted": False,
            "promoted_at_iter": -1,
        }
        _no_promote_kinds = ("verify",)  # verify 已经足够强,不升
        _is_already_think = "think" in (_h_spec.model or "") or _h_spec.reasoning in (
            "low", "high", "max"
        )
        _can_auto_promote = (
            kind in ("code", "coding", "edit", "verify")
            and kind not in _no_promote_kinds
            and not _is_already_think
        )

        def _auto_promote_callback(it: int, messages: list[dict]):
            """检测 helper 失败模式,触发档位升级。"""
            if not _can_auto_promote or _auto_promote_state["promoted"]:
                return None
            # 1) 累计 timeout
            try:
                from app.llm.client import _get_timeout_dict
                _timeout_count = _get_timeout_dict().get(task_id, 0)
            except Exception:
                _timeout_count = 0
            # 2) stuck detector
            _is_stuck = bool(stuck_detector.stuck)
            # 3) iter > 30 + 无产出文件
            _no_files_runaway = False
            if it > 30:
                try:
                    has_files = any(
                        not name.startswith(".") and
                        not name.endswith((".session_tag", ".helper_summary.txt"))
                        for name in os.listdir(helper_workspace)
                    )
                    if not has_files:
                        _no_files_runaway = True
                except OSError:
                    pass

            should_promote = (
                _timeout_count >= 3 or
                _is_stuck or
                _no_files_runaway
            )
            if not should_promote:
                return None

            # 决定升档目标
            try:
                from app.llm.model_pool import resolve_task
                if kind in ("code", "coding"):
                    new_spec = resolve_task("helper_full_legacy_hard")
                elif kind == "verify":
                    new_spec = resolve_task("helper_full_verify_think")
                else:
                    # edit: 升到 mid-pro disabled (已是 P0,再升没用)
                    return None
            except Exception:
                return None

            reason_parts = []
            if _timeout_count >= 3:
                reason_parts.append(f"timeout={_timeout_count}")
            if _is_stuck:
                reason_parts.append(f"stuck:{stuck_detector.stuck_reason}")
            if _no_files_runaway:
                reason_parts.append("iter>30+no_files")

            _auto_promote_state["promoted"] = True
            _auto_promote_state["promoted_at_iter"] = it
            debug.log(
                f"delegate.{task_id}.auto_promote",
                f"auto-promoted to think tier at iter {it} | "
                f"reason: {', '.join(reason_parts)}",
            )
            return {"model_spec": new_spec}

        # 按 kind 过滤工具集；edit 保留 python 轻量文本处理，但禁用 bash
        _tools_for_helper = _filter_tools_for_kind(kind, _HELPER_TOOLS)
        content, _msgs = await chat_with_tools_loop(
            msgs,
            tools=_tools_for_helper,
            dispatcher=_dispatch,
            model_spec=_h_spec,
            abort_event=abort_event,
            finalize_kind="text_summary",
            reasoning_callback=_auto_promote_callback if _can_auto_promote else None,
            task_id=task_id,  # 2026-05-03 v18.x:用于跨 timeout 累计降级
            helper_kind="legacy_hard" if (kind == "code" and mode == "hard") else kind,
            chunk_callback=_on_stream_chunk,  # 2026-05-05: API stall detection
            stream_event_cb=_on_stream_event,  # 2026-05-18 P173: gate stall monitor on stream open/close
        )

        # ── 清理 stall monitor ──
        if not _stall_monitor_task.done():
            _stall_monitor_task.cancel()
            try:
                await _stall_monitor_task
            except asyncio.CancelledError:
                pass
        report = content.strip() if content else "(无输出)"
        if (
            main_resource_request is None
            and not abort_event.is_set()
            and _needs_helper_report_format_self_repair(
            report,
            expected_outputs=expected_outputs,
            kind=kind,
            )
        ):
            try:
                repair_prompt = _build_helper_report_format_repair_prompt(
                    original_report=report,
                    expected_outputs=expected_outputs,
                )
                repair_msgs = list(_msgs or msgs)
                repair_msgs.append({"role": "user", "content": repair_prompt})
                debug.log(
                    f"delegate.{task_id}.report_format_repair",
                    "helper final report output-file declaration malformed; giving helper one local repair turn",
                )
                repaired_content, _repair_msgs = await chat_with_tools_loop(
                    repair_msgs,
                    tools=_tools_for_helper,
                    dispatcher=_dispatch,
                    model_spec=_h_spec,
                    abort_event=abort_event,
                    finalize_kind="text_summary",
                    reasoning_callback=None,
                    task_id=task_id,
                    helper_kind="legacy_hard" if (kind == "code" and mode == "hard") else kind,
                    chunk_callback=_on_stream_chunk,
                    stream_event_cb=_on_stream_event,
                )
                if repaired_content and repaired_content.strip():
                    report = repaired_content.strip()
                    _msgs = _repair_msgs
            except Exception as _repair_e:
                debug.log(
                    f"delegate.{task_id}.report_format_repair_failed",
                    f"helper report format repair turn failed: {type(_repair_e).__name__}: {_repair_e}",
                )
        _disk_declared_report_path = os.path.join(helper_workspace, f".helper_{task_id}_full_report.txt")
        if os.path.isfile(_disk_declared_report_path):
            try:
                with open(_disk_declared_report_path, "r", encoding="utf-8") as _f:
                    _disk_declared_report = _f.read()
                if _extract_declared_files(_disk_declared_report):
                    report = _disk_declared_report if report == "(无输出)" else report + "\n\n" + _disk_declared_report
            except OSError:
                pass
        interrupted = abort_event.is_set()
        # ── stuck 信息附加(让主线程能识别"反复失败,该换策略"的 helper)──
        was_stuck = stuck_detector.stuck
        stuck_reason = stuck_detector.stuck_reason if was_stuck else ""
        expected_outputs_for_copy = list(expected_outputs or [])
        if kind in {"read", "ocr"}:
            # read helper 的 .txt 是内部识别/提取材料,不是用户交付物。主线程可读取后重写,
            # 或派 edit helper 生成 docx/txt 终稿；不要把原始 OCR 报告自动提升给用户。
            expected_outputs_for_copy = []

        # 将 helper 生成的非源码文件复制回主工作区(即使中断也要尝试,
        # 那些可能是部分成果)
        # B1+B2 修复: 用 fork_snapshot diff 只复制真正新增/修改的;清单截顶 20 条
        _reported_declared = set(_extract_reported_output_files(report))
        declared = _reported_declared or _extract_declared_files(report)
        declared_for_copy = declared if kind not in {"read", "ocr"} else set()
        if not declared_for_copy:
            declared_for_copy = set(expected_outputs_for_copy)
        copied_back, copy_stats, _file_map = _copy_results_to_main(
            helper_workspace, main_workspace, task_id,
            fork_snapshot=fork_snapshot,
            declared_files=declared_for_copy if declared_for_copy else None,
            allow_shared_merge=not interrupted,
            expected_outputs=expected_outputs_for_copy,  # P73: 让 copy 知道哪些文件可以干净命名
            # 2026-05-21: 被竞速 abort 的 hard 副本(命名 *_hard + interrupted)无独立产物属正常,
            # 抑制 declared_missing 误报。
            is_race_aborted_twin=bool(
                interrupted and mode == "hard" and str(task_id).endswith("_hard")
            ),
            suppress_declared_missing=bool(main_resource_request),
            copy_unexpected_env_files=kind in {"read", "ocr"},
            helper_kind=kind,
        )
        if copied_back:
            shown = copied_back[:20]
            suffix = (
                f" and {len(copied_back)} files total"
                if len(copied_back) > 20 else ""
            )
            note_lines = [
                f"\n\n[Generated files copied to the main workspace: {', '.join(shown)}{suffix}]\n"
                f"生成文件已复制到主工作区。"
            ]
            if copy_stats.get("capped"):
                note_lines.append(
                    f"[Warning: helper copyback exceeded {_RESULT_COPY_BACK_MAX_FILES} files and was truncated. "
                    f"The main process should inspect task granularity; this often indicates runaway artifact growth "
                    f"or an overly broad helper boundary.]\n"
                    f"复制文件过多已截断；主进程应检查任务粒度。"
                )
            report += "".join(note_lines)

        # ── Bug H 修 (2026-05-02): 通知主线程哪些文件因太大没复制 ──
        # 实测 trace e4eeb133: helper 写 269MB bench_output.txt,被 50MB 上限 skip,
        # 但 report 里没说,主线程在主区找不到这个文件却以为存在,会困惑或重做。
        # 把 skipped_large 列表附在 report 里让主线程清楚"哪些产物没拿到"。
        skipped_large = copy_stats.get("skipped_large") or []
        if skipped_large:
            total_mb = sum(s for _, s in skipped_large) / 1024 / 1024
            shown_skip = ", ".join(
                f"{n} ({s/1024/1024:.0f}MB)"
                for n, s in skipped_large[:5]
            )
            more = (
                f" and {len(skipped_large)} files total"
                if len(skipped_large) > 5 else ""
            )
            report += (
                f"\n\n[Warning: these files exceeded the {_RESULT_COPY_BACK_MAX_SIZE//1024//1024}MB "
                f"copyback limit and were not copied to the main workspace "
                f"(total {total_mb:.0f}MB): {shown_skip}{more}. "
                f"If the main process needs this evidence, the helper should reduce it first: trim, summarize, "
                f"split, or extract key rows into a smaller file below the limit, then deliver that smaller file.]\n"
                f"大文件未回传；helper 应先裁剪、摘要或拆分成可用小文件。"
            )

        # ── 完整报告写入磁盘(替代旧 _REPORT_HARD_CAP 32KB 截断补丁) ──
        # 架构修复: 主线程拿结构化摘要,完整报告存磁盘按需读取。
        _full_report = report
        try:
            if main_workspace and os.path.isdir(main_workspace):
                main_full_path = os.path.join(
                    main_workspace, f".helper_{task_id}_full_report.txt"
                )
                with open(main_full_path, "w", encoding="utf-8") as f:
                    f.write(_full_report)
            # 同时写沙箱(供 resume 读取)
            full_path = os.path.join(helper_workspace, ".helper_summary.txt")
            with open(full_path, "w", encoding="utf-8") as f:
                f.write(_full_report)
            debug.log(
                f"delegate.{task_id}.full_report",
                f"full report ({len(_full_report)} chars) persisted to disk",
            )
        except OSError:
            log.exception("failed to persist full helper report")

        if interrupted:
            # 显式标注让主进程一眼看出可以续作
            report = (
                f"[This helper run was interrupted, and the workspace was preserved. "
                f"To continue the same work boundary, resume task_id={task_id} with resume=true.]\n"
                f"helper 中断摘要：工作区已保留，可同 task_id 续作。\n\n"
                + report
            )

        # ── stuck 时给主线程明确的 escalation 建议 ──
        if was_stuck:
            report = (
                f"[Warning: the helper stopped after a repeated-failure pattern.]\n"
                f"Reason: {stuck_reason}\n"
                f"Main process recovery options:\n"
                f"  - Resume the same task_id with a narrower prompt and a changed approach.\n"
                f"  - Escalate the same base kind to a stronger mode when the task boundary is already narrow.\n"
                f"  - Split the work into smaller helpers or request missing resources before continuing.\n"
                f"  - If recovery is not useful, report the verified partial result and the exact remaining gap.\n"
                f"helper 反复失败后停手；主进程应续作、升级、拆分或明确部分完成。\n\n"
                + report
            )
        if main_resource_request:
            report = (
                "[Helper frozen: main-process resource required]\n"
                f"Needed resource kind: '{main_resource_request.get('suggested_helper_kind')}'\n"
                f"Reason: {main_resource_request.get('blocked_reason') or main_resource_request.get('error')}\n"
                f"Suggested action: {main_resource_request.get('main_thread_action') or 'spawn the requested helper resource, then resume or respawn this helper'}\n"
                f"helper 已冻结等待资源；主进程提供或拒绝资源后再续作。\n\n"
                + report
            )

        # ── Opt 5: 持久化失败计数(供 resume 自动升级) ──
        _failed_this_run = bool(interrupted or was_stuck or main_resource_request)
        try:
            if _failed_this_run:
                _new_fail = _fail_count + 1
                with open(_fail_path, "w", encoding="utf-8") as _f:
                    _f.write(str(_new_fail))
            elif _fail_count > 0 and os.path.isfile(_fail_path):
                os.remove(_fail_path)  # 成功则清零
        except OSError:
            pass
        # 从 report 提取语义摘要 — helper 被要求在报告末尾写 ## 摘要 段
        _clean = report
        if _clean.startswith("[This helper run was interrupted") or _clean.startswith("[本次执行被打断"):
            _clean = _clean.split("\n\n", 1)[-1] if "\n\n" in _clean else _clean
        if _clean.startswith("[Warning: the helper stopped") or _clean.startswith("[⚠ helper 检测到反复失败"):
            _clean = _clean.split("\n\n", 1)[-1] if "\n\n" in _clean else _clean
        summary = ""
        for marker in ("## 摘要", "## 总结"):
            if marker in _clean:
                after = _clean.split(marker, 1)[-1].strip()
                cutoff = len(after)
                for delim in ("\n## ", "\n```"):
                    idx = after.find(delim)
                    if idx != -1 and idx < cutoff:
                        cutoff = idx
                summary = after[:cutoff].strip()
                break
        if not summary:
            for line in _clean.splitlines():
                s = line.strip()
                if s.startswith("VERDICT:") and len(s) > 10:
                    summary = s
                    break
        _report_verdict: str | None = None
        for line in _clean.splitlines()[:20]:
            s = line.strip()
            if not s.upper().startswith("VERDICT:"):
                continue
            value = s.split(":", 1)[1].strip().upper()
            for candidate in ("FAIL", "PARTIAL", "PASS"):
                if candidate in value:
                    _report_verdict = candidate
                    break
            break
        if not summary:
            # 2026-05-08 优化: 跳过 ```json {"files": [...]} ``` / 单纯的文件声明,
            # 那不是有意义的摘要。helper 可能把文件声明放在报告最前面。
            paras = [p.strip() for p in _clean.split("\n\n") if p.strip()]
            for p in paras:
                _stripped = p.strip("` \n\t")
                # 跳过 JSON 文件声明、纯 markdown 标题、产出文件清单
                if _stripped.startswith("{") and "files" in _stripped[:20]:
                    continue
                if _stripped.startswith("##") and len(_stripped) < 20:
                    continue
                if _stripped.startswith("产出文件:") or _stripped.startswith("## 产出"):
                    continue
                summary = p[:300]
                break
        summary = summary[:300] if len(summary) > 300 else summary
        # ── 构建紧凑 report(替代旧 _REPORT_HARD_CAP 32KB 截断) ──
        # 主线程默认看 summary(≤300 字),report 是 summary + 文件清单 + 磁盘指针,
        # 自然边界 < 2KB。需要完整内容时 read_file .helper_{task_id}_full_report.txt。
        _reported_declared = set(_extract_reported_output_files(_full_report))
        _declared = _reported_declared or _extract_declared_files(_full_report)
        _declared_set = set(_declared) if _declared else set()
        _registry_copied_back, _registry_visible_outputs = _registry_helper_outputs(main_workspace, task_id, kind)
        if _registry_copied_back:
            copied_back = sorted(set(copied_back or []) | set(_registry_copied_back))
            copy_stats.setdefault("registry_helper_outputs", _registry_copied_back)
        _copied_set = set(copied_back) if copied_back else set()
        _env_available_files = sorted(set(copy_stats.get("env_copied_files") or []))
        _workspace_copied_back = [
            path for path in copied_back
            if not _is_internal_helper_artifact(path)
            and not _is_shared_support_artifact(path)
        ]
        if _registry_visible_outputs:
            _workspace_copied_back = sorted(set(_workspace_copied_back) | set(_registry_visible_outputs))
        if kind in {"read", "ocr"}:
            _workspace_copied_back = []

        _visible_copied_back = [
            path for path in copied_back
            if not _is_internal_helper_artifact(path)
            and not _is_shared_support_artifact(path)
            and not str(path).replace("\\", "/").startswith(("_helpers_shared/", "_shared/"))
            and not str(path).replace("\\", "/").startswith("_env/")
        ]
        if _registry_visible_outputs:
            _visible_copied_back = sorted(
                path for path in (set(_visible_copied_back) | set(_registry_visible_outputs))
                if not str(path).replace("\\", "/").startswith(("_env/", "_helpers_shared/", "_shared/"))
            )
        if kind in {"read", "ocr"}:
            _visible_copied_back = []
        if kind not in {"read", "ocr"} and _env_available_files:
            _existing_basenames = {os.path.basename(p) for p in _workspace_copied_back}
            _deduped_env = [p for p in _env_available_files if os.path.basename(p) not in _existing_basenames]
            _workspace_copied_back = sorted(set(_workspace_copied_back) | set(_deduped_env))
        _visible_copied_set = set(_visible_copied_back)

        # 2026-05-15 修(trace 28d9525e):此 debug.log 原本放在 ~60 行之前的位置,
        # 引用了 _visible_copied_back,但变量定义在这里(几十行之后) —— Python
        # 编译器在函数内见到下方有 `_visible_copied_back = ...` 赋值就把整个名字
        # 标为 local,上面那次引用直接 raise UnboundLocalError, helper 已成功
        # 写完 docx 也被报为 "crashed",主线程重试又重试,最终绝望转人设拒绝。
        # 把日志移到变量定义之后即可。
        debug.log(
            f"delegate.{task_id}.done",
            f"report={len(report)} chars, files={_visible_copied_back}, "
            f"interrupted={interrupted}, stuck={was_stuck}",
            report[:500],
        )

        _internal_delivered = {
            path for path in _copied_set
            if _is_internal_helper_artifact(path)
        }
        _support_delivered = {
            path for path in _copied_set
            if _is_shared_support_artifact(path)
        }
        _env_project_staged_delivered = {
            path for path in _copied_set
            if str(path).replace("\\", "/").startswith("_env/")
        }
        _user_visible_delivered = _copied_set - _internal_delivered - _support_delivered - _env_project_staged_delivered
        if kind in {"read", "ocr"}:
            _user_visible_delivered = set()
        _user_visible_copied = [path for path in _visible_copied_back if path in _user_visible_delivered]
        _read_internal_evidence_files: list[str] = []
        if kind in {"read", "ocr"}:
            _read_internal_evidence_files = _collect_read_evidence_files(
                main_workspace=main_workspace,
                task_id=task_id,
                copied_paths=copied_back or [],
            )
        _shared_visible_delivered = {
            path for path in copied_back
            if path.startswith("_helpers_shared/")
        }
        _declared_to_copied = {
            str(m.get("helper_name")): str(m.get("main_name"))
            for m in (_file_map or [])
            if isinstance(m, dict) and m.get("helper_name") and m.get("main_name")
        }
        _shared_to_main = {
            str(m.get("shared_name")): str(m.get("main_name"))
            for m in (_file_map or [])
            if isinstance(m, dict) and m.get("shared_name") and m.get("main_name")
        }
        for _shared_name, _main_name in _shared_to_main.items():
            _declared_to_copied[_shared_name] = _main_name
        _visible_name_set = set(_visible_copied_set)
        _visible_name_set.update(_shared_to_main.keys())
        _visible_name_set.update(_shared_visible_delivered)
        _visible_name_set.update(_env_available_files)
        _declared_but_missing = sorted(
            name for name in _declared_set
            if _declared_to_copied.get(name, name) not in _visible_name_set
            and not _matches_declared_output_via_mapping(
                _declared_to_copied.get(name, name),
                {name},
                _file_map or [],
            )
        )
        _pending_declared_outputs = _declared_but_missing if main_resource_request else []
        if main_resource_request and not _pending_declared_outputs:
            _pending_declared_outputs = sorted(
                str(name).replace("\\", "/").strip()
                for name in (expected_outputs_for_copy or [])
                if str(name).strip()
                and str(name).replace("\\", "/").strip() not in _visible_name_set
            )
        if main_resource_request:
            _declared_but_missing = []
        _delivered_but_not_declared = sorted(
            path for path in _user_visible_delivered
            if not _matches_declared_output_via_mapping(path, _declared_set, _file_map or [])
        )
        _file_list = ", ".join(sorted(_declared)) if _declared else (
            ", ".join(copied_back[:8]) if copied_back else "(无)"
        )
        _disk_ptr = (
            f".helper_{task_id}_full_report.txt"
            f"({len(_full_report)} 字符)"
        )
        # L5-1 (2026-05-09): report 区分"实际交付"和"声明但未生成"
        _report_lines = [summary]
        _report_lines.append(f"已交付到主区: {', '.join(_user_visible_copied[:20]) if _user_visible_copied else '(无)'}")
        if _read_internal_evidence_files:
            _report_lines.append(
                f"内部读取证据: {', '.join(_read_internal_evidence_files[:12])}"
        )
        if _declared_but_missing:
            _report_lines.append(
                f"Declared but not produced: {', '.join(_declared_but_missing[:10])}\n"
                f"声明但未生成。"
            )
        if _pending_declared_outputs:
            _report_lines.append(
                f"Pending until requested resource is provided: {', '.join(_pending_declared_outputs[:10])}\n"
                f"等待资源后继续生成。"
            )
        if _delivered_but_not_declared:
            _report_lines.append(
                f"Delivered but not declared: {', '.join(_delivered_but_not_declared[:10])}\n"
                f"已交付但未声明。"
            )
        if interrupted and not copy_stats.get("shared_merge_allowed", True):
            _report_lines.append(
                "Shared handoff artifacts were not merged: `_helpers_shared/` remains only in the helper workspace.\n"
                "共享支撑产物未合并。"
            )
        _report_lines.append(f"Full report: {_disk_ptr}\n完整报告路径。")
        report = "\n".join(_report_lines)
        _elapsed = time.monotonic() - helper_started_at
        _pre_output_quality_warnings: list[dict] = []
        try:
            for _cw in text_mojibake_warnings(_full_report, file=f".helper_{task_id}_full_report.txt"):
                _pre_output_quality_warnings.append(_cw)
        except Exception:
            pass
        # 2026-05-08 优化: 主线程响应字段精简。
        # 旧版每个 result 含 workspace_dir(绝对路径,泄露)、summary(已并入 report)、
        # resumed_from(=输入参数 echo)、resumed_actually(False 时无意义)、stuck_reason(空字符串)。
        # 新版:summary 并入 report;workspace_dir 通过 task_id 间接寻址不暴露;
        # 仅在确实 stuck/interrupted/resumed 时才保留对应字段。
        _result = {
            "task_id": task_id,
            "ok": (
                not bool(main_resource_request)
                and _report_verdict not in ("FAIL", "PARTIAL")
            ),
            "report": report,
            "files": _workspace_copied_back,
            "workspace_files": _workspace_copied_back,
            "user_visible_files": _visible_copied_back,
            "declared_files": sorted(_declared_set),
            "declared_but_missing": _declared_but_missing,
            "delivered_but_not_declared": _delivered_but_not_declared,
            "file_map": _file_map,
            "copy_stats": copy_stats,
            "elapsed_sec": round(_elapsed, 1),
            # 2026-05-15 P105: 加 kind/mode 让 ledger 能记录, 后续 tier2b 看到不一致提醒主线程
            "kind": kind,
            "mode": mode,
        }
        if _report_verdict:
            _result["report_verdict"] = _report_verdict
        if copy_stats.get("capped"):
            _result["ok"] = False
            _result["terminal_reason"] = "copyback_blocked"
            _result["copyback_blocked"] = True
            if copy_stats.get("error_kind"):
                _result["error_kind"] = copy_stats.get("error_kind")
            if copy_stats.get("recovery_hint"):
                _result["report"] += "\n\n" + str(copy_stats.get("recovery_hint"))
                _result["retry_instruction"] = str(copy_stats.get("recovery_hint"))
            if copy_stats.get("suggested_next_action"):
                _result["next_action"] = copy_stats.get("suggested_next_action")
        if _env_available_files:
            _result["main_available_files"] = _env_available_files[:50]
            if kind in {"read", "ocr"}:
                _result["internal_evidence_files"] = sorted(set(
                    _read_internal_evidence_files + _env_available_files
                ))[:50]
            _result["post_helper_usage_hint"] = (
                "environment files have already been merged into the main workspace. "
                "Use/read/run these _env/... paths directly; do not copy from or refer to "
                ".temp/_delegate_* helper workspaces. For read helpers, these paths are "
                "internal evidence for the main process, not user-facing deliverables.\n\n"
                "环境文件已合并到主工作区；直接使用 _env 路径，勿引用 helper 临时目录。"
            )
        elif kind in {"read", "ocr"} and _read_internal_evidence_files:
            _result["internal_evidence_files"] = _read_internal_evidence_files[:50]
            _result["post_helper_usage_hint"] = (
                "Read-helper evidence files have been merged into the main workspace. "
                "Use these paths for synthesis and verification; they are internal evidence, "
                "not user-facing deliverables.\n\n"
                "读取/OCR 证据文件已在主工作区；用于综合与验收，不作为用户交付物。"
            )
        if kind in {"read", "ocr"}:
            _result["read_evidence_summary"] = _build_read_evidence_summary(
                main_workspace=main_workspace,
                task_id=task_id,
                evidence_files=_result.get("internal_evidence_files") or _read_internal_evidence_files,
                report=_full_report,
            )
        if interrupted:
            _result["interrupted"] = True
            if not copy_stats.get("shared_merge_allowed", True):
                _result["artifact_note"] = (
                    "The helper was interrupted. Regular artifacts were copied, but `_helpers_shared/` handoff files "
                    "were not merged. If those files are useful, resume the same task_id and let the helper organize "
                    "them into deliverable files before acceptance.\n\n"
                    "helper 被中断；共享区需同 task_id 续作整理后再交付。"
                )
        if was_stuck:
            _result["stuck"] = True
            if stuck_reason:
                _result["stuck_reason"] = stuck_reason
        if resumed_actually:
            _result["resumed"] = True
        if main_resource_request:
            _result["frozen"] = True
            _result["resource_required"] = main_resource_request
            if _pending_declared_outputs:
                _result["pending_declared_outputs"] = _pending_declared_outputs
            _result["suggested_retry_kind"] = main_resource_request.get("suggested_helper_kind")
            _result["suggested_retry_mode"] = "easy"

        # 2026-05-11 P2.1: 加 terminal_reason 字段 (参照 Claude Code transitions.ts)
        # 现有 boolean(ok/interrupted/stuck) 保留向后兼容; terminal_reason 是新增的
        # 权威单一 reason 字段, 主线程/日志/decision 优先看它, 不必拼接 boolean。
        # 枚举值含义:
        #   completed   = helper 自然结束 (LLM 不再调工具 → 出最终报告)
        #   interrupted = 主线程 kill / abort_event 触发 / wait_window 协作中断
        #   stuck       = StuckDetector 命中 (反复同错/同工具失败)
        #   crashed     = helper exception / 无法恢复
        #   api_stalled = API 60s 无 chunk → kill 重试
        # 注: 这是 helper "终态" enum, 跟 helper kind/task_id/elapsed 一起决定主线程下一步。
        if main_resource_request:
            _result["terminal_reason"] = "resource_required"
        elif _report_verdict in ("FAIL", "PARTIAL"):
            _result["terminal_reason"] = "failed"
        elif was_stuck:
            _result["terminal_reason"] = "stuck"
        elif interrupted:
            _result["terminal_reason"] = "interrupted"
        else:
            _result["terminal_reason"] = "completed"

        # 2026-05-11 Tier 1.G: delivered_summary 自动分类
        # helper 完成时按扩展名分组, 主线程一眼看出"给了什么类型的产物"
        # 不依赖 LLM 扫文件名列表
        _summary = {"code": [], "data": [], "docs": [], "binaries": [],
                    "images": [], "config": [], "other": []}
        for _f in ([] if kind == "ocr" else _user_visible_copied):
            _base = os.path.basename(_f).lower()
            _ext = _base.rsplit(".", 1)[-1] if "." in _base else ""
            if _ext in ("c", "cpp", "cc", "h", "hpp", "py", "js", "ts", "go", "rs", "java"):
                _summary["code"].append(_f)
            elif _ext in ("csv", "json", "tsv", "xml", "yaml", "yml", "parquet"):
                _summary["data"].append(_f)
            elif _ext in ("docx", "pdf", "md", "txt", "rst", "tex"):
                _summary["docs"].append(_f)
            elif _ext in ("exe", "dll", "so", "dylib", "o", "a", "lib"):
                _summary["binaries"].append(_f)
            elif _ext in ("png", "jpg", "jpeg", "gif", "svg", "bmp", "webp"):
                _summary["images"].append(_f)
            elif _ext in ("ini", "toml", "cfg", "conf") or _base.startswith("makefile"):
                _summary["config"].append(_f)
            else:
                _summary["other"].append(_f)
        # 只保留非空类
        _result["delivered_summary"] = {k: v for k, v in _summary.items() if v}

        # 2026-05-11 Tier 1.C: expected_outputs 系统比对
        # 主线程 spawn 时声明 ["abpt.c", "results_abpt.csv"], 完成时系统验收。
        # outputs_complete=true 意味着齐了, 主线程**绝对不应该再 resume**。
        # 2026-06-04 工程不变量: read/ocr helper 是内部 evidence 提取角色,
        # 永不产出用户可见 deliverable。这是工作流定位, 不接受 prompt 覆盖 ——
        # 用户没法决定内部 helper 的角色边界。expected_outputs 要走 edit/code/draw。
        if kind == "ocr" and expected_outputs:
            _result["outputs_check"] = {
                "expected": expected_outputs,
                "delivered_count": 0,
                "outputs_missing": list(expected_outputs),
                "outputs_complete": False,
                "quality_warnings": [{
                    "file": "<ocr-helper-output>",
                    "issue": "ocr_report_internal_only",
                    "details": (
                        "Read-helper `.txt` files are internal recognition/extraction evidence, not user-facing deliverables. "
                        "The main process should read and synthesize them, or delegate edit work for a Word/formal document.\n\n"
                        "read helper 的 txt 是内部证据，不是用户交付物。"
                    ),
                }],
            }
        elif expected_outputs:
            # 匹配方式: basename exact match OR endswith match (允许 helper_xxx_ 前缀)
            _delivered_basenames = {os.path.basename(f) for f in _visible_copied_back}
            _delivered_paths = set(_visible_copied_back)
            _all_copied_paths = {
                str(f or "").replace("\\", "/").strip()
                for f in (copied_back or [])
                if str(f or "").strip()
            }
            for _fm in (_file_map or []):
                if not isinstance(_fm, dict):
                    continue
                for _key in ("main_name", "shared_name"):
                    _mapped = str(_fm.get(_key) or "").replace("\\", "/").strip()
                    if _mapped:
                        _all_copied_paths.add(_mapped)
            _env_copied_paths = {
                str(path or "").replace("\\", "/").strip()
                for path in (copy_stats.get("env_copied_files") or [])
                if str(path or "").strip()
            }
            _missing = []
            _matched_files: list[tuple[str, str]] = []  # (expected, actual_path)
            _env_matched_files: list[tuple[str, str]] = []
            _shared_protocol_matches: list[dict[str, str]] = []

            def _matches_expected_output_path(actual_path: str, expected_norm: str, expected_base: str) -> bool:
                actual_norm = str(actual_path or "").replace("\\", "/").strip()
                if not actual_norm:
                    return False
                expected_dir = expected_norm.rstrip("/") + "/" if expected_norm.endswith("/") else ""
                if expected_dir:
                    if actual_norm.startswith(expected_dir) and actual_norm != expected_dir.rstrip("/"):
                        return True
                    if expected_dir.startswith("_shared/"):
                        shared_rel = expected_dir[len("_shared/"):]
                        writable_shared_dir = f"_helpers_shared/{shared_rel}"
                        if actual_norm.startswith(writable_shared_dir) and actual_norm != writable_shared_dir.rstrip("/"):
                            return True
                    return False
                actual_base = os.path.basename(actual_norm)
                if (
                    actual_norm == expected_norm
                    or actual_base == expected_base
                    or actual_norm.endswith(expected_norm)
                    or actual_norm.endswith("_" + expected_base)
                ):
                    return True
                if expected_norm.startswith("_shared/"):
                    shared_rel = expected_norm[len("_shared/"):]
                    writable_shared = f"_helpers_shared/{shared_rel}"
                    if actual_norm == writable_shared:
                        return True
                return False

            for _exp in expected_outputs:
                _exp_clean = _exp.strip().lstrip("./").lstrip("\\")
                _exp_norm = _exp_clean.replace("\\", "/")
                _exp_is_dir = _exp_norm.endswith("/")
                _exp_base = os.path.basename(_exp_norm.rstrip("/")) if _exp_is_dir else os.path.basename(_exp_clean)
                _found_path = None
                if _exp_is_dir:
                    _found_path = next(
                        (
                            f for f in sorted(_all_copied_paths)
                            if _matches_expected_output_path(f, _exp_norm, _exp_base)
                        ),
                        None,
                    )
                    if _found_path and _exp_norm.startswith("_shared/") and _found_path.startswith("_helpers_shared/"):
                        _shared_protocol_matches.append({
                            "expected_readonly": _exp_norm,
                            "delivered_writable_shared": os.path.dirname(_found_path).replace("\\", "/") + "/",
                        })
                elif _exp_norm == "_env" or _exp_norm.startswith("_env/"):
                    _env_abs = os.path.join(main_workspace, *_exp_norm.split("/"))
                    if _exp_norm in _env_copied_paths and os.path.isfile(_env_abs):
                        _env_matched_files.append((_exp, _exp_norm))
                        continue
                else:
                    _env_equivalent = f"_env/{_exp_norm}"
                    _env_abs = os.path.join(main_workspace, *_env_equivalent.split("/"))
                    if _env_equivalent in _env_copied_paths and os.path.isfile(_env_abs):
                        _env_matched_files.append((_exp, _env_equivalent))
                        continue
                if _found_path:
                    pass
                elif _exp_clean in _delivered_paths:
                    _found_path = _exp_clean
                elif _exp_base in _delivered_basenames:
                    _found_path = next(
                        (f for f in _visible_copied_back if os.path.basename(f) == _exp_base),
                        None
                    )
                else:
                    _found_path = next(
                        (f for f in _visible_copied_back
                         if f.endswith(_exp_clean) or f.endswith(_exp_base)
                         or f.endswith(_exp_norm) or f.endswith("_" + _exp_base)),
                        None
                    )
                if not _found_path:
                    _found_path = next(
                        (
                            f for f in sorted(_all_copied_paths)
                            if _matches_expected_output_path(f, _exp_norm, _exp_base)
                        ),
                        None,
                    )
                    if _found_path and _exp_norm.startswith("_shared/") and _found_path.startswith("_helpers_shared/"):
                        _shared_protocol_matches.append({
                            "expected_readonly": _exp_norm,
                            "delivered_writable_shared": _found_path,
                        })
                if _found_path:
                    _matched_files.append((_exp, _found_path))
                else:
                    _missing.append(_exp)
            if _env_matched_files:
                _matched_files.extend(_env_matched_files)

            # 2026-05-11 P11: 内容质量验收 (quality_warnings)
            # outputs_complete=true 只表示文件存在, 不验内容。
            # 这里对 PNG/CSV/docx 等做基础 sanity 检查, 发现"垃圾文件"给主线程警告。
            # 不阻断 outputs_complete(向后兼容), 但加 quality_warnings 字段。
            _quality_warnings: list[dict] = list(_pre_output_quality_warnings)
            for _exp, _actual_path in _matched_files:
                try:
                    _abs_path = (
                        _actual_path if os.path.isabs(_actual_path)
                        else os.path.join(main_workspace, _actual_path)
                    )
                    if not os.path.exists(_abs_path):
                        continue
                    _size = os.path.getsize(_abs_path)
                    _ext = _abs_path.lower().rsplit(".", 1)[-1] if "." in _abs_path else ""
                    if _ext in {"txt", "md", "markdown", "csv", "tsv", "json", "jsonl", "yaml", "yml", "xml"} and _size <= 2 * 1024 * 1024:
                        try:
                            with open(_abs_path, "r", encoding="utf-8", errors="replace") as _txt_f:
                                _text_for_encoding_check = _txt_f.read(120000)
                            for _cw in text_mojibake_warnings(_text_for_encoding_check, file=_exp):
                                _quality_warnings.append(_cw)
                        except OSError:
                            pass

                    # PNG/JPG: 太小可能是空图
                    if _ext in ("png", "jpg", "jpeg"):
                        if _size < 2000:  # < 2 KB 通常是空 PNG
                            _quality_warnings.append({
                                "file": _exp,
                                "issue": "image_too_small",
                                "size_bytes": _size,
                                "details": (
                                    f"PNG/JPG file is only {_size} bytes; it may be blank or failed to save.\n\n"
                                    "图片文件过小，可能为空图或保存失败。"
                                ),
                            })
                    # CSV/TSV: 太小可能是空数据
                    elif _ext in ("csv", "tsv"):
                        if _size < 100:  # < 100 B 至少没数据
                            _quality_warnings.append({
                                "file": _exp,
                                "issue": "data_file_empty",
                                "size_bytes": _size,
                                "details": (
                                    f"CSV/TSV file is only {_size} bytes; it may contain only a header or be empty.\n\n"
                                    "数据文件过小，可能只有表头或为空。"
                                ),
                            })
                        else:
                            # 数行数
                            try:
                                with open(_abs_path, "rb") as _f:
                                    _n_lines = sum(1 for _ in _f)
                                if _n_lines < 2:  # 至少 header + 1 行
                                    _quality_warnings.append({
                                        "file": _exp,
                                        "issue": "data_file_no_rows",
                                        "lines": _n_lines,
                                        "details": (
                                            f"CSV has only {_n_lines} line(s), so it may contain no data rows.\n\n"
                                            "CSV 行数过少，可能没有数据。"
                                        ),
                                    })
                            except OSError:
                                pass
                    # docx/pptx/xlsx: < 5 KB 通常是空文档
                    elif _ext in ("docx", "pptx", "xlsx"):
                        if _size < 5000:
                            _quality_warnings.append({
                                "file": _exp,
                                "issue": "document_too_small",
                                "size_bytes": _size,
                                "details": (
                                    f"{_ext} file is only {_size} bytes; it may be an empty template.\n\n"
                                    "文档文件过小，可能是空模板。"
                                ),
                            })
                        _doc_text_for_grounding = ""
                        if _ext == "pptx":
                            expected_slide_groups = _extract_expected_ppt_slide_token_groups(prompt)
                            if expected_slide_groups:
                                slide_texts = _pptx_slide_texts_for_quality_check(_abs_path)
                                _doc_text_for_grounding = "\n".join(slide_texts)
                                slide_missing: list[dict] = []
                                for slide_no, tokens in expected_slide_groups.items():
                                    if slide_no < 1 or slide_no > len(slide_texts):
                                        slide_missing.append({
                                            "slide": slide_no,
                                            "missing": tokens,
                                            "reason": "slide_absent",
                                        })
                                        continue
                                    slide_norm = _normalize_doc_text_for_match(slide_texts[slide_no - 1])
                                    missing_on_slide: list[str] = []
                                    for token in tokens:
                                        token_norm = _normalize_doc_text_for_match(token)
                                        if not token_norm:
                                            continue
                                        if "=" in token_norm:
                                            left, right = token_norm.split("=", 1)
                                            found = bool(left and right and left in slide_norm and right in slide_norm)
                                        else:
                                            found = token_norm in slide_norm
                                        if not found:
                                            missing_on_slide.append(token)
                                    if missing_on_slide:
                                        slide_missing.append({
                                            "slide": slide_no,
                                            "missing": missing_on_slide[:10],
                                        })
                                if slide_missing:
                                    _quality_warnings.append({
                                        "file": _exp,
                                        "issue": "pptx_expected_slide_order_mismatch",
                                        "missing_by_slide": slide_missing[:10],
                                        "details": (
                                            "Requested PPT content does not appear on the specified slide numbers. "
                                            "The slide order may be wrong; reorder or rebuild according to the requested page numbers before delivery.\n\n"
                                            "PPT 页序可能错误，需按指定页码重排或重建。"
                                        ),
                                    })
                        expected_text_tokens = _extract_expected_text_tokens_for_document(prompt)
                        if not _doc_text_for_grounding:
                            _doc_text_for_grounding = _document_text_for_quality_check(_abs_path, _ext)
                        if expected_text_tokens:
                            actual_text_norm = _normalize_doc_text_for_match(
                                _doc_text_for_grounding
                            )
                            missing_tokens: list[str] = []
                            for token in expected_text_tokens:
                                token_norm = _normalize_doc_text_for_match(token)
                                if not token_norm:
                                    continue
                                if "=" in token_norm:
                                    left, right = token_norm.split("=", 1)
                                    found = bool(left and right and left in actual_text_norm and right in actual_text_norm)
                                else:
                                    found = token_norm in actual_text_norm
                                if not found:
                                    missing_tokens.append(token)
                            if missing_tokens:
                                _quality_warnings.append({
                                    "file": _exp,
                                    "issue": "document_expected_text_missing",
                                    "missing": missing_tokens[:10],
                                    "details": (
                                        "The Office artifact is missing explicit text or numbers requested by the user. "
                                        "Read the produced file and repair it before delivery instead of relying only on the helper report.\n\n"
                                        "Office 产物缺少用户明确要求的文字或数字。"
                                    ),
                                })
                        for _cw in document_source_grounding_warnings(_doc_text_for_grounding):
                            _cw["file"] = _exp
                            _quality_warnings.append(_cw)
                        if _ext == "docx":
                            for _cw in academic_document_warnings(_doc_text_for_grounding):
                                _cw["file"] = _exp
                                _quality_warnings.append(_cw)
                            for _cw in docx_table_structure_warnings(_abs_path):
                                _cw["file"] = _exp
                                _quality_warnings.append(_cw)
                    # 代码文件: 过小
                    elif _ext in ("c", "cpp", "py", "go", "rs", "java"):
                        if _size < 100:
                            _quality_warnings.append({
                                "file": _exp,
                                "issue": "code_file_too_small",
                                "size_bytes": _size,
                                "details": (
                                    f"Code file `{_ext}` is only {_size} bytes; it may be empty or incomplete.\n\n"
                                    "代码文件过小，可能为空或未完成。"
                                ),
                            })

                    # 2026-05-11 P12.H: 深度内容验证(PNG 能 open / CSV 列合理)
                    # 不只看大小, 用对应 parser 验证文件实际可用
                    if _ext in ("png", "jpg", "jpeg"):
                        # PNG 头部 magic: \x89PNG\r\n\x1a\n
                        try:
                            with open(_abs_path, "rb") as _img_f:
                                _head = _img_f.read(16)
                            _is_valid_png = _head.startswith(b"\x89PNG\r\n\x1a\n")
                            _is_valid_jpg = _head[:2] == b"\xff\xd8"
                            if _ext == "png" and not _is_valid_png:
                                _quality_warnings.append({
                                    "file": _exp,
                                    "issue": "png_invalid_header",
                                    "details": (
                                        "PNG magic bytes do not match; the file may be corrupt.\n\n"
                                        "PNG 头部不匹配，文件可能损坏。"
                                    ),
                                })
                            elif _ext in ("jpg", "jpeg") and not _is_valid_jpg:
                                _quality_warnings.append({
                                    "file": _exp,
                                    "issue": "jpg_invalid_header",
                                    "details": (
                                        "JPG magic bytes do not match; the file may be corrupt.\n\n"
                                        "JPG 头部不匹配，文件可能损坏。"
                                    ),
                                })
                        except OSError:
                            pass

                    if _ext in ("csv", "tsv") and _size >= 100:
                        # CSV 验证: 至少 1 行 header + 1 行 data, 且每行列数一致
                        try:
                            with open(_abs_path, "r", encoding="utf-8-sig", errors="replace") as _csv_f:
                                _first_line = _csv_f.readline().strip()
                                _second_line = _csv_f.readline().strip()
                                if _first_line and _second_line:
                                    _delim = "," if "," in _first_line else "\t"
                                    _header_cols = _first_line.count(_delim) + 1
                                    _data_cols = _second_line.count(_delim) + 1
                                    if abs(_header_cols - _data_cols) > 1:
                                        _quality_warnings.append({
                                            "file": _exp,
                                            "issue": "csv_column_mismatch",
                                            "details": (
                                                f"CSV header has {_header_cols} column(s), but the first data row has "
                                                f"{_data_cols} column(s); the schema is inconsistent.\n\n"
                                                "CSV 表头与首行数据列数不一致。"
                                            )
                                        })
                        except OSError:
                            pass

                    # 2026-05-11 P12.F: benchmark 数据公平性 sanity check
                    # 检测某算法在某操作上是否比同类中位数快 100x+ (典型 tombstone 虚假快)
                    # 病因(实测 18:30 ABPT batch_delete tombstone): ABPT=6666 Mops/s,
                    # 其他算法=5-25 Mops/s, 差 332x → 论文里的图表都是这种虚假数据。
                    # 触发条件: CSV 文件 + 列含 algorithm/operation/throughput_mops 类
                    if _ext in ("csv", "tsv") and (
                        "result" in _exp.lower()
                        or "bench" in _exp.lower()
                        or "throughput" in _exp.lower()
                    ):
                        # 2026-05-21: 复杂度异常检测(单文件即可查 range_query 等随 N 平方增长)。
                        # 病因(trace c6e42ed6"1100 倍"): rbtree range_query O(n^2) 实现 → 论文夸张失真。
                        try:
                            for _cw in _detect_benchmark_complexity_anomaly(_abs_path):
                                _cw["file"] = _exp
                                _quality_warnings.append(_cw)
                        except Exception:
                            pass
                        # 2026-05-22: CSV schema 校验 + 内存/计时质量检测(交付前报警不丢弃)
                        try:
                            for _cw in _detect_benchmark_csv_schema_issues(_abs_path):
                                _cw["file"] = _exp
                                _quality_warnings.append(_cw)
                            for _cw in _detect_memory_non_monotonic(_abs_path):
                                _cw["file"] = _exp
                                _quality_warnings.append(_cw)
                            for _cw in _detect_timing_precision_loss(_abs_path):
                                _cw["file"] = _exp
                                _quality_warnings.append(_cw)
                        except Exception:
                            pass
                        try:
                            import csv as _csv_m
                            with open(_abs_path, "r", encoding="utf-8-sig", errors="replace") as _csv_f:
                                _reader = _csv_m.DictReader(_csv_f)
                                _rows = list(_reader)
                            if _rows and len(_rows) > 5:
                                # 找 throughput 列(常见命名: throughput_mops / throughput / tp)
                                _tp_col = None
                                for _c in _rows[0].keys():
                                    if _c and _c.lower().replace(' ','').replace('_','') in (
                                        "throughputmops", "throughput", "tp", "mops", "opsps"
                                    ):
                                        _tp_col = _c
                                        break
                                _algo_col = None
                                _op_col = None
                                for _c in _rows[0].keys():
                                    if _c and _c.lower() in ("algorithm", "algo", "algoname"):
                                        _algo_col = _c
                                    elif _c and _c.lower() in ("operation", "op", "opname"):
                                        _op_col = _c
                                if _tp_col and _algo_col and _op_col:
                                    # 按 (operation,) 分组, 找每个操作下最快算法的倍数
                                    from collections import defaultdict as _dd
                                    _by_op = _dd(lambda: _dd(list))  # op → algo → [tp...]
                                    for _r in _rows:
                                        try:
                                            _tp = float(_r.get(_tp_col, 0) or 0)
                                            _algo = _r.get(_algo_col, "")
                                            _op = _r.get(_op_col, "")
                                            if _tp > 0 and _algo and _op:
                                                _by_op[_op][_algo].append(_tp)
                                        except (ValueError, TypeError):
                                            continue

                                    # 对每个 operation, 看是否有算法超过中位数 100x
                                    for _op, _algo_tps in _by_op.items():
                                        if len(_algo_tps) < 3:  # 至少 3 个算法才能比
                                            continue
                                        # 取每算法平均 tp
                                        _algo_avg = {a: sum(v)/len(v) for a, v in _algo_tps.items()}
                                        _values = sorted(_algo_avg.values())
                                        _median = _values[len(_values)//2]
                                        _max_algo = max(_algo_avg, key=_algo_avg.get)
                                        _max_val = _algo_avg[_max_algo]
                                        if _median > 0 and _max_val / _median >= 100:
                                            _quality_warnings.append({
                                                "file": _exp,
                                                "issue": "benchmark_unfair_ratio",
                                                "details": (
                                                    f"Algorithm `{_max_algo}` has throughput={_max_val:.1f} for operation `{_op}`, "
                                                    f"while the median for other algorithms is {_median:.1f}; the ratio is "
                                                    f"{_max_val/_median:.0f}x. This is likely a measurement or implementation issue "
                                                    f"such as lazy deletion, inaccurate timing, or an incorrect benchmark.\n\n"
                                                    "性能数据异常偏离，需复核计时、实现和测试口径。"
                                                ),
                                                "ratio_vs_median": round(_max_val / _median, 1),
                                                "outlier_algorithm": _max_algo,
                                                "operation": _op,
                                            })
                        except (OSError, csv.Error, ImportError):
                            pass
                except OSError as _e:
                    _quality_warnings.append({
                        "file": _exp,
                        "issue": "stat_failed",
                        "details": (
                            f"Could not stat the file: {_e}\n\n"
                            "无法读取文件元数据。"
                        ),
                    })

            # 2026-05-11 P14.G: 空报告 / 短产出 helper 检测
            # 病因(实测 gen_paper 第 3 次 21:26): report=178 chars 通过验收, 但实际
            # helper 可能没真做事就 ok=true。系统层检测 report 过短 + 文件数过少 →
            # 警告主线程"这可能是假完成"。
            #
            # 2026-05-15 P74 修正(实测 sin_plot trace 20:48-20:50):
            # 病因: "画 sin 函数图" 这种简单单文件 draw 任务, helper 产出 1 个 PNG +
            #   113 字符简洁报告 (符合任务规模), 但被误触发 suspicious_short_completion
            #   → 主线程做 inspect/commit/copy/commit/locate 5 次工具调用 + 3 次 LLM
            #   round-trip 浪费 30s。
            # 修法: 收紧触发条件:
            #   (a) 文件数 == 0 (声称完成但没产物 — 真可疑)
            #   (b) 多 deliverable 任务 (expected_outputs >= 2) 但报告极短 (< 200 字)
            #       — 多产物任务通常需要解释每件做了什么
            #   (c) 多文件 (>= 2) 但报告 < 100 字 — 没说每件做了什么
            #   不触发: 单 expected_output + 1 文件 + 报告 50-300 字 (简单任务正常情况)
            _report_chars = len(report or "")
            _delivered_output_count = len(copied_back) + len(_env_matched_files)
            _n_expected = len(expected_outputs or [])
            _output_evidence_complete = (
                _n_expected > 0
                and not _missing
                and _delivered_output_count >= _n_expected
                and not main_resource_request
            )
            _suspicious = False
            _suspicious_reason = ""
            if _delivered_output_count == 0 and expected_outputs:
                _suspicious = True
                _suspicious_reason = (
                    "The helper claimed completion but delivered zero files.\n\n"
                    "helper 声称完成但没有交付文件。"
                )
            elif not _output_evidence_complete and _n_expected >= 2 and _report_chars < 200:
                _suspicious = True
                _suspicious_reason = (
                    f"Multi-output task expected {_n_expected} files, but the report is very short "
                    f"({_report_chars} chars) and does not explain each output.\n\n"
                    f"多产物任务报告过短，未说明每个产物。"
                )
            elif not _output_evidence_complete and _delivered_output_count >= 2 and _report_chars < 100:
                _suspicious = True
                _suspicious_reason = (
                    f"The helper delivered {_delivered_output_count} files, but the report is very short "
                    f"({_report_chars} chars) and does not explain the files.\n\n"
                    f"多文件交付报告过短，未说明文件用途。"
                )
            if _suspicious:
                _quality_warnings.append({
                    "file": "<helper-report>",
                    "issue": "suspicious_short_completion",
                    "report_chars": _report_chars,
                    "delivered_files": _delivered_output_count,
                    "details": _suspicious_reason,
                    "suggestion": (
                        "Inspect ledger.delivered_summary and the produced files against the original acceptance "
                        "points. If the same deliverable is incomplete or low quality, resume the same task_id with "
                        "resume=true, keep the base kind, and upgrade mode when needed. Use a verify helper only to "
                        "check an existing artifact; use a new task_id only for a genuinely different work boundary.\n"
                        "主线程按验收点检查产物；同一交付物修复用同 task_id + resume，必要时升 hard。"
                    ),
                })

            # 2026-05-11 P14.K: docx 深度内容验证
            # 病因(论文场景): outputs_check 只看大小 / PNG magic bytes, 不验段落数。
            # paper.docx 可能 500 KB 但只有 3 段 (helper 没写够内容)。
            for _exp, _actual_path in _matched_files:
                try:
                    if not _exp.lower().endswith(".docx"):
                        continue
                    _abs_path = (
                        _actual_path if os.path.isabs(_actual_path)
                        else os.path.join(main_workspace, _actual_path)
                    )
                    if not os.path.exists(_abs_path):
                        continue
                    from docx import Document as _Doc
                    _doc = _Doc(_abs_path)
                    _n_para = len(_doc.paragraphs)
                    # 非空段落数
                    _n_nonempty = sum(1 for p in _doc.paragraphs if (p.text or "").strip())
                    _total_chars = sum(len(p.text or "") for p in _doc.paragraphs)
                    # 阈值: 期望至少 5 个非空段落 + 200 字符
                    if _n_nonempty < 5:
                        _quality_warnings.append({
                            "file": _exp,
                            "issue": "docx_too_few_paragraphs",
                            "nonempty_paragraphs": _n_nonempty,
                            "total_chars": _total_chars,
                            "details": (
                                f"DOCX has only {_n_nonempty} non-empty paragraph(s), expected at least 5. "
                                f"The helper may not have completed the content.\n\n"
                                f"DOCX 段落过少，可能未写完。"
                            ),
                        })
                    elif _total_chars < 200:
                        _quality_warnings.append({
                            "file": _exp,
                            "issue": "docx_too_few_chars",
                            "total_chars": _total_chars,
                            "details": (
                                f"DOCX has {_n_nonempty} paragraph(s) but only {_total_chars} total characters, "
                                f"expected at least 200. The content may be placeholder or skeleton text.\n\n"
                                f"DOCX 字符数过少，可能只是占位或骨架。"
                            ),
                        })
                except (ImportError, Exception):  # docx 不可用 / 解析失败 静默
                    pass

            if main_resource_request:
                _quality_warnings.append({
                    "file": "<helper>",
                    "issue": "requires_main_resource",
                    "severity": "blocking",
                    "details": (
                        f"Helper is frozen and needs a main-process resource of kind "
                        f"{main_resource_request.get('suggested_helper_kind')}: "
                        f"{main_resource_request.get('main_thread_action')}\n\n"
                        f"helper 冻结，等待主进程提供资源。"
                    ),
                    "resource_required": main_resource_request,
                })

            _result["outputs_check"] = {
                "expected": expected_outputs,
                "delivered_count": len({
                    actual for _, actual in (_matched_files + _env_matched_files)
                    if str(actual or "").strip()
                }),
                "outputs_missing": _missing,
                "outputs_complete": (not _missing and not main_resource_request),
                # P11: quality_warnings 列表(空表示全部 OK)
                "quality_warnings": _quality_warnings,
            }
            if _shared_protocol_matches:
                _result["outputs_check"]["shared_protocol_matches"] = _shared_protocol_matches
                _result["post_helper_usage_hint"] = (
                    "Some expected outputs named `_shared/...`, which is read-only for helpers. "
                    "The helper correctly delivered matching files under `_helpers_shared/...`; "
                    "treat these as shared implementation candidates and integrate or verify them from the main process."
                    "\nhelper 对只读 `_shared/...` 的期望已由可写 `_helpers_shared/...` 产物满足；主线程负责整合和验证。"
                )
            elif _shared_visible_delivered and not _visible_copied_back and not _env_available_files:
                _result["post_helper_usage_hint"] = (
                    "This helper produced only `_helpers_shared/...` handoff files. They are available to the main "
                    "process and downstream helpers, but they are not user-facing deliverables and are not real "
                    "project files. Continue the chain: inspect or collect the shared files, assemble the requested "
                    "final artifact, and in environment mode apply verified `_env/...` or project-relative files "
                    "to the real project before final acceptance.\n\n"
                    "`_helpers_shared/...` 只是协作证据；继续汇总、生成最终产物，并在项目模式落到真实项目路径后再验收。"
                )
            if not _result["outputs_check"].get("outputs_complete"):
                _result["ok"] = False
                if _result.get("terminal_reason") == "completed":
                    _result["terminal_reason"] = "outputs_missing" if _missing else "resource_required"
                _result["outputs_check"]["outputs_complete"] = False
                if _missing and not _result["outputs_check"].get("outputs_missing"):
                    _result["outputs_check"]["outputs_missing"] = _missing
            _blocking_quality_warnings = blocking_quality_warnings(_quality_warnings)
            if _blocking_quality_warnings and _result.get("terminal_reason") == "completed":
                _result["ok"] = False
                _result["terminal_reason"] = "quality_blocked"
                _result["outputs_check"]["quality_blocked"] = True
                _result["outputs_check"]["blocking_quality_warnings"] = _blocking_quality_warnings[:10]
                _result["quality_blocked"] = True
                _result["next_action"] = {
                    "type": "resume_upgraded",
                    "rationale": (
                        "The helper produced files, but system quality checks found blocking issues. "
                        "Resume from the same workspace, inspect the existing artifacts and source evidence, "
                        "repair the listed issues, then rerun the relevant verification."
                    ),
                    "params": {
                        "action": "spawn",
                        "task_id": task_id,
                        "resume": True,
                        "kind": kind,
                        "mode": "hard" if mode != "hard" else mode,
                        "prompt": (
                            "Continue the same task from the preserved workspace. Fix the blocking "
                            "outputs_check.quality_warnings, especially missing requested text/numbers, "
                            "internal source labels, malformed data, or incomplete document structure. "
                            "Read the actual artifact and source evidence before reporting success."
                        ),
                    },
                }
                _result["retry_instruction"] = (
                    "This helper is not cleanly complete: outputs_check.quality_warnings contains "
                    "blocking issues. Resume the same task_id with resume=true, repair the existing "
                    "artifact from verified source evidence, and verify again before final delivery.\n\n"
                    "helper 未干净完成；同 task_id 恢复修复并重新验证。"
                )

            _terminal_converged = _converge_terminal_state_for_complete_outputs(
                _result,
                expected_outputs=expected_outputs,
                kind=kind,
                main_resource_request=main_resource_request,
            )

            # 2026-05-15 P78: 干净完成时给主线程明确"跳过验证, 直接出 JSON"信号。
            # 病因(实测 Task 2 sin_plot trace 124s 中 30s 浪费):
            #   helper outputs_complete=true, file_map 已显示文件在主区, quality_warnings 空,
            #   但主线程仍做 inspect_file + commit_to_main + run copy + commit_to_main + workspace.locate
            #   5 个工具调用 + 5 次 LLM round-trip。TEMPLATE 5.7 已写"快路径", 但 LLM 不一定看。
            # 修法: 系统检测到全清后, 在 result 顶层加 _post_helper_action="output_json_directly"
            #   + _verification_skipped 文字, LLM 看到直接知道不需要验证。
            _no_outputs_missing = not _missing
            _no_warnings = not _quality_warnings
            _has_expected = bool(expected_outputs) and kind not in {"read", "ocr"}
            _all_files_promoted = bool(_visible_copied_back or _env_available_files)
            _has_office_output = any(
                str(path).lower().endswith((".docx", ".pptx", ".xlsx"))
                for path in _visible_copied_back
            )
            if (_no_outputs_missing and _no_warnings and _has_expected
                    and _all_files_promoted and (not interrupted and not was_stuck or _terminal_converged)
                    and not main_resource_request and not _has_office_output):
                _result["_post_helper_action"] = "output_json_directly"
                # 2026-06-05 简化: 之前的提示太长 (~280 chars * 14 helpers = 4K context),
                # 而且每个 helper 完成都重复;改为 1 行核心信息。
                _result["_post_helper_hint"] = (
                    "Outputs complete. Put the clean filename in plan.deliverables "
                    "(pick one final if multiple candidates exist).\n"
                    "产物完整；deliverables 写干净文件名,多候选选最终版。"
                )
                if _env_available_files:
                    _result["_post_helper_hint"] += (
                        " _env/ paths are in main workspace; use those directly.\n"
                        "_env 产物已在主区,直接用。"
                    )

        # 2026-05-10 Patch 84: 对 interrupted / stuck / 长跑 helper 加 retry hint
        _is_long_running = _elapsed > 1200.0  # 20 min
        if (interrupted or was_stuck or _is_long_running) and not _result.get("_terminal_converged_from"):
            # 2026-05-11 fix: _trigger 优先级跟 terminal_reason 对齐
            # (was_stuck 优先于 interrupted, 因为 StuckDetector 先触发,
            # interrupted=True 是 stuck 后 abort 流程的副产物)
            # 旧版本 interrupted 优先 → 日志显示 terminal_reason="stuck"
            # 但 retry_instruction 文字写 "interrupted",LLM 困惑。
            _trigger = ("stuck" if was_stuck else
                        "interrupted" if interrupted else
                        "long_running")
            _result["suggested_retry_kind"] = kind
            _kind_for_retry = str(kind or "").lower()
            _code_hard_retry = was_stuck and _kind_for_retry in {"code", "coding"}
            _same_kind_hard_retry = bool(was_stuck)
            _retry_mode = "hard" if _same_kind_hard_retry else mode
            _result["suggested_retry_mode"] = _retry_mode

            # 2026-05-15 P104: 智能 prompt_hint - 根据 stuck_reason 给具体重试指引
            # 病因(实测 trace 16:24 压缩论文): 24 次 hard mode 重试, prompt_hint
            # 仅 "Copy original prompt; optionally prepend '已知失败原因: stuck after 400s'"。
            # 新 helper 拿同 prompt 容易卡同 bug。
            # 修法: 解析 stuck_reason, 生成针对性 hint, 让重试不只是"换更强模型"
            # 而是带具体新方向。
            _smart_prompt_hint = (
                f"Copy the original prompt only after adding a concise recovery note: "
                f"previous helper terminated as {_trigger} after {_elapsed:.0f}s. "
                f"State what will change in this retry.\n\n"
                f"重试提示需说明失败状态和新的处理方向。"
            )
            if was_stuck and stuck_reason:
                # 根据 stuck_reason 关键词推断具体问题, 生成针对性 hint
                _sr_lower = stuck_reason.lower()
                _retry_suggestions = []
                if "win_fatal" in _sr_lower or "heap_corruption" in _sr_lower or "p103" in _sr_lower:
                    # 2026-05-21: ASan 不可用环境(MinGW)不推 ASan,给替代手段。
                    try:
                        from app.llm.tools.workspace import has_asan as _has_asan
                        _asan_ok = _has_asan()
                    except Exception:
                        _asan_ok = False
                    if _asan_ok:
                        _retry_suggestions.append(
                            "Run a small debug build with `gcc -fsanitize=address -g -O0 your_code.c -o test && ./test`, "
                            "then use the ASan trace to locate the memory bug."
                        )
                        _retry_suggestions.append("Start with N=10 under ASan; increase input size only after the small case is clean.")
                    else:
                        _retry_suggestions.append(
                            "ASan is unavailable on this host, so avoid `-fsanitize=address`. Use UBSan when available, "
                            "assertions, boundary checks, and focused diagnostic output instead."
                        )
                        _retry_suggestions.append("Reproduce with N=10 first, add boundary assertions, then increase input size.")
                if "bash_fail_rate" in _sr_lower or "p70" in _sr_lower:
                    _retry_suggestions.append(
                        "The previous helper looped through edit/build/fail. Start with a progress_note that states the current diagnosis, "
                        "then consider rewriting the coherent module rather than repeatedly patching small spans."
                    )
                if "edit_same_file" in _sr_lower or "p69" in _sr_lower:
                    _retry_suggestions.append(
                        "The previous helper repeatedly edited the same file without passing checks. Re-read the relevant file or function first, "
                        "then rewrite the coherent function/module instead of another local patch."
                    )
                if "cumulative" in _sr_lower or "p14" in _sr_lower:
                    _retry_suggestions.append(
                        "The previous helper hit a repeated-error threshold. The retry prompt should name the known failure pattern, "
                        "require reading the error output, and locate the cause before editing."
                    )
                if "office_args_too_large" in _sr_lower:
                    _retry_suggestions.append(
                        "The previous Office call batched too many blocks. Limit each append to 2-8 blocks and split long sections across calls."
                    )
                if "long_no_delegate" in _sr_lower:
                    _retry_suggestions.append(
                        "The previous attempt did too much serial work. Split independent subtasks into helpers before fan-in."
                    )

                if _retry_suggestions:
                    _smart_prompt_hint = (
                        f"Previous helper stuck reason: {stuck_reason[:250]}\n"
                        f"The new prompt should include:\n"
                        + "\n".join(f"  - {s}" for s in _retry_suggestions)
                        + "\nUse a concrete revised direction instead of merely reusing the old prompt with a stuck label.\n\n"
                        + "重派 helper 时说明上次卡点，并给出新的具体方向。"
                    )
                else:
                    _smart_prompt_hint = (
                        f"Previous helper stuck reason: {stuck_reason[:250]}\n"
                        f"Recommended new prompt shape:\n"
                        f"  - State where the previous helper got stuck, using details from its summary or report.\n"
                        f"  - Give one concrete revised direction such as stronger verification, narrower tests, or a rewrite path.\n"
                        f"  - Preserve useful work, but make the next attempt meaningfully different.\n\n"
                        f"重派 helper 时保留已有证据，明确卡点，并换一个可产生新证据的方向。"
                    )

            # 结构化 next_action(主线程 LLM 直接复制 params 用 delegate)
            _result["next_action"] = {
                "type": "resume_upgraded",
                "rationale": (
                    f"helper terminated with '{_trigger}' after {_elapsed:.0f}s. "
                    + (
                        "Repeated code-helper failure may justify mode='hard' for the same narrow task after reading the failure evidence."
                        if _code_hard_retry else
                        "Repeated failure should first trigger root-cause repair: check kind, resources, paths, dependency order, scope, and acceptance evidence. If the same kind remains correct, mode='hard' is a stricter same-kind retry with richer evidence discipline, not new tool access."
                        if was_stuck else
                        "This may be a long healthy flow; reuse the same task_id and mode unless evidence shows a real capability or boundary problem."
                    )
                ),
                "params": {
                    "action": "spawn",
                    "task_id": task_id,
                    "resume": True,
                    "kind": kind,
                    "mode": _retry_mode,
                    "prompt_hint": _smart_prompt_hint,
                    "wait_window_sec": 600,
                },
            }
            _result["retry_instruction"] = (
                f"Helper `{task_id}` ended with `{_trigger}` after {_elapsed:.0f}s. Prefer continuing the same "
                f"task_id with resume=true after reading the failure evidence.\n"
                f"- If progress is still healthy, keep kind='{kind}' and mode='{mode}'.\n"
                f"- If it failed, diagnose kind, resources, paths, dependency order, scope, and acceptance evidence.\n"
                f"- If the same kind remains appropriate, use mode='{_retry_mode}' for a stricter continuation with "
                f"clearer evidence and verification, not a restart or a same-goal v2 task.\n"
                f"- Code hard mode strengthens implementation, debugging, build, and benchmark work. Other hard modes "
                f"strengthen evidence collection, shard verification, and reporting while keeping the tool boundary.\n"
                f"Suggested call: delegate(action='spawn', tasks=[{{'task_id': '{task_id}', 'resume': true, "
                f"'kind': '{kind}', 'mode': '{_retry_mode}', 'prompt': '<previous prompt plus revised direction>'}}], "
                f"wait_window_sec=600)\n\n"
                f"同一 helper 续作应保留已有证据，先修复具体失败点，再重新验收产物。"
            )

        # 2026-05-09 Patch 47: verify helper 完成时,解析 VERDICT 暴露给主线程
        # 病因(P45 引入):auto_verify 派出后,verify 报告里的 VERDICT 行只在
        # report 文本中,主线程模型可能漏读。修法:在 result dict 加结构化字段:
        #   - verify_verdict:解析出的 PASS / FAIL / PARTIAL(三选一,unknown 时为 None)
        #   - verify_target:被验证的父 task_id(从命名约定 {parent}_verify 反推)
        # 主线程模型在 round2 看 results 时直接看到 verify_verdict 字段,不需深读
        # report 文本就能判断。FAIL 时模型应该 spawn 修复 helper 或告诉用户产物有问题。
        #
        # 2026-05-09 Patch 48: verify FAIL/PARTIAL 时生成修复建议
        # 病因(P47 完成后留的洞):主线程拿到 verify_verdict=FAIL,但要自己**从零想**
        # 怎么写修复 prompt(读 verify 报告找 FAIL 项 + 拼成 prompt)。
        # 修法:在 verify 完成时同时生成 `repair_recommendation` 字段 — 一个建议性
        # 修复 prompt 模板,主线程**可直接复制**用于 delegate(task_id={父}, resume=true,
        # prompt=该 recommendation)。**建议性**:主线程可以无视(误报场景)或修改后用。
        # 设计原则:
        #   - 只在 FAIL/PARTIAL 时生成,PASS 不生成
        #   - 提取 verify report 中的 FAIL 项摘要(简短,≤ 800 字),不直接 dump 整个报告
        #   - 模板里明确说"基于现有产物 resume",不让模型推倒重来
        if kind == "verify":
            # 解析 VERDICT 行(verify report 第一行必须是判决行,P46-C 格式约束)
            _verify_verdict: str | None = None
            for _line in (report or "").splitlines()[:20]:  # 仅看前 20 行,防恶意长 prompt
                _ls = _line.strip()
                if _ls.startswith("VERDICT:"):
                    _v = _ls[len("VERDICT:"):].strip().upper()
                    # 兼容"VERDICT: PASS | FAIL | PARTIAL" 模板字面 / 多关键词同行
                    for _candidate in ("PASS", "FAIL", "PARTIAL"):
                        if _candidate in _v:
                            _verify_verdict = _candidate
                            break
                    break
            _result["verify_verdict"] = _verify_verdict  # None 表示未能解析(verify 报告不规范)
            # 反推被验证的父 task_id(命名约定 {parent}_verify)
            _verify_target: str | None = None
            if task_id.endswith("_verify") and len(task_id) > len("_verify"):
                _verify_target = task_id[: -len("_verify")]
                _result["verify_target"] = _verify_target

            # P48 修复建议生成
            if _verify_verdict in ("FAIL", "PARTIAL") and _verify_target:
                _fail_excerpt = _extract_verify_fail_excerpt(report or "")
                _result["repair_recommendation"] = (
                    f"## Repair Task (verify verdict: {_verify_verdict})\n"
                    f"Parent task_id `{_verify_target}` reported completion, but an independent verify helper found problems.\n"
                    f"Repair the existing artifacts and preserve correct delivered parts. Do not restart from scratch unless evidence shows the artifact is unusable.\n"
                    f"\n"
                    f"## Key FAIL Items From Verify Report\n"
                    f"{_fail_excerpt}\n"
                    f"\n"
                    f"## Required Work\n"
                    f"Fix each listed FAIL item, then rerun the same round-trip or test scenario. "
                    f"Report completion only after all checks pass. If a FAIL item is a verify false positive, "
                    f"state that with evidence.\n"
                    f"\n"
                    f"Suggested main-process call:\n"
                    f"  delegate(action='spawn', tasks=[{{\n"
                    f"    'task_id': '{_verify_target}',\n"
                    f"    'resume': True,\n"
                    f"    'kind': 'code',\n"
                    f"    'prompt': '<repair task text above>'\n"
                    f"  }}])\n\n"
                    f"验证失败时优先同 task_id 修复已有产物，并复跑相同验收。"
                )

        # 2026-05-10 Patch 55: 已删除原 P45 自动派 verify 块。
        # 主进程现在显式委派 verify(看 codehelper 报告里的 "建议验证" 提示决定)。
        try:
            _env_files_for_provenance = sorted(set(copy_stats.get("env_copied_files") or []))
            if _env_files_for_provenance:
                from app.llm.tools.environment import record_env_helper_outputs
                _outputs_check = _result.get("outputs_check") if isinstance(_result.get("outputs_check"), dict) else {}
                record_env_helper_outputs(
                    main_workspace,
                    task_id=task_id,
                    files=_env_files_for_provenance,
                    ok=bool(_result.get("ok")),
                    terminal_reason=str(_result.get("terminal_reason") or ""),
                    outputs_complete=_outputs_check.get("outputs_complete"),
                    kind=kind,
                    mode=mode,
                )
        except Exception as _prov_e:
            log.debug("environment helper provenance write failed: %r", _prov_e)

        # 2026-05-11 P5.2: SUBAGENT_STOP hook (正常路径)
        try:
            from app.core.hooks import dispatch_hook, HookEvent
            dispatch_hook(HookEvent.SUBAGENT_STOP, {
                "task_id": task_id, "kind": kind,
                "ok": _result.get("ok", False),
                "terminal_reason": _result.get("terminal_reason", "?"),
                "elapsed_sec": _result.get("elapsed_sec"),
                "files_count": len(_result.get("files", [])),
            })
        except Exception as _hook_e:
            log.debug("subagent_stop hook dispatch failed: %r", _hook_e)

        return _result
    except Exception as e:
        log.exception("delegate helper %s failed", task_id)
        _err_msg = f"执行失败：{type(e).__name__}: {e}"
        _elapsed = time.monotonic() - helper_started_at
        # 2026-05-08 优化: 同上,精简 result 字段
        _result = {
            "task_id": task_id,
            "ok": False,
            "report": _err_msg,
            "elapsed_sec": round(_elapsed, 1),
            # 2026-05-11 P2.1: helper 主入口异常也带 terminal_reason
            "terminal_reason": "crashed",
            "crash_type": type(e).__name__,
        }
        if resumed_actually:
            _result["resumed"] = True
        # 2026-05-11 P5.2: SUBAGENT_STOP hook (异常路径)
        try:
            from app.core.hooks import dispatch_hook, HookEvent
            dispatch_hook(HookEvent.SUBAGENT_STOP, {
                "task_id": task_id, "kind": kind,
                "ok": False,
                "terminal_reason": "crashed",
                "crash_type": type(e).__name__,
                "elapsed_sec": round(_elapsed, 1),
            })
        except Exception as _hook_e:
            log.debug("subagent_stop hook (except) dispatch failed: %r", _hook_e)
        return _result
    finally:
        # 通知 long-run observer 退出(无论 helper 怎么结束)
        try:
            long_run_observer_done.set()
            long_run_task.cancel()
            try:
                await long_run_task
            except (asyncio.CancelledError, Exception):
                pass
        except (NameError, UnboundLocalError):
            pass  # observer 可能在异常路径下还没创建
        # 还原 trace_id(虽然 ContextVar 是 task-local,理论上不需要,
        # 但若同一 task 内之后还有逻辑就靠这个 reset 保平安)
        debug.set_trace_id(parent_trace)
        # 还原 owner ContextVar
        try:
            reset_current_owner(owner_token)
        except (LookupError, NameError):
            pass
        # 还原 abort_event ContextVar
        try:
            reset_current_abort_event(abort_token)
        except (LookupError, NameError):
            pass
        # 2026-05-11 P9: 还原 helper_kind ContextVar
        try:
            reset_current_helper_kind(kind_token)
        except (LookupError, NameError):
            pass
        try:
            reset_current_helper_expected_outputs(expected_outputs_token)
        except (LookupError, NameError):
            pass
        try:
            reset_current_helper_write_scopes(write_scopes_token)
        except (LookupError, NameError):
            pass
        # 还原 helper_proc_id ContextVar (心跳汇报机制)
        if helper_proc_id_token is not None:
            try:
                reset_current_helper_proc_id(helper_proc_id_token)
            except (LookupError, NameError):
                pass
        # 还原 thread_ctx ContextVar(2026-05-03:recall_thread 工具用)
        try:
            _h_reset_thread_ctx(_helper_thread_token)
        except (LookupError, NameError):
            pass
        # _helper_fhc 已是 plain dict(非 ContextVar),clear() 在上方已调用,无需 reset
        # ── 取消 abort 桥接任务(Bug #6 修)──
        try:
            bridge_task.cancel()
            try:
                await bridge_task
            except (asyncio.CancelledError, Exception):
                pass
        except (NameError, UnboundLocalError):
            pass  # bridge_task 可能在很早就 raise 之前未创建
