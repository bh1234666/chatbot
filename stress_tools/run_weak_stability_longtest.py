from __future__ import annotations

import argparse
import asyncio
import json
import re
import shutil
import time
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Callable
from zipfile import ZipFile

try:
    from run_agent_longtest_no_group import _workflow_issue_summary
    from run_capability_regression import (
        EnvClient,
        PROJECT_ROOT,
        Recorder,
        response_issues,
        setup_projects,
        static_inventory,
        validate_app_clone,
        validate_engineering,
        validate_greenfield,
    )
    from run_focused_capability_regression import (
        scenario_consistency_issues,
        validate_ielts_extended,
    )
    from run_four_hour_longtest import start_service, stop_service, wait_health
except ModuleNotFoundError:
    from .run_agent_longtest_no_group import _workflow_issue_summary
    from .run_capability_regression import (
        EnvClient,
        PROJECT_ROOT,
        Recorder,
        response_issues,
        setup_projects,
        static_inventory,
        validate_app_clone,
        validate_engineering,
        validate_greenfield,
    )
    from .run_focused_capability_regression import (
        scenario_consistency_issues,
        validate_ielts_extended,
    )
    from .run_four_hour_longtest import start_service, stop_service, wait_health


RUNS_DIR = PROJECT_ROOT / "stress_tools" / "runs" / "weak_stability_longtest"


def _output_files(root: Path, patterns: tuple[str, ...]) -> list[str]:
    matches: set[str] = set()
    for pattern in patterns:
        for path in root.rglob(pattern):
            if path.is_file():
                matches.add(str(path.relative_to(root)).replace("\\", "/"))
    return sorted(matches)[:50]


def validate_ielts_weak(root: Path) -> dict[str, Any]:
    data = validate_ielts_extended(root)
    data["weak_outputs"] = _output_files(
        root,
        (
            "analysis_outputs/*",
            "_env/*ielts*",
            "_env/*雅思*",
            "_env/*evidence*",
            "_env/*report*",
            "_env/*整理*",
            "*.md",
            "*.html",
        ),
    )
    return data


def validate_engineering_weak(root: Path) -> dict[str, Any]:
    data = validate_engineering(root)
    data["weak_outputs"] = _output_files(
        root,
        (
            "analysis_outputs/*",
            "_env/*engineering*",
            "_env/*工程*",
            "_env/*管理*",
            "_env/*report*",
            "_env/*index*",
            "*.md",
            "*.html",
        ),
    )
    return data


def now_id() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S") + "_" + uuid.uuid4().hex[:6]


def _read_docx_text(path: Path) -> dict[str, Any]:
    try:
        from docx import Document

        doc = Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        table_cells: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        table_cells.append(text)
        text = "\n".join(paragraphs + table_cells)
        return {
            "ok": True,
            "method": "python-docx",
            "paragraphs": len(paragraphs),
            "tables": len(doc.tables),
            "chars": len(text),
            "text": text,
        }
    except Exception as exc:
        try:
            with ZipFile(path) as zf:
                parts = []
                for name in zf.namelist():
                    if name.startswith("word/") and name.endswith(".xml"):
                        parts.append(zf.read(name).decode("utf-8", errors="replace"))
                text = re.sub(r"<[^>]+>", "\n", "\n".join(parts))
                return {
                    "ok": True,
                    "method": "zip-xml",
                    "paragraphs": text.count("\n"),
                    "tables": text.count("tbl"),
                    "chars": len(text),
                    "text": text,
                    "python_docx_error": f"{type(exc).__name__}: {exc}",
                }
        except Exception as fallback_exc:
            return {
                "ok": False,
                "method": "failed",
                "error": f"{type(exc).__name__}: {exc}",
                "fallback_error": f"{type(fallback_exc).__name__}: {fallback_exc}",
            }


def validate_db_paper(root: Path) -> dict[str, Any]:
    inv = static_inventory(root)
    docx_files = sorted(
        [p for p in root.rglob("*.docx") if not p.name.startswith("~$")],
        key=lambda p: p.stat().st_size,
        reverse=True,
    )
    docs: list[dict[str, Any]] = []
    for path in docx_files[:5]:
        data = _read_docx_text(path)
        text = str(data.pop("text", ""))
        lower = text.lower()
        terms = {
            "red_black_tree": "红黑树" in text or "red-black" in lower,
            "skip_list": "跳表" in text or "skip list" in lower,
            "btree": "B树" in text or "b-tree" in lower or "B-Tree" in text,
            "bplus_tree": "B+树" in text or "b+ tree" in lower or "B+Tree" in text,
            "new_algorithm": any(
                marker in text
                for marker in ("新算法", "新型", "本文提出", "原创", "发明", "一种新的")
            )
            or "novel" in lower,
            "references_or_limitations": any(marker in text for marker in ("参考文献", "局限", "不足", "限制"))
            or any(marker in lower for marker in ("references", "limitations")),
        }
        leak_markers = [
            marker
            for marker in ("_delegate_", "_env", "trace_id", "tool_call", "workspace.", "helper_progress")
            if marker.lower() in lower
        ]
        docs.append(
            {
                "path": str(path.relative_to(root)).replace("\\", "/"),
                "bytes": path.stat().st_size,
                **data,
                "terms": terms,
                "missing_terms": [k for k, ok in terms.items() if not ok],
                "internal_leak_markers": leak_markers,
                "text_sample": text[:800],
            }
        )
    best = docs[0] if docs else {}
    return {
        "inventory": inv,
        "docx_files": [str(p.relative_to(root)).replace("\\", "/") for p in docx_files[:20]],
        "docx_count": len(docx_files),
        "best_docx": best,
        "has_nontrivial_docx": bool(best) and int(best.get("bytes") or 0) >= 20_000,
        "has_table": bool(best) and int(best.get("tables") or 0) >= 1,
        "has_required_terms": bool(best) and not best.get("missing_terms"),
        "has_internal_leak": bool(best.get("internal_leak_markers")),
    }


def db_paper_issues(text: str, after: dict[str, Any]) -> list[str]:
    issues: list[str] = []
    if not after.get("has_nontrivial_docx"):
        issues.append("db_docx_missing_or_too_small")
    if not after.get("has_table"):
        issues.append("db_docx_missing_tables")
    if not after.get("has_required_terms"):
        missing = (after.get("best_docx") or {}).get("missing_terms") or []
        issues.append("db_docx_missing_terms:" + ",".join(missing))
    if after.get("has_internal_leak"):
        issues.append("db_docx_internal_leak_markers")
    if ".docx" in text.lower() and not after.get("docx_count"):
        issues.append("claimed_docx_but_none_found")
    return issues


WEAK_TASKS: dict[str, list[str]] = {
    "ielts": [
        (
            "查看当前目录下的所有文件，整理雅思备考内容。分听力、阅读、写作、口语四科；"
            "按重要性和类别分类。词汇里独立零散出现的内容优先；作文、口语框架模板靠前。"
            "口语中只有关键词或关键句的答案要扩写为 6-6.5 分范例，作文可补充范文或段落，"
            "需要处做中英对照。先覆盖材料并说明已读、未读、需 OCR 或转写的文件，再输出报告。"
        ),
        (
            "继续上一轮雅思整理。检查是否遗漏文件和四科要求，补齐口语扩写、写作模板、"
            "零散高频词汇、中英对照和证据文件名。把最终结果写入 analysis_outputs/ielts_final_report.md。"
        ),
    ],
    "engineering": [
        (
            "分析当前工程管理目录的所有文件。先做目录与材料类型盘点，再按课程、任务、表格、"
            "文档、风险与待办分类整理成结构化报告。对 Office/PDF/图片等材料分批读取并形成证据，"
            "明确哪些文件已读、哪些无法读取以及理由。"
        ),
        (
            "继续上一轮工程管理材料整理。复核文件覆盖率，补齐课程任务、可复用模板、风险清单、"
            "进度建议和下一步行动，写入 analysis_outputs/engineering_final_report.md。"
        ),
    ],
    "db_paper": [
        (
            "比较红黑树、跳表、B树、B+树等数据库维护相关数据结构算法，并发明一种新的数据结构算法，"
            "为其编写一篇严谨的论文，Word 格式输出。先建立论文框架、算法比较标准、验收清单和文件命名，"
            "再把算法分析、复杂度表格、基准设计、论文写作和 Word 格式验证拆成可验证分片完成。"
        ),
        (
            "继续上一轮数据库算法论文任务。复核 Word 文档是否包含红黑树、跳表、B树、B+树、"
            "新算法定义、复杂度/适用场景表格、局限性与参考依据，修正格式问题并保留一个清晰命名的最终 docx。"
        ),
    ],
    "app": [
        (
            "在这个隔离 app 副本里检查文件读写与 helper 分工。请递归读取真实源码，说明 read/helper、"
            "edit/helper、code/helper 的边界和文件访问路径如何工作，找一个会影响复杂任务稳定性的真实问题，"
            "只在副本中做低风险改进并验证。"
        ),
        (
            "继续上一轮 app 副本维护。复核修改是否与提示词、工具描述和实际行为一致，补一个聚焦的回归测试，"
            "运行最小验证命令并报告结果。"
        ),
    ],
    "greenfield": [
        (
            "从空目录创建一个复杂但小型的本地 agent 工程：Python 核心、JavaScript 前端、一个 C 或 C++ 工具、"
            "测试、文档、fixtures 和 scripts/check_project.py。要求可本地运行且依赖轻。"
        ),
        (
            "继续这个新工程。增加一个跨多文件协作的功能，运行自检，修复发现的问题，并更新架构文档。"
        ),
    ],
}


def build_projects(run_dir: Path) -> dict[str, Path]:
    projects = setup_projects(run_dir)
    db_project = run_dir / "projects" / "db_index_paper_project"
    if db_project.exists():
        shutil.rmtree(db_project)
    db_project.mkdir(parents=True, exist_ok=True)
    projects["db_paper"] = db_project
    return projects


def scenario_plan() -> list[dict[str, Any]]:
    return [
        {
            "label": "ielts",
            "user_id": "weak_ielts",
            "tasks": WEAK_TASKS["ielts"],
            "validate": validate_ielts_weak,
            "must_reference_files": True,
        },
        {
            "label": "engineering",
            "user_id": "weak_engineering",
            "tasks": WEAK_TASKS["engineering"],
            "validate": validate_engineering_weak,
            "must_reference_files": True,
        },
        {
            "label": "db_paper",
            "user_id": "weak_db_paper",
            "tasks": WEAK_TASKS["db_paper"],
            "validate": validate_db_paper,
            "must_reference_files": False,
        },
        {
            "label": "app",
            "user_id": "weak_app",
            "tasks": WEAK_TASKS["app"],
            "validate": validate_app_clone,
            "must_reference_files": True,
        },
        {
            "label": "greenfield",
            "user_id": "weak_greenfield",
            "tasks": WEAK_TASKS["greenfield"],
            "validate": validate_greenfield,
            "must_reference_files": False,
        },
    ]


def consistency_issues(label: str, text: str, after: dict[str, Any]) -> list[str]:
    issues = scenario_consistency_issues(label, text, after)
    if label == "db_paper":
        issues.extend(db_paper_issues(text, after))
    return issues


async def monitor_loop(client: EnvClient, recorder: Recorder, stop: asyncio.Event, interval: float) -> None:
    while not stop.is_set():
        try:
            r = await client.client.get(f"{client.base_url}/v1/environment/active")
            active = r.json() if r.status_code == 200 else {"status_code": r.status_code, "text": r.text[:500]}
        except Exception as exc:
            active = {"error": f"{type(exc).__name__}: {exc}"}
        await recorder.record({"kind": "monitor", "active": active})
        try:
            await asyncio.wait_for(stop.wait(), timeout=interval)
        except asyncio.TimeoutError:
            pass


async def write_report(run_dir: Path, projects: dict[str, Path], *, partial: bool) -> None:
    events_path = run_dir / "events.jsonl"
    events = (
        [json.loads(line) for line in events_path.read_text(encoding="utf-8").splitlines() if line.strip()]
        if events_path.exists()
        else []
    )
    results = [e for e in events if e.get("kind") == "scenario_result"]
    validations = {
        "ielts": validate_ielts_weak(projects["ielts"]),
        "engineering": validate_engineering_weak(projects["engineering"]),
        "db_paper": validate_db_paper(projects["db_paper"]),
        "app": validate_app_clone(projects["app"]),
        "greenfield": validate_greenfield(projects["greenfield"]),
    }
    summary = {
        "run_dir": str(run_dir),
        "partial": partial,
        "events": len(events),
        "scenario_results": len(results),
        "issue_results": sum(1 for e in results if e.get("issues") or not e.get("ok")),
        "projects": {k: str(v) for k, v in projects.items()},
        "final_validation": validations,
    }
    (run_dir / "summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = [
        "# Weak Stability Long Test",
        "",
        f"- Partial: {partial}",
        f"- Events: {summary['events']}",
        f"- Scenario results: {summary['scenario_results']}",
        f"- Results with issues: {summary['issue_results']}",
        "",
        "## Scenario Results",
        "",
    ]
    for e in results[-80:]:
        text = str(e.get("text") or "").replace("\n", " ")
        lines.append(
            f"- `{e.get('label')}` turn={e.get('turn')} ok={e.get('ok')} "
            f"latency={e.get('latency_sec')}s issues={e.get('issues')} text={text[:700]}"
        )
    lines.extend(["", "## Final Validation", "", "```json", json.dumps(validations, ensure_ascii=False, indent=2), "```"])
    (run_dir / "report.md").write_text("\n".join(lines), encoding="utf-8")


async def run(args: argparse.Namespace) -> Path:
    run_dir = RUNS_DIR / ("weak_" + now_id())
    run_dir.mkdir(parents=True, exist_ok=True)
    base_url = f"http://127.0.0.1:{args.port}"
    service = start_service(run_dir, args.port) if args.start_service else None
    client: EnvClient | None = None
    projects: dict[str, Path] = {}
    try:
        await wait_health(base_url, args.health_timeout_sec)
        projects = build_projects(run_dir)
        recorder = Recorder(run_dir)
        client = EnvClient(base_url)
        await recorder.record(
            {
                "kind": "run_start",
                "base_url": base_url,
                "duration_min": args.duration_min,
                "scope": "weak_stability_environment_agent_no_group_chat",
                "projects": {k: str(v) for k, v in projects.items()},
                "initial_inventory": {k: static_inventory(v) for k, v in projects.items()},
            }
        )
        stop = asyncio.Event()
        monitor = asyncio.create_task(monitor_loop(client, recorder, stop, args.monitor_interval_sec))
        plan = scenario_plan()
        if args.scenario:
            selected = list(dict.fromkeys(args.scenario))
            available = {item["label"]: item for item in plan}
            unknown = set(selected).difference(available)
            if unknown:
                raise ValueError(f"unknown scenario(s): {', '.join(sorted(unknown))}")
            plan = [available[label] for label in selected]
        end_at = time.monotonic() + args.duration_min * 60
        report_lock = asyncio.Lock()

        async def flush_report(partial: bool = True) -> None:
            async with report_lock:
                await write_report(run_dir, projects, partial=partial)

        async def run_scenario(item: dict[str, Any]) -> None:
            label = item["label"]
            turns = 0
            in_flight: dict[str, Any] | None = None
            while time.monotonic() < end_at and turns < args.max_turns_per_scenario:
                task_list = item["tasks"]
                turn = turns
                message = (
                    task_list[turn % len(task_list)]
                    + f"\n\nStability cycle: {turn // len(task_list) + 1}. "
                    "Continue from the actual current state. Prefer correct verified results over speed. "
                    "For broad reading or broad writing, delegate by clear slices and verify concrete files before completion."
                )
                root = projects[label]
                before = item["validate"](root)
                await recorder.record(
                    {
                        "kind": "scenario_request",
                        "label": label,
                        "turn": turn,
                        "root": str(root),
                        "before": before,
                        "message": message,
                    }
                )
                in_flight = {"label": label, "turn": turn, "root": str(root), "started": time.monotonic()}
                started = time.monotonic()
                try:
                    result = await asyncio.wait_for(
                        client.ask(
                            user_id=item["user_id"],
                            current_dir=root,
                            message=message,
                            label=label,
                            turn=turn,
                        ),
                        timeout=args.scenario_timeout_sec,
                    )
                    after = item["validate"](root)
                    issues = response_issues(
                        result.get("text") or "",
                        must_reference_files=item["must_reference_files"],
                    )
                    issues.extend(consistency_issues(label, result.get("text") or "", after))
                    issues.extend(_workflow_issue_summary(result))
                except asyncio.TimeoutError:
                    after = item["validate"](root)
                    result = {
                        "ok": False,
                        "status_code": 0,
                        "latency_sec": round(time.monotonic() - started, 3),
                        "text": "",
                        "error": f"timeout after {args.scenario_timeout_sec}s",
                    }
                    issues = ["scenario_timeout"]
                except asyncio.CancelledError:
                    after = item["validate"](root)
                    await recorder.record(
                        {
                            "kind": "scenario_result",
                            "label": label,
                            "turn": turn,
                            "root": str(root),
                            "after": after,
                            "issues": ["scenario_cancelled"],
                            "ok": False,
                            "status_code": 0,
                            "latency_sec": round(time.monotonic() - started, 3),
                            "text": "",
                            "error": "scenario cancelled by runner time limit",
                        }
                    )
                    await flush_report(partial=True)
                    raise
                except Exception as exc:
                    after = item["validate"](root)
                    result = {
                        "ok": False,
                        "status_code": 0,
                        "latency_sec": round(time.monotonic() - started, 3),
                        "text": "",
                        "error": f"{type(exc).__name__}: {exc}",
                    }
                    issues = ["scenario_exception"]
                await recorder.record(
                    {
                        "kind": "scenario_result",
                        "label": label,
                        "turn": turn,
                        "root": str(root),
                        "after": after,
                        "issues": issues,
                        **result,
                    }
                )
                in_flight = None
                turns += 1
                await flush_report(partial=True)
                if time.monotonic() < end_at and turns < args.max_turns_per_scenario:
                    await asyncio.sleep(args.gap_sec)
            if in_flight is not None:
                await recorder.record(
                    {
                        "kind": "scenario_result",
                        "label": in_flight["label"],
                        "turn": in_flight["turn"],
                        "root": in_flight["root"],
                        "after": item["validate"](projects[label]),
                        "issues": ["scenario_cancelled_before_result"],
                        "ok": False,
                        "status_code": 0,
                        "latency_sec": round(time.monotonic() - in_flight["started"], 3),
                        "text": "",
                        "error": "scenario task ended before a result was recorded",
                    }
                )

        try:
            scenario_tasks = [asyncio.create_task(run_scenario(item)) for item in plan]
            try:
                await asyncio.wait_for(
                    asyncio.gather(*scenario_tasks, return_exceptions=True),
                    timeout=args.duration_min * 60 + args.drain_sec,
                )
            except asyncio.TimeoutError:
                await recorder.record({"kind": "run_timeout"})
                for task in scenario_tasks:
                    task.cancel()
                await asyncio.gather(*scenario_tasks, return_exceptions=True)
        finally:
            stop.set()
            await monitor
            await flush_report(partial=False)
    finally:
        if client is not None:
            await client.close()
        if args.stop_service:
            stop_service(service)
    return run_dir


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--duration-min", type=float, default=120.0)
    parser.add_argument("--port", type=int, default=8076)
    parser.add_argument("--start-service", action=argparse.BooleanOptionalAction, default=True)
    parser.add_argument("--stop-service", action=argparse.BooleanOptionalAction, default=True)
    parser.add_argument("--health-timeout-sec", type=float, default=120.0)
    parser.add_argument("--scenario-timeout-sec", type=float, default=2400.0)
    parser.add_argument("--drain-sec", type=float, default=1800.0)
    parser.add_argument(
        "--scenario",
        action="append",
        choices=["ielts", "engineering", "db_paper", "app", "greenfield"],
    )
    parser.add_argument("--max-turns-per-scenario", type=int, default=2)
    parser.add_argument("--gap-sec", type=float, default=5.0)
    parser.add_argument("--monitor-interval-sec", type=float, default=20.0)
    return parser.parse_args()


if __name__ == "__main__":
    print(asyncio.run(run(parse_args())))
